"""
智能文档解析器 (Document Parser)
==============================

支持 PDF、Word、图片等格式的智能解析，提取环评关键信息。
"""

import asyncio
import json
import os
import re
import hashlib
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from typing import Any, Callable, Optional


class DocumentType(Enum):
    """文档类型"""
    FEASIBILITY = "feasibility"  # 可行性研究报告
    LAYOUT = "layout"            # 总平面布置图
    EQUIPMENT = "equipment"      # 设备清单
    ENVIRONMENTAL = "environmental"  # 环境影响报告
    OTHER = "other"


@dataclass
class ParsedDocument:
    """解析后的文档"""
    document_id: str
    doc_type: DocumentType
    file_path: str
    file_name: str
    page_count: int = 0
    extracted_data: "ExtractedData" = None
    raw_text: str = ""
    tables: list = field(default_factory=list)
    images: list = field(default_factory=list)
    parsed_at: datetime = field(default_factory=datetime.now)
    metadata: dict = field(default_factory=dict)


@dataclass
class ExtractedData:
    """提取的数据"""
    # 项目基本信息
    project_name: str = ""
    location: str = ""
    industry: str = ""
    scale: str = ""
    investment: str = ""

    # 工艺信息
    processes: list = field(default_factory=list)  # 工艺单元列表
    materials: list = field(default_factory=list)  # 原辅材料列表

    # 设备信息
    equipment_list: list = field(default_factory=list)  # 设备清单

    # 建筑信息
    buildings: list = field(default_factory=list)  # 建筑物列表

    # 污染源信息
    pollution_sources: list = field(default_factory=list)  # 污染源

    # 公用工程
    utilities: dict = field(default_factory=dict)  # 水、电、汽等


class DocumentParser:
    """
    智能文档解析器

    用法:
        parser = DocumentParser()

        # 解析文档
        result = await parser.parse("可研报告.pdf", "feasibility")

        # 提取的数据
        print(result.extracted_data.project_name)
        print(result.extracted_data.processes)
    """

    def __init__(self, data_dir: str = "./data/eia"):
        self.data_dir = data_dir
        self._extractors: dict[DocumentType, Callable] = {}
        self._ai_handler: Optional[Callable] = None

        self._register_default_extractors()

    def register_ai_handler(self, handler: Callable) -> None:
        """注册 AI 处理器"""
        self._ai_handler = handler

    def _register_default_extractors(self) -> None:
        """注册默认提取器"""
        self._extractors[DocumentType.FEASIBILITY] = self._extract_feasibility
        self._extractors[DocumentType.LAYOUT] = self._extract_layout
        self._extractors[DocumentType.EQUIPMENT] = self._extract_equipment

    async def parse(self, file_path: str, doc_type: str = "auto") -> ParsedDocument:
        """
        解析文档

        Args:
            file_path: 文件路径
            doc_type: 文档类型 (feasibility/layout/equipment/auto)

        Returns:
            ParsedDocument: 解析结果
        """
        # 自动检测文档类型
        if doc_type == "auto":
            doc_type = self._detect_doc_type(file_path)

        doc_type_enum = DocumentType(doc_type)

        # 生成文档 ID
        doc_id = hashlib.sha256(f"{file_path}{datetime.now().isoformat()}".encode()).hexdigest()[:12]

        parsed = ParsedDocument(
            document_id=doc_id,
            doc_type=doc_type_enum,
            file_path=file_path,
            file_name=os.path.basename(file_path)
        )

        try:
            # 读取文件内容
            raw_content = await self._read_file(file_path)
            parsed.raw_text = raw_content
            parsed.page_count = self._estimate_page_count(raw_content)

            # 提取数据
            extractor = self._extractors.get(doc_type_enum, self._extract_generic)
            parsed.extracted_data = await extractor(raw_content, file_path)

            # 如果有 AI 处理器，用它增强提取
            if self._ai_handler:
                parsed.extracted_data = await self._ai_handler_enhance(
                    parsed.extracted_data,
                    doc_type_enum
                )

        except Exception as e:
            parsed.metadata["error"] = str(e)

        return parsed

    async def _read_file(self, file_path: str) -> str:
        """读取文件内容"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        elif ext == ".pdf":
            # 简化处理：实际应使用 PDF 解析库
            return await self._extract_pdf_text(file_path)

        elif ext in [".docx", ".doc"]:
            # 简化处理：实际应使用 python-docx
            return await self._extract_docx_text(file_path)

        elif ext in [".jpg", ".jpeg", ".png", ".bmp"]:
            # 图片：OCR 或 AI 识别
            return await self._extract_image_text(file_path)

        else:
            raise ValueError(f"Unsupported file type: {ext}")

    async def _extract_pdf_text(self, file_path: str) -> str:
        """提取 PDF 文本（简化版）"""
        # 实际应使用 pypdf 或 pdfplumber
        try:
            import pypdf
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
            ...
            return "[PDF解析需要 pypdf 库]"

    async def _extract_docx_text(self, file_path: str) -> str:
        """提取 Word 文本（简化版）"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return text
        except ImportError:
            return "[Word解析需要 python-docx库]"

    async def _extract_image_text(self, file_path: str) -> str:
        """从图片提取文本（需要 OCR）"""
        # 简化处理
        return "[图片需要 OCR 识别]"

    def _estimate_page_count(self, text: str) -> int:
        """估算页数"""
        # 假设每页约 3000 字符
        return max(1, len(text) // 3000)

    def _detect_doc_type(self, file_path: str) -> str:
        """自动检测文档类型"""
        file_name = os.path.basename(file_path).lower()

        if any(kw in file_name for kw in ["可研", "可行性", "project", "feasibility"]):
            return "feasibility"
        elif any(kw in file_name for kw in ["平面", "布置", "layout", "plan"]):
            return "layout"
        elif any(kw in file_name for kw in ["设备", "equipment", "清单"]):
            return "equipment"
        elif any(kw in file_name for kw in ["环评", "环境", "报告", "environmental"]):
            return "environmental"
        else:
            return "other"

    async def _extract_feasibility(self, text: str, file_path: str) -> ExtractedData:
        """提取可研报告关键信息"""
        data = ExtractedData()

        # 项目名称
        name_patterns = [
            r"项目名称[：:]\s*([^\n]+)",
            r"项目名称\s+([^\n]+)",
            r"^#\s*(.+?)$",  # Markdown标题
        ]
        for pattern in name_patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                data.project_name = match.group(1).strip()
                break

        # 地点
        location_patterns = [
            r"建设地点[：:]\s*([^\n]+)",
            r"项目地点[：:]\s*([^\n]+)",
            r"厂址[：:]\s*([^\n]+)",
        ]
        for pattern in location_patterns:
            match = re.search(pattern, text)
            if match:
                data.location = match.group(1).strip()
                break

        # 行业
        industry_patterns = [
            r"行业类别[：:]\s*([^\n]+)",
            r"所属行业[：:]\s*([^\n]+)",
        ]
        for pattern in industry_patterns:
            match = re.search(pattern, text)
            if match:
                data.industry = match.group(1).strip()
                break

        # 规模
        scale_patterns = [
            r"建设规模[：:]\s*([^\n]+)",
            r"生产规模[：:]\s*([^\n]+)",
            r"规模[：:]\s*([^\n]+)",
        ]
        for pattern in scale_patterns:
            match = re.search(pattern, text)
            if match:
                data.scale = match.group(1).strip()
                break

        # 工艺单元
        data.processes = self._extract_processes(text)

        # 原辅材料
        data.materials = self._extract_materials(text)

        return data

    def _extract_processes(self, text: str) -> list[dict]:
        """提取工艺单元"""
        processes = []

        # 工艺关键词
        process_keywords = [
            "破碎", "研磨", "筛分", "混合", "搅拌", "反应", "蒸馏",
            "萃取", "结晶", "干燥", "包装", "储存", "输送"
        ]

        # 简化：查找包含工艺关键词的句子
        lines = text.split("\n")
        for line in lines:
            for keyword in process_keywords:
                if keyword in line:
                    # 提取工艺名称
                    name = line.strip()[:50]  # 取前50字符
                    processes.append({
                        "name": name,
                        "type": keyword,
                        "description": line.strip()
                    })
                    break

        return processes[:20]  # 限制数量

    def _extract_materials(self, text: str) -> list[dict]:
        """提取原辅材料"""
        materials = []

        # 材料关键词
        material_pattern = r"([\u4e00-\u9fa5a-zA-Z0-9]+)\s*[\（(]?\s*([0-9.]+\s*[吨万千克个])\s*[\）)]?"

        for match in re.finditer(material_pattern, text):
            name = match.group(1)
            amount = match.group(2)

            # 过滤非材料词
            exclude_words = ["项目", "工程", "设备", "工艺", "建设", "生产"]
            if name and not any(ex in name for ex in exclude_words):
                materials.append({
                    "name": name,
                    "amount": amount,
                    "unit": "吨/年" if "吨" in amount else "kg/年"
                })

        return materials[:30]

    async def _extract_layout(self, text: str, file_path: str) -> ExtractedData:
        """提取总平面布置图信息"""
        data = ExtractedData()

        # 建筑物
        building_keywords = ["车间", "厂房", "仓库", "办公楼", "宿舍", "配电", "锅炉", "罐区", "池"]

        lines = text.split("\n")
        for line in lines:
            for keyword in building_keywords:
                if keyword in line:
                    # 提取尺寸
                    size_match = re.search(r"(\d+(?:\.\d+)?)\s*[m米×x]\s*(\d+(?:\.\d+)?)", line)
                    size = f"{size_match.group(1)}m × {size_match.group(2)}m" if size_match else "未知"

                    data.buildings.append({
                        "name": line.strip()[:30],
                        "type": keyword,
                        "size": size
                    })
                    break

        return data

    async def _extract_equipment(self, text: str, file_path: str) -> ExtractedData:
        """提取设备清单"""
        data = ExtractedData()

        # 设备列表
        equipment_list = []

        # 常见设备关键词
        equipment_keywords = [
            "破碎机", "球磨机", "筛分机", "输送机", "风机", "泵",
            "反应釜", "蒸馏塔", "萃取器", "干燥机", "蒸发器", "换热器",
            "储罐", "计量罐", "混合机", "搅拌器", "除尘器", "过滤器"
        ]

        lines = text.split("\n")
        for line in lines:
            for keyword in equipment_keywords:
                if keyword in line:
                    # 提取数量
                    qty_match = re.search(r"(\d+)\s*[台套个部]", line)
                    qty = qty_match.group(1) if qty_match else "1"

                    equipment_list.append({
                        "name": keyword,
                        "quantity": int(qty) if qty.isdigit() else 1,
                        "description": line.strip()[:50]
                    })
                    break

        data.equipment_list = equipment_list
        return data

    async def _extract_generic(self, text: str, file_path: str) -> ExtractedData:
        """通用提取"""
        data = ExtractedData()
        data.project_name = os.path.basename(file_path)
        return data

    async def _ai_handler_enhance(self, data: ExtractedData, doc_type: DocumentType) -> ExtractedData:
        """使用 AI 增强提取"""
        if not self._ai_handler:
            return data

        # 使用 AI 进一步提取
        enhanced = await self._ai_handler({
            "type": "document_extraction",
            "doc_type": doc_type.value,
            "current_data": data.__dict__
        })

        # 更新数据
        for key, value in enhanced.items():
            if hasattr(data, key):
                setattr(data, key, value)

        return data

    def extract_tables(self, text: str) -> list[list]:
        """提取表格"""
        tables = []

        # 简单的表格识别（制表符或对齐的空格）
        lines = text.split("\n")
        current_table = []
        in_table = False

        for line in lines:
            # 检测表格行（包含多个分隔符）
            if "\t" in line or "│" in line or "|" in line:
                in_table = True
                # 分割单元格
                if "\t" in line:
                    cells = [c.strip() for c in line.split("\t")]
                else:
                    cells = [c.strip() for c in line.split("|")]
                current_table.append(cells)
            elif in_table and current_table:
                tables.append(current_table)
                current_table = []
                in_table = False

        if current_table:
            tables.append(current_table)

        return tables


def create_document_parser(data_dir: str = "./data/eia") -> DocumentParser:
    """创建文档解析器实例"""
    return DocumentParser(data_dir=data_dir)