"""
Document Parser - 文档解析器
============================

支持多种文档格式的解析，提取文本、表格、图像等结构化内容。
"""

import re
import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Tuple
from enum import Enum
from pathlib import Path


class BlockType(Enum):
    """文档块类型"""
    TEXT = "text"           # 普通文本段落
    HEADING = "heading"     # 标题
    TABLE = "table"         # 表格
    IMAGE = "image"         # 图像
    CODE = "code"           # 代码块
    LIST = "list"           # 列表
    FOOTNOTE = "footnote"   # 脚注
    PAGE_BREAK = "page_break"  # 分页符


@dataclass
class TextBlock:
    """文本块"""
    text: str
    block_type: BlockType = BlockType.TEXT
    page_num: int = 1
    bbox: Optional[Tuple[float, float, float, float]] = None  # 边界框 (x0, y0, x1, y1)
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    is_bold: bool = False
    is_italic: bool = False
    language: Optional[str] = None  # 检测到的语言
    confidence: float = 1.0  # 置信度

    def __post_init__(self):
        if self.bbox and len(self.bbox) != 4:
            self.bbox = None


@dataclass
class TableBlock:
    """表格块"""
    rows: List[List[str]]  # 表格数据
    headers: Optional[List[str]] = None
    page_num: int = 1
    bbox: Optional[Tuple[float, float, float, float]] = None
    style: Optional[str] = None  # grid, bordered, etc.

    def to_markdown(self) -> str:
        """转换为 Markdown 表格格式"""
        if not self.rows:
            return ""

        lines = []
        if self.headers:
            lines.append("| " + " | ".join(self.headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(self.headers)) + " |")

        for row in self.rows:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")

        return "\n".join(lines)


@dataclass
class ParsedDocument:
    """解析后的文档"""
    file_path: str
    file_name: str
    total_pages: int = 1
    text_blocks: List[TextBlock] = field(default_factory=list)
    tables: List[TableBlock] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    raw_text: str = ""  # 纯文本内容

    @property
    def total_words(self) -> int:
        """总词数"""
        return len(self.raw_text.split())

    @property
    def total_chars(self) -> int:
        """总字符数"""
        return len(self.raw_text)

    def get_text_by_page(self, page_num: int) -> List[TextBlock]:
        """获取指定页的文本块"""
        return [b for b in self.text_blocks if b.page_num == page_num]

    def to_plain_text(self) -> str:
        """转换为纯文本"""
        lines = []
        for block in self.text_blocks:
            if block.block_type == BlockType.PAGE_BREAK:
                lines.append("\n--- Page Break ---\n")
            elif block.block_type == BlockType.TABLE:
                # 查找对应表格
                for table in self.tables:
                    if abs(table.page_num - block.page_num) < 1:
                        lines.append(table.to_markdown())
            else:
                lines.append(block.text)
        return "\n".join(lines)


class BaseParser(ABC):
    """解析器基类"""

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """解析文档"""
        pass

    @abstractmethod
    def is_supported(self, file_path: str) -> bool:
        """是否支持该文件"""
        pass


class TextParser(BaseParser):
    """纯文本解析器 (TXT)"""

    def is_supported(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.txt', '.text', '.log']

    def parse(self, file_path: str) -> ParsedDocument:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # 移除 BOM
        if content.startswith('\ufeff'):
            content = content[1:]

        blocks = []
        for i, line in enumerate(content.split('\n')):
            if line.strip():
                block_type = BlockType.HEADING if self._is_heading(line) else BlockType.TEXT
                blocks.append(TextBlock(
                    text=line,
                    block_type=block_type,
                    page_num=i // 50 + 1  # 估计页码
                ))

        file_name = Path(file_path).name
        return ParsedDocument(
            file_path=file_path,
            file_name=file_name,
            text_blocks=blocks,
            raw_text=content,
            metadata={'format': 'txt'}
        )

    def _is_heading(self, line: str) -> bool:
        """判断是否为标题"""
        # 以 # 开头或者是全大写且较短
        if line.strip().startswith('#'):
            return True
        stripped = line.strip()
        if 3 < len(stripped) < 100 and stripped.isupper():
            return True
        return False


class MarkdownParser(BaseParser):
    """Markdown 解析器"""

    def is_supported(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.md', '.markdown', '.mdown']

    def parse(self, file_path: str) -> ParsedDocument:
        """解析 Markdown 文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        blocks = []
        lines = content.split('\n')
        page_num = 1

        for line in lines:
            if line.strip() == '':
                continue

            block_type = self._detect_block_type(line)
            blocks.append(TextBlock(
                text=line,
                block_type=block_type,
                page_num=page_num
            ))

            # 估计页码 (约50行一页)
            if len(blocks) % 50 == 0:
                page_num += 1

        file_name = Path(file_path).name
        return ParsedDocument(
            file_path=file_path,
            file_name=file_name,
            text_blocks=blocks,
            raw_text=content,
            metadata={'format': 'markdown'}
        )

    def _detect_block_type(self, line: str) -> BlockType:
        """检测块类型"""
        stripped = line.strip()
        if stripped.startswith('#'):
            return BlockType.HEADING
        elif stripped.startswith('```'):
            return BlockType.CODE
        elif stripped.startswith('|'):
            return BlockType.TABLE
        elif stripped.startswith('- ') or stripped.startswith('* '):
            return BlockType.LIST
        return BlockType.TEXT


class PDFParser(BaseParser):
    """PDF 解析器 (可选 PyMuPDF)"""

    def __init__(self):
        self._pymupdf = None
        self._pdfminer = None

    def _try_import(self):
        """尝试导入依赖"""
        if self._pymupdf is None:
            try:
                import fitz  # PyMuPDF
                self._pymupdf = fitz
            except ImportError:
                pass

    def is_supported(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext == '.pdf'

    def parse(self, file_path: str) -> ParsedDocument:
        """解析 PDF 文件"""
        self._try_import()

        if self._pymupdf is None:
            # 如果没有 PyMuPDF，返回基础解析
            return self._fallback_parse(file_path)

        doc = self._pymupdf.open(file_path)
        blocks = []
        tables = []
        total_pages = len(doc)

        for page_num, page in enumerate(doc, 1):
            # 获取文本块
            text_dict = page.get_text("dict")

            for block in text_dict.get("blocks", []):
                if block.get("type") == 0:  # 文本块
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text:
                                blocks.append(TextBlock(
                                    text=text,
                                    block_type=self._detect_text_type(span),
                                    page_num=page_num,
                                    bbox=tuple(block.get("bbox", [])),
                                    font_name=span.get("font"),
                                    font_size=span.get("size"),
                                    is_bold=span.get("bold", False),
                                    is_italic=span.get("italic", False)
                                ))
                elif block.get("type") == 1:  # 图像块
                    blocks.append(TextBlock(
                        text="[IMAGE]",
                        block_type=BlockType.IMAGE,
                        page_num=page_num,
                        bbox=tuple(block.get("bbox", [])) if block.get("bbox") else None
                    ))

            # 检测表格 (简化实现)
            tables.extend(self._extract_tables(page, page_num))

        doc.close()

        # 合并纯文本
        raw_text = "\n".join(b.text for b in blocks if b.block_type != BlockType.IMAGE)

        file_name = Path(file_path).name
        return ParsedDocument(
            file_path=file_path,
            file_name=file_name,
            total_pages=total_pages,
            text_blocks=blocks,
            tables=tables,
            raw_text=raw_text,
            metadata={'format': 'pdf', 'engine': 'pymupdf'}
        )

    def _detect_text_type(self, span: Dict) -> BlockType:
        """检测文本类型"""
        text = span.get("text", "")
        size = span.get("size", 12)
        font = span.get("font", "").lower()

        if "heading" in font or "title" in font or size > 16:
            return BlockType.HEADING
        if "code" in font or text.startswith("```"):
            return BlockType.CODE
        if text.startswith(("•", "- ", "* ", "1.", "2.", "3.")):
            return BlockType.LIST
        return BlockType.TEXT

    def _extract_tables(self, page, page_num: int) -> List[TableBlock]:
        """提取表格 (简化实现)"""
        # 实际应使用表格检测算法，这里返回空
        return []

    def _fallback_parse(self, file_path: str) -> ParsedDocument:
        """回退解析方案"""
        return TextParser().parse(file_path)


class DOCXParser(BaseParser):
    """Word 文档解析器 (可选 python-docx)"""

    def __init__(self):
        self._docx = None

    def _try_import(self):
        if self._docx is None:
            try:
                import docx
                self._docx = docx
            except ImportError:
                pass

    def is_supported(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.docx', '.doc']

    def parse(self, file_path: str) -> ParsedDocument:
        """解析 Word 文档"""
        self._try_import()

        if self._docx is None:
            return self._fallback_parse(file_path)

        doc = self._docx.Document(file_path)
        blocks = []
        page_num = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检测样式
            style_name = para.style.name if para.style else "Normal"
            block_type = self._detect_style_type(style_name, text)

            blocks.append(TextBlock(
                text=text,
                block_type=block_type,
                page_num=page_num,
                is_bold=para.runs[0].bold if para.runs else False,
                is_italic=para.runs[0].italic if para.runs else False
            ))

        # 提取表格
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(TableBlock(rows=table_data))

        raw_text = "\n".join(b.text for b in blocks)

        file_name = Path(file_path).name
        return ParsedDocument(
            file_path=file_path,
            file_name=file_name,
            text_blocks=blocks,
            tables=tables,
            raw_text=raw_text,
            metadata={'format': 'docx', 'engine': 'python-docx'}
        )

    def _detect_style_type(self, style_name: str, text: str) -> BlockType:
        """检测样式类型"""
        name_lower = style_name.lower()
        if "heading" in name_lower or "title" in name_lower:
            return BlockType.HEADING
        if "code" in name_lower:
            return BlockType.CODE
        if "list" in name_lower:
            return BlockType.LIST
        return BlockType.TEXT

    def _fallback_parse(self, file_path: str) -> ParsedDocument:
        """回退解析"""
        return TextParser().parse(file_path)


class DocumentParser:
    """文档解析器主类"""

    def __init__(self):
        self.parsers: List[BaseParser] = [
            PDFParser(),
            DOCXParser(),
            MarkdownParser(),
            TextParser(),
        ]

    def parse(self, file_path: str) -> ParsedDocument:
        """解析文档"""
        file_path = str(Path(file_path).resolve())

        for parser in self.parsers:
            if parser.is_supported(file_path):
                return parser.parse(file_path)

        # 默认使用纯文本解析
        return TextParser().parse(file_path)

    def supported_formats(self) -> List[str]:
        """支持的文件格式"""
        return ['.pdf', '.docx', '.doc', '.md', '.markdown', '.txt', '.text']

    @classmethod
    def get_info(cls, file_path: str) -> Dict[str, Any]:
        """获取文档基本信息（不解析内容）"""
        path = Path(file_path)
        return {
            'file_name': path.name,
            'file_size': path.stat().st_size if path.exists() else 0,
            'extension': path.suffix.lower(),
            'exists': path.exists()
        }
