# -*- coding: utf-8 -*-
"""
文档/网址自动分析入库 - Document Analyzer
==========================================

功能：
1. 自动识别网址或文档内容
2. 网页内容抓取和解析
3. 数据去重和清洗
4. 自动入库到知识库

复用模块：
- KnowledgeBaseVectorStore (知识存储)
- WritingDataCollector (数据采集)
- SmartWritingEvolutionEngine (自进化)

Author: Hermes Desktop Team
"""

import logging
import re
import hashlib
import json
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from urllib.parse import urlparse, urljoin
import asyncio

logger = logging.getLogger(__name__)


class DocumentSource(Enum):
    """文档来源"""
    URL = "url"
    FILE = "file"
    CLIPBOARD = "clipboard"
    USER_INPUT = "user_input"


class ContentType(Enum):
    """内容类型"""
    HTML = "html"
    PDF = "pdf"
    DOC = "doc"
    MARKDOWN = "markdown"
    TEXT = "text"
    UNKNOWN = "unknown"


@dataclass
class AnalyzedDocument:
    """分析后的文档"""
    source: DocumentSource
    original_url: Optional[str] = None
    original_content: str = ""
    cleaned_content: str = ""
    title: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    entities: Dict[str, Any] = field(default_factory=dict)
    sections: List[Dict] = field(default_factory=list)
    content_hash: str = ""
    is_duplicate: bool = False
    quality_score: float = 0.0
    
    def to_dict(self) -> Dict:
        return {
            "source": self.source.value,
            "original_url": self.original_url,
            "title": self.title,
            "cleaned_content": self.cleaned_content[:500] + "..." if len(self.cleaned_content) > 500 else self.cleaned_content,
            "metadata": self.metadata,
            "entities": self.entities,
            "sections_count": len(self.sections),
            "content_hash": self.content_hash,
            "is_duplicate": self.is_duplicate,
            "quality_score": self.quality_score,
        }


@dataclass
class DataCleaner:
    """数据清洗配置"""
    remove_html_tags: bool = True
    remove_urls: bool = False
    remove_emails: bool = True
    remove_phones: bool = True
    normalize_whitespace: bool = True
    remove_special_chars: bool = False
    min_content_length: int = 50  # 最小内容长度
    max_content_length: int = 500000  # 最大内容长度（字符）


class DocumentAnalyzer:
    """
    文档/网址自动分析器
    
    使用示例：
    ```python
    analyzer = DocumentAnalyzer()
    
    # 分析网址
    result = await analyzer.analyze_url("https://example.com/report.pdf")
    
    # 分析本地文件
    result = analyzer.analyze_file("/path/to/document.pdf")
    
    # 分析文本内容
    result = analyzer.analyze_content("这是一段文本内容...")
    
    # 批量分析
    results = await analyzer.analyze_batch(["url1", "url2", "file1.pdf"])
    ```
    """
    
    def __init__(self, cleaner_config: Optional[DataCleaner] = None):
        self.cleaner_config = cleaner_config or DataCleaner()
        self._kb = None
        self._evolution = None
        self._url_cache: Dict[str, str] = {}  # URL -> content hash
        self._content_cache: Dict[str, str] = {}  # hash -> content
        
    @property
    def knowledge_base(self):
        """延迟加载知识库"""
        if self._kb is None:
            try:
                from core.knowledge_vector_db import KnowledgeBaseVectorStore
                self._kb = KnowledgeBaseVectorStore()
            except ImportError:
                logger.warning("KnowledgeBaseVectorStore 未安装")
        return self._kb
    
    @property
    def evolution_engine(self):
        """延迟加载进化引擎"""
        if self._evolution is None:
            try:
                from core.smart_writing.self_evolution import get_evolution_engine
                self._evolution = get_evolution_engine()
            except ImportError:
                logger.warning("EvolutionEngine 未安装")
        return self._evolution
    
    async def analyze(self, input_data: str) -> AnalyzedDocument:
        """
        智能分析入口
        
        Args:
            input_data: 网址、文件路径或文本内容
        
        Returns:
            AnalyzedDocument: 分析结果
        """
        # 自动识别类型
        input_type = self._identify_input_type(input_data)
        
        if input_type == DocumentSource.URL:
            return await self.analyze_url(input_data)
        elif input_type == DocumentSource.FILE:
            return self.analyze_file(input_data)
        else:
            return self.analyze_content(input_data)
    
    def _identify_input_type(self, input_data: str) -> DocumentSource:
        """识别输入类型"""
        input_data = input_data.strip()
        
        # URL识别
        if input_data.startswith(("http://", "https://", "ftp://")):
            return DocumentSource.URL
        
        # 文件路径识别
        if input_data.startswith(("C:\\", "/", "./", "../")) or \
           input_data.endswith((".pdf", ".doc", ".docx", ".txt", ".md", ".html")):
            return DocumentSource.FILE
        
        return DocumentSource.USER_INPUT
    
    async def analyze_url(self, url: str) -> AnalyzedDocument:
        """
        分析网址
        
        Args:
            url: 网址
        
        Returns:
            AnalyzedDocument: 分析结果
        """
        logger.info(f"开始分析URL: {url}")
        
        doc = AnalyzedDocument(source=DocumentSource.URL, original_url=url)
        
        try:
            # 1. 抓取内容
            content, content_type = await self._fetch_url(url)
            doc.original_content = content
            
            # 2. 确定内容类型
            detected_type = self._detect_content_type(url, content, content_type)
            doc.metadata["content_type"] = detected_type.value
            
            # 3. 提取标题
            doc.title = self._extract_title(content, detected_type)
            doc.metadata["url_title"] = doc.title
            
            # 4. 清洗数据
            doc.cleaned_content = self.clean_data(content, detected_type)
            
            # 5. 生成哈希
            doc.content_hash = self._generate_hash(doc.cleaned_content)
            
            # 6. 检查重复
            doc.is_duplicate = self._check_duplicate(doc.content_hash)
            
            # 7. 提取实体
            doc.entities = self._extract_entities(doc.cleaned_content)
            
            # 8. 提取章节
            doc.sections = self._extract_sections(doc.cleaned_content)
            
            # 9. 质量评分
            doc.quality_score = self._score_quality(doc)
            
            logger.info(f"URL分析完成: {url}, 质量={doc.quality_score:.2f}, 重复={doc.is_duplicate}")
            
        except Exception as e:
            logger.error(f"URL分析失败: {url}, {e}")
            doc.metadata["error"] = str(e)
            
        return doc
    
    async def _fetch_url(self, url: str) -> Tuple[str, str]:
        """抓取URL内容"""
        try:
            import httpx
            from bs4 import BeautifulSoup
            
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            }
            
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url, headers=headers)
                response.raise_for_status()
                
                content_type = response.headers.get("Content-Type", "")
                
                # HTML处理
                if "text/html" in content_type:
                    soup = BeautifulSoup(response.text, "html.parser")
                    
                    # 移除脚本和样式
                    for tag in soup(["script", "style", "nav", "footer", "header"]):
                        tag.decompose()
                    
                    # 提取正文
                    main_content = soup.find("main") or soup.find("article") or soup.find("body")
                    text = main_content.get_text(separator="\n", strip=True) if main_content else soup.get_text(separator="\n", strip=True)
                    
                    return text, content_type
                
                # PDF/其他直接返回
                return response.text[:10000], content_type
                
        except ImportError:
            # 降级：使用基础urllib
            import urllib.request
            with urllib.request.urlopen(url, timeout=30) as response:
                return response.read().decode("utf-8", errors="ignore"), ""
        except Exception as e:
            logger.error(f"URL抓取失败: {e}")
            raise
    
    def _detect_content_type(self, url: str, content: str, content_type: str) -> ContentType:
        """检测内容类型"""
        if ".pdf" in url.lower():
            return ContentType.PDF
        elif ".doc" in url.lower():
            return ContentType.DOC
        elif ".md" in url.lower():
            return ContentType.MARKDOWN
        elif "<html" in content.lower() or "<!doctype" in content.lower():
            return ContentType.HTML
        else:
            return ContentType.TEXT
    
    def _extract_title(self, content: str, content_type: ContentType) -> str:
        """提取标题"""
        if content_type == ContentType.HTML:
            # 从HTML中提取标题
            title_match = re.search(r"<title[^>]*>([^<]+)</title>", content, re.IGNORECASE)
            if title_match:
                return title_match.group(1).strip()
        
        # 从文本中提取第一行作为标题
        lines = content.split("\n")
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 5 and len(line) < 100:
                return line
                
        return "未命名文档"
    
    def analyze_file(self, file_path: str) -> AnalyzedDocument:
        """
        分析本地文件
        
        Args:
            file_path: 文件路径
        
        Returns:
            AnalyzedDocument: 分析结果
        """
        logger.info(f"开始分析文件: {file_path}")
        
        doc = AnalyzedDocument(source=DocumentSource.FILE, original_url=file_path)
        
        try:
            import os
            
            if not os.path.exists(file_path):
                doc.metadata["error"] = "文件不存在"
                return doc
            
            # 根据扩展名处理
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == ".pdf":
                doc = self._analyze_pdf(file_path, doc)
            elif ext in [".doc", ".docx"]:
                doc = self._analyze_doc(file_path, doc)
            elif ext == ".md":
                doc = self._analyze_markdown(file_path, doc)
            elif ext == ".txt":
                doc = self._analyze_text_file(file_path, doc)
            else:
                doc.metadata["error"] = f"不支持的文件类型: {ext}"
                
        except Exception as e:
            logger.error(f"文件分析失败: {file_path}, {e}")
            doc.metadata["error"] = str(e)
            
        return doc
    
    def _analyze_pdf(self, file_path: str, doc: AnalyzedDocument) -> AnalyzedDocument:
        """分析PDF文件"""
        try:
            import PyPDF2
            
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                doc.metadata["page_count"] = len(reader.pages)
                
                # 提取文本
                text_parts = []
                for page in reader.pages[:50]:  # 限制页数
                    text_parts.append(page.extract_text())
                    
                doc.original_content = "\n".join(text_parts)
                doc.title = os.path.basename(file_path).replace(".pdf", "")
                
        except ImportError:
            # 降级：简单读取
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                doc.original_content = f.read()[:10000]
                
        doc.cleaned_content = self.clean_data(doc.original_content, ContentType.PDF)
        doc.content_hash = self._generate_hash(doc.cleaned_content)
        doc.is_duplicate = self._check_duplicate(doc.content_hash)
        doc.entities = self._extract_entities(doc.cleaned_content)
        doc.sections = self._extract_sections(doc.cleaned_content)
        doc.quality_score = self._score_quality(doc)
        
        return doc
    
    def _analyze_doc(self, file_path: str, doc: AnalyzedDocument) -> AnalyzedDocument:
        """分析Word文档"""
        try:
            from docx import Document
            
            doc_obj = Document(file_path)
            paragraphs = [p.text for p in doc_obj.paragraphs]
            doc.original_content = "\n".join(paragraphs)
            doc.title = os.path.basename(file_path).replace(".docx", "").replace(".doc", "")
            
        except ImportError:
            doc.metadata["error"] = "python-docx 未安装"
            return doc
            
        doc.cleaned_content = self.clean_data(doc.original_content, ContentType.DOC)
        doc.content_hash = self._generate_hash(doc.cleaned_content)
        doc.is_duplicate = self._check_duplicate(doc.content_hash)
        doc.entities = self._extract_entities(doc.cleaned_content)
        doc.sections = self._extract_sections(doc.cleaned_content)
        doc.quality_score = self._score_quality(doc)
        
        return doc
    
    def _analyze_markdown(self, file_path: str, doc: AnalyzedDocument) -> AnalyzedDocument:
        """分析Markdown文件"""
        import os
        
        with open(file_path, "r", encoding="utf-8") as f:
            doc.original_content = f.read()
            
        doc.title = os.path.basename(file_path).replace(".md", "")
        doc.cleaned_content = self.clean_data(doc.original_content, ContentType.MARKDOWN)
        doc.content_hash = self._generate_hash(doc.cleaned_content)
        doc.is_duplicate = self._check_duplicate(doc.content_hash)
        doc.entities = self._extract_entities(doc.cleaned_content)
        doc.sections = self._extract_sections(doc.cleaned_content)
        doc.quality_score = self._score_quality(doc)
        
        return doc
    
    def _analyze_text_file(self, file_path: str, doc: AnalyzedDocument) -> AnalyzedDocument:
        """分析文本文件"""
        import os
        
        with open(file_path, "r", encoding="utf-8") as f:
            doc.original_content = f.read()
            
        doc.title = os.path.basename(file_path)
        doc.cleaned_content = self.clean_data(doc.original_content, ContentType.TEXT)
        doc.content_hash = self._generate_hash(doc.cleaned_content)
        doc.is_duplicate = self._check_duplicate(doc.content_hash)
        doc.entities = self._extract_entities(doc.cleaned_content)
        doc.sections = self._extract_sections(doc.cleaned_content)
        doc.quality_score = self._score_quality(doc)
        
        return doc
    
    def analyze_content(self, content: str, title: str = "") -> AnalyzedDocument:
        """
        分析文本内容
        
        Args:
            content: 文本内容
            title: 标题（可选）
        
        Returns:
            AnalyzedDocument: 分析结果
        """
        doc = AnalyzedDocument(
            source=DocumentSource.USER_INPUT,
            original_content=content,
            title=title or self._extract_title(content, ContentType.TEXT)
        )
        
        doc.cleaned_content = self.clean_data(content, ContentType.TEXT)
        doc.content_hash = self._generate_hash(doc.cleaned_content)
        doc.is_duplicate = self._check_duplicate(doc.content_hash)
        doc.entities = self._extract_entities(doc.cleaned_content)
        doc.sections = self._extract_sections(doc.cleaned_content)
        doc.quality_score = self._score_quality(doc)
        
        return doc
    
    def clean_data(self, content: str, content_type: ContentType = ContentType.TEXT) -> str:
        """
        数据清洗
        
        Args:
            content: 原始内容
            content_type: 内容类型
        
        Returns:
            str: 清洗后的内容
        """
        cfg = self.cleaner_config
        text = content
        
        # 1. HTML标签移除
        if cfg.remove_html_tags and content_type in [ContentType.HTML, ContentType.UNKNOWN]:
            text = re.sub(r"<[^>]+>", "", text)
            
        # 2. URL移除
        if cfg.remove_urls:
            text = re.sub(r"https?://\S+", "", text)
            
        # 3. 邮箱移除
        if cfg.remove_emails:
            text = re.sub(r"\S+@\S+\.\S+", "", text)
            
        # 4. 电话号码移除
        if cfg.remove_phones:
            text = re.sub(r"\d{3,4}-?\d{7,8}", "", text)
            text = re.sub(r"1[3-9]\d{9}", "", text)
            
        # 5. 空白规范化
        if cfg.normalize_whitespace:
            text = re.sub(r"\s+", " ", text)
            text = re.sub(r"\n+", "\n", text)
            
        # 6. 特殊字符处理
        if cfg.remove_special_chars:
            text = re.sub(r"[^\u4e00-\u9fa5a-zA-Z0-9\s\n.,;:!?，。；：！？、]", "", text)
            
        # 7. 长度过滤
        if len(text) < cfg.min_content_length:
            return ""
        if len(text) > cfg.max_content_length:
            text = text[:cfg.max_content_length]
            
        return text.strip()
    
    def _generate_hash(self, content: str) -> str:
        """生成内容哈希"""
        return hashlib.md5(content.encode("utf-8")).hexdigest()
    
    def _check_duplicate(self, content_hash: str) -> bool:
        """检查重复"""
        # 检查本地缓存
        if content_hash in self._content_cache:
            return True
            
        # 检查知识库
        if self.knowledge_base:
            try:
                results = self.knowledge_base.search(content_hash, top_k=1)
                if results and results[0].score > 0.95:
                    return True
            except Exception:
                pass
                
        return False
    
    def _extract_entities(self, content: str) -> Dict[str, List[str]]:
        """提取实体"""
        entities = {
            "locations": [],      # 地点
            "organizations": [],   # 组织/公司
            "numbers": [],        # 数字
            "dates": [],          # 日期
            "standards": [],      # 标准规范
            "keywords": [],       # 关键词
        }
        
        # 地点提取
        locations = re.findall(r"[省市县区镇](?:[^省市县区镇]+)?(?:市|县|区|镇|街|路|村)", content)
        entities["locations"] = list(set(locations))[:20]
        
        # 标准规范提取
        standards = re.findall(r"(GB\s*\d+(?:\.\d+)*|HJ\s*\d+(?:\.\d+)*|AQ\s*\d+(?:\.\d+)*|YD\s*\d+(?:\.\d+)*)", content, re.IGNORECASE)
        entities["standards"] = list(set(standards))[:10]
        
        # 数字提取（金额、百分比等）
        numbers = re.findall(r"\d+(?:\.\d+)?\s*(?:万|亿|元|%|吨|千瓦|平方米|公顷)", content)
        entities["numbers"] = list(set(numbers))[:20]
        
        # 日期提取
        dates = re.findall(r"\d{4}年\d{1,2}月\d{1,2}日|\d{4}-\d{2}-\d{2}|\d{4}/\d{2}/\d{2}", content)
        entities["dates"] = list(set(dates))[:10]
        
        # 关键词提取（出现频率高的词）
        words = re.findall(r"[\u4e00-\u9fa5]{2,4}", content)
        word_freq = {}
        for word in words:
            word_freq[word] = word_freq.get(word, 0) + 1
            
        # 取高频词作为关键词
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        entities["keywords"] = [w[0] for w in sorted_words if len(w[0]) >= 3][:30]
        
        return entities
    
    def _extract_sections(self, content: str) -> List[Dict]:
        """提取章节结构"""
        sections = []
        lines = content.split("\n")
        
        current_section = None
        current_content = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # 标题识别
            is_title = False
            
            # 模式1: 第X章/节
            chapter_match = re.match(r"^第([一二三四五六七八九十\d]+)[章节]", line)
            if chapter_match:
                is_title = True
                title_type = "chapter"
                
            # 模式2: # 标题
            hash_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if hash_match:
                is_title = True
                title_type = f"h{len(hash_match.group(1))}"
                
            # 模式3: X.X 标题
            num_match = re.match(r"^(\d+\.)+\s*(.+)$", line)
            if num_match:
                is_title = True
                title_type = "numbered"
                
            # 模式4: 【标题】
            bracket_match = re.match(r"^【(.+?)】$", line)
            if bracket_match:
                is_title = True
                title_type = "bracket"
                
            if is_title:
                # 保存上一个章节
                if current_section:
                    current_section["content"] = "\n".join(current_content[:10])  # 取前10行作为摘要
                    sections.append(current_section)
                    
                current_section = {
                    "title": line,
                    "title_type": title_type,
                    "line_start": i,
                    "content": "",
                }
                current_content = []
            elif current_section:
                current_content.append(line)
                
        # 保存最后一个章节
        if current_section:
            current_section["content"] = "\n".join(current_content[:10])
            sections.append(current_section)
            
        return sections[:50]  # 最多50个章节
    
    def _score_quality(self, doc: AnalyzedDocument) -> float:
        """质量评分"""
        score = 0.5
        
        # 内容长度加分
        length = len(doc.cleaned_content)
        if length > 1000:
            score += 0.15
        elif length > 5000:
            score += 0.25
            
        # 章节结构加分
        if len(doc.sections) >= 3:
            score += 0.1
        if len(doc.sections) >= 5:
            score += 0.1
            
        # 实体提取加分
        entity_count = sum(len(v) for v in doc.entities.values())
        if entity_count > 10:
            score += 0.1
            
        return min(1.0, score)
    
    async def store_to_knowledge_base(self, doc: AnalyzedDocument) -> bool:
        """
        存储到知识库
        
        Args:
            doc: 分析后的文档
        
        Returns:
            bool: 是否存储成功
        """
        if doc.is_duplicate:
            logger.info(f"跳过重复文档: {doc.content_hash}")
            return False
            
        if not doc.cleaned_content:
            logger.warning("内容为空，跳过存储")
            return False
            
        # 更新缓存
        self._content_cache[doc.content_hash] = doc.cleaned_content
        if doc.original_url:
            self._url_cache[doc.original_url] = doc.content_hash
            
        # 存储到知识库
        if self.knowledge_base:
            try:
                self.knowledge_base.add_document(
                    text=doc.cleaned_content,
                    metadata={
                        "type": "uploaded_document",
                        "source": doc.source.value,
                        "url": doc.original_url,
                        "title": doc.title,
                        "content_hash": doc.content_hash,
                        "entities": doc.entities,
                        "sections_count": len(doc.sections),
                        "quality_score": doc.quality_score,
                        "created_at": datetime.now().isoformat(),
                    }
                )
                logger.info(f"文档已入库: {doc.title}")
                
                # 更新进化引擎
                if self.evolution_engine:
                    self.evolution_engine.learn_from_generation(
                        requirement=doc.title,
                        doc_type=doc.metadata.get("detected_type", "unknown"),
                        content=doc.cleaned_content,
                        quality_score=doc.quality_score,
                        metadata=doc.to_dict()
                    )
                    
                return True
                
            except Exception as e:
                logger.error(f"知识库存储失败: {e}")
                
        return False
    
    async def analyze_batch(self, inputs: List[str]) -> List[AnalyzedDocument]:
        """
        批量分析
        
        Args:
            inputs: 输入列表
        
        Returns:
            List[AnalyzedDocument]: 分析结果列表
        """
        results = []
        
        for inp in inputs:
            try:
                result = await self.analyze(inp)
                results.append(result)
                
                # 自动入库（非重复）
                if not result.is_duplicate:
                    await self.store_to_knowledge_base(result)
                    
            except Exception as e:
                logger.error(f"批量分析失败: {inp}, {e}")
                
        return results


import os


# 全局实例
_analyzer: Optional[DocumentAnalyzer] = None


def get_document_analyzer() -> DocumentAnalyzer:
    """获取全局文档分析器"""
    global _analyzer
    if _analyzer is None:
        _analyzer = DocumentAnalyzer()
    return _analyzer


async def quick_analyze(input_data: str) -> Dict[str, Any]:
    """快速分析接口"""
    analyzer = get_document_analyzer()
    result = await analyzer.analyze(input_data)
    
    # 自动入库
    await analyzer.store_to_knowledge_base(result)
    
    return result.to_dict()
