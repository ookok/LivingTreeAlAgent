# -*- coding: utf-8 -*-
"""
AI OS Office 集成示例脚本
========================

此脚本演示如何通过 Python 调用 AI OS，实现 WPS/Word/Excel 的自动化处理。

依赖安装：
    pip install requests python-docx openpyxl

使用方法：
    python examples/office_integration.py
"""

import requests
import json
import time
from typing import Optional, List, Dict, Any

# ============== 配置 ==============

API_BASE_URL = "http://127.0.0.1:8898/api/v1"
TIMEOUT = 30  # 请求超时（秒）


# ============== API 客户端 ==============

class AIOSClient:
    """AI OS API 客户端"""

    def __init__(self, base_url: str = API_BASE_URL):
        self.base_url = base_url.rstrip('/')
        self.session = requests.Session()
        self.session.headers.update({"Content-Type": "application/json"})

    def _call(self, endpoint: str, data: Dict) -> Dict:
        """调用 API"""
        url = f"{self.base_url}/{endpoint}"
        try:
            response = self.session.post(url, json=data, timeout=TIMEOUT)
            response.raise_for_status()
            result = response.json()
            if not result.get('success'):
                raise Exception(result.get('error', 'Unknown error'))
            return result['data']
        except requests.exceptions.ConnectionError:
            raise Exception(f"无法连接到 AI OS 服务 (http://127.0.0.1:8898)\n请先运行: python -m core.external_integration")
        except requests.exceptions.Timeout:
            raise Exception("请求超时，请检查网络连接")

    def query(self, text: str, context: str = None) -> str:
        """知识库查询"""
        result = self._call("query", {"text": text, "context": context})
        return result.get('answer', '')

    def summarize(self, text: str) -> Dict:
        """文档摘要"""
        return self._call("summarize", {"text": text})

    def polish(self, text: str, style: str = "formal") -> Dict:
        """文本润色"""
        return self._call("polish", {"text": text, "options": {"style": style}})

    def translate(self, text: str, target_lang: str = "en") -> str:
        """翻译"""
        result = self._call("translate", {
            "text": text,
            "options": {"target_lang": target_lang}
        })
        return result.get('translated', '')

    def correct(self, text: str) -> str:
        """错别字纠正"""
        result = self._call("correct", {"text": text})
        return result.get('corrected', text)

    def analyze(self, text: str) -> Dict:
        """分析"""
        return self._call("analyze", {"text": text})

    def batch(self, texts: List[str], operation: str = "query") -> List[Dict]:
        """批量处理"""
        result = self._session.post(
            f"{self.base_url}/batch",
            json={"texts": texts, "operation": operation},
            timeout=TIMEOUT * 2
        )
        return result.json()['data']['results']


# ============== Word 处理 ==============

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("⚠️ 未安装 python-docx，请运行: pip install python-docx")


class AIOSWordProcessor:
    """AI OS Word 文档处理器"""

    def __init__(self, client: Optional[AIOSClient] = None):
        self.client = client or AIOSClient()

    def process_document(
        self,
        input_file: str,
        output_file: str,
        operations: List[str] = None
    ):
        """
        处理 Word 文档

        Args:
            input_file: 输入文件
            output_file: 输出文件
            operations: 操作列表 ["summarize", "polish", "correct"]
        """
        if not HAS_DOCX:
            raise Exception("需要安装 python-docx: pip install python-docx")

        doc = Document(input_file)
        operations = operations or ["correct"]

        # 统计
        stats = {
            "total_paragraphs": 0,
            "processed": 0,
            "summaries": 0,
            "polished": 0,
            "corrected": 0,
        }

        print(f"\n📄 开始处理: {input_file}")

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text or len(text) < 10:
                continue

            stats["total_paragraphs"] += 1

            try:
                for op in operations:
                    if op == "summarize" and len(text) > 200:
                        result = self.client.summarize(text)
                        if result.get('summary'):
                            para.text = result['summary']
                            stats["summaries"] += 1
                            stats["processed"] += 1
                            print(f"  [摘要] 段落 {i+1}")

                    elif op == "polish":
                        result = self.client.polish(text)
                        if result.get('polished'):
                            para.text = result['polished']
                            stats["polished"] += 1
                            stats["processed"] += 1
                            print(f"  [润色] 段落 {i+1}")

                    elif op == "correct":
                        result = self.client.correct(text)
                        if result != text:
                            para.text = result
                            stats["corrected"] += 1
                            stats["processed"] += 1
                            print(f"  [纠正] 段落 {i+1}")

            except Exception as e:
                print(f"  [错误] 段落 {i+1}: {e}")

            # 避免请求过快
            time.sleep(0.5)

        # 保存
        doc.save(output_file)
        print(f"\n✅ 处理完成: {output_file}")
        print(f"   总段落: {stats['total_paragraphs']}")
        print(f"   处理: {stats['processed']}")
        print(f"   摘要: {stats['summaries']}")
        print(f"   润色: {stats['polished']}")
        print(f"   纠正: {stats['corrected']}")

        return stats

    def create_summary_document(self, input_file: str, output_file: str):
        """创建摘要文档"""
        if not HAS_DOCX:
            raise Exception("需要安装 python-docx")

        doc = Document(input_file)
        summary_doc = Document()

        # 标题
        summary_doc.add_heading('📋 文档摘要', 0)
        summary_doc.add_paragraph(f'来源: {input_file}')
        summary_doc.add_paragraph(f'生成时间: {time.strftime("%Y-%m-%d %H:%M:%S")}')
        summary_doc.add_paragraph()

        # 摘要内容
        summary_doc.add_heading('内容摘要', 1)

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text or len(text) < 50:
                continue

            try:
                result = self.client.summarize(text)
                summary = result.get('summary', text)

                summary_doc.add_heading(f'段落 {i+1}', 2)
                summary_doc.add_paragraph(summary)

                # 要点
                key_points = result.get('key_points', [])
                if key_points:
                    summary_doc.add_paragraph('要点:')
                    for point in key_points:
                        summary_doc.add_paragraph(f'• {point}')

                print(f"  [处理] 段落 {i+1}")

            except Exception as e:
                print(f"  [跳过] 段落 {i+1}: {e}")

            time.sleep(0.5)

        summary_doc.save(output_file)
        print(f"\n✅ 摘要文档已生成: {output_file}")


# ============== Excel 处理 ==============

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    print("⚠️ 未安装 openpyxl，请运行: pip install openpyxl")


class AIOSExcelProcessor:
    """AI OS Excel 处理器"""

    def __init__(self, client: Optional[AIOSClient] = None):
        self.client = client or AIOSClient()

    def analyze_spreadsheet(self, input_file: str, output_file: str = None):
        """分析电子表格"""
        if not HAS_OPENPYXL:
            raise Exception("需要安装 openpyxl: pip install openpyxl")

        wb = openpyxl.load_workbook(input_file)
        ws = wb.active

        print(f"\n📊 开始分析: {input_file}")

        # 收集文本数据
        text_data = []
        for row in ws.iter_rows(max_row=50):  # 限制前50行
            for cell in row:
                if cell.value and isinstance(cell.value, str) and len(cell.value) > 10:
                    text_data.append(cell.value)

        if not text_data:
            print("未找到可分析的文本数据")
            return

        # 合并分析
        combined_text = "\n---\n".join(text_data[:10])  # 限制10个单元格

        try:
            result = self.client.analyze(combined_text)
            analysis = result.get('analysis', '')

            # 输出结果
            print("\n" + "="*60)
            print("📈 分析结果")
            print("="*60)
            print(analysis)
            print("="*60)

            # 保存到新sheet
            if output_file:
                result_sheet = wb.create_sheet("AI分析结果")
                result_sheet['A1'] = "AI 分析结果"
                result_sheet['A2'] = analysis
                wb.save(output_file)
                print(f"\n✅ 结果已保存: {output_file}")

        except Exception as e:
            print(f"分析失败: {e}")

    def batch_translate_sheet(
        self,
        input_file: str,
        sheet_name: str,
        column: int,
        output_file: str,
        target_lang: str = "en"
    ):
        """批量翻译工作表"""
        if not HAS_OPENPYXL:
            raise Exception("需要安装 openpyxl")

        wb = openpyxl.load_workbook(input_file)

        if sheet_name not in wb.sheetnames:
            print(f"未找到工作表: {sheet_name}")
            return

        ws = wb[sheet_name]

        print(f"\n🌐 开始翻译工作表: {sheet_name}")

        translations = []
        row_count = 0

        for row in ws.iter_rows(min_col=column, max_col=column):
            for cell in row:
                if cell.value and isinstance(cell.value, str) and len(cell.value) > 2:
                    try:
                        result = self.client.translate(cell.value, target_lang)
                        translations.append((cell.row, cell.value, result))
                        row_count += 1
                        print(f"  [翻译] 行 {cell.row}: {cell.value[:20]}...")

                        if row_count >= 50:  # 限制50条
                            break
                    except Exception as e:
                        print(f"  [跳过] 行 {cell.row}: {e}")

                    time.sleep(0.3)

            if row_count >= 50:
                break

        # 写入翻译结果
        target_col = column + 1
        for row, original, translated in translations:
            ws.cell(row=row, column=target_col, value=translated)

        wb.save(output_file)
        print(f"\n✅ 翻译完成: {output_file}")
        print(f"   翻译条目: {len(translations)}")


# ============== 演示 ==============

def demo():
    """演示函数"""
    print("=" * 60)
    print("🤖 AI OS Office 集成演示")
    print("=" * 60)

    client = AIOSClient()

    # 健康检查
    print("\n1️⃣ 检查服务状态...")
    try:
        health = requests.get(f"{client.base_url}/health", timeout=5)
        print(f"   ✅ 服务正常: {health.json()}")
    except Exception as e:
        print(f"   ❌ 服务未运行: {e}")
        print("\n请先启动服务:")
        print("   python -m core.external_integration")
        return

    # 能力列表
    print("\n2️⃣ 获取 AI OS 能力...")
    caps = requests.get(f"{client.base_url}/capabilities")
    print(f"   支持的操作:")
    for cap in caps.json()['capabilities']:
        print(f"   • {cap['name']}: {cap['description']}")

    # 示例操作
    print("\n3️⃣ 测试操作...")

    test_text = "人工智能是计算机科学的一个分支，它企图了解智能的实质，并生产出一种新的能以人类智能相似的方式做出反应的智能机器。"

    print(f"\n   原文: {test_text[:50]}...")

    # 摘要
    print("\n   [摘要测试]")
    result = client.summarize(test_text)
    print(f"   结果: {result.get('summary', '')[:100]}...")

    # 润色
    print("\n   [润色测试]")
    result = client.polish("请帮我写一份报告")
    print(f"   结果: {result.get('polished', '')}")

    # 翻译
    print("\n   [翻译测试]")
    result = client.translate("Hello, how are you?", "zh")
    print(f"   结果: {result}")

    # 纠正
    print("\n   [纠正测试]")
    result = client.correct("今天天气很好")
    print(f"   结果: {result}")

    print("\n" + "=" * 60)
    print("✅ 演示完成!")
    print("=" * 60)


# ============== 主入口 ==============

if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        if sys.argv[1] == "--demo":
            demo()
        elif sys.argv[1] == "--help":
            print("""
AI OS Office 集成示例

用法:
  python office_integration.py --demo        # 运行演示
  python office_integration.py --help        # 显示帮助

示例脚本中包含:
  - AIOSClient: API 客户端
  - AIOSWordProcessor: Word 文档处理器
  - AIOSExcelProcessor: Excel 处理器

请参考 docs/office_integration_guide.md 获取详细文档
            """)
    else:
        demo()
