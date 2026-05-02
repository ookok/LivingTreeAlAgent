"""
文档生成器 - 核心实现

基于JSON Schema生成Word/PDF/Markdown文档。
"""
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Any, Optional

from .json_schema import (
    ReportSchema,
    ReportSection,
    ReportType,
    ContentType,
    SchemaValidator,
    ReportTemplates,
)

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    logger.warning("python-docx not installed, Word generation disabled")


@dataclass
class DocumentGenerationResult:
    """文档生成结果"""
    success: bool
    output_path: Optional[str] = None
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    generated_sections: int = 0


class DocumentGenerator:
    """
    文档生成器
    
    核心能力：
    1. 根据Schema生成Word文档
    2. 生成Markdown文档
    3. 填充数据到模板
    """
    
    def __init__(self):
        self.validator = SchemaValidator()
    
    def generate_from_schema(
        self,
        schema: ReportSchema,
        data: Dict[str, Any],
        output_path: str,
        format: str = "docx"
    ) -> DocumentGenerationResult:
        """
        从Schema生成文档
        
        Args:
            schema: 报告Schema
            data: 数据字典
            output_path: 输出路径
            format: 输出格式 (docx/md)
        
        Returns:
            DocumentGenerationResult
        """
        logger.info(f"Generating document: {schema.title}")
        
        # 验证Schema和数据
        validation_errors = self.validator.validate_schema(schema)
        validation_errors.extend(self.validator.validate_data(data, schema))
        
        if validation_errors:
            return DocumentGenerationResult(
                success=False,
                errors=validation_errors
            )
        
        # 根据格式生成文档
        if format.lower() == "docx":
            return self._generate_docx(schema, data, output_path)
        elif format.lower() == "md":
            return self._generate_markdown(schema, data, output_path)
        else:
            return DocumentGenerationResult(
                success=False,
                errors=[f"不支持的格式: {format}"]
            )
    
    def _generate_docx(self, schema: ReportSchema, data: Dict[str, Any], output_path: str) -> DocumentGenerationResult:
        """生成Word文档"""
        if not HAS_PYTHON_DOCX:
            return DocumentGenerationResult(
                success=False,
                errors=["python-docx未安装，无法生成Word文档"]
            )
        
        try:
            doc = Document()
            
            # 设置文档属性
            doc.core_properties.title = schema.title
            doc.core_properties.version = schema.version
            
            # 添加标题
            title_para = doc.add_heading(schema.title, 0)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 遍历章节
            generated_count = 0
            for section in schema.sections:
                generated_count += self._add_section_to_docx(doc, section, data, 1)
            
            # 保存文档
            doc.save(output_path)
            
            logger.info(f"Document saved to: {output_path}")
            return DocumentGenerationResult(
                success=True,
                output_path=output_path,
                generated_sections=generated_count
            )
        
        except Exception as e:
            logger.error(f"Word generation failed: {e}")
            return DocumentGenerationResult(
                success=False,
                errors=[str(e)]
            )
    
    def _add_section_to_docx(self, doc, section: ReportSection, data: Dict[str, Any], level: int) -> int:
        """添加章节到Word文档"""
        count = 1
        
        # 添加标题
        heading = doc.add_heading(section.title, level)
        
        # 根据内容类型处理
        if section.content_type == ContentType.TEXT:
            # 填充内容
            content = self._fill_content(section.content, data) if section.content else ""
            if content:
                doc.add_paragraph(content)
        
        elif section.content_type == ContentType.TABLE:
            if section.table_spec:
                self._add_table_to_docx(doc, section.table_spec, data)
        
        elif section.content_type == ContentType.CHART:
            # 图表占位符
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("[图表] " + (section.chart_spec.title if section.chart_spec else ""))
            run.italic = True
        
        elif section.content_type == ContentType.CODE:
            if section.code:
                code_content = self._fill_content(section.code, data)
                paragraph = doc.add_paragraph(code_content)
                paragraph.style = "Code"
        
        elif section.content_type == ContentType.LIST:
            for item in section.items:
                doc.add_paragraph(self._fill_content(item, data), style='List Bullet')
        
        # 递归处理子章节
        for child in section.children:
            count += self._add_section_to_docx(doc, child, data, level + 1)
        
        return count
    
    def _add_table_to_docx(self, doc, table_spec, data: Dict[str, Any]):
        """添加表格到Word文档"""
        # 获取表格数据
        table_data = []
        if table_spec.data_source and table_spec.data_source in data:
            table_data = data[table_spec.data_source]
        
        # 创建表格
        if table_data:
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else len(table_spec.headers)
        else:
            rows = 3
            cols = len(table_spec.headers)
        
        table = doc.add_table(rows=rows + 1, cols=cols)
        table.style = 'Table Grid'
        
        # 添加表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(table_spec.headers[:cols]):
            hdr_cells[i].text = header
            # 设置表头样式
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # 添加数据行
        for i, row in enumerate(table_data[:rows]):
            row_cells = table.rows[i + 1].cells
            for j, cell in enumerate(row[:cols]):
                row_cells[j].text = str(cell)
        
        # 添加标题
        if table_spec.caption:
            paragraph = doc.add_paragraph()
            paragraph.add_run(table_spec.caption).italic = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def _generate_markdown(self, schema: ReportSchema, data: Dict[str, Any], output_path: str) -> DocumentGenerationResult:
        """生成Markdown文档"""
        try:
            content = []
            
            # 添加标题
            content.append(f"# {schema.title}")
            content.append(f"> 版本: {schema.version}")
            content.append("")
            
            # 遍历章节
            generated_count = 0
            for section in schema.sections:
                generated_count += self._generate_markdown_section(content, section, data, 1)
            
            # 写入文件
            with open(output_path, "w", encoding="utf-8") as f:
                f.write("\n".join(content))
            
            logger.info(f"Markdown saved to: {output_path}")
            return DocumentGenerationResult(
                success=True,
                output_path=output_path,
                generated_sections=generated_count
            )
        
        except Exception as e:
            logger.error(f"Markdown generation failed: {e}")
            return DocumentGenerationResult(
                success=False,
                errors=[str(e)]
            )
    
    def _generate_markdown_section(self, content: List[str], section: ReportSection, data: Dict[str, Any], level: int) -> int:
        """生成Markdown章节内容"""
        count = 1
        
        # 添加标题
        content.append("#" * level + " " + section.title)
        
        # 根据内容类型处理
        if section.content_type == ContentType.TEXT:
            content_text = self._fill_content(section.content, data) if section.content else ""
            if content_text:
                content.append(content_text)
                content.append("")
        
        elif section.content_type == ContentType.TABLE:
            if section.table_spec:
                content.extend(self._generate_markdown_table(section.table_spec, data))
                content.append("")
        
        elif section.content_type == ContentType.CHART:
            title = section.chart_spec.title if section.chart_spec else "图表"
            content.append(f"![{title}]()")
            content.append("")
        
        elif section.content_type == ContentType.CODE:
            if section.code:
                code_content = self._fill_content(section.code, data)
                content.append(f"```python")
                content.append(code_content)
                content.append("```")
                content.append("")
        
        elif section.content_type == ContentType.LIST:
            for item in section.items:
                content.append(f"- {self._fill_content(item, data)}")
            content.append("")
        
        # 递归处理子章节
        for child in section.children:
            count += self._generate_markdown_section(content, child, data, level + 1)
        
        return count
    
    def _generate_markdown_table(self, table_spec, data: Dict[str, Any]) -> List[str]:
        """生成Markdown表格"""
        lines = []
        
        # 获取表格数据
        table_data = []
        if table_spec.data_source and table_spec.data_source in data:
            table_data = data[table_spec.data_source]
        
        # 添加标题
        if table_spec.caption:
            lines.append(f"*{table_spec.caption}*")
        
        # 添加表头
        lines.append("| " + " | ".join(table_spec.headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(table_spec.headers)) + " |")
        
        # 添加数据行
        for row in table_data:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
        
        return lines
    
    def _fill_content(self, template: str, data: Dict[str, Any]) -> str:
        """使用数据填充模板内容"""
        if not template:
            return ""
        
        result = template
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            result = result.replace(placeholder, str(value))
        
        return result
    
    def generate_eia_report(self, data: Dict[str, Any], output_path: str) -> DocumentGenerationResult:
        """生成环评报告"""
        schema = ReportTemplates.get_eia_report_schema()
        return self.generate_from_schema(schema, data, output_path)
    
    def generate_feasibility_study(self, data: Dict[str, Any], output_path: str) -> DocumentGenerationResult:
        """生成可行性研究报告"""
        schema = ReportTemplates.get_feasibility_study_schema()
        return self.generate_from_schema(schema, data, output_path)


# 便捷函数
def generate_document(
    report_type: str,
    data: Dict[str, Any],
    output_path: str,
    format: str = "docx"
) -> DocumentGenerationResult:
    """便捷函数：生成文档"""
    generator = DocumentGenerator()
    
    if report_type.lower() == "eia":
        return generator.generate_eia_report(data, output_path)
    elif report_type.lower() == "feasibility":
        return generator.generate_feasibility_study(data, output_path)
    else:
        # 创建自定义Schema
        schema = ReportSchema(
            report_type=ReportType.CUSTOM,
            title=data.get("title", "报告"),
            sections=[
                ReportSection(
                    id="content",
                    title="内容",
                    level=1,
                    content_type=ContentType.TEXT,
                    content=data.get("content", ""),
                ),
            ],
        )
        return generator.generate_from_schema(schema, data, output_path, format)