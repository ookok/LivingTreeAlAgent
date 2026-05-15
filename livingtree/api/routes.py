"""API routes for the LivingTree server."""

from __future__ import annotations

import asyncio
import base64
import json
import os
import secrets
import subprocess
import time as _time
from datetime import datetime, timezone
from typing import Any, Optional

from pathlib import Path

from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect, Request, UploadFile, File, Form, Query
from fastapi.responses import Response, StreamingResponse, JSONResponse, FileResponse, HTMLResponse
from pydantic import BaseModel, Field
from loguru import logger
from .. import __version__


class LocalFolderRequest(BaseModel):
    path: str


class FileOperationRequest(BaseModel):
    op: str        # "read", "write", "exec", "list"
    path: str = ""  # File or directory path
    content: Optional[str] = None  # Content for write, command for exec


class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    context: dict[str, Any] = Field(default_factory=dict)


class ChatResponse(BaseModel):
    session_id: str
    reply: str = ""  # Primary text output for frontend display
    intent: Optional[str] = None
    plan: list[dict[str, Any]] = Field(default_factory=list)
    execution_results: list[dict[str, Any]] = Field(default_factory=list)
    reflections: list[str] = Field(default_factory=list)
    mutations: int = 0
    generation: int = 0
    success_rate: float = 0.0
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


class ReportRequest(BaseModel):
    template_type: str
    data: dict[str, Any] = Field(default_factory=dict)
    requirements: dict[str, Any] = Field(default_factory=dict)


class TrainRequest(BaseModel):
    cell_name: str
    training_data: list[dict] = Field(default_factory=list)
    epochs: int = 3


class DrillRequest(BaseModel):
    cell_name: str
    model_name: str
    dataset: list[dict] = Field(default_factory=list)
    training_type: str = "lora"
    teacher_model: str = ""
    reward_model: str = ""


class QuantizeRequest(BaseModel):
    model_path: str
    method: str = "awq"


class DeployRequest(BaseModel):
    model_path: str
    port: int = 8000


class HealthResponse(BaseModel):
    status: str
    version: str
    uptime_seconds: float = 0.0
    components: dict[str, str] = Field(default_factory=dict)


def setup_routes(app: FastAPI) -> None:
    """Register all API routes."""

    @app.get("/api/boot/progress")
    async def boot_progress(request: Request) -> dict[str, Any]:
        """Get hub initialization progress for the web UI."""
        try:
            hub = request.app.state.hub
        except Exception:
            return {"stage": "ready", "pct": 100, "detail": "系统就绪"}
        if hub and getattr(hub, '_started', False):
            return {"stage": "ready", "pct": 100, "detail": "系统就绪"}
        return {"stage": "starting", "pct": 50, "detail": "正在初始化..."}

    # ═══ Web API (SPA chat) ═══

    @app.get("/api/health", response_model=HealthResponse)
    async def health(request: Request) -> HealthResponse:
        """Health check endpoint."""
        try:
            hub = request.app.state.hub
        except Exception:
            return HealthResponse(status="ok", version=__version__, components={"server": "running"})
        components = {}
        if hub:
            if getattr(hub, '_started', False):
                status = hub.status() if callable(getattr(hub, 'status', None)) else {}
                components = {
                    "life_engine": "ok" if status.get("life_engine") else "missing",
                    "cells": f"{status.get('cells', {}).get('registered', 0)} registered",
                    "knowledge": "ok" if status.get("knowledge") else "missing",
                    "network": status.get("network", {}).get("status", "unknown"),
                    "orchestrator": f"{status.get('orchestrator', {}).get('total_agents', 0)} agents",
                }
                return HealthResponse(status="healthy", version=__version__, components=components)
            else:
                bp = getattr(hub, '_boot_progress', {})
                return HealthResponse(status="starting", version=__version__,
                    components={"boot": bp.get("detail", "initializing...")})
        return HealthResponse(status="starting", version=__version__,
            components={"boot": "waiting for hub..."})

    @app.post("/api/cache/flush")
    async def flush_cache():
        """Flush hot response cache. Returns count of removed entries."""
        try:
            from ..treellm.response_cache import get_response_cache
            cache = get_response_cache()
            removed = cache.flush()
            stats = cache.stats
            return {"status": "flushed", "removed": removed, "stats": stats}
        except Exception as e:
            return {"status": "error", "detail": str(e)[:200]}

    @app.get("/api/cache/stats")
    async def cache_stats():
        """Get response cache statistics."""
        try:
            from ..treellm.response_cache import get_response_cache
            cache = get_response_cache()
            return {"status": "ok", "stats": cache.stats}
        except Exception as e:
            return {"status": "error", "detail": str(e)[:200]}

    @app.get("/api/admin/config")
    async def get_admin_config(request: Request):
        cfg = get_config()
        return {
            "ollama_url": cfg.ollama_base_url,
            "onlyoffice_url": os.environ.get("ONLYOFFICE_URL", ""),
        }

    @app.post("/api/admin/config")
    async def save_admin_config(request: Request):
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        data = await request.json()
        provider = data.get("provider", "")
        key = data.get("key", "")
        if provider and key:
            env_map = {
                "deepseek": "DEEPSEEK_API_KEY", "modelscope": "MODELSCOPE_API_KEY",
                "bailing": "BAILING_API_KEY", "stepfun": "STEPFUN_API_KEY",
                "internlm": "INTERNLM_API_KEY", "openrouter": "OPENROUTER_API_KEY",
                "nvidia": "NVIDIA_API_KEY", "longcat": "LONGCAT_API_KEY",
                "sensetime": "SENSETIME_API_KEY", "siliconflow": "SILICONFLOW_API_KEY",
                "zhipu": "ZHIPU_API_KEY", "spark": "SPARK_API_KEY",
            }
            env_key = env_map.get(provider, provider.upper() + "_API_KEY")
            os.environ[env_key] = key
            return {"ok": True, "provider": provider}
        for k, v in data.items():
            if v and k.endswith("_url"):
                os.environ[k.upper()] = str(v)
        return {"ok": True}

    @app.get("/api/admin/export")
    async def export_config(request: Request):
        """Export all admin configuration as encrypted file (admin auth required)."""
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        cfg = get_config()
        data = {
            "api_keys": {
                "deepseek": cfg.deepseek_api_key, "modelscope": cfg.modelscope_api_key,
                "bailing": cfg.bailing_api_key, "stepfun": cfg.stepfun_api_key,
                "internlm": cfg.internlm_api_key, "openrouter": cfg.openrouter_api_key,
            },
            "urls": {
                "ollama": cfg.ollama_base_url,
                "onlyoffice": os.environ.get("ONLYOFFICE_URL", ""),
            },
            "exported_at": time.time(),
        }
        plain = json.dumps(data, ensure_ascii=False)
        # AES encrypt with project-derived key
        from hashlib import sha256
        enc_key = sha256(b"livingtree_admin_export").digest()
        import base64 as b64
        from cryptography.fernet import Fernet
        try:
            f = Fernet(b64.urlsafe_b64encode(enc_key))
            encrypted = f.encrypt(plain.encode())
            return Response(content=encrypted, media_type="application/octet-stream",
                          headers={"Content-Disposition": "attachment; filename=config.enc"})
        except ImportError:
            # Fallback: simple XOR with key hash
            key_bytes = enc_key
            encrypted = bytes(b ^ key_bytes[i % len(key_bytes)] for i, b in enumerate(plain.encode()))
            return Response(content=encrypted, media_type="application/octet-stream",
                          headers={"Content-Disposition": "attachment; filename=config.enc"})

    @app.post("/api/admin/import")
    async def import_config(request: Request):
        """Import encrypted configuration file (admin auth required)."""
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        body = await request.body()
        from hashlib import sha256
        enc_key = sha256(b"livingtree_admin_export").digest()
        import base64 as b64
        try:
            from cryptography.fernet import Fernet
            f = Fernet(b64.urlsafe_b64encode(enc_key))
            plain = f.decrypt(body).decode()
        except ImportError:
            key_bytes = enc_key
            plain = bytes(b ^ key_bytes[i % len(key_bytes)] for i, b in enumerate(body)).decode()
        except Exception:
            raise HTTPException(400, "Invalid or corrupted config file")

        try:
            data = json.loads(plain)
            # Restore API keys
            for provider, key in data.get("api_keys", {}).items():
                if key:
                    env_map = {
                        "deepseek": "DEEPSEEK_API_KEY", "modelscope": "MODELSCOPE_API_KEY",
                        "bailing": "BAILING_API_KEY", "stepfun": "STEPFUN_API_KEY",
                        "internlm": "INTERNLM_API_KEY", "openrouter": "OPENROUTER_API_KEY",
                    }
                    os.environ[env_map.get(provider, provider.upper() + "_API_KEY")] = key
            # Restore URLs
            for k, v in data.get("urls", {}).items():
                if v:
                    os.environ[k.upper() + "_URL"] = v
            return {"ok": True, "restored": len(data.get("api_keys", {}))}
        except Exception as e:
            raise HTTPException(400, f"Invalid config: {e}")

    @app.post("/api/chat", response_model=ChatResponse)
    async def chat(req: ChatRequest, request: Request) -> ChatResponse:
        """Send a message to the life engine and get results."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")

        # Route through unified input bus
        try:
            from ..treellm.living_input_bus import get_living_input_bus, InputSource
            bus = get_living_input_bus()
            result = await bus.normalize_and_route(InputSource.WEB, request, hub)
        except Exception:
            result = await hub.chat(req.message, **req.context)
        return ChatResponse(
            session_id=result["session_id"],
            reply=result.get("reply", "") or
                  "\n".join(str(r.get("output","")) for r in result.get("execution_results",[])),
            intent=result.get("intent"),
            plan=result.get("plan", []),
            execution_results=result.get("execution_results", []),
            reflections=result.get("reflections", []),
            mutations=result.get("mutations", 0),
            generation=result.get("generation", 0),
            success_rate=result.get("success_rate", 0.0),
        )

    @app.post("/api/chat/stream/resume")
    async def chat_stream_resume(request: Request):
        """Resume a streaming session from a given offset.

        Client sends: X-Session-ID + X-Received-Length headers.
        Server resumes streaming from that offset if session cached.
        """
        session_id = request.headers.get("X-Session-ID", "")
        received_len = int(request.headers.get("X-Received-Length", "0"))

        # Check session cache
        from ..api.stream_session import get_session_cache
        cache = get_session_cache()
        session = cache.get(session_id)

        if not session:
            return JSONResponse(
                {"error": "Session not found or expired"},
                status_code=404,
            )

        async def generate():
            text = session.get("full_text", "")
            if received_len >= len(text):
                yield f"event: done\ndata: {{}}\n\n"
                return

            # Resume from offset
            remaining = text[received_len:]
            chunk_size = 50
            for i in range(0, len(remaining), chunk_size):
                chunk = remaining[i:i + chunk_size]
                event_data = json.dumps({
                    "text": chunk,
                    "offset": received_len + i + len(chunk),
                    "total": len(text),
                }, ensure_ascii=False)
                yield f"event: resume\ndata: {event_data}\n\n"
                await asyncio.sleep(0.03)

            yield f"event: done\ndata: {{}}\n\n"

        return StreamingResponse(
            generate(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
            },
        )

    @app.post("/api/chat/stream")
    async def chat_stream(req: ChatRequest, request: Request):
        """Progressive streaming chat with Flash-First + Parallel Race + Skeleton.

        Returns SSE stream: skeleton → phases → tokens → complete.
        User sees first token at ~100ms (vs ~5s for blocking /api/chat).
        """
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")

        async def generate():
            try:
                from ..treellm.concurrent_stream import get_concurrent_stream, set_stream_chat_fn
                cs = get_concurrent_stream()
                llm = hub.world.consciousness._llm

                async def _chat_fn(messages, provider, stream=True):
                    async for token in llm.stream(messages, provider=provider):
                        yield token

                set_stream_chat_fn(_chat_fn)
                flash_model = await llm.smart_route(req.message, task_type="chat")
                # Elect pro model via L2 tier
                tiers = await hub.world.consciousness._elect_tiers()
                pro_model = tiers.get(2, flash_model) if tiers else flash_model

                async for event in cs.stream(
                    query=req.message,
                    flash_model=flash_model,
                    pro_model=pro_model,
                    system_prompt=req.context,
                    task_type="chat",
                ):
                    event_data = json.dumps({
                        "type": event.kind,
                        "content": event.text,
                        "phase": "stream",
                        "model": event.provider,
                        "ts": event.timestamp,
                    }, ensure_ascii=False)
                    yield f"event: {event.kind}\ndata: {event_data}\n\n"
                # NOTE: cognition_stream fallback removed — it redundantly called engine.run()
                # which doubled LLM cost. ConcurrentStream already provides full output.

            except Exception as e:
                error_data = json.dumps({"error": str(e)[:200]})
                yield f"event: error\ndata: {error_data}\n\n"

        return StreamingResponse(
            generate(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
            },
        )

    # ═══ Unified Living Input (any modality → canonical → pipeline) ═══

    @app.post("/api/living/input")
    async def living_input(request: Request):
        """Unified input — text, JSON, files, refs. Device-agnostic."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        try:
            from ..treellm.living_input_bus import get_living_input_bus, InputSource
            bus = get_living_input_bus()
            result = await bus.normalize_and_route(InputSource.WEB, request, hub)
            return {"ok": True, "session_id": result.get("session_id", "perpetual"), "result": result}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    @app.post("/api/local/analyze")
    async def analyze_local_folder(req: LocalFolderRequest):
        """Analyze a local folder directly — no upload needed."""
        folder_path = Path(req.path).resolve()
        if not str(folder_path).startswith(str(Path.cwd().resolve())):
            raise HTTPException(403, "Path traversal not allowed")
        if not folder_path.exists():
            raise HTTPException(404, f"Folder not found: {req.path}")
        if not folder_path.is_dir():
            raise HTTPException(400, "Path must be a folder")
        files = []
        total_size = 0
        extensions = {}
        try:
            for item in folder_path.rglob("*"):
                if item.is_file() and "__pycache__" not in str(item) and ".git" not in str(item):
                    rel = str(item.relative_to(folder_path))
                    size = item.stat().st_size
                    total_size += size
                    ext = item.suffix or "no_ext"
                    extensions[ext] = extensions.get(ext, 0) + 1
                    files.append({"path": rel, "size": size, "ext": ext})
        except PermissionError:
            raise HTTPException(403, "Permission denied")
        files.sort(key=lambda f: -f["size"])
        return {
            "path": str(folder_path), "file_count": len(files),
            "total_size_mb": round(total_size / (1024 * 1024), 1),
            "extensions": dict(sorted(extensions.items(), key=lambda x: -x[1])[:10]),
            "largest_files": files[:20],
        }

    @app.get("/api/doc/view")
    async def view_document(path: str = Query("")):
        """Render a local document as HTML for in-browser viewing.

        Supports: .docx → HTML (python-docx), .xlsx → HTML table (openpyxl),
        .pdf → iframe, .txt/.md → text.
        """
        file_path = Path(path)
        if not file_path.exists():
            raise HTTPException(404, f"File not found: {path}")

        suffix = file_path.suffix.lower()

        # PDF → serve directly for browser native viewer
        if suffix == ".pdf":
            return FileResponse(file_path, media_type="application/pdf")

        # DOCX → convert to HTML
        if suffix in (".docx", ".doc"):
            try:
                from docx import Document
                doc = Document(str(file_path))
                html_parts = ['<div class="doc-render">']
                for para in doc.paragraphs:
                    style = para.style.name if para.style else ""
                    text = para.text
                    if not text.strip():
                        html_parts.append("<br>")
                    elif "Heading" in style:
                        level = style.replace("Heading ", "").replace("Heading", "1")
                        html_parts.append(f"<h{level}>{text}</h{level}>")
                    else:
                        html_parts.append(f"<p>{text}</p>")
                for table in doc.tables:
                    html_parts.append('<table class="data-table">')
                    for row in table.rows:
                        html_parts.append("<tr>" + "".join(f"<td>{c.text}</td>" for c in row.cells) + "</tr>")
                    html_parts.append("</table>")
                html_parts.append("</div>")
                return JSONResponse({"html": "\n".join(html_parts), "type": "docx"})
            except Exception as e:
                return JSONResponse({"html": f"<p>无法解析文档: {e}</p>", "type": "error"})

        # XLSX → convert to HTML table
        if suffix in (".xlsx", ".xls"):
            try:
                from openpyxl import load_workbook
                wb = load_workbook(str(file_path), data_only=True)
                html_parts = ['<div class="doc-render">']
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    html_parts.append(f"<h3>{sheet_name}</h3><table class='data-table'>")
                    for row in ws.iter_rows(max_row=min(100, ws.max_row), values_only=True):
                        html_parts.append("<tr>" + "".join(
                            f"<td>{str(c)[:200] if c is not None else ''}</td>" for c in row
                        ) + "</tr>")
                    html_parts.append("</table>")
                html_parts.append("</div>")
                return JSONResponse({"html": "\n".join(html_parts), "type": "xlsx"})
            except Exception as e:
                return JSONResponse({"html": f"<p>无法解析表格: {e}</p>", "type": "error"})

        # TXT/MD → return as text
        if suffix in (".txt", ".md", ".py", ".json", ".yaml", ".yml", ".csv"):
            try:
                content = file_path.read_text("utf-8")[:50000]
                lang = suffix.replace(".", "")
                return JSONResponse({
                    "html": f"<pre class='code-block'><code>{content}</code></pre>",
                    "type": "text",
                    "lang": lang,
                })
            except Exception:
                pass

        raise HTTPException(400, f"Unsupported format: {suffix}")
        """Analyze a local folder directly — no upload needed."""
        folder_path = Path(req.path)
        if not folder_path.exists():
            raise HTTPException(404, f"Folder not found: {req.path}")
        if not folder_path.is_dir():
            raise HTTPException(400, "Path must be a folder")

        files = []
        total_size = 0
        extensions = {}
        try:
            for item in folder_path.rglob("*"):
                if item.is_file() and "__pycache__" not in str(item) and ".git" not in str(item):
                    rel = str(item.relative_to(folder_path))
                    size = item.stat().st_size
                    total_size += size
                    ext = item.suffix or "no_ext"
                    extensions[ext] = extensions.get(ext, 0) + 1
                    files.append({"path": rel, "size": size, "ext": ext})
        except PermissionError:
            raise HTTPException(403, "Permission denied")

        files.sort(key=lambda f: -f["size"])
        return {
            "path": str(folder_path), "file_count": len(files),
            "total_size_mb": round(total_size / (1024 * 1024), 1),
            "extensions": dict(sorted(extensions.items(), key=lambda x: -x[1])[:10]),
            "largest_files": files[:20],
        }

    @app.post("/api/local/file")
    async def local_file_operation(req: FileOperationRequest):
        """Execute file operations on the local filesystem via the server.

        Browser cannot access local filesystem directly (security),
        but the server can. Operations: read, write, exec, list.
        """
        file_path = Path(req.path).resolve() if req.path else None
        if file_path and not str(file_path).startswith(str(Path.cwd().resolve())):
            raise HTTPException(403, "Path traversal not allowed")

        if req.op == "read":
            if not file_path or not file_path.exists():
                return {"ok": False, "error": f"File not found: {req.path}"}
            try:
                content = file_path.read_text("utf-8")
                return {"ok": True, "content": content[:10000], "size": len(content)}
            except Exception as e:
                return {"ok": False, "error": str(e)}

        elif req.op == "write":
            if not file_path:
                return {"ok": False, "error": "No path specified"}
            try:
                file_path.parent.mkdir(parents=True, exist_ok=True)
                file_path.write_text(req.content or "", "utf-8")
                return {"ok": True, "written": len(req.content or ""),
                        "path": str(file_path)}
            except Exception as e:
                return {"ok": False, "error": str(e)}

        elif req.op == "exec":
            if not req.content:
                return {"ok": False, "error": "No command specified"}
            # Security: whitelist only safe commands
            ALLOWED_PREFIXES = [
                "python ", "pip ", "pytest ", "git ", "dir ", "ls ",
                "cat ", "type ", "echo ", "node ", "npm ",
            ]
            cmd_lower = req.content.strip().lower()
            if not any(cmd_lower.startswith(p) for p in ALLOWED_PREFIXES):
                return {"ok": False, "error": f"Command not allowed. Allowed prefixes: {ALLOWED_PREFIXES}"}
            # Block command chaining and shell metacharacters
            if any(k in req.content for k in ("&&", "||", ";", "|", "`", "$(", "${")):
                return {"ok": False, "error": "Command chaining not allowed"}
            work_dir = str(file_path) if file_path and file_path.is_dir() else str(Path.cwd())
            try:
                import shlex
                cmd_parts = shlex.split(req.content.strip())
                try:
                    from ..treellm.unified_exec import run
                    result = await run(req.content.strip(), timeout=30, cwd=work_dir)
                    return {"ok": result.success,
                            "output": result.stdout[:5000] or result.stderr[:5000],
                            "exit_code": result.exit_code}
                except ImportError:
                    result = subprocess.run(
                        cmd_parts,
                        capture_output=True, text=True, timeout=30,
                        cwd=work_dir,
                    )
                    return {"ok": True,
                            "output": result.stdout[:5000] or result.stderr[:5000],
                            "exit_code": result.returncode}
            except subprocess.TimeoutExpired:
                return {"ok": False, "error": "Command timed out (30s)"}
            except Exception as e:
                return {"ok": False, "error": str(e)}

        elif req.op == "list":
            if not file_path or not file_path.is_dir():
                return {"ok": False, "error": "Not a directory"}
            try:
                items = []
                for item in sorted(file_path.iterdir()):
                    items.append({
                        "name": item.name,
                        "type": "dir" if item.is_dir() else "file",
                        "size": item.stat().st_size if item.is_file() else 0,
                    })
                return {"ok": True, "items": items[:100]}
            except Exception as e:
                return {"ok": False, "error": str(e)}

        return {"ok": False, "error": f"Unknown operation: {req.op}"}

    # ═══ Web API (SPA chat) ═══

    @app.get("/api/status")
    async def status(request: Request) -> dict[str, Any]:
        """Get the current status of the life form."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        try:
            if hasattr(hub, 'status'):
                return hub.status()
            return hub.get_status() if hasattr(hub, 'get_status') else {"version": __version__, "online": True}
        except Exception as e:
            return {"error": str(e), "version": __version__, "status": "initializing"}

    @app.get("/api/config/tianditu_key")
    async def tianditu_key(request: Request) -> dict[str, str]:
        """Serve Tianditu API key from encrypted vault (exposed to frontend for map tiles)."""
        try:
            from ..config.secrets import get_secret_vault
            key = get_secret_vault().get("tianditu_key", "")
            # Also try env var override
            if not key:
                key = os.environ.get("LT_TIANDITU_KEY", "")
            return {"key": key}
        except Exception:
            return {"key": ""}

    @app.get("/api/config/tencent_map_key")
    async def tencent_map_key(request: Request) -> dict[str, str]:
        """Serve Tencent Map API key from encrypted vault (exposed to frontend for map tiles)."""
        try:
            from ..config.secrets import get_secret_vault
            key = get_secret_vault().get("tencent_map_key", "")
            if not key:
                key = os.environ.get("LT_TENCENT_MAP_KEY", "")
            return {"key": key}
        except Exception:
            return {"key": ""}

    @app.get("/api/config/baidu_map_key")
    async def baidu_map_key(request: Request) -> dict[str, str]:
        """Serve Baidu Map browser Key from encrypted vault (for frontend map.html)."""
        try:
            from ..config.secrets import get_secret_vault
            key = get_secret_vault().get("baidu_map_browser_key", "")
            if not key:
                key = os.environ.get("LT_BAIDU_MAP_BROWSER_KEY", "")
            return {"key": key}
        except Exception:
            return {"key": ""}

    @app.get("/api/tools")
    async def list_tools(request: Request) -> list[dict[str, Any]]:
        """List all registered tools."""
        hub = request.app.state.hub
        if not hub:
            return []
        try:
            tools = hub.tool_market.discover_tools()
            return [t.model_dump() for t in tools]
        except Exception:
            return []

    @app.get("/api/metrics")
    async def metrics(request: Request) -> dict[str, Any]:
        """Get runtime metrics."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return hub.metrics.get_snapshot() if hub.metrics else {}

    @app.post("/api/report/generate")
    async def generate_report(req: ReportRequest, request: Request) -> dict[str, Any]:
        """Generate an industrial report."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return await hub.generate_report(
            template_type=req.template_type,
            data=req.data,
            requirements=req.requirements,
        )

    @app.post("/api/cell/train")
    async def train_cell(req: TrainRequest, request: Request) -> dict[str, Any]:
        """Train an AI cell."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return await hub.train_cell(
            cell_name=req.cell_name,
            training_data=req.training_data,
            epochs=req.epochs,
        )

    @app.post("/api/drill/train")
    async def drill_train(req: DrillRequest, request: Request) -> dict[str, Any]:
        """Train a cell using MS-SWIFT automated drill pipeline."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return await hub.drill_train(
            cell_name=req.cell_name,
            model_name=req.model_name,
            dataset=req.dataset,
            training_type=req.training_type,
            teacher_model=req.teacher_model,
            reward_model=req.reward_model,
        )

    @app.get("/api/drill/models")
    async def list_drill_models(request: Request) -> list[str]:
        """List MS-SWIFT supported models."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return hub.drill.SUPPORTED_MODELS if hub.drill else []

    @app.get("/api/drill/system")
    async def drill_system_info(request: Request) -> dict[str, Any]:
        """Get system info for training optimization."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return hub.drill._system_info if hub.drill else {}

    @app.post("/api/drill/evaluate")
    async def drill_evaluate(request: Request, model_path: str = "", benchmarks: str = "") -> dict[str, Any]:
        """Evaluate a trained model."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        bench_list = benchmarks.split(",") if benchmarks else None
        return await hub.drill_evaluate(model_path, bench_list)

    @app.post("/api/drill/quantize")
    async def drill_quantize(req: QuantizeRequest, request: Request) -> dict[str, Any]:
        """Quantize a trained model."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return await hub.drill_quantize(req.model_path, req.method)

    @app.post("/api/drill/deploy")
    async def drill_deploy(req: DeployRequest, request: Request) -> dict[str, Any]:
        """Deploy a trained model as API server."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return await hub.drill_deploy(req.model_path, req.port)

    @app.get("/api/drill/queue")
    async def drill_queue(request: Request) -> list[dict[str, Any]]:
        jobs = hub.drill.get_queue() if hub.drill else []
        return [{"model_name": j.model_name, "training_type": j.training_type} for j in jobs]

    @app.get("/api/hitl/pending")
    async def hitl_pending(request: Request) -> list[dict[str, Any]]:
        """List pending HITL approval requests."""
        hub = request.app.state.hub
        if not hub or not hub.world.hitl:
            return []
        return [{"id": r.id, "task": r.task_name, "question": r.question,
                 "status": r.status, "created": r.created} for r in hub.world.hitl.get_pending()]

    @app.post("/api/hitl/deny")
    async def hitl_deny(request: Request) -> dict[str, Any]:
        hub = request.app.state.hub
        body = await request.json()
        if hub and hub.world.hitl:
            ok = hub.world.hitl.deny(body.get("id", ""), body.get("reason", ""))
            return {"denied": ok}
        return {"denied": False}

    @app.get("/api/cost/status")
    async def cost_status(request: Request) -> dict[str, Any]:
        """Get cost/budget status."""
        hub = request.app.state.hub
        if not hub or not hub.world.cost_aware:
            return {"error": "CostAware not available"}
        st = hub.world.cost_aware.status()
        return {"daily_limit": st.daily_limit, "used": st.used_today,
                "remaining": st.remaining, "usage_pct": st.usage_pct,
                "cost_yuan": round(st.cost_yuan, 4), "degraded": st.degraded}

    @app.get("/api/checkpoint/sessions")
    async def checkpoint_sessions(request: Request) -> list[str]:
        hub = request.app.state.hub
        if not hub or not hub.world.checkpoint:
            return []
        return await hub.world.checkpoint.list_sessions()

    # ── Swarm: Direct P2P collaboration endpoints ──

    @app.get("/api/swarm/status")
    async def swarm_status(request: Request) -> dict:
        from ..network.swarm_coordinator import get_swarm
        return get_swarm().status()

    @app.post("/api/swarm/ping")
    async def swarm_ping(request: Request) -> Response:
        """Binary health probe."""
        from ..network.message_bus import decode_health_report, encode_health_report
        raw = await request.body()
        decoded = decode_health_report(raw)
        resp_bin = encode_health_report("ok", "ok", 1.0, [])
        return Response(content=resp_bin, media_type="application/octet-stream")

    @app.post("/api/swarm/cell/receive")
    async def swarm_receive_cell(request: Request) -> Response:
        from ..network.swarm_coordinator import get_swarm
        from ..network.message_bus import encode_task_response
        raw = await request.body()
        result = await get_swarm().receive_cell(raw)
        resp_bin = encode_task_response(
            "cell", "completed" if result.get("cell_name") else "failed",
            str(result.get("cell_name", result.get("error", ""))),
        )
        return Response(content=resp_bin, media_type="application/octet-stream")

    @app.post("/api/swarm/knowledge/receive")
    async def swarm_receive_knowledge(request: Request) -> Response:
        from ..network.swarm_coordinator import get_swarm
        from ..network.message_bus import encode_task_response
        raw = await request.body()
        result = await get_swarm().receive_knowledge(raw)
        resp_bin = encode_task_response("knowledge", "completed", str(result.get("received", 0)))
        return Response(content=resp_bin, media_type="application/octet-stream")

    @app.post("/api/swarm/task/execute")
    async def swarm_execute_task(request: Request) -> Response:
        """Execute a subtask from binary protobuf TaskRequest."""
        from ..network.message_bus import decode_task_distribute, encode_task_response
        hub = request.app.state.hub
        raw = await request.body()
        decoded = decode_task_distribute(raw)
        if not decoded or not hub:
            resp_bin = encode_task_response("task", "failed", "no task or hub")
            return Response(content=resp_bin, media_type="application/octet-stream")
        try:
            subtask = decoded.get("subtask", decoded.get("goal", ""))
            result = await hub.chat(subtask)
            output = str(result.get("intent", result.get("reflections", ["done"])[0]))[:500]
            resp_bin = encode_task_response(decoded.get("task_id", "task"), "completed", output)
        except Exception as e:
            resp_bin = encode_task_response(decoded.get("task_id", "task"), "failed", str(e)[:200])
        return Response(content=resp_bin, media_type="application/octet-stream")

    @app.post("/api/swarm/task/distribute")
    async def swarm_distribute_task(request: Request) -> dict:
        """Distribute a complex task. Accepts JSON for human-facing API."""
        from ..network.swarm_coordinator import get_swarm
        body = await request.json()
        goal = body.get("goal", "")
        peers = body.get("peers", [])
        return await get_swarm().distribute_task(goal, peers)

    @app.delete("/api/checkpoint/{session_id}")
    async def checkpoint_delete(session_id: str, request: Request) -> dict:
        hub = request.app.state.hub
        if hub and hub.world.checkpoint:
            await hub.world.checkpoint.delete(session_id)
            return {"deleted": True}
        return {"deleted": False}

    @app.get("/api/cells")
    async def list_cells(request: Request) -> list[dict[str, Any]]:
        """List all registered AI cells."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        cells = hub.cell_registry.discover()
        return [c.model_dump() if hasattr(c, 'model_dump') else {"name": getattr(c, 'name', 'unknown')} for c in cells]

    @app.post("/api/phage/absorb")
    async def absorb_repo(request: Request, repo_url: str = "") -> dict[str, Any]:
        """Absorb code patterns from a GitHub repository."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        if not repo_url:
            raise HTTPException(status_code=400, detail="repo_url is required")
        return await hub.absorb_github(repo_url)

    @app.get("/api/peers")
    async def list_peers(request: Request) -> list[dict[str, Any]]:
        """List discovered P2P peers."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")
        return await hub.discover_peers()

    @app.post("/api/web/chat")
    async def web_chat(request: Request):
        """SSE streaming endpoint for web frontend. Accepts {messages: [...]}."""
        hub = request.app.state.hub
        if not hub:
            raise HTTPException(status_code=503, detail="Hub not initialized")

        try:
            body = await request.json()
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid JSON")

        messages = body.get("messages", [])
        if not messages:
            raise HTTPException(status_code=400, detail="messages required")

        last_msg = messages[-1].get("content", "") if messages else ""

        # Compress tool outputs in messages to save tokens (9Router RTK-inspired)
        try:
            from livingtree.core.token_compressor import compress
            total_saved = 0
            for msg in messages:
                content = msg.get("content", "")
                if isinstance(content, str) and len(content) > 500:
                    result = compress(content)
                    if result["saved_pct"] > 0:
                        msg["content"] = result["compressed"]
                        total_saved += result["saved_pct"]
            if total_saved > 0:
                logger.debug(f"Token compressor: saved ~{total_saved / len(messages):.0f}% avg across {len(messages)} msgs")
        except Exception:
            pass

        # Inject memory context (non-blocking fire-and-forget)
        try:
            from livingtree.core.session_memory import agent_memory
            context = await agent_memory.get_context_injection(last_msg)
            if context:
                last_msg = f"{context}\n\n---\n用户消息:\n{last_msg}"
        except Exception:
            pass

        from fastapi.responses import StreamingResponse
        import json as _json, time as _time, asyncio as _asyncio, httpx

        async def generate():
            sid = f"web_{int(_time.time()*1000)}"
            thinking = body.get("thinking", False)

            if thinking:
                try:
                    llm = hub.world.consciousness._llm
                    flash_model = await llm.smart_route(last_msg, task_type="chat")
                    if flash_model:
                        async for token in llm.stream(
                            [{"role": "user", "content": last_msg}],
                            provider=flash_model, max_tokens=4096,
                        ):
                            yield f"data: {_json.dumps({'type':'content','content':token,'session_id':sid})}\n\n"
                        yield f"data: {_json.dumps({'type':'done','content':'','session_id':sid})}\n\n"
                        return
                except Exception:
                    pass

            # Fallback: non-streaming via hub chat
            try:
                result = await hub.chat(last_msg)
                reply = result.get("reflections", [""])[0] if result.get("reflections") else ""
                if reply:
                    for i in range(0, len(reply), 2):
                        chunk = reply[i:i+2]
                        yield f"data: {_json.dumps({'type':'content','content':chunk,'session_id':sid})}\n\n"
                        await _asyncio.sleep(0.012)
                    yield f"data: {_json.dumps({'type':'done','content':'','session_id':sid})}\n\n"
                    return
            except Exception:
                pass
            yield f"data: {_json.dumps({'type':'done','content':'模型暂时不可用','session_id':sid})}\n\n"

        return StreamingResponse(generate(), media_type="text/event-stream")

    @app.websocket("/ws")
    async def websocket_endpoint(websocket: WebSocket):
        """WebSocket endpoint for real-time communication."""
        await websocket.accept()
        hub = getattr(websocket.app.state, 'hub', None)

        if not hub:
            await websocket.send_json({"type": "error", "message": "Hub not initialized"})
            await websocket.close()
            return

        await websocket.send_json({
            "type": "connected",
            "node_id": hub.node.info.id if hub.node else "unknown",
            "version": __version__,
        })

        try:
            while True:
                data = await websocket.receive_json()
                msg_type = data.get("type", "chat")

                if msg_type == "chat":
                    message = data.get("message", "")
                    if not message:
                        await websocket.send_json({"type": "error", "message": "Empty message"})
                        continue

                    result = await hub.chat(message)
                    await websocket.send_json({
                        "type": "chat_response",
                        **result,
                    })

                elif msg_type == "status":
                    status = hub.status() if callable(getattr(hub, 'status', None)) else {}
                    await websocket.send_json({"type": "status", **status})

                elif msg_type == "ping":
                    await websocket.send_json({"type": "pong", "timestamp": datetime.now(timezone.utc).isoformat()})

                elif msg_type == "generate_report":
                    template = data.get("template_type", "环评报告")
                    report_data = data.get("data", {})
                    requirements = data.get("requirements", {})
                    result = await hub.generate_report(template, report_data, requirements)
                    await websocket.send_json({"type": "report_response", **result})

                else:
                    await websocket.send_json({"type": "error", "message": f"Unknown message type: {msg_type}"})

        except WebSocketDisconnect:
            logger.info("WebSocket client disconnected")
        except Exception as e:
            logger.error(f"WebSocket error: {e}")
            try:
                await websocket.send_json({"type": "error", "message": str(e)})
            except Exception:
                pass

    # ── IM WebSocket: Instant messaging, voice/video calls, virtual meetings ──

    @app.websocket("/ws/im")
    async def im_websocket(websocket: WebSocket):
        """IM WebSocket — handles DM, group chat, typing, presence, calls, file transfer."""
        await websocket.accept()

        from ..network.im_core import get_im, IMMessage, MessageType, OnlineStatus
        im = get_im()

        params = dict(websocket.query_params or {})
        user_id = params.get("user_id", params.get("uid", ""))
        user_name = params.get("name", user_id[:8] if user_id else "anonymous")

        if not user_id:
            await websocket.send_json({"type": "error", "message": "user_id required"})
            await websocket.close()
            return

        im.connect(user_id, websocket)
        await websocket.send_json({
            "type": "connected", "user_id": user_id,
            "online_count": len(im._connections),
        })

        try:
            while True:
                data = await websocket.receive_json()
                msg_type_str = data.get("type", "")

                if msg_type_str == "dm":
                    msg = IMMessage(
                        msg_id=f"dm_{int(_time.time()*1000)}_{secrets.token_hex(3)}",
                        msg_type=MessageType.DM,
                        sender_id=user_id, sender_name=data.get("sender_name", user_name),
                        receiver_id=data.get("to", data.get("receiver_id", "")),
                        content=data.get("content", data.get("message", "")),
                    )
                    await im.send_message(msg)

                elif msg_type_str == "group_msg":
                    msg = IMMessage(
                        msg_id=f"grp_{int(_time.time()*1000)}_{secrets.token_hex(3)}",
                        msg_type=MessageType.GROUP,
                        sender_id=user_id, sender_name=user_name,
                        receiver_id=data.get("group_id", ""),
                        content=data.get("content", ""),
                    )
                    await im.send_message(msg)

                elif msg_type_str == "typing":
                    msg = IMMessage(
                        msg_id=f"typ_{int(_time.time()*1000)}",
                        msg_type=MessageType.TYPING,
                        sender_id=user_id, receiver_id=data.get("to", ""),
                    )
                    await im.send_message(msg)

                elif msg_type_str == "presence":
                    st = OnlineStatus(data.get("status", "online"))
                    im.set_presence(user_id, st)

                elif msg_type_str == "call_signal":
                    await websocket.send_json({
                        "type": "call_signal_relay",
                        "signal": data.get("signal", {}),
                        "from": user_id,
                    })

                elif msg_type_str == "get_history":
                    msgs = im.get_history(
                        user_id,
                        with_user=data.get("with", ""),
                        group_id=data.get("group_id", ""),
                        limit=data.get("limit", 50),
                    )
                    await websocket.send_json({"type": "history", "messages": msgs})

                elif msg_type_str == "get_contacts":
                    await websocket.send_json({
                        "type": "contacts",
                        "contacts": [c.to_dict() for c in im.get_contacts(user_id)],
                    })

                elif msg_type_str == "add_contact":
                    from ..network.im_core import Contact
                    im.add_contact(user_id, Contact(
                        user_id=data.get("contact_id", ""),
                        name=data.get("contact_name", ""),
                    ))
                    await websocket.send_json({"type": "contact_added", "ok": True})

                elif msg_type_str == "create_group":
                    g = im.create_group(
                        name=data.get("name", "New Group"),
                        owner_id=user_id,
                        members=data.get("members", []),
                    )
                    await websocket.send_json({"type": "group_created", "group": g.to_dict()})

                elif msg_type_str == "get_groups":
                    groups = [g.to_dict() for g in im.get_user_groups(user_id)]
                    await websocket.send_json({"type": "groups", "groups": groups})

                elif msg_type_str == "nearby":
                    nearby = im.get_nearby_users(exclude_user_id=user_id)
                    await websocket.send_json({"type": "nearby", "users": nearby})

                elif msg_type_str == "file_chunk":
                    im.receive_file_chunk(
                        file_id=data.get("file_id", ""),
                        chunk_index=data.get("chunk_index", 0),
                        data=data.get("data", "").encode() if isinstance(data.get("data"), str) else b"",
                    )

                elif msg_type_str == "file_meta":
                    im.start_file_receive(
                        file_id=data.get("file_id", ""),
                        total_chunks=data.get("total_chunks", 1),
                        filename=data.get("filename", "file"),
                        file_size=data.get("file_size", 0),
                    )

                elif msg_type_str == "create_meeting":
                    meeting = await im.create_meeting(
                        topic=data.get("topic", "会议"),
                        host_id=user_id,
                        ai_participants=data.get("ai_participants"),
                    )
                    if hasattr(im, '_hub'):
                        im._hub = getattr(websocket.app.state, 'hub', None)
                    await websocket.send_json({"type": "meeting_created", "meeting": meeting})

                elif msg_type_str == "ping":
                    await websocket.send_json({"type": "pong"})

        except WebSocketDisconnect:
            logger.info(f"IM: {user_id} disconnected")
        except Exception as e:
            logger.error(f"IM WS error: {e}")
        finally:
            im.disconnect(user_id)

    # ── WebRTC Signaling: Browser ↔ Server P2P remote ops ──

    @app.websocket("/ws/rtc-signal")
    async def rtc_signaling(websocket: WebSocket):
        """WebRTC signaling for remote ops. Browser creates offer → server answers."""
        await websocket.accept()

        from ..network.webrtc_remote import get_remote_hub, RemoteOpsSession
        hub = get_remote_hub()
        session_id = f"rtc_{int(_time.time() * 1000)}"
        pc = None

        try:
            data = await websocket.receive_json()
            if "sdp" not in data and "ice" not in data:
                await websocket.close()
                return

            if "sdp" in data:
                sdp_dict = data["sdp"]
                try:
                    import json as _j
                    sdp_dict = _j.loads(sdp_dict) if isinstance(sdp_dict, str) else sdp_dict
                except Exception:
                    pass
                await websocket.send_json({"sdp": {
                    "type": "answer",
                    "sdp": "v=0\r\no=- 0 0 IN IP4 127.0.0.1\r\ns=-\r\nt=0 0\r\na=group:BUNDLE data\r\nm=application 9 DTLS/SCTP 5000\r\nc=IN IP4 0.0.0.0\r\na=ice-pwd:livingtree\r\na=ice-ufrag:ltree\r\na=fingerprint:sha-256 " + "00"*32 + "\r\na=setup:active\r\na=mid:data\r\na=sctp-port:5000\r\n"
                }})
            elif "ice" in data:
                await websocket.send_json({"ice": {"candidate": "candidate:1 1 UDP 2122252543 127.0.0.1 9999 typ host", "sdpMid": "data"}})

            raw = await websocket.receive_text()
            try:
                msg = _json.loads(raw)
            except Exception:
                msg = {"op": "ping"}

            # Create session and handle ops
            class FakeChannel:
                def __init__(self, ws): self._ws = ws
                def send(self, data): asyncio.create_task(self._ws.send_text(data))
            channel = FakeChannel(websocket)
            session = hub.create_session(session_id, channel)

            await session.handle_message(_json.dumps(msg))

            while True:
                raw = await websocket.receive_text()
                try:
                    msg_data = _json.loads(raw)
                    await session.handle_message(_json.dumps(msg_data))
                except Exception:
                    await session.handle_message(raw)

        except WebSocketDisconnect:
            logger.debug(f"RTC signal disconnected: {session_id}")
        except Exception as e:
            logger.debug(f"RTC signal error: {e}")
        finally:
            hub.close_session(session_id)

    # ── Reach Gateway: Cross-device AI sensory extension ──

    @app.websocket("/ws/reach")
    async def reach_websocket(websocket: WebSocket):
        """Reach Gateway WebSocket — cross-device AI sensory network.

        Devices connect here. AI pushes sensor requests (photo, QR, GPS, etc.),
        devices execute and respond. This is how the AI "reaches through" to
        mobile devices for capabilities it doesn't have.
        """
        await websocket.accept()

        from ..network.reach_gateway import get_reach_gateway, DeviceType
        reach = get_reach_gateway()

        hub = getattr(websocket.app.state, 'hub', None)
        if hub:
            reach.set_hub(hub)

        params = dict(websocket.query_params or {})
        device_type = params.get("device_type", "unknown")
        device_name = params.get("device_name", "")
        session_id = params.get("session_id", "")
        ua = websocket.headers.get("user-agent", "")

        device_id = None
        try:
            device = await reach.register_device(
                websocket, device_type, device_name, session_id, ua,
            )
            device_id = device.device_id

            await websocket.send_json({
                "type": "reach_connected",
                "device_id": device_id,
                "device_type": device.device_type.value,
                "capabilities": device.capabilities,
                "message": f"已连接。AI 可以在需要时向你发送任务。",
            })

            while True:
                data = await websocket.receive_json()
                msg_type = data.get("type", "")

                if msg_type == "ping":
                    await websocket.send_json({"type": "pong"})
                    device.last_seen = _time.time()

                elif msg_type == "sensor_response":
                    await reach.receive_response(device_id, data)
                    await websocket.send_json({"type": "ack", "request_id": data.get("request_id", "")})

                elif msg_type == "capabilities":
                    device.capabilities = data.get("capabilities", device.capabilities)
                    await websocket.send_json({"type": "caps_updated"})

                elif msg_type == "task_card_result":
                    card_id = data.get("card_id", "")
                    card = reach._pending_cards.get(card_id)
                    if card:
                        card.status = "completed"
                        card.result = data.get("result", {})
                    await websocket.send_json({"type": "card_ack", "card_id": card_id})

                else:
                    await websocket.send_json({"type": "error", "message": f"Unknown: {msg_type}"})

        except WebSocketDisconnect:
            logger.info(f"Reach: device {device_id or '?'} disconnected")
        except Exception as e:
            logger.error(f"Reach WS error: {e}")
        finally:
            if device_id:
                reach.unregister_device(device_id)

    # ── Voice Call WebSocket ──

    @app.websocket("/ws/voice")
    async def voice_call_websocket(websocket: WebSocket):
        """Voice call WebSocket — bilateral real-time voice conversation.

        Client sends browser-STT text, 小树 responds with streaming TTS audio.
        Audio is NEVER stored to disk. Voice persona: 活跃可爱温暖女声.
        """
        from ..core.voice_call import get_voice_call_engine
        engine = get_voice_call_engine()
        await engine.handle(websocket)

    # ── OpenAI-compatible relay for external tools (opencode, Claude Code, etc.) ──

    @app.post("/v1/chat/completions")
    async def openai_chat_completions(request: Request):
        """OpenAI-compatible endpoint. Routes through LivingTree's auto-election."""
        hub = getattr(request.app.state, 'hub', None)
        if not hub:
            from fastapi import HTTPException
            raise HTTPException(status_code=503, detail="Hub not initialized")

        try:
            body = await request.json()
        except Exception:
            from fastapi import HTTPException
            raise HTTPException(status_code=400, detail="Invalid JSON")

        messages = body.get("messages", [])
        if not messages:
            from fastapi import HTTPException
            raise HTTPException(status_code=400, detail="messages required")

        last_msg = messages[-1].get("content", "") if messages else ""
        stream = body.get("stream", False)

        if stream:
            from fastapi.responses import StreamingResponse
            import json as _json

            async def generate():
                try:
                    async for token in hub.world.consciousness.stream_of_thought(last_msg):
                        chunk = _json.dumps({
                            "choices": [{"delta": {"content": token}, "index": 0}],
                            "object": "chat.completion.chunk",
                        })
                        yield f"data: {chunk}\n\n"
                    yield "data: [DONE]\n\n"
                except Exception as e:
                    error_chunk = _json.dumps({"error": str(e)})
                    yield f"data: {error_chunk}\n\n"

            return StreamingResponse(generate(), media_type="text/event-stream")

        result = await hub.chat(last_msg)
        reply = str(result.get("intent", "") or result.get("reflections", [""])[0] or "")
        return {
            "id": f"livingtree-{result.get('session_id', '')}",
            "object": "chat.completion",
            "created": int(__import__('time').time()),
            "model": "livingtree-dual",
            "choices": [{
                "index": 0,
                "message": {"role": "assistant", "content": reply},
                "finish_reason": "stop",
            }],
            "usage": {
                "prompt_tokens": len(last_msg),
                "completion_tokens": len(reply),
                "total_tokens": len(last_msg) + len(reply),
            },
        }

    @app.get("/v1/models")
    async def openai_list_models(request: Request):
        return {
            "object": "list",
            "data": [
                {"id": "livingtree-dual", "object": "model"},
                {"id": "livingtree-flash", "object": "model"},
                {"id": "longcat-flash-lite", "object": "model"},
                {"id": "longcat-flash-chat", "object": "model"},
            ],
        }

    # ═══ Scinet Service ═══

    @app.get("/api/scinet/status")
    async def scinet_status(request: Request):
        from ..network.scinet_service import get_scinet
        s = get_scinet().get_status()
        return {
            "running": s.running,
            "port": s.port,
            "pac_url": s.pac_url,
            "uptime_seconds": s.uptime_seconds,
            "total_requests": s.total_requests,
            "success_requests": s.success_requests,
            "failed_requests": s.failed_requests,
            "bandwidth_mb": round(s.bandwidth_bytes / 1024 / 1024, 2),
        }

    @app.post("/api/scinet/start")
    async def scinet_start(request: Request):
        from ..network.scinet_service import get_scinet
        s = await get_scinet().start()
        return {"status": "started", "port": s.port, "pac_url": s.pac_url}

    @app.post("/api/scinet/stop")
    async def scinet_stop(request: Request):
        from ..network.scinet_service import get_scinet
        s = await get_scinet().stop()
        return {"status": "stopped", "uptime_seconds": s.uptime_seconds, "total_requests": s.total_requests}

    @app.get("/api/scinet/pac")
    async def scinet_pac(request: Request):
        from ..network.scinet_service import get_scinet
        return Response(
            content=get_scinet().generate_pac(),
            media_type="application/x-ns-proxy-autoconfig",
        )

    @app.get("/api/scinet/test")
    async def scinet_test(request: Request):
        from ..network.scinet_service import get_scinet
        import aiohttp, time
        results = {}
        for url in ["https://github.com", "https://huggingface.co", "https://stackoverflow.com"]:
            try:
                async with aiohttp.ClientSession() as session:
                    start = time.time()
                    async with session.get(url, timeout=aiohttp.ClientTimeout(total=5)) as resp:
                        results[url] = f"OK ({resp.status}, {(time.time()-start)*1000:.0f}ms)"
            except Exception as e:
                results[url] = f"FAIL ({e})"
        return results

    # ── Register doc routes (LT-Office integration) ──
    from .doc_routes import setup_doc_routes
    setup_doc_routes(app)

    # ── Memory management routes ──

    @app.get("/api/memory/status")
    async def memory_status(request: Request) -> dict[str, Any]:
        """Get memory engine status."""
        try:
            from livingtree.core.session_memory import agent_memory, _cognee_available
            await agent_memory._ensure_init()
            entries = agent_memory._load_fallback()
            return {
                "cognee_available": _cognee_available,
                "fallback_entries": len(entries),
                "status": "ok",
            }
        except Exception as e:
            return {"status": "error", "error": str(e)}

    @app.post("/api/memory/remember")
    async def memory_remember(request: Request) -> dict[str, Any]:
        """Store a memory entry."""
        try:
            body = await request.json()
            from livingtree.core.session_memory import agent_memory
            ok = await agent_memory.remember(
                body.get("content", ""),
                user_id=body.get("user_id", ""),
                project=body.get("project", ""),
                workspace_id=body.get("workspace_id", ""),
            )
            return {"ok": ok}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.post("/api/memory/recall")
    async def memory_recall(request: Request) -> dict[str, Any]:
        """Recall relevant memories."""
        try:
            body = await request.json()
            from livingtree.core.session_memory import agent_memory
            results = await agent_memory.recall(
                body.get("query", ""),
                user_id=body.get("user_id", ""),
                project=body.get("project", ""),
                workspace_id=body.get("workspace_id", ""),
                limit=body.get("limit", 5),
            )
            return {"ok": True, "results": results}
        except Exception as e:
            return {"ok": False, "error": str(e), "results": []}

    @app.post("/api/memory/forget")
    async def memory_forget(request: Request) -> dict[str, Any]:
        """Forget memories."""
        try:
            body = await request.json()
            from livingtree.core.session_memory import agent_memory
            ok = await agent_memory.forget(
                user_id=body.get("user_id", ""),
                project=body.get("project", ""),
            )
            return {"ok": ok}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ── Skill management routes ──

    @app.get("/api/skills")
    async def list_all_skills(
        request: Request,
        workspace_id: str = Query(default=""),
        tag: str = Query(default=""),
        search: str = Query(default=""),
    ) -> list[dict]:
        """List skills for current user or workspace."""
        user_id = _get_user_id_from_request(request)
        from livingtree.core.skills import list_skills
        return list_skills(
            user_id=user_id if not workspace_id else "",
            workspace_id=workspace_id,
            tag=tag,
            search=search,
        )

    @app.get("/api/skills/{name}")
    async def get_single_skill(
        name: str,
        request: Request,
        workspace_id: str = Query(default=""),
    ) -> dict:
        """Get a single skill by name."""
        user_id = _get_user_id_from_request(request)
        from livingtree.core.skills import get_skill
        skill = get_skill(name, user_id=user_id if not workspace_id else "", workspace_id=workspace_id)
        if not skill:
            raise HTTPException(status_code=404, detail="技能不存在")
        return skill

    @app.post("/api/skills")
    async def create_or_update_skill(request: Request) -> dict:
        """Create or update a skill."""
        try:
            body = await request.json()
            user_id = _get_user_id_from_request(request)
            from livingtree.core.skills import create_or_update_skill
            return create_or_update_skill(
                name=body.get("name", ""),
                body=body.get("body", ""),
                user_id=user_id if not body.get("workspace_id") else "",
                workspace_id=body.get("workspace_id", ""),
                description=body.get("description", ""),
                tags=body.get("tags", []),
                source_project=body.get("source_project", ""),
                source_user=body.get("source_user", user_id),
            )
        except Exception as e:
            raise HTTPException(status_code=400, detail=str(e))

    @app.delete("/api/skills/{name}")
    async def delete_skill_by_name(
        name: str,
        request: Request,
        workspace_id: str = Query(default=""),
    ) -> dict:
        """Delete a skill."""
        user_id = _get_user_id_from_request(request)
        from livingtree.core.skills import delete_skill
        ok = delete_skill(name, user_id=user_id if not workspace_id else "", workspace_id=workspace_id)
        return {"ok": ok}

    @app.post("/api/skills/{name}/touch")
    async def touch_skill_usage(
        name: str,
        request: Request,
        workspace_id: str = Query(default=""),
    ) -> dict:
        """Increment skill usage count."""
        user_id = _get_user_id_from_request(request)
        from livingtree.core.skills import touch_skill
        ok = touch_skill(name, user_id=user_id if not workspace_id else "", workspace_id=workspace_id)
        return {"ok": ok}

    @app.get("/api/skills/suggestions")
    async def skill_suggestions(
        request: Request,
        workspace_id: str = Query(default=""),
    ) -> list[dict]:
        """Get skill suggestions from recorded sessions."""
        user_id = _get_user_id_from_request(request)
        from livingtree.core.skills import suggest_skills_from_sessions
        return suggest_skills_from_sessions(
            user_id=user_id if not workspace_id else "",
            workspace_id=workspace_id,
        )

    @app.get("/api/skills/dedup")
    async def skill_deduplication(
        request: Request,
        workspace_id: str = Query(default=""),
    ) -> dict:
        """Find duplicate skills."""
        user_id = _get_user_id_from_request(request)
        from livingtree.core.skills import deduplicate_skills
        return deduplicate_skills(
            user_id=user_id if not workspace_id else "",
            workspace_id=workspace_id,
        )


    @app.get("/api/docs/templates")
    async def list_doc_templates() -> list[dict]:
        """List available Kami document templates."""
        from livingtree.core.doc_renderer import list_templates
        return list_templates()

    @app.post("/api/docs/render")
    async def render_document_endpoint(request: Request) -> dict:
        """Render Markdown content to Kami-styled HTML/PDF."""
        try:
            body = await request.json()
            from livingtree.core.doc_renderer import render_document
            return render_document(
                content=body.get("content", ""),
                template=body.get("template", "long_doc"),
                title=body.get("title", ""),
                output_path=body.get("output_path", ""),
            )
        except Exception as e:
            return {"ok": False, "error": str(e)}


async def _handle_relay_config(msg: str, ml: str):
    """Handle relay server configuration commands via admin chat."""
    import re as _regex
    import aiohttp
    actions, reply = [], ""

    relay_url = os.environ.get("LT_RELAY_URL", "http://127.0.0.1:8899")
    admin_pwd = os.environ.get("LT_RELAY_ADMIN_PWD", "")
    if not admin_pwd:
        try:
            from ..config.secrets import get_secret_vault
            admin_pwd = get_secret_vault().get("relay_admin_pwd", "")
        except Exception:
            pass

    async def _call_relay(method: str, path: str, data: dict | None = None) -> dict:
        try:
            async with aiohttp.ClientSession() as session:
                if method == "GET":
                    async with session.get(f"{relay_url}{path}", timeout=aiohttp.ClientTimeout(total=10)) as resp:
                        return await resp.json() if resp.content_type == "application/json" else {"status": resp.status}
                elif method == "POST":
                    async with session.post(f"{relay_url}{path}", json=data or {}, timeout=aiohttp.ClientTimeout(total=10)) as resp:
                        return await resp.json() if resp.content_type == "application/json" else {"status": resp.status}
        except Exception as e:
            return {"error": str(e)}

    if "添加账户" in msg or "创建用户" in msg or "注册账户" in msg or "add account" in ml:
        username = ""
        password = ""
        um = _regex.search(r'(?:用户|账户|账号|用户名)[:：\s]*(\S+)', msg)
        pm = _regex.search(r'(?:密码|口令)[:：\s]*(\S+)', msg)
        if um:
            username = um.group(1)
            if pm:
                password = pm.group(1)
            else:
                import secrets
                password = secrets.token_hex(6)
                actions.append(f"已生成随机密码: {password}")

            # Login as admin first
            r = await _call_relay("POST", "/admin/login", {"username": "admin", "password": admin_pwd})
            if not r.get("error"):
                r2 = await _call_relay("POST", "/admin/accounts/add", {"username": username, "password": password, "display_name": username})
                if r2.get("status") == 200:
                    actions.append(f"✅ 中继账户已创建: {username}")
                    reply = f"中继账户 {username} 已创建。密码: {password}"
                else:
                    reply = f"创建失败: {r2}"
                    actions.append(str(r2))
            else:
                reply = "中继服务器连接失败。确认中继已启动: python relay_server.py"
        else:
            reply = "请提供用户名，如: 创建用户 alice 密码 123456"

    elif "查看账户" in msg or "账户列表" in msg or "list" in ml:
        r = await _call_relay("GET", "/status")
        reply = f"中继状态: {r.get('peers','?')} 个节点在线"
        actions.append(f"中继地址: {relay_url}")

    elif "添加订阅" in msg or "添加模型" in msg or "subscription" in ml:
        base_url = ""
        api_key = ""
        model = ""
        provider = _regex.search(r'(?:提供方|provider|名称)[:：\s]*(\S+)', msg)
        bm = _regex.search(r'(?:地址|url|base)[:：\s]*(https?://[^\s]+)', msg)
        km = _regex.search(r'(?:key|密钥|api)[:：\s]*(sk-\S+)', msg)
        mm = _regex.search(r'(?:模型|model)[:：\s]*(\S+)', msg)
        if bm:
            base_url = bm.group(1)
            api_key = km.group(1) if km else ""
            model = mm.group(1) if mm else "default"
            name = provider.group(1) if provider else base_url.split("//")[1].split(".")[0][:10]
            r = await _call_relay("POST", "/admin/subscriptions/add", {
                "name": name, "base_url": base_url, "api_key": api_key,
                "model": model, "user": "admin",
            })
            actions.append(f"✅ 订阅已添加: {name} ({model})")
            reply = f"中继订阅 {name} 已配置: {base_url}"
        else:
            reply = "请提供订阅信息，如: 添加订阅 名称 deepseek 地址 https://api.deepseek.com/v1 key sk-xxx 模型 deepseek-chat"

    elif "查看节点" in msg or "节点列表" in msg or "在线" in msg or "peers" in ml:
        r = await _call_relay("GET", "/peers/discover")
        peers = r if isinstance(r, list) else r.get("peers", []) if isinstance(r, dict) else []
        count = len(peers) if isinstance(peers, list) else 0
        reply = f"中继在线节点: {count} 个"
        for p in (peers if isinstance(peers, list) else [])[:5]:
            actions.append(f"📡 {p.get('peer_id','?')[:12]} — {p.get('metadata',{}).get('platform','?')}")

    elif "状态" in msg or "status" in ml:
        r = await _call_relay("GET", "/health")
        reply = f"中继运行中: {r.get('status','?')} · {r.get('peers',0)} 节点 · {r.get('uptime',0):.0f}s"
        r2 = await _call_relay("GET", "/status")
        if isinstance(r2, dict):
            actions.append(f"请求: {r2.get('request_count',0)} · 用户: {r2.get('username','?')}")

    elif "密码" in msg and "relay" in ml:
        reply = f"中继管理员默认密码: {admin_pwd}。修改: POST /admin/password"

    else:
        r = await _call_relay("GET", "/health")
        if r.get("status") == "ok":
            reply = f"中继服务器运行正常 ({relay_url})。可用命令: 创建用户/查看节点/添加订阅"
        else:
            reply = f"中继服务器未响应 ({relay_url})。启动: python relay_server.py --port 8899"

    return actions, reply


def _get_user_id_from_request(request: Request) -> str:
    """Extract user_id from JWT token in request headers."""
    try:
        from livingtree.api.auth import verify_token
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
            payload = verify_token(token)
            if payload:
                return payload.get("user_id", "")
    except Exception:
        pass
    return ""

    # ── Document parsing routes ──

    @app.post("/api/docs/parse")
    async def parse_uploaded_doc(
        request: Request,
        file: UploadFile = File(...),
        project: str = Form(default=""),
        workspace_id: str = Form(default=""),
    ) -> dict:
        """Upload a document and parse it with MinerU. Results fed into agent memory."""
        from pathlib import Path
        import time as _upload_time
        user_id = _get_user_id_from_request(request)
        uploads_dir = Path(__file__).resolve().parent.parent.parent / "data" / "uploads"
        uploads_dir.mkdir(parents=True, exist_ok=True)
        tmp_path = uploads_dir / f"{int(_upload_time.time())}_{file.filename}"
        try:
            content = await file.read()
            tmp_path.write_bytes(content)
            from livingtree.core.session_memory import ingest_document
            result = await ingest_document(
                str(tmp_path), user_id=user_id, project=project, workspace_id=workspace_id)
            if result.get("ok"):
                tmp_path.unlink(missing_ok=True)
            return result
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.post("/api/docs/parse-file")
    async def parse_local_doc(request: Request) -> dict:
        """Parse a local file path with MinerU."""
        try:
            body = await request.json()
            user_id = _get_user_id_from_request(request)
            from livingtree.core.session_memory import ingest_document
            return await ingest_document(
                body.get("path", ""),
                user_id=user_id,
                project=body.get("project", ""),
                workspace_id=body.get("workspace_id", ""),
            )
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ═══ Admin API ═══

    @app.get("/api/admin")
    async def admin_page():
        from pathlib import Path
        template_dir = Path(__file__).resolve().parent.parent / "templates"
        from fastapi.responses import HTMLResponse
        return HTMLResponse((template_dir / "admin.html").read_text(encoding="utf-8"))

    @app.get("/api/admin/status")
    async def admin_status(request: Request):
        from ..core.admin_manager import get_admin
        admin = get_admin()
        hub = getattr(request.app.state, "hub", None)
        config_info = {}
        if hub and hasattr(hub, "config"):
            cfg = hub.config
            config_info = {
                "flash_model": getattr(getattr(cfg, "model", None), "flash_model", ""),
                "pro_model": getattr(getattr(cfg, "model", None), "pro_model", ""),
                "api_port": getattr(getattr(cfg, "api", None), "port", 8100) if hasattr(cfg, "api") else 8100,
            }
        return {**admin.status(), "config": config_info}

    class AdminPassword(BaseModel):
        password: str

    @app.post("/api/admin/init")
    async def admin_init(req: AdminPassword):
        from ..core.admin_manager import get_admin
        ok = get_admin().initialize(req.password)
        return {"ok": ok, "error": "" if ok else "已初始化或密码过短(至少6位)"}

    @app.post("/api/admin/login")
    async def admin_login(req: AdminPassword):
        from ..core.admin_manager import get_admin
        a = get_admin()
        if not a.is_initialized:
            return {"token": "", "error": "未初始化"}
        if a.verify_password(req.password):
            return {"token": a.create_admin_token()}
        return {"token": "", "error": "密码错误"}

    class ChangePassword(BaseModel):
        old_password: str
        new_password: str

    @app.post("/api/admin/change-password")
    async def admin_change_password(req: ChangePassword, request: Request):
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        return {"ok": a.change_password(req.old_password, req.new_password)}

    class CredentialReq(BaseModel):
        key: str
        value: str = ""

    @app.get("/api/admin/credentials")
    async def admin_list_credentials(request: Request):
        from fastapi.responses import HTMLResponse
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            return HTMLResponse('<div style="color:var(--err)">未授权</div>')
        keys = a.list_credential_keys()
        rows = ""
        for k in keys[:30]:
            rows += (
                f'<div class="cred-row"><span>{k}</span>'
                f'<span class="masked">●●●●●●●●</span>'
                f'<button class="danger" onclick="deleteCredential(\'{k}\')" style="font-size:10px;padding:2px 8px">删除</button></div>'
            )
        return HTMLResponse(rows or '<div style="color:var(--dim);font-size:12px">暂无存储的密钥</div>')

    @app.post("/api/admin/credentials")
    async def admin_add_credential(req: CredentialReq, request: Request):
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        return {"ok": a.store_credential(req.key, req.value)}

    @app.delete("/api/admin/credentials/{key}")
    async def admin_delete_credential(key: str, request: Request):
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        return {"ok": a.delete_credential(key)}

    class ConfigChatReq(BaseModel):
        message: str

    @app.post("/api/admin/config/chat")
    async def admin_config_chat(req: ConfigChatReq, request: Request):
        import re as _regex
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)

        hub = request.app.state.hub
        if not hub:
            return {"reply": "系统未就绪", "actions": []}

        msg = req.message
        actions = []
        reply = ""
        world = hub.world
        consc = getattr(world, "consciousness", None) if world else None
        ml = msg.lower().replace(" ", "")

        if "apikey" in ml or "api_key" in ml or "密钥" in msg or "sk-" in msg:
            km = _regex.search(r'(?:api[_ ]?key|密钥|key)\s*[:：=]\s*["\']?(\S+)', msg, _regex.IGNORECASE)
            if not km:
                km = _regex.search(r'(sk-[a-zA-Z0-9]+)', msg)
            if km:
                key_val = km.group(1)
                pv = "deepseek"
                for p in ["deepseek", "openai", "longcat", "qwen", "zhipu", "siliconflow"]:
                    if p in ml:
                        pv = p
                        break
                a.store_credential(f"{pv}_api_key", key_val)
                actions.append(f"已加密保存 {pv}_api_key")
                reply = f"✅ API密钥已加密保存。提供方: {pv}"
            else:
                reply = "请提供完整的API密钥(如 api_key=sk-xxx)"

        elif "扫描" in msg or "scan" in ml or "发现" in msg:
            from ..core.universal_scanner import get_universal_scanner
            scanner = get_universal_scanner()
            scanner._hub = hub
            found = await scanner.discover_from_description(msg)
            if found and found.is_alive:
                reply = f"发现: {found.name} ({found.protocol})"
                if found.models:
                    reply += f" — {len(found.models)}个模型"
            else:
                reply = "未发现。请提供具体地址"

        elif "模型" in msg or "model" in ml or "默认" in msg:
            if consc:
                reply = f"Flash: {getattr(consc, 'flash_model', '?')}\nPro: {getattr(consc, 'pro_model', '?')}"

        elif "预算" in msg or "budget" in ml or "token" in ml or "费用" in msg:
            ca = getattr(world, "cost_aware", None)
            if ca:
                st = ca.status()
                reply = f"已用 {st.used_today:,} tokens · ¥{st.cost_yuan:.4f}"
            else:
                reply = "费用模块未就绪"

        elif "宪法" in msg or "spec" in ml or "价值观" in msg or "行为准则" in msg:
            from ..core.model_spec import get_spec
            spec = get_spec()
            if "修改" in msg or "更新" in msg or "update" in ml:
                content = msg.split("修改", 1)[-1].strip() if "修改" in msg else msg
                if len(content) > 10:
                    spec.update_spec(content if content.startswith("#") else f"# Updated Spec\n\n{content}")
                    actions.append(f"宪法已更新至 v{spec._version}")
                    reply = f"✅ Model Spec 已更新 (v{spec._version})"
                else:
                    reply = "请提供具体的宪法内容"
            else:
                info = spec.get_spec_for_admin()
                reply = f"Model Spec v{info['version']} — {info['values_count']} 个价值观 · {info['size_chars']} 字符"
                actions.append(f"Hash: {info['hash']}")

        elif "中继" in msg or "relay" in ml or "节点" in msg:
            actions, reply = await _handle_relay_config(msg, ml)
            if not reply:
                reply = "中继配置完成"

        else:
            if consc:
                try:
                    rt = await consc.chain_of_thought(
                        f"系统配置助手。管理员: {msg}\n返回JSON: {{'reply':'回复','actions':[]}}", steps=1,
                    )
                    txt = rt if isinstance(rt, str) else str(rt)
                    try:
                        import json as _j
                        d = _j.loads(txt[txt.find("{"):txt.rfind("}") + 1])
                        reply = d.get("reply", "已处理")
                        actions = d.get("actions", [])
                    except Exception:
                        reply = txt[:400]
                except Exception:
                    reply = "配置助手暂不可用"

        return {"reply": reply or "已处理", "actions": actions}

    class SvcDiscover(BaseModel):
        description: str

    @app.post("/api/admin/services/discover")
    async def admin_svc_discover(req: SvcDiscover, request: Request):
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        from ..core.universal_scanner import get_universal_scanner
        sc = get_universal_scanner()
        sc._hub = request.app.state.hub
        svc = await sc.discover_from_description(req.description)
        return {"service": svc.to_dict() if svc else None}

    class SvcRegister(BaseModel):
        url: str
        api_key: str = ""

    @app.post("/api/admin/services/register")
    async def admin_svc_register(req: SvcRegister, request: Request):
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        from ..core.universal_scanner import get_universal_scanner
        sc = get_universal_scanner()
        sc._hub = request.app.state.hub
        svc = sc._discovered.get(req.url)
        if not svc:
            return {"ok": False, "error": "服务未发现"}
        return await sc.auto_register_service(svc, api_key=req.api_key)

    @app.get("/api/admin/services/status")
    async def admin_svc_status(request: Request):
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        from ..core.universal_scanner import get_universal_scanner
        return get_universal_scanner().status()

    @app.get("/api/admin/services/scan")
    async def admin_svc_scan(request: Request):
        from ..core.admin_manager import get_admin
        a = get_admin()
        t = (request.headers.get("Authorization", "")).replace("Bearer ", "")
        if not a.verify_admin_token(t):
            raise HTTPException(status_code=401)
        from ..core.universal_scanner import get_universal_scanner
        sc = get_universal_scanner()
        sc._hub = request.app.state.hub
        discovered = await sc.scan_network(max_ports=50)
        return {"discovered": len(discovered), "services": [s.to_dict() for s in discovered]}

    # ═══ HITL Bridge ═══

    @app.post("/api/hitl/approve")
    async def hitl_approve(request: Request):
        from ..core.agent_qa import get_hitl_bridge
        from ..core.dpo_prefs import get_dpo_hooks
        body = await request.json()
        ok = get_hitl_bridge().approve(body.get("id", ""), body.get("response", "approved"))
        get_dpo_hooks().on_hitl_decision(body.get("id", ""), True, body.get("context", ""))
        return {"ok": ok}

    @app.post("/api/hitl/reject")
    async def hitl_reject(request: Request):
        from ..core.agent_qa import get_hitl_bridge
        from ..core.dpo_prefs import get_dpo_hooks
        body = await request.json()
        ok = get_hitl_bridge().reject(body.get("id", ""), body.get("reason", "rejected"))
        get_dpo_hooks().on_hitl_decision(body.get("id", ""), False, body.get("context", ""))
        return {"ok": ok}

    # ═══ Chrome DevTools MCP ═══

    @app.post("/api/chrome/start")
    async def chrome_start(request: Request):
        from ..core.chrome_dual import get_chrome_dual
        body = await request.json()
        headless = body.get("headless", True)
        bridge = get_chrome_dual()
        await bridge.probe()
        result = await bridge.start(headless=headless)
        return result

    @app.post("/api/chrome/stop")
    async def chrome_stop(request: Request):
        from ..core.chrome_dual import get_chrome_dual
        return await get_chrome_dual().stop()

    @app.post("/api/chrome/screenshot")
    async def chrome_screenshot(request: Request):
        from ..core.chrome_dual import get_chrome_dual
        body = await request.json()
        url = body.get("url", "")
        bridge = get_chrome_dual()
        if url:
            await bridge.navigate(url)
            await asyncio.sleep(1)
        return await bridge.screenshot(
            selector=body.get("selector", ""),
            full_page=body.get("full_page", False),
        )

    @app.post("/api/chrome/eval")
    async def chrome_eval(request: Request):
        from ..core.chrome_dual import get_chrome_dual
        body = await request.json()
        return await get_chrome_dual().eval_js(body.get("expression", "document.title"))

    @app.post("/api/chrome/navigate")
    async def chrome_navigate(request: Request):
        from ..core.chrome_dual import get_chrome_dual
        body = await request.json()
        return await get_chrome_dual().navigate(body.get("url", "about:blank"))

    @app.post("/api/chrome/audit")
    async def chrome_audit(request: Request):
        from ..core.chrome_dual import get_chrome_dual
        return await get_chrome_dual().accessibility_audit()

    @app.post("/api/chrome/dom")
    async def chrome_dom(request: Request):
        from ..core.chrome_dual import get_chrome_dual
        body = await request.json()
        return await get_chrome_dual().dom_query(body.get("selector", "body"))

    # ═══ Voice / Speech API ═══

    @app.post("/api/speech/transcribe")
    async def speech_transcribe(request: Request):
        from ..core.unified_speech import get_speech_pipeline
        body = await request.json()
        audio_b64 = body.get("audio", "")
        audio_format = body.get("format", "webm")

        if not audio_b64:
            return {"ok": False, "error": "no audio data", "text": ""}

        pipeline = get_speech_pipeline()
        if not pipeline._probed:
            await pipeline.probe()

        audio_bytes = base64.b64decode(audio_b64)
        return await pipeline.transcribe_direct(audio_bytes, format=audio_format)

    # ═══ Knowledge Ingest — pseudo-upload (no disk storage) ═══

    @app.post("/api/knowledge/ingest")
    async def knowledge_ingest(request: Request):
        """Pseudo-upload: extract text from files/audio/video in memory, store in KB.

        Original file bytes are NEVER saved to disk. Only the extracted text is
        stored in the knowledge base. Supports: txt/md/pdf/docx/csv/html +
        audio (STT) + video (extract audio → STT).
        """
        from ..core.inline_parser import get_inline_parser
        body = await request.json()
        data_b64 = body.get("data", "")
        filename = body.get("filename", "untitled")
        mime_type = body.get("mime_type", "")
        domain = body.get("domain", "user_upload")

        if not data_b64:
            return {"ok": False, "error": "no file data"}

        try:
            data = base64.b64decode(data_b64)
        except Exception:
            return {"ok": False, "error": "invalid base64"}

        parser = get_inline_parser()
        result = await parser.parse(data, filename, mime_type)

        if not result.ok:
            return {"ok": False, "error": result.error}

        doc_id = ""
        try:
            from ..knowledge.knowledge_base import Document, get_knowledge_base
            kb = get_knowledge_base()
            doc = Document(
                title=result.title,
                content=result.text,
                domain=domain,
                source="pseudo-upload",
                author="user",
                metadata={
                    "source_format": result.source_format,
                    "word_count": result.word_count,
                    "parse_time_ms": result.parse_time_ms,
                    "original_filename": filename,
                    "mime_type": mime_type,
                },
            )
            doc_id = kb.add_document(doc)
        except Exception as e:
            logger.warning(f"Knowledge base store failed: {e}")
            doc_id = ""

        return {
            "ok": True,
            "doc_id": doc_id,
            "title": result.title,
            "text_preview": result.text[:500],
            "word_count": result.word_count,
            "source_format": result.source_format,
            "parse_time_ms": round(result.parse_time_ms, 1),
        }

    # ═══ Video Search API ═══

    # ═══ Vitals API (Living Pot hardware) ═══

    @app.get("/api/status/vitals")
    async def status_vitals(request: Request):
        from ..core.vitals import get_vitals
        return get_vitals().measure()

    @app.get("/api/status/vitals/hardware")
    async def status_vitals_hardware(request: Request):
        from ..core.vitals import get_vitals
        return get_vitals().hardware_json()

    # ═══ City MCP API ═══

    @app.get("/api/city/tools")
    async def city_mcp_tools(request: Request):
        from ..core.city_mcp import get_city_mcp
        return {"ok": True, "tools": get_city_mcp().list_tools()}

    @app.post("/api/city/call")
    async def city_mcp_call(request: Request):
        from ..core.city_mcp import get_city_mcp
        body = await request.json()
        tool = body.get("tool", "")
        params = body.get("params", {})
        return await get_city_mcp().call_tool(tool, params)

    # ═══ Unified MCP API (aggregates all MCP tool sources) ═══

    @app.get("/api/mcp/tools")
    async def mcp_tools(request: Request):
        """List all MCP tools from all sources (server + chrome + city)."""
        tools = []
        try:
            from ..mcp.server import MCPServer
            tools.extend(MCPServer.TOOLS)
        except Exception:
            pass
        try:
            from ..core.chrome_mcp import CHROME_MCP_TOOLS
            tools.extend(CHROME_MCP_TOOLS)
        except Exception:
            pass
        try:
            from ..core.city_mcp import CITY_MCP_TOOLS
            tools.extend(CITY_MCP_TOOLS)
        except Exception:
            pass
        return {"ok": True, "tools": [{"name": t.get("name",""), "description": t.get("description","")} for t in tools], "count": len(tools)}

    @app.post("/api/mcp/call")
    async def mcp_call(request: Request):
        """Call an MCP tool by name, routed to the correct handler."""
        body = await request.json()
        tool = body.get("tool", "")
        params = body.get("params", {})
        try:
            from ..mcp.server import MCPServer
            for t in MCPServer.TOOLS:
                if t.get("name") == tool:
                    handler = t.get("handler")
                    if handler:
                        result = handler(**params) if callable(handler) else str(handler)
                        return {"ok": True, "result": result}
        except Exception:
            pass
        try:
            from ..core.chrome_mcp import get_chrome_mcp
            for t in getattr(get_chrome_mcp(), 'CHROME_MCP_TOOLS', []):
                if t.get("name") == tool:
                    return await get_chrome_mcp().execute(tool, params)
        except Exception:
            pass
        try:
            from ..core.city_mcp import get_city_mcp
            return await get_city_mcp().call_tool(tool, params)
        except Exception:
            pass
        return {"ok": False, "error": f"Tool '{tool}' not found"}

    # ═══ Green Scheduler API ═══

    @app.get("/api/scheduler/status")
    async def scheduler_status(request: Request):
        from ..core.green_scheduler import get_green_scheduler
        return get_green_scheduler().stats()

    # ═══ Prompt Shield API ═══

    @app.get("/api/shield/status")
    async def shield_status(request: Request):
        from ..core.prompt_shield import get_shield
        return get_shield().stats()

    # ═══ Virtual FS API ═══

    @app.post("/api/vfs/execute")
    async def vfs_execute(request: Request):
        """Execute a command on the Virtual File System."""
        body = await request.json()
        command = body.get("command", "")
        if not command:
            return {"error": "no command"}
        try:
            from ..capability.virtual_fs import get_virtual_fs
            vfs = get_virtual_fs()
            result = await vfs.execute(command)
            return {"ok": True, "output": result[:5000]}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ═══ Wiki Search API ═══

    @app.post("/api/wiki/search")
    async def wiki_search(request: Request):
        body = await request.json()
        query = body.get("query", "")
        path = body.get("path", "")
        section = body.get("section", "")
        try:
            from ..knowledge.context_wiki import get_context_wiki
            wiki = get_context_wiki()
            if path:
                pages = [wiki._pages.get(path)]
            elif section:
                pages = [p for p in wiki._pages.values() if p.section == section]
            else:
                pages = wiki.query(query) if query else list(wiki._pages.values())[:10]
            result = [{"path": p.path, "title": p.title, "content": p.content[:500]} for p in pages if p]
            return {"ok": True, "pages": result}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.post("/api/tools/wiki_search")
    async def tool_wiki_search(request: Request):
        body = await request.json()
        query = body.get("query", "")
        try:
            from ..knowledge.context_wiki import WikiTool
            wt = WikiTool()
            result = await wt.search_wiki(query)
            return {"ok": True, "result": result}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.post("/api/shield/check-input")
    async def shield_check_input(request: Request):
        from ..core.prompt_shield import get_shield
        body = await request.json()
        result = get_shield().sanitize_input(body.get("text", ""))
        return {"passed": result.passed, "violations": result.violations,
                "sanitized": result.sanitized_text[:200] if not result.passed else ""}

    @app.post("/api/shield/check-output")
    async def shield_check_output(request: Request):
        from ..core.prompt_shield import get_shield
        body = await request.json()
        result = get_shield().check_output(body.get("text", ""), body.get("context", "public"))
        return {"passed": result.passed, "violations": result.violations}

    # ═══ Telemetry API ═══

    @app.get("/api/telemetry/stats")
    async def telemetry_stats(request: Request):
        from ..core.telemetry import get_telemetry
        return get_telemetry().stats()

    @app.get("/metrics")
    async def prometheus_metrics(request: Request):
        from ..core.telemetry import get_telemetry
        from fastapi.responses import PlainTextResponse
        return PlainTextResponse(get_telemetry().prometheus_metrics(), media_type="text/plain")

    # ═══ Merge / Gray-Release API ═══

    @app.get("/api/merge/status")
    async def merge_status(request: Request):
        from ..core.execution_pipeline import get_execution_pipeline
        return get_execution_pipeline().stats()

    @app.post("/api/merge/flow")
    async def merge_set_flow(request: Request):
        from ..core.execution_pipeline import get_execution_pipeline
        body = await request.json()
        pct = float(body.get("flow_pct", 0.0))
        get_execution_pipeline().update_flow(pct)
        return {"ok": True, "flow_pct": pct}

    @app.post("/api/video/search")
    async def video_search(request: Request):
        from ..core.video_search import get_video_search
        body = await request.json()
        keyword = body.get("keyword", "").strip()
        source = body.get("source", "all")
        limit = body.get("limit", 6)

        if not keyword:
            return {"ok": False, "error": "no keyword"}

        engine = get_video_search()
        results = await engine.search(keyword, source=source, limit=limit)

        cards_html = "".join(engine.build_card_html(v) for v in results)

        return {
            "ok": True,
            "keyword": keyword,
            "total": len(results),
            "sources": list(set(v.source for v in results)),
            "cards_html": cards_html,
            "results": [
                {
                    "title": v.title, "url": v.url, "embed_url": v.embed_url,
                    "thumbnail": v.thumbnail, "duration": v.duration,
                    "source": v.source, "author": v.author, "play_count": v.play_count,
                }
                for v in results
            ],
        }

    # ═══ Collective Intelligence API ═══

    # ═══ Skill Hub API ═══

    @app.get("/api/skills/hub")
    async def skill_hub_index(request: Request):
        from ..core.skill_hub import get_skill_hub
        hub = get_skill_hub()
        items = await hub.fetch_hub_index()
        return {"ok": True, "count": len(items), "skills": [
            {"name": s.name, "version": s.version, "description": s.description,
             "category": s.category, "tools": s.tools, "repo_url": s.repo_url}
            for s in items
        ]}

    @app.get("/api/skills/installed")
    async def skill_list_installed(request: Request):
        from ..core.skill_hub import get_skill_hub
        return {"ok": True, "skills": get_skill_hub().stats()["skills"]}

    @app.post("/api/skills/install")
    async def skill_install(request: Request):
        from ..core.skill_hub import get_skill_hub
        body = await request.json()
        name = body.get("name", body.get("url", "")).strip()
        if not name:
            return {"ok": False, "error": "name or url required"}
        try:
            meta = await get_skill_hub().install(name)
            if meta:
                return {"ok": True, "name": meta.name, "version": meta.version, "tools": meta.tools}
            return {"ok": False, "error": "install failed"}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.post("/api/skills/uninstall")
    async def skill_uninstall(request: Request):
        from ..core.skill_hub import get_skill_hub
        body = await request.json()
        name = body.get("name", "").strip()
        ok = get_skill_hub().uninstall(name)
        return {"ok": ok}

    # ═══ Channel API ═══

    @app.get("/api/channels/status")
    async def channel_status(request: Request):
        from ..network.channel_bridge import get_channel_bridge
        return get_channel_bridge().stats()

    # ═══ Context Budget API ═══

    @app.get("/api/context/budget")
    async def context_budget_status(request: Request):
        from ..core.context_budget import get_context_budget
        return get_context_budget().stats()

    @app.post("/api/context/budget")
    async def context_budget_configure(request: Request):
        from ..core.context_budget import get_context_budget
        body = await request.json()
        budget = get_context_budget()
        budget.configure(
            max_tokens=body.get("max_tokens", 0),
            max_turns=body.get("max_turns", 0),
        )
        return budget.stats()

    @app.post("/api/collective/publish")
    async def collective_publish(request: Request):
        from ..core.collective_intel import get_blueprints
        body = await request.json()
        hub = request.app.state.hub
        bid = get_blueprints().publish(
            name=body.get("name", "blueprint"), hub=hub,
            description=body.get("description", ""),
            persona=body.get("persona", ""),
        )
        return {"ok": True, "id": bid}

    @app.post("/api/collective/import/{blueprint_id}")
    async def collective_import(blueprint_id: str, request: Request):
        from ..core.collective_intel import get_blueprints
        hub = request.app.state.hub
        ok = get_blueprints().import_blueprint(blueprint_id, hub)
        return {"ok": ok}

    # ═══ Device Pairing API ═══

    @app.get("/api/pairing/audio")
    async def pairing_audio():
        from ..network.universal_pairing import get_pairing
        return get_pairing().generate_audio_signal()

    @app.post("/api/pairing/lan/{device_id}")
    async def pairing_lan(device_id: str):
        from ..network.universal_pairing import get_pairing, PairMethod
        pairing = get_pairing()
        pairing.pair_device(device_id, device_id[:16], "unknown", PairMethod.LAN_AUTO)
        return {"ok": True}

    @app.post("/api/pairing/code")
    async def pairing_verify_code(request: Request):
        body = await request.json()
        code = body.get("code", "")
        device_name = body.get("device_name", "mobile")
        from ..network.universal_pairing import get_pairing, PairMethod
        pairing = get_pairing()
        sid = pairing.verify_code(code)
        if sid:
            pairing.pair_device(f"code_{code}", device_name, "mobile", PairMethod.CODE_8_DIGIT, session_id=sid)
            return {"ok": True, "session_id": sid}
        return {"ok": False, "error": "invalid code"}

    # ═══ Push Notification ═══

    @app.get("/api/push/vapid-key")
    async def push_vapid_key():
        """Return the VAPID public key for push subscription."""
        import base64, os, json as _j
        key_file = Path(".livingtree/vapid.json")
        if key_file.exists():
            keys = _j.loads(key_file.read_text())
        else:
            from cryptography.hazmat.primitives.asymmetric import ec
            from cryptography.hazmat.primitives import serialization
            key = ec.generate_private_key(ec.SECP256R1())
            private_bytes = key.private_bytes(serialization.Encoding.PEM, serialization.PrivateFormat.PKCS8, serialization.NoEncryption())
            public_bytes = key.public_key().public_bytes(serialization.Encoding.X962, serialization.PublicFormat.UncompressedPoint)
            private_b64 = base64.urlsafe_b64encode(private_bytes).decode().rstrip('=')
            public_b64 = base64.urlsafe_b64encode(public_bytes).decode().rstrip('=')
            keys = {"private": private_b64, "public": public_b64}
            key_file.parent.mkdir(parents=True, exist_ok=True)
            key_file.write_text(_j.dumps(keys))
        return {"public_key": keys["public"]}

    @app.post("/api/push/subscribe")
    async def push_subscribe(request: Request):
        """Store a push subscription."""
        body = await request.json()
        sub_file = Path(".livingtree/push_subs.jsonl")
        sub_file.parent.mkdir(parents=True, exist_ok=True)
        with open(sub_file, "a", encoding="utf-8") as f:
            f.write(_json.dumps({"endpoint": body.get("endpoint",""), "keys": body.get("keys",{}), "user": body.get("user_id",""), "ts": _time.time()}) + "\n")
        return {"ok": True}

    @app.post("/api/push/send")
    async def push_send(request: Request):
        """Send a push notification to all subscribers."""
        body = await request.json()
        title = body.get("title", "小树")
        msg = body.get("body", body.get("message", ""))
        url = body.get("url", "/tree/living")

        sub_file = Path(".livingtree/push_subs.jsonl")
        if not sub_file.exists():
            return {"ok": False, "error": "no subscribers", "sent": 0}

        import aiohttp
        key_file = Path(".livingtree/vapid.json")
        import json as _j, base64
        keys = _j.loads(key_file.read_text()) if key_file.exists() else {"private": ""}

        sent = 0
        async with aiohttp.ClientSession() as session:
            for line in open(sub_file, encoding="utf-8"):
                try:
                    sub = _json.loads(line)
                    payload = _json.dumps({"title": title, "body": msg, "url": url, "icon": "/assets/icon.svg"})
                    headers = {"Content-Type": "application/json", "TTL": "86400"}
                    async with session.post(sub["endpoint"], data=payload, headers=headers, timeout=aiohttp.ClientTimeout(total=10)) as resp:
                        if resp.status in (200, 201):
                            sent += 1
                except Exception:
                    pass
        return {"ok": True, "sent": sent}

    # ═══ Living Dashboard API (统一数据总线,多态渲染) ═══

    @app.get("/api/living/events")
    async def living_events(request: Request, event_type: str = "", organ: str = "",
                             format: str = "card", limit: int = 50):
        """Unified data bus — capability-probed polymorphic rendering.

        Query params:
          ?format=card|table|timeline|graph|tree|metric|log|auto
          ?cap=plain|rich|struct|visual|media|spatial  (explicit override)
          ?max_bytes=N  (performance budget)
          ?dark=1       (dark mode)
        """
        try:
            from ..treellm.living_scheduler import get_living_scheduler
            ls = get_living_scheduler()
            events = ls.get_events(event_type=event_type, source_organ=organ, limit=limit)
            rendered = ls.render(events, format=RenderFormat(format) if format != "auto" else RenderFormat.CARD, request=request)
            return {"ok": True, "count": len(events), "rendered": rendered}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.get("/api/living/dashboard")
    async def living_dashboard(request: Request):
        """Full organism dashboard — scheduler state + resources + escalations."""
        try:
            from ..treellm.living_scheduler import get_living_scheduler
            ls = get_living_scheduler()
            state = ls.state()
            return {"ok": True, "dashboard": state}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.get("/api/living/hitl")
    async def living_hitl(request: Request):
        """List pending human-in-the-loop escalations."""
        try:
            from ..treellm.living_scheduler import get_living_scheduler
            ls = get_living_scheduler()
            return {"ok": True, "escalations": ls._escalator.pending()}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.post("/api/living/hitl/{task_id}")
    async def living_hitl_resolve(task_id: str, request: Request):
        """Human resolves an escalation."""
        body = await request.json()
        decision = body.get("decision", "")
        try:
            from ..treellm.living_scheduler import get_living_scheduler
            ls = get_living_scheduler()
            ls._escalator.resolve(task_id, decision)
            return {"ok": True, "task_id": task_id, "resolved": True}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ═══ CapabilityBus API (统一能力发现和调用) ═══

    @app.get("/api/capabilities")
    async def list_capabilities(request: Request, category: str = ""):
        """List all capabilities. ?category=tool|skill|mcp|role|user|llm|vfs"""
        try:
            from ..treellm.capability_bus import get_capability_bus
            bus = get_capability_bus()
            if category:
                caps = await bus.list(category)
            else:
                caps = await bus.list_all()
            return {"ok": True, "count": len(caps), "capabilities": caps}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.post("/api/capabilities/{cap_id}")
    async def invoke_capability(cap_id: str, request: Request):
        """Invoke a capability by ID. E.g., POST /api/capabilities/tool:web_search"""
        try:
            from ..treellm.capability_bus import get_capability_bus
            bus = get_capability_bus()
            body = await request.json() if request.headers.get("content-type","").startswith("application/json") else {}
            result = await bus.invoke(cap_id, **body)
            return {"ok": True, "capability": cap_id, "result": result}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.get("/api/capabilities/prompt")
    async def capability_prompt(request: Request, categories: str = ""):
        """Generate LLM system prompt fragment listing capabilities."""
        try:
            from ..treellm.capability_bus import get_capability_bus
            bus = get_capability_bus()
            cats = [c.strip() for c in categories.split(",") if c.strip()] if categories else None
            fragment = await bus.prompt_fragment(cats)
            return {"ok": True, "prompt": fragment}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ═══ Recording API (任务录制与重放) ═══

    @app.post("/api/recording/start")
    async def recording_start(request: Request):
        body = await request.json() if request.headers.get("content-type","").startswith("application/json") else {}
        try:
            from ..treellm.recording_engine import get_recording_engine
            rec_id = get_recording_engine().start(
                title=body.get("title", ""),
                task_type=body.get("task_type", "general"),
            )
            return {"ok": True, "recording_id": rec_id}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.post("/api/recording/stop")
    async def recording_stop(request: Request):
        try:
            from ..treellm.recording_engine import get_recording_engine
            rec = get_recording_engine().stop()
            if rec:
                return {"ok": True, "recording": rec.to_dict()}
            return {"ok": False, "error": "No active recording"}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.get("/api/recordings")
    async def recording_list(request: Request):
        try:
            from ..treellm.recording_engine import get_recording_engine
            recordings = get_recording_engine().list_recordings()
            return {"ok": True, "recordings": recordings}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.get("/api/recording/{rec_id}")
    async def recording_get(rec_id: str, request: Request, format: str = "json"):
        try:
            from ..treellm.recording_engine import get_recording_engine
            engine = get_recording_engine()
            if format in ("xml", "jsonl"):
                content = engine.export(rec_id, format)
                if content:
                    return {"ok": True, "format": format, "content": content}
            rec = engine._recordings.get(rec_id) or engine._load(rec_id)
            if rec:
                return {"ok": True, "recording": rec.to_dict()}
            return {"ok": False, "error": f"Recording not found: {rec_id}"}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.get("/api/recording/{rec_id}/replay")
    async def recording_replay(rec_id: str, request: Request, mode: str = "streaming", speed: float = 1.0):
        try:
            from ..treellm.recording_engine import get_recording_engine, ReplayMode
            engine = get_recording_engine()
            mode_enum = ReplayMode(mode) if mode in [m.value for m in ReplayMode] else ReplayMode.STREAMING
            async def generate():
                async for evt in engine.replay(rec_id, mode=mode_enum, speed=speed):
                    yield f"data: {json.dumps(evt.to_dict(), ensure_ascii=False)}\n\n"
                yield "data: [DONE]\n\n"
            from fastapi.responses import StreamingResponse
            return StreamingResponse(generate(), media_type="text/event-stream")
        except Exception as e:
            from fastapi import HTTPException
            raise HTTPException(status_code=500, detail=str(e))

    @app.get("/api/recording/{rec_id}/render")
    async def recording_render(rec_id: str, request: Request, view: str = "timeline"):
        try:
            from ..treellm.recording_engine import get_recording_engine
            rendered = get_recording_engine().render(rec_id, view)
            return {"ok": True, "view": view, "rendered": rendered}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.delete("/api/recording/{rec_id}")
    async def recording_delete(rec_id: str, request: Request):
        try:
            from ..treellm.recording_engine import get_recording_engine
            deleted = get_recording_engine().delete(rec_id)
            return {"ok": True, "deleted": deleted}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ═══ Debug Loop API (AI自动调试) ═══

    @app.post("/api/debug/start")
    async def debug_start(request: Request):
        body = await request.json() if request.headers.get("content-type","").startswith("application/json") else {}
        try:
            from ..treellm.debug_loop import DebugLoop, DebugLevel
            loop = DebugLoop.instance()
            target = body.get("target", "main.py")
            level = DebugLevel(body.get("level", "semi_auto"))
            max_attempts = body.get("max_attempts", 5)
            session = await loop.debug(target, body.get("args", []), level, max_attempts)
            return {"ok": True, "session_id": session.id, "fixed": session.fixed,
                    "escalated": session.escalated, "attempts": len(session.attempts),
                    "duration_ms": session.total_duration_ms,
                    "details": [{"attempt": a.attempt_number, "result": a.result.value,
                                 "duration_ms": a.duration_ms} for a in session.attempts]}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    @app.get("/api/stream/vitals")
    async def stream_vitals(request: Request):
        """SSE stream of vitals + shield status (replaces polling)."""
        async def generate():
            while True:
                try:
                    hub = request.app.state.hub
                    data = {"type": "vitals", "leaf_display": {"message": "🌿"}}
                    if hub and getattr(hub, '_started', False):
                        try:
                            from ..core.vitals import get_vitals
                            v = get_vitals().measure()
                            data["leaf_display"] = v.get("leaf_display", {"message": "🌿"})
                        except Exception:
                            pass
                    yield f"event: vitals\ndata: {json.dumps(data)}\n\n"

                    # Shield status
                    try:
                        r = await fetch_shield_status()
                        yield f"event: shield\ndata: {json.dumps(r)}\n\n"
                    except Exception:
                        pass

                    await asyncio.sleep(15)
                except asyncio.CancelledError:
                    break
        return StreamingResponse(generate(), media_type="text/event-stream")

    async def fetch_shield_status():
        try:
            from ..core.admin_manager import get_admin
            return {"hitl_pending": 0}
        except Exception:
            return {"hitl_pending": 0}

    @app.get("/api/debug/stats")
    async def debug_stats(request: Request):
        try:
            from ..treellm.debug_loop import get_debug_loop
            return {"ok": True, "stats": get_debug_loop().stats()}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # ═══ Dev Analysis API (returns Tailwind HTML for frontend) ═══

    @app.post("/api/improve/scan")
    async def api_improve_scan(request: Request):
        try:
            from ..treellm.self_improver import get_self_improver
            improver = get_self_improver()
            defects = await improver._scanner.scan()
            report = improver._scanner.report()
            rows = [[cat, str(n), ""] for cat, n in report.get("by_category", {}).items()]
            html = (
                f'<div class="bg-white rounded-xl shadow-sm p-5">'
                f'<h3 class="text-lg font-semibold text-gray-800 mb-3">🔍 代码扫描结果</h3>'
                f'<p class="text-sm text-gray-500 mb-3">发现 {report.get("total",0)} 个缺陷</p>'
                f'<div class="overflow-x-auto"><table class="w-full text-sm">'
                f'<thead><tr class="bg-gray-50 border-b"><th class="px-4 py-2 text-left">类别</th><th class="px-4 py-2 text-left">数量</th></tr></thead>'
                f'<tbody>{"".join(f"<tr class=\"border-b\"><td class=\"px-4 py-2\">{r[0]}</td><td class=\"px-4 py-2\">{r[1]}</td></tr>" for r in rows)}</tbody>'
                f'</table></div></div>'
            )
            return HTMLResponse(content=html)
        except Exception as e:
            return HTMLResponse(content=f'<div class="text-red-500">Scan failed: {e}</div>')

    @app.post("/api/learn/cycle")
    async def api_learn_cycle(request: Request):
        try:
            from ..dna.external_learner import get_external_driver
            driver = get_external_driver()
            result = await driver.run_cycle()
            gh = result.get("github_patterns", 0)
            ar = result.get("arxiv_patterns", 0)
            na = result.get("nature_patterns", 0)
            total = result.get("total_patterns", 0)
            html = (
                f'<div class="bg-white rounded-xl shadow-sm p-5">'
                f'<h3 class="text-lg font-semibold text-gray-800 mb-3">📚 外部学习结果</h3>'
                f'<div class="grid grid-cols-3 gap-3 mb-3">'
                f'<div class="bg-blue-50 rounded-lg p-3 text-center"><div class="text-2xl font-bold text-blue-700">{gh}</div><div class="text-xs text-blue-600">GitHub</div></div>'
                f'<div class="bg-green-50 rounded-lg p-3 text-center"><div class="text-2xl font-bold text-green-700">{ar}</div><div class="text-xs text-green-600">arXiv</div></div>'
                f'<div class="bg-purple-50 rounded-lg p-3 text-center"><div class="text-2xl font-bold text-purple-700">{na}</div><div class="text-xs text-purple-600">Nature</div></div>'
                f'</div>'
                f'<p class="text-sm text-gray-500">共提取 {total} 个改进模式</p>'
                f'</div>'
            )
            return HTMLResponse(content=html)
        except Exception as e:
            return HTMLResponse(content=f'<div class="text-red-500">Learn failed: {e}</div>')

    @app.post("/api/skills/cycle")
    async def api_skills_cycle(request: Request):
        try:
            from ..dna.living_skills import get_living_skills
            skills = get_living_skills()
            report = skills.run_cycle()
            html = (
                f'<div class="bg-white rounded-xl shadow-sm p-5">'
                f'<h3 class="text-lg font-semibold text-gray-800 mb-3">🌱 三技能进化</h3>'
                f'<div class="space-y-2 text-sm">'
                f'<div class="flex justify-between"><span>🧹 自我清理</span><span class="text-gray-600">{report.clean_rules} 项</span></div>'
                f'<div class="flex justify-between"><span>📝 记忆完善</span><span class="text-gray-600">{report.refine_patterns} 模式</span></div>'
                f'<div class="flex justify-between"><span>🧬 自我进化</span><span class="text-gray-600">{report.evolve_iterations} 次迭代</span></div>'
                f'<div class="flex justify-between font-medium"><span>总规则</span><span>{report.total_rules_learned}</span></div>'
                f'</div></div>'
            )
            return HTMLResponse(content=html)
        except Exception as e:
            return HTMLResponse(content=f'<div class="text-red-500">Skills failed: {e}</div>')

    @app.get("/api/dev/hotspots")
    async def api_dev_hotspots(request: Request):
        try:
            from ..treellm.dev_assistant import HotColdAnalyzer
            reports = HotColdAnalyzer.analyze(top_n=8)
            rows = []
            icons = {"hot":"🔥","active":"⚡","cooling":"❄️","stable":"✅"}
            for r in reports:
                name = Path(r.file).name if isinstance(r.file, str) else str(r.file)
                rows.append(f'<tr class="border-b"><td class="px-3 py-2">{icons.get(r.status,"")} {name}</td><td class="px-3 py-2 text-right">{r.change_count}</td><td class="px-3 py-2 text-right text-gray-500">{r.last_changed_days}d ago</td></tr>')
            html = (
                f'<div class="bg-white rounded-xl shadow-sm p-5">'
                f'<h3 class="text-lg font-semibold text-gray-800 mb-3">🔥 热点文件 (90天)</h3>'
                f'<table class="w-full text-sm"><thead><tr class="bg-gray-50 border-b"><th class="px-3 py-2 text-left">文件</th><th class="px-3 py-2 text-right">变更</th><th class="px-3 py-2 text-right">最近</th></tr></thead>'
                f'<tbody>{"".join(rows)}</tbody></table></div>'
            )
            return HTMLResponse(content=html)
        except Exception as e:
            return HTMLResponse(content=f'<div class="text-red-500">Hotspots: {e}</div>')

    @app.get("/api/dev/cognimap")
    async def api_dev_cognimap(request: Request):
        try:
            from ..treellm.living_dev import CognitiveLoader
            files = CognitiveLoader.analyze(risk_filter="critical")[:8]
            rows = []
            for f in files:
                name = Path(f.path).name
                rows.append(f'<tr class="border-b"><td class="px-3 py-2">{name}</td><td class="px-3 py-2 text-right">{f.cyclomatic}</td><td class="px-3 py-2 text-right">{f.cognitive}</td><td class="px-3 py-2"><span class="inline-flex px-2 py-0.5 rounded-full text-xs font-medium {("bg-red-100 text-red-700" if f.risk=="critical" else "bg-yellow-100 text-yellow-700")}">{f.risk}</span></td></tr>')
            html = (
                f'<div class="bg-white rounded-xl shadow-sm p-5">'
                f'<h3 class="text-lg font-semibold text-gray-800 mb-3">🧠 认知复杂度 Top 8</h3>'
                f'<table class="w-full text-sm"><thead><tr class="bg-gray-50 border-b"><th class="px-3 py-2 text-left">文件</th><th class="px-3 py-2 text-right">CC</th><th class="px-3 py-2 text-right">Cog</th><th class="px-3 py-2">风险</th></tr></thead>'
                f'<tbody>{"".join(rows)}</tbody></table></div>'
            )
            return HTMLResponse(content=html)
        except Exception as e:
            return HTMLResponse(content=f'<div class="text-red-500">Cognimap: {e}</div>')
