"""
Word 模板学习模块
=================

自动学习 Word 模板的结构和格式，用于后续报告生成。

采用自我学习策略：
1. 解析 Word 模板（.docx）
2. 提取章节结构、样式、表格格式
3. 存储到知识库（向量数据库）
4. Agent 生成报告时参考学习到的模板

Author: LivingTreeAI Agent
Date: 2026-04-27
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WordTemplateLearner] 警告: python-docx 未安装")

from client.src.business.tools.tool_definition import ToolDefinition
from client.src.business.tools.base_tool import BaseTool
from client.src.business.tools.tool_result import ToolResult

logger = logging.getLogger(__name__)


@dataclass
class TemplateStyle:
    """模板样式"""
    name: str
    font_name: str = "宋体"
    font_size: int = 12
    bold: bool = False
    italic: bool = False
    color: str = "000000"  # RGB hex
    alignment: str = "LEFT"
    

@dataclass
class TemplateSection:
    """模板章节"""
    title: str
    level: int  # 标题级别 1-4
    content_template: str = ""  # 内容模板（如 "本章描述了{project_name}的..."）
    required: bool = True  # 是否必需章节
    children: List['TemplateSection'] = field(default_factory=list)
    

@dataclass
class TemplateTable:
    """模板表格格式"""
    name: str
    rows: int
    cols: int
    headers: List[str]
    style: str = "Table Grid"
    

class WordTemplateLearner:
    """
    Word 模板学习器
    
    功能：
    1. 解析 Word 模板（.docx）
    2. 提取章节结构、样式、表格
    3. 存储到知识库
    4. 生成模板描述（供 Agent 参考）
    """
    
    def __init__(self, template_path: Optional[str] = None):
        """
        初始化模板学习器
        
        Args:
            template_path: Word 模板文件路径（.docx）
        """
        self.template_path = template_path
        self.sections: List[TemplateSection] = []
        self.styles: Dict[str, TemplateStyle] = {}
        self.tables: List[TemplateTable] = []
        self.metadata: Dict[str, Any] = {}
        
    def load_template(self, template_path: str) -> bool:
        """
        加载 Word 模板
        
        Args:
            template_path: 模板文件路径
            
        Returns:
            是否加载成功
        """
        if not HAS_DOCX:
            logger.error("[WordTemplateLearner] python-docx 未安装，无法解析模板")
            return False
        
        if not os.path.exists(template_path):
            logger.error(f"[WordTemplateLearner] 模板文件不存在: {template_path}")
            return False
        
        self.template_path = template_path
        
        try:
            doc = Document(template_path)
            
            # 1. 提取元数据
            self._extract_metadata(doc)
            
            # 2. 提取章节结构
            self._extract_sections(doc)
            
            # 3. 提取样式
            self._extract_styles(doc)
            
            # 4. 提取表格
            self._extract_tables(doc)
            
            logger.info(f"[WordTemplateLearner] 模板加载成功: {template_path}")
            logger.info(f"  章节数: {len(self.sections)}")
            logger.info(f"  样式数: {len(self.styles)}")
            logger.info(f"  表格数: {len(self.tables)}")
            
            return True
            
        except Exception as e:
            logger.error(f"[WordTemplateLearner] 加载模板失败: {e}")
            return False
    
    def _extract_metadata(self, doc: Document):
        """提取文档元数据"""
        self.metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            "paragraphs_count": len(doc.paragraphs),
            "tables_count": len(doc.tables),
        }
        
        # 提取文档主体字体设置
        if doc.styles['Normal']:
            style = doc.styles['Normal']
            if style.font:
                self.metadata["default_font"] = {
                    "name": style.font.name or "宋体",
                    "size": style.font.size.pt if style.font.size else 12,
                }
    
    def _extract_sections(self, doc: Document):
        """提取章节结构（从标题段落）"""
        self.sections = []
        current_section = None
        
        for para in doc.paragraphs:
            # 检查是否是标题
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                title = para.text.strip()
                
                section = TemplateSection(
                    title=title,
                    level=level,
                    content_template=self._extract_content_template(para),
                )
                
                # 构建章节树
                if level == 1:
                    self.sections.append(section)
                    current_section = section
                elif current_section:
                    parent = self._find_parent(current_section, level)
                    if parent:
                        parent.children.append(section)
                
                logger.debug(f"[WordTemplateLearner] 发现章节: {title} (级别 {level})")
    
    def _extract_content_template(self, para) -> str:
        """提取章节内容模板（从标题后的段落）"""
        # 简化实现：返回标题本身
        return para.text.strip()
    
    def _find_parent(self, section: TemplateSection, level: int) -> Optional[TemplateSection]:
        """查找父章节"""
        if level <= 1:
            return None
        
        # 简化实现：假设当前章节的父章节是最近的高一级章节
        # 实际应该维护一个栈
        return None  # 暂时返回 None
    
    def _extract_styles(self, doc: Document):
        """提取文档样式"""
        self.styles = {}
        
        for style in doc.styles:
            if style.type == 1:  # 段落样式
                ts = TemplateStyle(
                    name=style.name,
                    font_name=style.font.name or "宋体",
                    font_size=int(style.font.size.pt) if style.font.size else 12,
                    bold=style.font.bold or False,
                    italic=style.font.italic or False,
                )
                
                # 提取颜色
                if style.font.color and style.font.color.rgb:
                    ts.color = str(style.font.color.rgb)
                
                # 提取对齐方式
                if style.paragraph_format:
                    alignment = style.paragraph_format.alignment
                    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        ts.alignment = "CENTER"
                    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                        ts.alignment = "RIGHT"
                    else:
                        ts.alignment = "LEFT"
                
                self.styles[style.name] = ts
                logger.debug(f"[WordTemplateLearner] 发现样式: {style.name}")
    
    def _extract_tables(self, doc: Document):
        """提取表格格式"""
        self.tables = []
        
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            
            # 提取表头（第一行）
            headers = []
            if rows > 0:
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip())
            
            tt = TemplateTable(
                name=f"Table_{i+1}",
                rows=rows,
                cols=cols,
                headers=headers,
            )
            
            self.tables.append(tt)
            logger.debug(f"[WordTemplateLearner] 发现表格: {tt.name} ({rows}x{cols})")
    
    def generate_template_description(self) -> str:
        """
        生成模板描述（供 Agent 参考）
        
        Returns:
            模板描述文本（Markdown 格式）
        """
        desc = "# Word 模板结构描述\n\n"
        
        # 元数据
        desc += "## 元数据\n"
        desc += f"- 标题: {self.metadata.get('title', '未知')}\n"
        desc += f"- 作者: {self.metadata.get('author', '未知')}\n"
        desc += f"- 段落数: {self.metadata.get('paragraphs_count', 0)}\n"
        desc += f"- 表格数: {self.metadata.get('tables_count', 0)}\n\n"
        
        # 章节结构
        desc += "## 章节结构\n"
        for section in self.sections:
            desc += self._format_section(section, level=0)
        
        # 样式
        desc += "\n## 样式\n"
        for style_name, style in self.styles.items():
            desc += f"- **{style_name}**: {style.font_name} {style.font_size}pt"
            if style.bold:
                desc += " **加粗**"
            if style.italic:
                desc += " *斜体*"
            desc += f" [颜色: #{style.color}]\n"
        
        # 表格
        desc += "\n## 表格格式\n"
        for table in self.tables:
            desc += f"- **{table.name}**: {table.rows}行 x {table.cols}列\n"
            if table.headers:
                desc += f"  - 表头: {', '.join(table.headers)}\n"
        
        return desc
    
    def _format_section(self, section: TemplateSection, level: int) -> str:
        """格式化章节（递归）"""
        indent = "  " * level
        result = f"{indent}- **{section.title}** (级别 {section.level})\n"
        
        for child in section.children:
            result += self._format_section(child, level + 1)
        
        return result
    
    def save_to_knowledge_base(self, kb_name: str = "template_knowledge") -> bool:
        """
        将学习到的模板存储到知识库
        
        Args:
            kb_name: 知识库名称
            
        Returns:
            是否保存成功
        """
        try:
            from client.src.business.knowledge_vector_db import VectorDatabase
            
            db = VectorDatabase()
            
            # 生成模板描述
            description = self.generate_template_description()
            
            # 存储到向量数据库
            db.add_document(
                text=description,
                metadata={
                    "type": "word_template",
                    "template_path": self.template_path or "",
                    "sections_count": len(self.sections),
                    "tables_count": len(self.tables),
                }
            )
            
            logger.info(f"[WordTemplateLearner] 模板已保存到知识库: {kb_name}")
            return True
            
        except Exception as e:
            logger.error(f"[WordTemplateLearner] 保存到知识库失败: {e}")
            return False
    
    def get_section_names(self) -> List[str]:
        """获取所有章节名称（扁平列表）"""
        names = []
        for section in self.sections:
            names.extend(self._collect_section_names(section))
        return names
    
    def _collect_section_names(self, section: TemplateSection) -> List[str]:
        """递归收集章节名称"""
        names = [section.title]
        for child in section.children:
            names.extend(self._collect_section_names(child))
        return names


# ============================================================
# 工具接口（注册到 ToolRegistry）
# ============================================================

class LearnWordTemplateTool(BaseTool):
    """学习 Word 模板工具"""
    
    def __init__(self):
        definition = ToolDefinition(
            name="learn_word_template",
            description="学习 Word 模板的结构和格式，用于后续报告生成",
            category="document",
            tags=["word", "template", "learning", "ei"],
        )
        super().__init__(definition)
    
    def execute(self, template_path: str, save_to_kb: bool = True) -> ToolResult:
        """
        学习 Word 模板
        
        Args:
            template_path: Word 模板文件路径（.docx）
            save_to_kb: 是否保存到知识库
            
        Returns:
            ToolResult
        """
        try:
            learner = WordTemplateLearner()
            
            if not learner.load_template(template_path):
                return ToolResult(
                    success=False,
                    error="加载模板失败"
                )
            
            # 生成模板描述
            description = learner.generate_template_description()
            
            # 保存到知识库
            if save_to_kb:
                learner.save_to_knowledge_base()
            
            return ToolResult(
                success=True,
                data={
                    "template_path": template_path,
                    "sections_count": len(learner.sections),
                    "tables_count": len(learner.tables),
                    "styles_count": len(learner.styles),
                    "description": description,
                    "section_names": learner.get_section_names(),
                }
            )
            
        except Exception as e:
            logger.error(f"[LearnWordTemplateTool] 执行失败: {e}")
            return ToolResult(success=False, error=str(e))


class AnalyzeTemplateStructureTool(BaseTool):
    """分析模板结构工具（仅分析，不保存）"""
    
    def __init__(self):
        definition = ToolDefinition(
            name="analyze_template_structure",
            description="分析 Word 模板的章节结构（不保存到知识库）",
            category="document",
            tags=["word", "template", "analysis", "ei"],
        )
        super().__init__(definition)
    
    def execute(self, template_path: str) -> ToolResult:
        """
        分析模板结构
        
        Args:
            template_path: Word 模板文件路径（.docx）
            
        Returns:
            ToolResult
        """
        try:
            learner = WordTemplateLearner()
            
            if not learner.load_template(template_path):
                return ToolResult(
                    success=False,
                    error="加载模板失败"
                )
            
            # 生成模板描述
            description = learner.generate_template_description()
            
            return ToolResult(
                success=True,
                data={
                    "template_path": template_path,
                    "description": description,
                    "sections": learner.get_section_names(),
                    "metadata": learner.metadata,
                }
            )
            
        except Exception as e:
            logger.error(f"[AnalyzeTemplateStructureTool] 执行失败: {e}")
            return ToolResult(success=False, error=str(e))


def register_word_template_tools():
    """注册 Word 模板学习工具到 ToolRegistry"""
    from client.src.business.tools.tool_registry import ToolRegistry
    
    registry = ToolRegistry.get_instance()
    
    registry.register_tool(LearnWordTemplateTool())
    registry.register_tool(AnalyzeTemplateStructureTool())
    
    logger.info("[WordTemplateLearner] Word 模板学习工具已注册")
