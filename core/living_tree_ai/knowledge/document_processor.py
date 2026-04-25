"""
文档处理器（增强版）

支持更多文档格式：PDF、Word、Excel、文本等
"""

import os
import re
from typing import List, Dict, Optional, Any
from dataclasses import dataclass

try:
    import fitz  # PyMuPDF
    PYMPDF_AVAILABLE = True
except ImportError:
    PYMPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


def _get_openai_api_key() -> Optional[str]:
    """获取 OpenAI API Key（兼容统一配置）"""
    try:
        from core.config.unified_config import get_config
        key = get_config().get_api_key("openai")
        if key:
            return key
    except Exception:
        pass
    return os.getenv("OPENAI_API_KEY")


@dataclass
class DocumentChunk:
    """文档块"""
    id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(
        self,
        embedding_model: str = "text-embedding-ada-002",
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ):
        """
        初始化文档处理器
        
        Args:
            embedding_model: 嵌入模型名称
            chunk_size: 文本分块大小
            chunk_overlap: 分块重叠大小
        """
        self.embedding_model = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._init_text_splitter()
        self._init_embeddings()
    
    def _init_text_splitter(self):
        """初始化文本分割器"""
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
            length_function=len
        )
    
    def _init_embeddings(self):
        """初始化嵌入模型"""
        from langchain.embeddings import OpenAIEmbeddings, HuggingFaceEmbeddings
        
        try:
            # 优先使用 OpenAI 嵌入（通过统一配置）
            api_key = _get_openai_api_key()
            if api_key:
                self.embeddings = OpenAIEmbeddings(
                    model=self.embedding_model,
                    api_key=api_key
                )
                return
        except Exception:
            pass
        
        # 回退到 HuggingFace 嵌入
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    
    def process_file(self, file_path: str) -> List[DocumentChunk]:
        """
        处理文件，自动识别格式
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档块列表
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self.process_pdf(file_path)
        elif ext == '.docx':
            return self.process_docx(file_path)
        elif ext == '.xlsx':
            return self.process_xlsx(file_path)
        elif ext == '.txt':
            return self.process_text_file(file_path)
        else:
            return self.process_text_file(file_path)
    
    def process_pdf(self, pdf_path: str) -> List[DocumentChunk]:
        """
        处理 PDF 文档
        
        Args:
            pdf_path: PDF 文件路径
            
        Returns:
            文档块列表
        """
        chunks = []
        
        if not PYMPDF_AVAILABLE:
            print(f"[DocumentProcessor] PyMuPDF 不可用，无法处理 PDF: {pdf_path}")
            return chunks
        
        try:
            doc = fitz.open(pdf_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                text += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
            
            doc.close()
            
            # 分割文本
            text_chunks = self.text_splitter.split_text(text)
            
            for i, chunk in enumerate(text_chunks):
                chunk_id = f"chunk_{i}"
                metadata = {
                    "source": pdf_path,
                    "type": "pdf",
                    "chunk_id": chunk_id,
                    "index": i
                }
                
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    content=chunk,
                    metadata=metadata
                ))
                
        except Exception as e:
            print(f"[DocumentProcessor] 处理 PDF 失败: {e}")
        
        return chunks
    
    def process_docx(self, docx_path: str) -> List[DocumentChunk]:
        """
        处理 Word 文档
        
        Args:
            docx_path: Word 文件路径
            
        Returns:
            文档块列表
        """
        chunks = []
        
        if not DOCX_AVAILABLE:
            print(f"[DocumentProcessor] python-docx 不可用，无法处理 Word: {docx_path}")
            return chunks
        
        try:
            doc = DocxDocument(docx_path)
            text = ""
            
            # 提取段落
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text += f"\n{para.text}"
            
            # 提取表格
            for i, table in enumerate(doc.tables):
                text += f"\n\n--- Table {i + 1} ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += f"{row_text}\n"
            
            # 分割文本
            text_chunks = self.text_splitter.split_text(text)
            
            for i, chunk in enumerate(text_chunks):
                chunk_id = f"chunk_{i}"
                metadata = {
                    "source": docx_path,
                    "type": "docx",
                    "chunk_id": chunk_id,
                    "index": i
                }
                
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    content=chunk,
                    metadata=metadata
                ))
                
        except Exception as e:
            print(f"[DocumentProcessor] 处理 Word 文档失败: {e}")
        
        return chunks
    
    def process_xlsx(self, xlsx_path: str) -> List[DocumentChunk]:
        """
        处理 Excel 文档
        
        Args:
            xlsx_path: Excel 文件路径
            
        Returns:
            文档块列表
        """
        chunks = []
        
        if not OPENPYXL_AVAILABLE:
            print(f"[DocumentProcessor] openpyxl 不可用，无法处理 Excel: {xlsx_path}")
            return chunks
        
        try:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            
            text = ""
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"\n\n--- Sheet: {sheet_name} ---\n"
                
                # 提取表格数据
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([
                        str(cell) if cell is not None else ""
                        for cell in row
                    ])
                    if row_text.strip():
                        text += f"{row_text}\n"
                
                # 提取命名范围
                if sheet.defined_names:
                    text += f"\n命名范围: {list(sheet.defined_names.definedName)}\n"
            
            wb.close()
            
            # 分割文本
            text_chunks = self.text_splitter.split_text(text)
            
            for i, chunk in enumerate(text_chunks):
                chunk_id = f"chunk_{i}"
                metadata = {
                    "source": xlsx_path,
                    "type": "xlsx",
                    "chunk_id": chunk_id,
                    "index": i
                }
                
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    content=chunk,
                    metadata=metadata
                ))
                
        except Exception as e:
            print(f"[DocumentProcessor] 处理 Excel 文档失败: {e}")
        
        return chunks
    
    def process_text(self, text: str, source: str = "text") -> List[DocumentChunk]:
        """
        处理文本
        
        Args:
            text: 文本内容
            source: 来源
            
        Returns:
            文档块列表
        """
        chunks = []
        
        try:
            # 分割文本
            text_chunks = self.text_splitter.split_text(text)
            
            for i, chunk in enumerate(text_chunks):
                chunk_id = f"chunk_{i}"
                metadata = {
                    "source": source,
                    "type": "text",
                    "chunk_id": chunk_id,
                    "index": i
                }
                
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    content=chunk,
                    metadata=metadata
                ))
                
        except Exception as e:
            print(f"[DocumentProcessor] 处理文本失败: {e}")
        
        return chunks
    
    def process_text_file(self, file_path: str) -> List[DocumentChunk]:
        """
        处理文本文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档块列表
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return self.process_text(text, source=file_path)
        except Exception as e:
            print(f"[DocumentProcessor] 处理文本文件失败: {e}")
            return []
    
    def generate_embeddings(self, chunks: List[DocumentChunk]) -> List[DocumentChunk]:
        """
        生成嵌入向量
        
        Args:
            chunks: 文档块列表
            
        Returns:
            带嵌入向量的文档块列表
        """
        try:
            # 提取文本
            texts = [chunk.content for chunk in chunks]
            
            # 生成嵌入
            embeddings = self.embeddings.embed_documents(texts)
            
            # 赋值嵌入向量
            for i, chunk in enumerate(chunks):
                if i < len(embeddings):
                    chunk.embedding = embeddings[i]
                    
        except Exception as e:
            print(f"[DocumentProcessor] 生成嵌入失败: {e}")
        
        return chunks
    
    def create_vector_store(self, chunks: List[DocumentChunk]):
        """
        创建向量存储
        
        Args:
            chunks: 文档块列表
            
        Returns:
            FAISS 向量存储
        """
        try:
            # 提取文本和元数据
            texts = [chunk.content for chunk in chunks]
            metadatas = [chunk.metadata for chunk in chunks]
            
            # 创建向量存储
            from langchain.vectorstores import FAISS
            vector_store = FAISS.from_texts(
                texts=texts,
                embedding=self.embeddings,
                metadatas=metadatas
            )
            
            return vector_store
            
        except Exception as e:
            print(f"[DocumentProcessor] 创建向量存储失败: {e}")
            return None
    
    def save_vector_store(self, vector_store, path: str):
        """
        保存向量存储
        
        Args:
            vector_store: 向量存储
            path: 保存路径
        """
        try:
            vector_store.save_local(path)
        except Exception as e:
            print(f"[DocumentProcessor] 保存向量存储失败: {e}")
    
    def load_vector_store(self, path: str):
        """
        加载向量存储
        
        Args:
            path: 加载路径
            
        Returns:
            FAISS 向量存储
        """
        try:
            from langchain.vectorstores import FAISS
            vector_store = FAISS.load_local(
                path=path,
                embeddings=self.embeddings,
                allow_dangerous_deserialization=True
            )
            return vector_store
        except Exception as e:
            print(f"[DocumentProcessor] 加载向量存储失败: {e}")
            return None
