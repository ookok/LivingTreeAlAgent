"""
文档导出引擎

支持将文档导出为多种格式：Word、PDF、HTML、Markdown、Excel等。
"""

import json
import asyncio
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Any, Callable
from enum import Enum
from datetime import datetime
from pathlib import Path


# ==================== 数据模型 ====================

class ExportFormat(Enum):
    """导出格式"""
    WORD = "word"
    PDF = "pdf"
    HTML = "html"
    MARKDOWN = "markdown"
    EXCEL = "excel"
    JSON = "json"
    XML = "xml"


@dataclass
class DocumentExporterConfig:
    """导出配置"""
    format: ExportFormat = ExportFormat.WORD
    include_metadata: bool = True
    include_toc: bool = True
    include_header: bool = True
    include_footer: bool = True
    page_size: str = "A4"
    font_family: str = "宋体"
    font_size: int = 12
    line_spacing: float = 1.5


@dataclass
class ExportResult:
    """导出结果"""
    success: bool
    format: ExportFormat
    file_path: Optional[str] = None
    content: Optional[bytes] = None
    error: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)


# ==================== 导出器基类 ====================

class BaseExporter:
    """导出器基类"""

    def __init__(self, config: DocumentExporterConfig):
        self.config = config

    async def export(self, document_data: Dict, template: Dict) -> ExportResult:
        """
        导出文档

        Args:
            document_data: 文档数据
            template: 模板定义

        Returns:
            ExportResult: 导出结果
        """
        raise NotImplementedError

    def _build_document(self, document_data: Dict, template: Dict) -> str:
        """构建文档内容"""
        raise NotImplementedError


# ==================== HTML导出器 ====================

class HTMLExporter(BaseExporter):
    """HTML导出器"""

    async def export(self, document_data: Dict, template: Dict) -> ExportResult:
        """导出为HTML"""
        try:
            html_content = self._build_document(document_data, template)

            return ExportResult(
                success=True,
                format=ExportFormat.HTML,
                content=html_content.encode('utf-8'),
                metadata={
                    "generated_at": datetime.now().isoformat(),
                    "template": template.get("name", "未知模板"),
                    "sections_count": len(document_data.get("sections", []))
                }
            )
        except Exception as e:
            return ExportResult(
                success=False,
                format=ExportFormat.HTML,
                error=str(e)
            )

    def _build_document(self, document_data: Dict, template: Dict) -> str:
        """构建HTML文档"""
        sections = document_data.get("sections", [])

        html_parts = ['''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>''', document_data.get("title", template.get("name", "文档")), '''</title>
    <style>
        body {
            font-family: "''', self.config.font_family, '''", "SimSun", serif;
            font-size: 14px;
            line-height: 1.8;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            color: #333;
        }
        h1 {
            font-size: 24px;
            text-align: center;
            margin-bottom: 30px;
            color: #1a1a1a;
        }
        h2 {
            font-size: 18px;
            margin-top: 30px;
            border-bottom: 2px solid #667eea;
            padding-bottom: 8px;
            color: #2d3748;
        }
        .section {
            margin-bottom: 25px;
        }
        .field {
            margin-bottom: 12px;
        }
        .field-label {
            font-weight: bold;
            color: #4a5568;
            margin-bottom: 4px;
        }
        .field-value {
            padding: 8px 12px;
            background: #f7fafc;
            border-left: 3px solid #667eea;
            min-height: 20px;
        }
        .table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        .table th, .table td {
            border: 1px solid #e2e8f0;
            padding: 8px 12px;
            text-align: left;
        }
        .table th {
            background: #f7fafc;
            font-weight: bold;
        }
        .metadata {
            font-size: 12px;
            color: #718096;
            margin-top: 30px;
            padding-top: 15px;
            border-top: 1px solid #e2e8f0;
        }
    </style>
</head>
<body>
    <h1>''', document_data.get("title", template.get("name", "文档")), '''</h1>
''']

        for section in sections:
            section_id = section.get("id", "")
            section_title = section.get("title", "")
            section_data = document_data.get(section_id, {})

            html_parts.append(f'    <div class="section">\n')
            html_parts.append(f'        <h2>{section_title}</h2>\n')

            for field_def in section.get("fields", []):
                field_id = field_def.get("id", "")
                field_label = field_def.get("label", field_id)
                field_value = ""

                # 获取字段值
                if isinstance(section_data, dict):
                    field_value = section_data.get(field_id, "")
                elif isinstance(section_data, str):
                    field_value = section_data

                field_type = field_def.get("type", "text")

                if field_type == "table" and isinstance(field_value, list):
                    html_parts.append(self._render_table(field_label, field_def, field_value))
                else:
                    html_parts.append(f'''
        <div class="field">
            <div class="field-label">{field_label}</div>
            <div class="field-value">{field_value or "—"}</div>
        </div>
''')

            html_parts.append('    </div>\n')

        # 元数据
        if self.config.include_metadata:
            html_parts.append(f'''
    <div class="metadata">
        <p>生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
        <p>模板: {template.get("name", "未知模板")} v{template.get("version", "1.0")}</p>
    </div>
''')

        html_parts.append('''</body>\n</html>''')

        return ''.join(html_parts)

    def _render_table(self, label: str, field_def: Dict, rows: List) -> str:
        """渲染表格"""
        columns = field_def.get("columns", [])

        header_html = ''.join([
            f'<th>{col.get("label", col.get("id", ""))}</th>'
            for col in columns
        ])

        body_html = ''
        for row in rows:
            cells = []
            for col in columns:
                col_id = col.get("id", "")
                cell_value = row.get(col_id, "") if isinstance(row, dict) else ""
                cells.append(f'<td>{cell_value}</td>')
            body_html += f'<tr>{"".join(cells)}</tr>'

        return f'''
        <div class="field">
            <div class="field-label">{label}</div>
            <table class="table">
                <thead><tr>{header_html}</tr></thead>
                <tbody>{body_html}</tbody>
            </table>
        </div>
'''


# ==================== Markdown导出器 ====================

class MarkdownExporter(BaseExporter):
    """Markdown导出器"""

    async def export(self, document_data: Dict, template: Dict) -> ExportResult:
        """导出为Markdown"""
        try:
            md_content = self._build_document(document_data, template)

            return ExportResult(
                success=True,
                format=ExportFormat.MARKDOWN,
                content=md_content.encode('utf-8'),
                metadata={
                    "generated_at": datetime.now().isoformat(),
                    "template": template.get("name", "未知模板")
                }
            )
        except Exception as e:
            return ExportResult(
                success=False,
                format=ExportFormat.MARKDOWN,
                error=str(e)
            )

    def _build_document(self, document_data: Dict, template: Dict) -> str:
        """构建Markdown文档"""
        sections = document_data.get("sections", [])

        md_parts = [
            f"# {document_data.get('title', template.get('name', '文档'))}\n\n"
        ]

        if self.config.include_metadata:
            md_parts.append(
                f"*生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
            )

        for section in sections:
            section_id = section.get("id", "")
            section_title = section.get("title", "")
            section_data = document_data.get(section_id, {})

            md_parts.append(f"## {section_title}\n\n")

            for field_def in section.get("fields", []):
                field_id = field_def.get("id", "")
                field_label = field_def.get("label", field_id)
                field_value = ""

                if isinstance(section_data, dict):
                    field_value = section_data.get(field_id, "")
                elif isinstance(section_data, str):
                    field_value = section_data

                field_type = field_def.get("type", "text")

                if field_type == "table" and isinstance(field_value, list):
                    md_parts.append(self._render_table(field_label, field_def, field_value))
                else:
                    md_parts.append(f"**{field_label}**: {field_value or '—'}\n\n")

        return ''.join(md_parts)

    def _render_table(self, label: str, field_def: Dict, rows: List) -> str:
        """渲染Markdown表格"""
        columns = field_def.get("columns", [])

        header = '| ' + ' | '.join([col.get("label", col.get("id", "")) for col in columns]) + ' |'
        separator = '| ' + ' | '.join(['---' for _ in columns]) + ' |'

        body_rows = []
        for row in rows:
            cells = []
            for col in columns:
                col_id = col.get("id", "")
                cell_value = row.get(col_id, "") if isinstance(row, dict) else ""
                cells.append(str(cell_value))
            body_rows.append('| ' + ' | '.join(cells) + ' |')

        return f"**{label}**\n\n{header}\n{separator}\n{'  \n'.join(body_rows)}\n\n"


# ==================== JSON导出器 ====================

class JSONExporter(BaseExporter):
    """JSON导出器"""

    async def export(self, document_data: Dict, template: Dict) -> ExportResult:
        """导出为JSON"""
        try:
            # 构建结构化数据
            structured_data = {
                "metadata": {
                    "title": document_data.get("title", template.get("name", "文档")),
                    "template": template.get("name", ""),
                    "version": template.get("version", "1.0"),
                    "exported_at": datetime.now().isoformat()
                },
                "content": document_data.get("sections", []),
                "raw_data": document_data
            }

            json_content = json.dumps(structured_data, ensure_ascii=False, indent=2)

            return ExportResult(
                success=True,
                format=ExportFormat.JSON,
                content=json_content.encode('utf-8'),
                metadata={
                    "generated_at": datetime.now().isoformat(),
                    "template": template.get("name", "未知模板")
                }
            )
        except Exception as e:
            return ExportResult(
                success=False,
                format=ExportFormat.JSON,
                error=str(e)
            )


# ==================== Word导出器 ====================

class WordExporter(BaseExporter):
    """Word文档导出器（需要python-docx）"""

    async def export(self, document_data: Dict, template: Dict) -> ExportResult:
        """导出为Word"""
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置页面
            sections = doc.sections
            for section in sections:
                section.page_width = Cm(21)
                section.page_height = Cm(29.7)
                section.left_margin = Cm(3)
                section.right_margin = Cm(2)

            # 标题
            title = doc.add_heading(document_data.get("title", template.get("name", "文档")), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 章节
            for section in document_data.get("sections", []):
                section_title = section.get("title", "")
                doc.add_heading(section_title, level=1)

                section_id = section.get("id", "")
                section_data = document_data.get(section_id, {})

                for field_def in section.get("fields", []):
                    field_id = field_def.get("id", "")
                    field_label = field_def.get("label", field_id)
                    field_value = ""

                    if isinstance(section_data, dict):
                        field_value = section_data.get(field_id, "")
                    elif isinstance(section_data, str):
                        field_value = section_data

                    field_type = field_def.get("type", "text")

                    if field_type == "table" and isinstance(field_value, list):
                        self._add_table_to_doc(doc, field_def, field_value)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"{field_label}: ").bold = True
                        p.add_run(str(field_value) if field_value else "—")

            # 保存到字节
            import io
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return ExportResult(
                success=True,
                format=ExportFormat.WORD,
                content=buffer.read(),
                metadata={
                    "generated_at": datetime.now().isoformat(),
                    "template": template.get("name", "未知模板")
                }
            )

        except ImportError:
            return ExportResult(
                success=False,
                format=ExportFormat.WORD,
                error="需要安装python-docx库: pip install python-docx"
            )
        except Exception as e:
            return ExportResult(
                success=False,
                format=ExportFormat.WORD,
                error=str(e)
            )

    def _add_table_to_doc(self, doc, field_def: Dict, rows: List):
        """添加表格到Word文档"""
        columns = field_def.get("columns", [])

        table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
        table.style = 'Table Grid'

        # 表头
        for i, col in enumerate(columns):
            cell = table.rows[0].cells[i]
            cell.text = col.get("label", col.get("id", ""))

        # 数据行
        for row_idx, row in enumerate(rows):
            for col_idx, col in enumerate(columns):
                col_id = col.get("id", "")
                cell_value = row.get(col_id, "") if isinstance(row, dict) else ""
                table.rows[row_idx + 1].cells[col_idx].text = str(cell_value)


# ==================== 文档导出引擎 ====================

class DocumentExporter:
    """
    文档导出引擎

    支持多种格式的文档导出。
    """

    def __init__(self, config: DocumentExporterConfig = None):
        self.config = config or DocumentExporterConfig()

        # 注册导出器
        self._exporters: Dict[ExportFormat, BaseExporter] = {
            ExportFormat.HTML: HTMLExporter(self.config),
            ExportFormat.MARKDOWN: MarkdownExporter(self.config),
            ExportFormat.JSON: JSONExporter(self.config),
            ExportFormat.WORD: WordExporter(self.config),
        }

    def set_format(self, format: ExportFormat):
        """设置导出格式"""
        self.config.format = format

    async def export(
        self,
        document_data: Dict,
        template: Dict,
        format: ExportFormat = None
    ) -> ExportResult:
        """
        导出文档

        Args:
            document_data: 文档数据
            template: 模板定义
            format: 导出格式

        Returns:
            ExportResult: 导出结果
        """
        export_format = format or self.config.format

        exporter = self._exporters.get(export_format)
        if not exporter:
            return ExportResult(
                success=False,
                format=export_format,
                error=f"不支持的导出格式: {export_format.value}"
            )

        return await exporter.export(document_data, template)

    def get_supported_formats(self) -> List[ExportFormat]:
        """获取支持的导出格式"""
        return list(self._exporters.keys())


# ==================== 便捷函数 ====================

_exporter_instance: Optional[DocumentExporter] = None


def get_exporter(config: DocumentExporterConfig = None) -> DocumentExporter:
    """获取导出器单例"""
    global _exporter_instance
    if _exporter_instance is None:
        _exporter_instance = DocumentExporter(config)
    return _exporter_instance


async def export_document_async(
    document_data: Dict,
    template: Dict,
    format: str = "html"
) -> ExportResult:
    """
    导出文档的便捷函数

    Args:
        document_data: 文档数据
        template: 模板定义
        format: 导出格式 (word/pdf/html/markdown/json)

    Returns:
        ExportResult: 导出结果
    """
    exporter = get_exporter()
    export_format = ExportFormat(format)
    return await exporter.export(document_data, template, export_format)
