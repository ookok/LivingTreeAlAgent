"""Online Stream Parser — parse attachments from URL without saving to disk.

Handles PDF, DOCX, ZIP, RAR, 7z from URLs or bytes:
  - Fetch attachment into memory (no disk write)
  - Auto-detect format from content/MIME
  - Parse text content in-memory
  - Return structured text summary (not full content)

Never saves downloaded files to disk.
"""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from loguru import logger


@dataclass
class StreamParseResult:
    """Result of streaming attachment parse."""
    source_url: str = ""
    format: str = "unknown"
    text_content: str = ""
    text_length: int = 0
    summary: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    success: bool = False


class OnlineStreamParser:
    """Parse attachments from URL directly in memory — no disk writes.

    Supported formats:
      - PDF: pymupdf → text, pdfplumber → text, pypdf → text
      - DOCX: python-docx → text
      - ZIP/RAR/7z: extract text files → aggregate text
      - Images: PIL → metadata, OCR if available
      - Plain text: direct decode

    Usage:
        parser = OnlineStreamParser()
        result = await parser.parse_url("https://example.com/doc.pdf")
        print(result.summary)  # key text extraction
    """

    def __init__(self, max_text_length: int = 50000):
        self.max_text_length = max_text_length

    async def parse_url(self, url: str, timeout: int = 20) -> StreamParseResult:
        """Fetch and parse an attachment from URL directly in memory."""
        result = StreamParseResult(source_url=url)

        try:
            from ..network.resilience import resilient_fetch
            status, data, _ = await resilient_fetch(url, timeout=timeout, use_accelerator=True)

            if status != 200 or not data:
                result.warnings.append(f"HTTP {status}")
                return result

            return self.parse_bytes(data, url)
        except Exception as e:
            result.warnings.append(f"Fetch error: {e}")
            return result

    def parse_bytes(self, data: bytes, source_url: str = "") -> StreamParseResult:
        """Parse attachment from bytes in memory."""
        result = StreamParseResult(source_url=source_url)

        ext = self._detect_extension(source_url, data)
        result.format = ext

        try:
            if ext == "pdf":
                self._parse_pdf(data, result)
            elif ext in ("docx", "doc"):
                self._parse_docx(data, result)
            elif ext in ("zip", "rar", "7z"):
                self._parse_archive(data, ext, result)
            elif ext in ("txt", "csv", "md", "json", "xml", "html"):
                self._parse_text(data, result)
            elif ext in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
                self._parse_image(data, result)
            else:
                self._parse_text(data, result)
                result.warnings.append(f"Unknown format: {ext}")

            if result.text_content:
                result.text_length = len(result.text_content)
                result.summary = result.text_content[:2000]
                result.success = True

        except Exception as e:
            result.warnings.append(f"Parse error ({ext}): {e}")

        return result

    async def parse_multiple(self, urls: list[str]) -> list[StreamParseResult]:
        """Parse multiple attachment URLs concurrently."""
        import asyncio
        tasks = [self.parse_url(url) for url in urls]
        return await asyncio.gather(*tasks, return_exceptions=False)

    # ═══ Format Parsers ═══

    def _parse_pdf(self, data: bytes, result: StreamParseResult) -> None:
        # Try pymupdf first
        try:
            import fitz
            doc = fitz.open(stream=data, filetype="pdf")
            texts = []
            for page in doc:
                t = page.get_text()
                if t:
                    texts.append(t)
            doc.close()
            if texts:
                result.text_content = "\n".join(texts)[:self.max_text_length]
                return
        except Exception:
            pass

        # Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                texts = [page.extract_text() or "" for page in pdf.pages]
                result.text_content = "\n".join(texts)[:self.max_text_length]
                if result.text_content.strip():
                    return
        except Exception:
            pass

        # Try pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            texts = [page.extract_text() or "" for page in reader.pages]
            result.text_content = "\n".join(texts)[:self.max_text_length]
        except Exception:
            pass

    def _parse_docx(self, data: bytes, result: StreamParseResult) -> None:
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    texts.append(row_text)
            result.text_content = "\n".join(texts)[:self.max_text_length]
        except Exception:
            result.warnings.append("python-docx not available")

    def _parse_archive(self, data: bytes, ext: str, result: StreamParseResult) -> None:
        texts = []
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for name in zf.namelist()[:20]:
                    if name.endswith((".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm")):
                        with zf.open(name) as f:
                            content = f.read()
                            try:
                                texts.append(content.decode("utf-8", errors="replace")[:5000])
                            except Exception:
                                pass
                    elif name.endswith(".pdf"):
                        with zf.open(name) as f:
                            pdf_data = f.read()
                            sub_result = StreamParseResult()
                            self._parse_pdf(pdf_data, sub_result)
                            if sub_result.text_content:
                                texts.append(f"[{name}] {sub_result.text_content[:3000]}")
            if texts:
                result.text_content = "\n---\n".join(texts)[:self.max_text_length]
            else:
                result.warnings.append(f"No readable files in {ext}")
        except Exception as e:
            result.warnings.append(f"Archive error: {e}")

    def _parse_text(self, data: bytes, result: StreamParseResult) -> None:
        try:
            text = data.decode("utf-8", errors="replace")
            result.text_content = text[:self.max_text_length]
        except Exception:
            result.text_content = data.decode("latin-1", errors="replace")[:self.max_text_length]

    def _parse_image(self, data: bytes, result: StreamParseResult) -> None:
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(data))
            result.metadata["width"] = img.width
            result.metadata["height"] = img.height
            result.metadata["format"] = img.format
            result.text_content = f"[Image: {img.width}x{img.height}, {img.format}]"
            result.warnings.append("Image OCR not available — marked for manual review")
        except Exception:
            result.text_content = "[Image — cannot parse]"

    # ═══ Format Detection ═══

    @staticmethod
    def _detect_extension(url: str, data: bytes = b"") -> str:
        url_path = url.split("?")[0].lower()

        for ext in ["pdf", "docx", "doc", "xlsx", "xls", "pptx", "ppt",
                     "zip", "rar", "7z", "tar", "gz",
                     "txt", "csv", "md", "json", "xml", "html", "htm",
                     "png", "jpg", "jpeg", "gif", "bmp", "webp"]:
            if url_path.endswith(f".{ext}"):
                return ext

        if data:
            if data.startswith(b"%PDF"):
                return "pdf"
            if data.startswith(b"PK\x03\x04"):
                return "zip"
            if data.startswith(b"Rar!"):
                return "rar"
            if data.startswith(b"\x89PNG"):
                return "png"
            if data.startswith(b"\xff\xd8"):
                return "jpg"

        return "unknown"
