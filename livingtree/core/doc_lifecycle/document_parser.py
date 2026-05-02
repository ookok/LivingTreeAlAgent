"""
文档解析引擎
DocLifecycle 文档解析引擎 - 支持多种格式文档解析
"""

import hashlib
import logging
import os
import subprocess
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from .models import DocumentInfo, DocumentType

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析文档，返回(文本内容, 元数据)"""
        pass
    
    @abstractmethod
    def get_type(self) -> DocumentType:
        """获取支持的文档类型"""
        pass
    
    def can_parse(self, file_path: str) -> bool:
        """检查是否支持解析此文件"""
        return True


class TextParser(DocumentParser):
    """纯文本解析器"""
    
    def get_type(self) -> DocumentType:
        return DocumentType.TEXT
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.txt', '.log', '.md', '.rst', '.text']
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析纯文本文件"""
        try:
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # 如果都失败，使用二进制读取
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='ignore')
            
            metadata = {
                "lines": len(content.splitlines()),
                "characters": len(content),
                "words": len(content.split())
            }
            
            return content, metadata
        except Exception as e:
            logger.error(f"Failed to parse text file {file_path}: {e}")
            raise


class MarkdownParser(DocumentParser):
    """Markdown解析器"""
    
    def get_type(self) -> DocumentType:
        return DocumentType.MARKDOWN
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.md', '.markdown', '.mdown']
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 统计信息
            lines = content.splitlines()
            metadata = {
                "lines": len(lines),
                "characters": len(content),
                "words": len(content.split()),
                "headings": sum(1 for line in lines if line.startswith('#')),
                "code_blocks": content.count('```'),
                "links": content.count('[')
            }
            
            return content, metadata
        except Exception as e:
            logger.error(f"Failed to parse markdown file {file_path}: {e}")
            raise


class CSVParser(DocumentParser):
    """CSV解析器"""
    
    def get_type(self) -> DocumentType:
        return DocumentType.CSV
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.csv', '.tsv']
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析CSV文件"""
        try:
            import csv
            
            delimiter = ',' if file_path.endswith('.csv') else '\t'
            rows = []
            
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                reader = csv.reader(f, delimiter=delimiter)
                for row in reader:
                    rows.append(row)
            
            if rows:
                header = rows[0] if rows else []
                data_rows = rows[1:] if len(rows) > 1 else []
                
                content = '\n'.join([','.join(row) for row in rows])
                
                metadata = {
                    "rows": len(data_rows),
                    "columns": len(header),
                    "headers": header[:10],  # 只保存前10个列名
                    "has_header": True
                }
                
                return content, metadata
            
            return "", {"rows": 0, "columns": 0}
        except Exception as e:
            logger.error(f"Failed to parse CSV file {file_path}: {e}")
            raise


class JSONParser(DocumentParser):
    """JSON解析器"""
    
    def get_type(self) -> DocumentType:
        return DocumentType.JSON
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.json', '.jsonl']
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析JSON文件"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                if file_path.endswith('.jsonl'):
                    data = [json.loads(line) for line in f if line.strip()]
                else:
                    data = json.load(f)
            
            # 转为字符串
            content = json.dumps(data, ensure_ascii=False, indent=2)
            
            metadata = {
                "type": type(data).__name__,
                "keys": list(data.keys()) if isinstance(data, dict) else [],
                "length": len(data) if isinstance(data, (list, dict)) else 0
            }
            
            return content, metadata
        except Exception as e:
            logger.error(f"Failed to parse JSON file {file_path}: {e}")
            raise


class PDFParser(DocumentParser):
    """PDF解析器"""
    
    def get_type(self) -> DocumentType:
        return DocumentType.PDF
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext == '.pdf'
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析PDF文件"""
        try:
            # 尝试使用PyPDF2
            try:
                import PyPDF2
                
                text_parts = []
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata = {
                        "pages": len(reader.pages),
                        "title": reader.metadata.get('/Title', '') if reader.metadata else ''
                    }
                    
                    for page_num, page in enumerate(reader.pages):
                        text = page.extract_text()
                        text_parts.append(f"=== Page {page_num + 1} ===\n{text}")
                
                content = '\n\n'.join(text_parts)
                return content, metadata
                
            except ImportError:
                # 尝试使用pdfplumber
                try:
                    import pdfplumber
                    
                    text_parts = []
                    metadata = {"pages": 0}
                    
                    with pdfplumber.open(file_path) as pdf:
                        metadata["pages"] = len(pdf.pages)
                        for page_num, page in enumerate(pdf.pages):
                            text = page.extract_text() or ""
                            text_parts.append(f"=== Page {page_num + 1} ===\n{text}")
                    
                    content = '\n\n'.join(text_parts)
                    return content, metadata
                    
                except ImportError:
                    # 尝试使用命令行工具
                    return self._parse_with_pdftotext(file_path)
                    
        except Exception as e:
            logger.error(f"Failed to parse PDF file {file_path}: {e}")
            raise
    
    def _parse_with_pdftotext(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """使用pdftotext命令行工具解析"""
        try:
            result = subprocess.run(
                ['pdftotext', '-layout', file_path, '-'],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0:
                return result.stdout, {"pages": "unknown", "method": "pdftotext"}
            else:
                raise Exception(f"pdftotext failed: {result.stderr}")
        except FileNotFoundError:
            raise Exception("pdftotext not found. Please install poppler-utils")
        except Exception as e:
            raise Exception(f"pdftotext parsing failed: {e}")


class WordParser(DocumentParser):
    """Word文档解析器"""
    
    def get_type(self) -> DocumentType:
        return DocumentType.DOCX
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.docx', '.doc']
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析Word文档"""
        try:
            # 优先使用python-docx
            try:
                from docx import Document as DocxDocument
                
                doc = DocxDocument(file_path)
                
                # 提取文本
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                content = '\n\n'.join(paragraphs)
                
                metadata = {
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections)
                }
                
                # 提取表格内容
                for i, table in enumerate(doc.tables):
                    content += f"\n\n[Table {i+1}]\n"
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        content += ' | '.join(row_data) + '\n'
                
                return content, metadata
                
            except ImportError:
                # 尝试使用antiword (用于.doc)
                if file_path.endswith('.doc'):
                    return self._parse_with_antiword(file_path)
                else:
                    raise Exception("python-docx not installed")
                    
        except Exception as e:
            logger.error(f"Failed to parse Word file {file_path}: {e}")
            raise
    
    def _parse_with_antiword(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """使用antiword命令行工具解析.doc文件"""
        try:
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode == 0:
                return result.stdout, {"method": "antiword"}
            else:
                raise Exception(f"antiword failed: {result.stderr}")
        except FileNotFoundError:
            raise Exception("antiword not found. Please install antiword")
        except Exception as e:
            raise Exception(f"antiword parsing failed: {e}")


class ExcelParser(DocumentParser):
    """Excel解析器"""
    
    def get_type(self) -> DocumentType:
        return DocumentType.XLSX
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.xlsx', '.xls']
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析Excel文件"""
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            
            all_text = []
            metadata = {
                "sheets": wb.sheetnames,
                "sheet_count": len(wb.sheetnames),
                "total_rows": 0,
                "total_cols": 0
            }
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                all_text.append(f"=== Sheet: {sheet_name} ===")
                
                for row_num, row in enumerate(ws.iter_rows(values_only=True), 1):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        all_text.append(row_text)
                    metadata["total_rows"] += 1
                    metadata["total_cols"] = max(metadata["total_cols"], len(row))
            
            content = '\n'.join(all_text)
            return content, metadata
            
        except ImportError:
            raise Exception("openpyxl not installed")
        except Exception as e:
            logger.error(f"Failed to parse Excel file {file_path}: {e}")
            raise


class HTMLParser(DocumentParser):
    """HTML解析器"""
    
    def get_type(self) -> DocumentType:
        return DocumentType.HTML
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in ['.html', '.htm', '.xhtml']
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析HTML文件"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除脚本和样式
            for tag in soup(['script', 'style']):
                tag.decompose()
            
            # 提取文本
            content = soup.get_text(separator='\n', strip=True)
            
            metadata = {
                "title": soup.title.string if soup.title else '',
                "headers": len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])),
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img')),
                "tables": len(soup.find_all('table'))
            }
            
            return content, metadata
            
        except ImportError:
            raise Exception("beautifulsoup4 not installed")
        except Exception as e:
            logger.error(f"Failed to parse HTML file {file_path}: {e}")
            raise


class CodeParser(DocumentParser):
    """代码文件解析器"""
    
    CODE_EXTENSIONS = {
        '.py': 'python',
        '.js': 'javascript',
        '.ts': 'typescript',
        '.java': 'java',
        '.c': 'c',
        '.cpp': 'cpp',
        '.h': 'c_header',
        '.hpp': 'cpp_header',
        '.cs': 'csharp',
        '.go': 'go',
        '.rs': 'rust',
        '.rb': 'ruby',
        '.php': 'php',
        '.swift': 'swift',
        '.kt': 'kotlin',
        '.scala': 'scala',
        '.sql': 'sql',
        '.sh': 'bash',
        '.bat': 'batch',
        '.ps1': 'powershell',
        '.r': 'r',
        '.lua': 'lua',
        '.pl': 'perl',
        '.xml': 'xml',
        '.yaml': 'yaml',
        '.yml': 'yaml',
        '.toml': 'toml',
        '.ini': 'ini',
        '.cfg': 'config',
        '.conf': 'config',
        '.properties': 'properties',
    }
    
    def get_type(self) -> DocumentType:
        return DocumentType.CODE
    
    def can_parse(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.CODE_EXTENSIONS or ext.startswith('.mcp')
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """解析代码文件"""
        try:
            ext = Path(file_path).suffix.lower()
            language = self.CODE_EXTENSIONS.get(ext, 'unknown')
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            lines = content.splitlines()
            
            metadata = {
                "language": language,
                "lines": len(lines),
                "characters": len(content),
                "code_lines": sum(1 for line in lines if line.strip() and not line.strip().startswith('//') and not line.strip().startswith('#')),
                "comment_lines": sum(1 for line in lines if line.strip().startswith('//') or line.strip().startswith('#') or line.strip().startswith('/*') or line.strip().startswith('*')),
                "blank_lines": sum(1 for line in lines if not line.strip())
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Failed to parse code file {file_path}: {e}")
            raise


class DocumentParserEngine:
    """文档解析引擎"""
    
    _instance = None
    
    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        if self._initialized:
            return
        
        self._parsers: List[DocumentParser] = []
        self._register_default_parsers()
        self._initialized = True
        
        logger.info(f"DocumentParserEngine initialized with {len(self._parsers)} parsers")
    
    def _register_default_parsers(self):
        """注册默认解析器"""
        self._parsers = [
            TextParser(),
            MarkdownParser(),
            CSVParser(),
            JSONParser(),
            PDFParser(),
            WordParser(),
            ExcelParser(),
            HTMLParser(),
            CodeParser()
        ]
    
    def register_parser(self, parser: DocumentParser):
        """注册自定义解析器"""
        self._parsers.append(parser)
        logger.info(f"Registered parser for type: {parser.get_type()}")
    
    def parse(self, file_path: str) -> Tuple[DocumentInfo, str, Dict[str, Any]]:
        """解析文档"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # 创建文档信息
        stat = path.stat()
        
        # 确定文档类型
        doc_type = self._detect_type(file_path)
        
        doc_info = DocumentInfo(
            doc_id=self._generate_doc_id(file_path),
            file_path=str(path.absolute()),
            file_name=path.name,
            file_type=doc_type,
            file_size=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            accessed_at=datetime.fromtimestamp(stat.st_atime),
            hash_value=self._calculate_hash(file_path)
        )
        
        # 查找合适的解析器
        parser = self._find_parser(file_path)
        
        if parser:
            try:
                content, metadata = parser.parse(file_path)
                doc_info.metadata = metadata
                return doc_info, content, metadata
            except Exception as e:
                logger.error(f"Parser failed for {file_path}: {e}")
                # 返回基本信息，不返回内容
                return doc_info, "", {}
        
        # 无可用解析器，返回基本信息
        return doc_info, "", {}
    
    def _detect_type(self, file_path: str) -> DocumentType:
        """检测文档类型"""
        ext = Path(file_path).suffix.lower()
        
        type_map = {
            '.txt': DocumentType.TEXT,
            '.log': DocumentType.TEXT,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN,
            '.pdf': DocumentType.PDF,
            '.doc': DocumentType.DOC,
            '.docx': DocumentType.DOCX,
            '.xls': DocumentType.XLS,
            '.xlsx': DocumentType.XLSX,
            '.csv': DocumentType.CSV,
            '.json': DocumentType.JSON,
            '.jsonl': DocumentType.JSON,
            '.xml': DocumentType.XML,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.zip': DocumentType.ARCHIVE,
            '.rar': DocumentType.ARCHIVE,
            '.7z': DocumentType.ARCHIVE,
            '.tar': DocumentType.ARCHIVE,
            '.gz': DocumentType.ARCHIVE,
        }
        
        return type_map.get(ext, DocumentType.UNKNOWN)
    
    def _find_parser(self, file_path: str) -> Optional[DocumentParser]:
        """查找合适的解析器"""
        for parser in self._parsers:
            if parser.can_parse(file_path):
                return parser
        return None
    
    def _generate_doc_id(self, file_path: str) -> str:
        """生成文档ID"""
        return hashlib.md5(str(file_path).encode()).hexdigest()[:16]
    
    def _calculate_hash(self, file_path: str) -> str:
        """计算文件哈希"""
        hasher = hashlib.sha256()
        try:
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(8192), b''):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception:
            return ""
    
    def batch_parse(self, file_paths: List[str]) -> List[Tuple[DocumentInfo, str, Dict[str, Any]]]:
        """批量解析文档"""
        results = []
        for path in file_paths:
            try:
                result = self.parse(path)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to parse {path}: {e}")
                results.append(None)
        return [r for r in results if r is not None]


# 全局实例
_parser_engine: Optional[DocumentParserEngine] = None


def get_parser_engine() -> DocumentParserEngine:
    """获取文档解析引擎实例"""
    global _parser_engine
    if _parser_engine is None:
        _parser_engine = DocumentParserEngine()
    return _parser_engine
