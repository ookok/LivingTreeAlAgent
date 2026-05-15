"""LearningEngine — Meta-learning for self-evolving systems.

Replaces hardcoded templates with learned knowledge:
- TemplateLearner: dynamically generates task templates from KB + Distillation
- SkillDiscoverer: auto-discovers tools from codebase via Phage AST
- RoleGenerator: generates agent roles from task descriptions
- PatternExtractor: extracts reusable patterns from successful executions

All knowledge is stored in KnowledgeBase for future recall.
"""

from __future__ import annotations

import asyncio
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Optional

from loguru import logger


@dataclass
class LearnedTemplate:
    """A task template learned from experience or Distillation."""
    domain: str
    sections: list[str]
    source: str  # "distillation", "execution", "format_discovery", "kb_merge"
    learned_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    success_rate: float = 0.0
    use_count: int = 0

    def to_plan_steps(self) -> list[dict[str, Any]]:
        return [
            {"name": section, "action": "execute", "roles": ["analyst"],
             "description": f"Learned from {self.source}"}
            for section in self.sections
        ]


class TemplateLearner:
    """Learns task templates through Distillation + KB + successful executions.

    No hardcoded templates. Every template is either:
    1. Distilled from expert model on first use
    2. Extracted from existing documents via FormatDiscovery
    3. Merged from similar previously-learned templates in KB
    4. Cached after learning for fast recall
    """

    def __init__(self, kb: Any = None, distillation: Any = None, expert_config: Any = None):
        self.kb = kb
        self.distillation = distillation
        self.expert_config = expert_config
        self._cache: dict[str, LearnedTemplate] = {}
        self._load_cache()

    async def get_template(self, domain: str, goal: str = "") -> list[dict[str, Any]]:
        """Get or learn a task template for a domain.

        Priority:
        1. Cached (already learned)
        2. From KnowledgeBase (previously stored)
        3. From Distillation (ask expert model)
        4. From FormatDiscovery (scan documents)
        5. Merge similar domains
        6. Minimal fallback
        """
        # 1. Cache hit
        if domain in self._cache:
            tpl = self._cache[domain]
            tpl.use_count += 1
            return tpl.to_plan_steps()

        # 2. KB lookup
        if self.kb:
            try:
                kb_results = self.kb.search(f"template {domain} sections", top_k=3)
                for doc in kb_results:
                    if "sections" in doc.metadata:
                        tpl = LearnedTemplate(
                            domain=domain, sections=doc.metadata["sections"],
                            source="kb_retrieval",
                        )
                        self._cache[domain] = tpl
                        return tpl.to_plan_steps()
            except Exception:
                pass

        # 3. Distillation
        tpl = await self._learn_from_expert(domain, goal)
        if tpl and tpl.sections:
            self._cache[domain] = tpl
            self._save(tpl)
            return tpl.to_plan_steps()

        # 4. FormatDiscovery scan
        if self.kb:
            tpl = await self._learn_from_documents(domain)
            if tpl and tpl.sections:
                self._cache[domain] = tpl
                self._save(tpl)
                return tpl.to_plan_steps()

        # 5. Merge similar
        tpl = self._merge_similar(domain)
        if tpl and tpl.sections:
            self._cache[domain] = tpl
            return tpl.to_plan_steps()

        # 6. Minimal
        return [{"name": f"Execute: {goal or domain}", "action": "execute",
                 "description": f"Handle {domain} task"}]

    async def _learn_from_expert(self, domain: str, goal: str) -> Optional[LearnedTemplate]:
        if not self.distillation or not self.expert_config:
            return None
        try:
            prompt = (
                f"For a '{domain}' task{f' ({goal})' if goal else ''}, "
                f"what are the standard steps and sections?\n"
                f"Return exactly: a JSON array of step names, one per line.\n"
                f'Example: ["总论","工程分析","环境现状","结论"]'
            )
            response = await self.distillation.query_expert(prompt, self.expert_config)
            import json, re
            match = re.search(r'\[.*?\]', response, re.DOTALL)
            if match:
                sections = json.loads(match.group())
                return LearnedTemplate(domain=domain, sections=sections, source="distillation")
        except Exception as e:
            logger.debug(f"Template distillation failed: {e}")
        return None

    async def _learn_from_documents(self, domain: str) -> Optional[LearnedTemplate]:
        if not self.kb:
            return None
        try:
            docs = self.kb.search(domain, top_k=10, as_of=datetime.utcnow())
            headings: set[str] = set()
            for doc in docs:
                for line in doc.content.split("\n"):
                    line = line.strip()
                    if line.startswith("#") or (line and line[0].isdigit() and len(line) > 3):
                        headings.add(line.lstrip("#0123456789. ").strip()[:80])
            if len(headings) >= 3:
                return LearnedTemplate(
                    domain=domain, sections=sorted(headings)[:20],
                    source="format_discovery",
                )
        except Exception:
            pass
        return None

    def _merge_similar(self, domain: str) -> Optional[LearnedTemplate]:
        """Merge sections from similar domains in cache."""
        merged: list[str] = []
        for key, tpl in self._cache.items():
            if any(word in key for word in domain.split()) or any(word in domain for word in key.split()):
                for s in tpl.sections:
                    if s not in merged:
                        merged.append(s)
        if len(merged) >= 3:
            return LearnedTemplate(domain=domain, sections=merged, source="merge")
        return None

    def record_success(self, domain: str, success_rate: float) -> None:
        """Update template success rate after execution."""
        if domain in self._cache:
            self._cache[domain].success_rate = success_rate

    def _save(self, tpl: LearnedTemplate) -> None:
        if self.kb:
            try:
                from ..knowledge.knowledge_base import Document
                doc = Document(
                    title=f"template:{tpl.domain}",
                    content="\n".join(tpl.sections),
                    domain=tpl.domain,
                    metadata={"sections": tpl.sections, "source": tpl.source},
                    source="template_learner",
                )
                self.kb.add_knowledge(doc)
            except Exception:
                pass

    def _load_cache(self) -> None:
        if not self.kb:
            return
        try:
            for doc in self.kb.search("template:", top_k=50):
                if "sections" in doc.metadata:
                    domain = doc.metadata.get("domain", doc.domain)
                    if domain:
                        self._cache[domain] = LearnedTemplate(
                            domain=domain,
                            sections=doc.metadata["sections"],
                            source="kb_load",
                        )
        except Exception:
            pass


class SkillDiscoverer:
    """Auto-discovers tools from codebase and skills.

    No hardcoded tool list. Tools are discovered:
    1. From Phage AST scan of the codebase
    2. From SkillFactory registered skills
    3. From KnowledgeBase (previously discovered)
    """

    def __init__(self, phage: Any = None, skill_factory: Any = None, ast_parser: Any = None, kb: Any = None):
        self.phage = phage
        self.skill_factory = skill_factory
        self.ast_parser = ast_parser
        self.kb = kb
        self._discovered: dict[str, Any] = {}

    async def discover(self, codebase_path: str = ".") -> dict[str, Any]:
        """Discover all available tools from all sources."""
        tools: dict[str, Any] = {}

        # From SkillFactory
        if self.skill_factory:
            for name in self.skill_factory.discover_skills():
                tools[name] = {"name": name, "type": "skill", "source": "skill_factory"}

        # From Phage AST scan
        if self.phage and self.ast_parser:
            try:
                scan = await self.phage.scan_directory(codebase_path)
                for func_name in scan.get("top_functions", [])[:20]:
                    name = func_name.get("name", "")
                    if name and name not in tools:
                        tools[name] = {
                            "name": name, "type": "function",
                            "source": "phage_scan",
                            "file": func_name.get("file", ""),
                            "connections": func_name.get("connections", 0),
                        }
            except Exception:
                pass

        # From KB
        if self.kb:
            try:
                kb_tools = self.kb.search("tool:function", top_k=50)
                for doc in kb_tools:
                    name = doc.metadata.get("name", doc.title)
                    if name and name not in tools:
                        tools[name] = {
                            "name": name, "type": "kb_learned",
                            "source": "knowledge_base",
                        }
            except Exception:
                pass

        self._discovered = tools
        return tools

    def get_tool(self, name: str) -> Optional[dict]:
        return self._discovered.get(name)


class RoleGenerator:
    """Generates agent roles dynamically from task descriptions.

    No hardcoded role list. Roles are:
    1. Generated via Distillation for the domain
    2. Merged from similar domains
    3. Cached for reuse
    """

    def __init__(self, distillation: Any = None, expert_config: Any = None, kb: Any = None):
        self.distillation = distillation
        self.expert_config = expert_config
        self.kb = kb
        self._cache: dict[str, list[dict]] = {}

    async def generate_roles(self, domain: str, task_description: str = "") -> list[dict[str, Any]]:
        """Generate appropriate agent roles for a domain task."""
        if domain in self._cache:
            return self._cache[domain]

        if self.distillation and self.expert_config:
            try:
                prompt = (
                    f"For a '{domain}' task, what specialized agent roles are needed?\n"
                    'Return JSON array: [{"name":"role_name","capabilities":["cap1","cap2"]},...]\n'
                    "Use brief English role names."
                )
                response = await self.distillation.query_expert(prompt, self.expert_config)
                import json, re
                match = re.search(r'\[.*?\]', response, re.DOTALL)
                if match:
                    roles = json.loads(match.group())
                    self._cache[domain] = roles
                    return roles
            except Exception:
                pass

        # Generic fallback
        roles = [
            {"name": "analyst", "capabilities": ["analysis", "reasoning"]},
            {"name": "executor", "capabilities": ["execution", "tool_use"]},
        ]
        self._cache[domain] = roles
        return roles


# ═══════════════════════════════════════════════════════════════════════════════
# Merged from auto_knowledge_miner.py
# ═══════════════════════════════════════════════════════════════════════════════

import json
import re
import time
from collections import Counter, defaultdict
from pathlib import Path

MINER_DIR = Path(".livingtree/auto_miner")
CONFIDENCE_HIGH = 0.8


@dataclass
class ExtractedTemplate:
    name: str
    sections: list[str] = field(default_factory=list)
    boilerplate: dict[str, str] = field(default_factory=dict)
    structure_skeleton: str = ""
    source_docs: list[str] = field(default_factory=list)
    confidence: float = 0.0


@dataclass
class AutoMinedTerm:
    term: str
    category: str = ""
    frequency: int = 0
    contexts: list[str] = field(default_factory=list)


@dataclass
class CodePattern:
    name: str
    pattern_type: str = ""
    description: str = ""
    examples: list[str] = field(default_factory=list)
    frequency: int = 0


class AutoKnowledgeMiner:
    """Autonomous knowledge extraction from existing enterprise artifacts.

    Point it at any directory. It scans, parses, extracts, links — fully autonomous.
    """

    def __init__(self, hub=None):
        self._hub = hub
        self._templates: dict[str, ExtractedTemplate] = {}
        self._terms: dict[str, AutoMinedTerm] = {}
        self._code_patterns: dict[str, CodePattern] = {}
        self._cross_refs: dict[str, list[str]] = defaultdict(list)
        self._stats = {"docs_parsed": 0, "projects_scanned": 0, "terms_mined": 0, "templates_extracted": 0}
        self._load()

    async def mine_directory(self, root: str | Path) -> dict:
        root = Path(root)
        t0 = time.time()

        doc_files = list(root.rglob("*.docx")) + list(root.rglob("*.pdf")) + list(root.rglob("*.md"))
        if doc_files:
            await self._mine_documents(doc_files)

        code_indicators = list(root.glob("*.py")) + list(root.glob("package.json")) + list(root.glob("Cargo.toml"))
        if code_indicators:
            self._mine_code_project(root)

        self._build_cross_refs()
        self._synthesize_templates()

        self._stats["_last_mining_duration"] = time.time() - t0
        self._save()
        return self._stats

    async def _mine_documents(self, files: list[Path]):
        all_sections: dict[str, Counter] = defaultdict(Counter)
        all_boilerplate: dict[str, list[str]] = defaultdict(list)
        term_counter: Counter = Counter()
        term_contexts: dict[str, list[str]] = defaultdict(list)

        for path in files[:200]:
            try:
                text = self._extract_text(path)
                if not text or len(text) < 100:
                    continue

                sections = re.findall(r'^#+\s*(.+)$|^第[一二三四五六七八九十百千]+[章节条]\s*(.+)$|^(\d+[\.\、])\s*(.+)$', text, re.MULTILINE)
                for m in sections:
                    heading = next(s for s in m if s).strip()
                    all_sections[path.suffix][heading] += 1
                    start = text.find(heading)
                    if start >= 0:
                        content = text[start + len(heading):start + len(heading) + 300]
                        all_boilerplate[heading].append(content.strip())

                terms = self._extract_terms(text)
                for term, cat in terms:
                    term_counter[(term, cat)] += 1
                    idx = text.find(term)
                    if idx >= 0:
                        ctx = text[max(0, idx - 30):idx + len(term) + 30]
                        term_contexts[term].append(ctx)

                self._stats["docs_parsed"] += 1
            except Exception as e:
                logger.debug(f"Mine doc {path}: {e}")

        for (term, cat), freq in term_counter.most_common(500):
            if freq >= 2:
                self._terms[term] = AutoMinedTerm(
                    term=term, category=cat, frequency=freq,
                    contexts=term_contexts.get(term, [])[:5],
                )
        self._stats["terms_mined"] = len(self._terms)

        self._raw_sections = all_sections
        self._raw_boilerplate = all_boilerplate

    def _extract_terms(self, text: str) -> list[tuple[str, str]]:
        terms = []
        patterns = [
            (r'GB/?T?\s*\d+[\.-]\d+', "standard"),
            (r'[A-Z]{2,8}(?:\s*\d+)?', "abbreviation"),
            (r'\d+\.?\d*\s*(?:μg/m³|mg/m³|mg/L|dB|t/a|m³/h)', "metric"),
            (r'《([^》]+)》', "legal"),
            (r'(?:环评|环境影响|排放标准|总量控制|清洁生产|风险评价)', "technical"),
        ]
        for pat, cat in patterns:
            for m in re.finditer(pat, text):
                terms.append((m.group(0), cat))
        return terms

    def _mine_code_project(self, root: Path):
        patterns = []

        tech = []
        if (root / "pyproject.toml").exists():
            tech.append("Python")
        if (root / "package.json").exists():
            tech.append("Node.js")
        if (root / "go.mod").exists():
            tech.append("Go")
        if (root / "Cargo.toml").exists():
            tech.append("Rust")

        if tech:
            patterns.append(CodePattern(
                name="tech_stack", pattern_type="project_meta",
                description="Tech stack: " + ", ".join(tech),
            ))

        py_files = list(root.rglob("*.py"))
        if py_files:
            imports = Counter()
            for f in py_files[:50]:
                try:
                    for line in f.read_text(errors="replace").splitlines()[:100]:
                        m = re.match(r'^(?:from|import)\s+(\w+)', line.strip())
                        if m:
                            imports[m.group(1)] += 1
                except Exception:
                    continue

            top_imports = imports.most_common(10)
            if top_imports:
                patterns.append(CodePattern(
                    name="dependency_cluster", pattern_type="import_cluster",
                    description="Top imports: " + ", ".join(f"{n}({c})" for n, c in top_imports[:5]),
                ))

            func_names = Counter()
            class_names = Counter()
            for f in py_files[:30]:
                try:
                    for line in f.read_text(errors="replace").splitlines()[:50]:
                        fm = re.match(r'def\s+([a-z_]\w*)', line)
                        if fm: func_names[fm.group(1)] += 1
                        cm = re.match(r'class\s+([A-Z]\w*)', line)
                        if cm: class_names[cm.group(1)] += 1
                except Exception:
                    continue

            if func_names:
                patterns.append(CodePattern(
                    name="function_naming", pattern_type="naming_convention",
                    description=f"Functions ({len(func_names)} unique): " + ", ".join(n for n, _ in func_names.most_common(5)),
                ))

        self._code_patterns = {p.name: p for p in patterns}
        self._stats["projects_scanned"] += 1

    def _build_cross_refs(self):
        for term_name, term_obj in self._terms.items():
            for std in ["GB3095", "GB3096", "GB3838", "GB/T3840"]:
                if std in term_name or any(std in ctx for ctx in term_obj.contexts):
                    self._cross_refs[std].append(term_name)

        for pattern in self._code_patterns.values():
            if pattern.pattern_type == "import_cluster":
                for imp in re.findall(r'(\w+)\(\d+\)', pattern.description):
                    self._cross_refs["code"].append(imp)

    def _synthesize_templates(self):
        raw_sec = getattr(self, '_raw_sections', {})
        raw_bp = getattr(self, '_raw_boilerplate', {})

        if not raw_sec or not raw_bp:
            return

        for fmt, section_counts in raw_sec.items():
            total_docs = max(max(section_counts.values()), 1)
            common = {s: c for s, c in section_counts.items() if c / total_docs >= 0.5}
            if not common:
                continue

            name = f"auto-template-{fmt.strip('.')}"
            template = ExtractedTemplate(
                name=name,
                sections=list(common.keys()),
                source_docs=[f"{fmt} x{self._stats['docs_parsed']}"],
                confidence=min(1.0, len(common) / 10.0),
            )

            for section in common:
                examples = raw_bp.get(section, [])
                if examples:
                    template.boilerplate[section] = examples[0][:200]

            template.structure_skeleton = "\n".join(
                f"## {s}\n{template.boilerplate.get(s, '...')}" for s in common
            )

            self._templates[name] = template
        self._stats["templates_extracted"] = len(self._templates)

    def _extract_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".md":
            return path.read_text(errors="replace")[:50000]
        elif suffix == ".docx":
            try:
                from docx import Document
                doc = Document(str(path))
                return "\n".join(p.text for p in doc.paragraphs if p.text)[:50000]
            except ImportError:
                return ""
        elif suffix == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(path))
                return "\n".join(p.extract_text() or "" for p in reader.pages[:20])
            except ImportError:
                return ""
        elif suffix in (".txt", ".py", ".json", ".yaml", ".xml"):
            return path.read_text(errors="replace")[:50000]
        return ""

    def get_templates(self) -> list[ExtractedTemplate]:
        return sorted(self._templates.values(), key=lambda t: -t.confidence)

    def get_terms(self, category: str = "") -> list[AutoMinedTerm]:
        terms = list(self._terms.values())
        if category:
            terms = [t for t in terms if t.category == category]
        return sorted(terms, key=lambda t: -t.frequency)

    def get_patterns(self) -> list[CodePattern]:
        return list(self._code_patterns.values())

    def get_cross_refs(self, entity: str) -> list[str]:
        return self._cross_refs.get(entity, [])

    def get_stats(self) -> dict:
        return dict(self._stats)

    def _save(self):
        try:
            MINER_DIR.mkdir(parents=True, exist_ok=True)
            (MINER_DIR / "knowledge.json").write_text(json.dumps({
                "templates": {n: {"name": t.name, "sections": t.sections,
                                  "structure_skeleton": t.structure_skeleton,
                                  "confidence": t.confidence}
                              for n, t in self._templates.items()},
                "terms": {n: {"term": t.term, "category": t.category, "frequency": t.frequency}
                          for n, t in self._terms.items()},
                "code_patterns": {n: {"name": p.name, "pattern_type": p.pattern_type,
                                      "description": p.description}
                                  for n, p in self._code_patterns.items()},
                "cross_refs": dict(self._cross_refs),
                "stats": dict(self._stats),
            }, indent=2, ensure_ascii=False))
        except Exception:
            pass

    def _load(self):
        try:
            path = MINER_DIR / "knowledge.json"
            if path.exists():
                data = json.loads(path.read_text())
                for n, d in data.get("templates", {}).items():
                    self._templates[n] = ExtractedTemplate(**d)
                for n, d in data.get("terms", {}).items():
                    self._terms[n] = AutoMinedTerm(**d)
                for n, d in data.get("code_patterns", {}).items():
                    self._code_patterns[n] = CodePattern(**d)
                self._cross_refs = defaultdict(list, data.get("cross_refs", {}))
                self._stats = data.get("stats", {})
        except Exception:
            pass


_miner: AutoKnowledgeMiner | None = None


def get_miner(hub=None) -> AutoKnowledgeMiner:
    global _miner
    if _miner is None:
        _miner = AutoKnowledgeMiner(hub)
    return _miner


__all__ = [
    "LearnedTemplate",
    "TemplateLearner",
    "SkillDiscoverer",
    "RoleGenerator",
    "ExtractedTemplate",
    "AutoMinedTerm",
    "CodePattern",
    "AutoKnowledgeMiner",
    "get_miner",
]
