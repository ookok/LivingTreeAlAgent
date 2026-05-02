"""
样式学习器 - 从用户上传的文档中自动学习格式

核心功能：
1. 解析用户上传的Word文档
2. 提取样式定义（标题、正文、表格等）
3. 建立样式映射表
4. 支持增量学习和样式进化
"""
import json
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class StyleDefinition:
    """样式定义"""
    name: str
    type: str  # paragraph, character, table
    font_name: str = ""
    font_size: float = 12.0
    font_bold: bool = False
    font_italic: bool = False
    font_color: str = "000000"
    alignment: str = "left"
    line_spacing: float = 1.5
    space_before: float = 0.0
    space_after: float = 0.0
    indent_first_line: float = 0.0
    based_on: str = ""
    next_style: str = ""
    usage_count: int = 0


@dataclass
class TableStyleDefinition:
    """表格样式定义"""
    name: str
    alignment: str = "center"
    cell_padding: float = 5.0
    has_header: bool = True
    header_background: str = "E0E0E0"
    border_style: str = "single"
    border_width: float = 1.0
    border_color: str = "000000"


class StyleLearner:
    """
    样式学习器
    
    从用户上传的文档中自动学习格式样式，支持：
    1. 批量学习多个文档
    2. 样式融合（学习多个文档后融合成统一样式）
    3. 增量学习（持续学习新文档）
    4. 样式进化（根据使用反馈优化）
    """
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装")
        
        self.paragraph_styles: Dict[str, StyleDefinition] = {}
        self.character_styles: Dict[str, StyleDefinition] = {}
        self.table_styles: Dict[str, TableStyleDefinition] = {}
        self.learning_history: List[Dict[str, Any]] = []
        
        # 样式类型映射
        self.style_type_map = {
            'heading': ['heading', '标题', '章', '节'],
            'normal': ['normal', '正文', '普通'],
            'code': ['code', '代码', '程序'],
            'table': ['table', '表格'],
            'list': ['list', '列表'],
        }
    
    def learn_from_document(self, docx_path: str, document_type: str = "unknown") -> Dict[str, Any]:
        """
        从单个文档学习样式
        
        Args:
            docx_path: Word文档路径
            document_type: 文档类型（eia_report, feasibility_study等）
        
        Returns:
            学习结果摘要
        """
        if not os.path.exists(docx_path):
            return {'success': False, 'error': '文档不存在'}
        
        try:
            doc = Document(docx_path)
            
            # 学习段落样式
            self._learn_paragraph_styles(doc)
            
            # 学习表格样式
            self._learn_table_styles(doc)
            
            # 学习字符样式
            self._learn_character_styles(doc)
            
            # 记录学习历史
            self.learning_history.append({
                'document_path': docx_path,
                'document_type': document_type,
                'timestamp': self._get_timestamp(),
                'styles_learned': len(self.paragraph_styles)
            })
            
            return {
                'success': True,
                'document_path': docx_path,
                'paragraph_styles': len(self.paragraph_styles),
                'table_styles': len(self.table_styles),
                'character_styles': len(self.character_styles)
            }
        
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _learn_paragraph_styles(self, doc):
        """学习段落样式"""
        for paragraph in doc.paragraphs:
            style = paragraph.style
            
            if style.name in self.paragraph_styles:
                # 已存在，更新使用次数
                self.paragraph_styles[style.name].usage_count += 1
                continue
            
            # 解析样式属性
            font = style.font
            paragraph_format = style.paragraph_format
            
            style_def = StyleDefinition(
                name=style.name,
                type='paragraph',
                font_name=font.name or '',
                font_size=font.size.pt if font.size else 12.0,
                font_bold=font.bold or False,
                font_italic=font.italic or False,
                font_color=self._rgb_to_hex(font.color.rgb) if font.color and font.color.rgb else '000000',
                alignment=self._get_alignment_name(paragraph_format.alignment),
                line_spacing=paragraph_format.line_spacing or 1.5,
                space_before=paragraph_format.space_before.pt if paragraph_format.space_before else 0.0,
                space_after=paragraph_format.space_after.pt if paragraph_format.space_after else 0.0,
                indent_first_line=paragraph_format.first_line_indent.pt if paragraph_format.first_line_indent else 0.0,
                based_on=style.base_style.name if style.base_style else '',
                next_style=style.next_style.name if style.next_style else '',
                usage_count=1
            )
            
            self.paragraph_styles[style.name] = style_def
    
    def _learn_table_styles(self, doc):
        """学习表格样式"""
        for table in doc.tables:
            # 使用表格在文档中的位置作为标识
            table_key = f"Table_{id(table)}"
            
            if table_key in self.table_styles:
                self.table_styles[table_key].alignment = self._get_table_alignment(table)
                continue
            
            table_style = TableStyleDefinition(
                name=table_key,
                alignment=self._get_table_alignment(table),
                has_header=True,  # 默认有表头
                usage_count=1
            )
            
            self.table_styles[table_key] = table_style
    
    def _learn_character_styles(self, doc):
        """学习字符样式"""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.style and run.style.name:
                    if run.style.name in self.character_styles:
                        self.character_styles[run.style.name].usage_count += 1
                        continue
                    
                    font = run.font
                    char_style = StyleDefinition(
                        name=run.style.name,
                        type='character',
                        font_name=font.name or '',
                        font_size=font.size.pt if font.size else 12.0,
                        font_bold=font.bold or False,
                        font_italic=font.italic or False,
                        font_color=self._rgb_to_hex(font.color.rgb) if font.color and font.color.rgb else '000000',
                        usage_count=1
                    )
                    
                    self.character_styles[run.style.name] = char_style
    
    def learn_from_multiple_documents(self, doc_paths: List[str], document_type: str = "unknown") -> Dict[str, Any]:
        """
        从多个文档学习样式
        
        Args:
            doc_paths: Word文档路径列表
            document_type: 文档类型
        
        Returns:
            汇总学习结果
        """
        results = []
        total_styles = 0
        
        for path in doc_paths:
            result = self.learn_from_document(path, document_type)
            results.append(result)
            if result['success']:
                total_styles += result.get('paragraph_styles', 0)
        
        return {
            'success': all(r['success'] for r in results),
            'documents_processed': len(doc_paths),
            'total_styles_learned': total_styles,
            'details': results
        }
    
    def fuse_styles(self, target_styles: List[str] = None) -> Dict[str, StyleDefinition]:
        """
        融合多个样式成统一样式
        
        Args:
            target_styles: 目标样式名称列表，None表示全部融合
        
        Returns:
            融合后的样式字典
        """
        fused = {}
        
        if target_styles is None:
            target_styles = list(self.paragraph_styles.keys())
        
        for style_name in target_styles:
            # 找到所有相似的样式
            similar_styles = self._find_similar_styles(style_name)
            
            if len(similar_styles) == 1:
                fused[style_name] = similar_styles[0]
            elif len(similar_styles) > 1:
                fused[style_name] = self._average_styles(similar_styles)
        
        return fused
    
    def _find_similar_styles(self, style_name: str) -> List[StyleDefinition]:
        """找到相似的样式"""
        similar = []
        
        for name, style in self.paragraph_styles.items():
            # 检查是否属于同一类型
            if any(keyword in name.lower() for keyword in self.style_type_map.get('heading', [])):
                if any(keyword in style_name.lower() for keyword in self.style_type_map.get('heading', [])):
                    similar.append(style)
            elif any(keyword in name.lower() for keyword in self.style_type_map.get('normal', [])):
                if any(keyword in style_name.lower() for keyword in self.style_type_map.get('normal', [])):
                    similar.append(style)
        
        return similar
    
    def _average_styles(self, styles: List[StyleDefinition]) -> StyleDefinition:
        """对多个样式取平均值"""
        if not styles:
            return StyleDefinition(name='Unknown')
        
        first = styles[0]
        
        # 计算平均值
        avg_font_size = sum(s.font_size for s in styles) / len(styles)
        avg_line_spacing = sum(s.line_spacing for s in styles) / len(styles)
        avg_space_before = sum(s.space_before for s in styles) / len(styles)
        avg_space_after = sum(s.space_after for s in styles) / len(styles)
        
        # 投票决定字体名称（选择使用最频繁的）
        font_counts = {}
        for s in styles:
            font_counts[s.font_name] = font_counts.get(s.font_name, 0) + s.usage_count
        
        most_common_font = max(font_counts, key=font_counts.get, default=first.font_name)
        
        return StyleDefinition(
            name=f"Fused_{first.name}",
            type=first.type,
            font_name=most_common_font,
            font_size=round(avg_font_size, 1),
            font_bold=any(s.font_bold for s in styles),
            font_italic=any(s.font_italic for s in styles),
            font_color=first.font_color,
            alignment=first.alignment,
            line_spacing=round(avg_line_spacing, 1),
            space_before=round(avg_space_before, 1),
            space_after=round(avg_space_after, 1),
            usage_count=sum(s.usage_count for s in styles)
        )
    
    def get_style_mapping(self) -> Dict[str, str]:
        """获取样式映射表"""
        mapping = {}
        
        for name, style in self.paragraph_styles.items():
            # 根据样式特征判断用途
            if style.font_bold and style.font_size > 14:
                mapping[name] = 'heading_1'
            elif style.font_bold and style.font_size > 12:
                mapping[name] = 'heading_2'
            elif style.font_bold:
                mapping[name] = 'heading_3'
            else:
                mapping[name] = 'normal_text'
        
        return mapping
    
    def export_styles(self, output_path: str):
        """导出学习到的样式到JSON文件"""
        data = {
            'paragraph_styles': {k: self._style_to_dict(v) for k, v in self.paragraph_styles.items()},
            'table_styles': {k: self._table_style_to_dict(v) for k, v in self.table_styles.items()},
            'character_styles': {k: self._style_to_dict(v) for k, v in self.character_styles.items()},
            'style_mapping': self.get_style_mapping(),
            'learning_history': self.learning_history
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        print(f"✅ 样式已导出到: {output_path}")
    
    def import_styles(self, input_path: str):
        """从JSON文件导入样式"""
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        self.paragraph_styles = {k: self._dict_to_style(v) for k, v in data.get('paragraph_styles', {}).items()}
        self.table_styles = {k: self._dict_to_table_style(v) for k, v in data.get('table_styles', {}).items()}
        self.character_styles = {k: self._dict_to_style(v) for k, v in data.get('character_styles', {}).items()}
        self.learning_history = data.get('learning_history', [])
        
        print(f"✅ 样式已从 {input_path} 导入")
    
    def _rgb_to_hex(self, rgb) -> str:
        """将RGB颜色转换为十六进制"""
        if rgb:
            return f"{rgb.red:02X}{rgb.green:02X}{rgb.blue:02X}"
        return "000000"
    
    def _get_alignment_name(self, alignment) -> str:
        """获取对齐方式名称"""
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return 'center'
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return 'right'
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return 'justify'
        return 'left'
    
    def _get_table_alignment(self, table) -> str:
        """获取表格对齐方式"""
        alignment = table.alignment
        if alignment == WD_TABLE_ALIGNMENT.CENTER:
            return 'center'
        elif alignment == WD_TABLE_ALIGNMENT.RIGHT:
            return 'right'
        return 'left'
    
    def _get_timestamp(self) -> str:
        """获取时间戳"""
        from datetime import datetime
        return datetime.now().isoformat()
    
    def _style_to_dict(self, style: StyleDefinition) -> Dict[str, Any]:
        """样式转字典"""
        return {
            'name': style.name,
            'type': style.type,
            'font_name': style.font_name,
            'font_size': style.font_size,
            'font_bold': style.font_bold,
            'font_italic': style.font_italic,
            'font_color': style.font_color,
            'alignment': style.alignment,
            'line_spacing': style.line_spacing,
            'space_before': style.space_before,
            'space_after': style.space_after,
            'indent_first_line': style.indent_first_line,
            'based_on': style.based_on,
            'next_style': style.next_style,
            'usage_count': style.usage_count
        }
    
    def _dict_to_style(self, data: Dict[str, Any]) -> StyleDefinition:
        """字典转样式"""
        return StyleDefinition(
            name=data.get('name', ''),
            type=data.get('type', 'paragraph'),
            font_name=data.get('font_name', ''),
            font_size=data.get('font_size', 12.0),
            font_bold=data.get('font_bold', False),
            font_italic=data.get('font_italic', False),
            font_color=data.get('font_color', '000000'),
            alignment=data.get('alignment', 'left'),
            line_spacing=data.get('line_spacing', 1.5),
            space_before=data.get('space_before', 0.0),
            space_after=data.get('space_after', 0.0),
            indent_first_line=data.get('indent_first_line', 0.0),
            based_on=data.get('based_on', ''),
            next_style=data.get('next_style', ''),
            usage_count=data.get('usage_count', 0)
        )
    
    def _table_style_to_dict(self, style: TableStyleDefinition) -> Dict[str, Any]:
        """表格样式转字典"""
        return {
            'name': style.name,
            'alignment': style.alignment,
            'cell_padding': style.cell_padding,
            'has_header': style.has_header,
            'header_background': style.header_background,
            'border_style': style.border_style,
            'border_width': style.border_width,
            'border_color': style.border_color
        }
    
    def _dict_to_table_style(self, data: Dict[str, Any]) -> TableStyleDefinition:
        """字典转表格样式"""
        return TableStyleDefinition(
            name=data.get('name', ''),
            alignment=data.get('alignment', 'center'),
            cell_padding=data.get('cell_padding', 5.0),
            has_header=data.get('has_header', True),
            header_background=data.get('header_background', 'E0E0E0'),
            border_style=data.get('border_style', 'single'),
            border_width=data.get('border_width', 1.0),
            border_color=data.get('border_color', '000000')
        )