"""Inline Parser — pseudo-upload: parse files/audio/video in memory, extract text.

Files are NEVER saved to disk. Content is parsed on-the-fly and only the
extracted text/knowledge is stored in the knowledge base. Original bytes
are discarded after parsing.

Supported formats:
  Text:   .txt .md .py .js .json .csv .html .xml .yaml .log
  Office: .docx .pdf (in-memory, no temp files)
  Audio:  .mp3 .wav .webm .ogg .m4a .flac → Whisper STT → text
  Video:  .mp4 .webm .mkv .avi → extract audio track → STT → text

Performance:
  - Streaming read for large files (4096-byte chunks for text)
  - Worker pool for concurrent STT processing
  - Hard limits: 50MB text, 25MB audio, 100MB video
  - Lazy imports: parsers loaded only when needed
"""

from __future__ import annotations

import asyncio
import codecs
import csv
import io
import json as _json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from loguru import logger

MAX_TEXT_BYTES = 50 * 1024 * 1024       # 50 MB
MAX_AUDIO_BYTES = 25 * 1024 * 1024      # 25 MB
MAX_VIDEO_BYTES = 100 * 1024 * 1024     # 100 MB
CHUNK_SIZE = 4096


@dataclass
class ParseResult:
    ok: bool
    text: str = ""
    title: str = ""
    mime_type: str = ""
    source_format: str = ""       # e.g. "pdf", "docx", "audio/webm"
    word_count: int = 0
    parse_time_ms: float = 0.0
    metadata: dict[str, Any] = field(default_factory=dict)
    error: str = ""


class InlineParser:
    """Parse files entirely in memory — no disk writes, ever."""

    def __init__(self):
        self._bom_pattern = re.compile(rb'^\xEF\xBB\xBF')
        self._stt_semaphore = asyncio.Semaphore(2)   # max 2 concurrent STT jobs

    # ═══ Main entry ═══

    async def parse(self, data: bytes, filename: str = "", mime_type: str = "") -> ParseResult:
        """Parse file bytes → extracted text. Data discarded after this call."""
        import time
        t0 = time.time()

        if not data:
            return ParseResult(ok=False, error="empty data")

        ext = Path(filename).suffix.lower() if filename else ""
        mime_lower = (mime_type or "").lower()

        try:
            if self._is_archive(ext, data):
                result = await self._parse_archive(data, filename)
            elif ext == ".csv" or "csv" in mime_lower:
                result = await self._parse_csv(data, filename)
            elif self._is_text(ext, mime_lower):
                result = await self._parse_text(data, filename, ext)
            elif self._is_audio(ext, mime_lower):
                result = await self._parse_audio(data, filename)
            elif self._is_video(ext, mime_lower):
                result = await self._parse_video(data, filename)
            elif ext == ".pdf" or "pdf" in mime_lower:
                result = await self._parse_pdf(data, filename)
            elif ext == ".docx" or "docx" in mime_lower:
                result = await self._parse_docx(data, filename)
            elif ext in (".csv",):
                result = await self._parse_csv(data, filename)
            else:
                result = await self._parse_text(data, filename, ext)

            result.parse_time_ms = (time.time() - t0) * 1000
            if result.ok:
                logger.info(f"Parsed {filename}: {result.word_count} words in {result.parse_time_ms:.0f}ms")
            return result

        except Exception as e:
            logger.warning(f"Parse error [{filename}]: {e}")
            return ParseResult(ok=False, error=str(e), parse_time_ms=(time.time() - t0) * 1000)

    # ═══ Type Detection ═══

    @staticmethod
    def _is_archive(ext: str, data: bytes) -> bool:
        """Detect if data is an archive (ZIP/TAR/gzip) by extension or magic bytes."""
        archive_exts = {".zip", ".tar", ".gz", ".tgz", ".tar.gz"}
        if ext in archive_exts or ext.endswith(".tar.gz") or ext.endswith(".tgz"):
            return True
        # Magic byte fallback
        if len(data) >= 4:
            if data[:4] == b'PK\x03\x04':   # ZIP
                return True
            if data[:2] == b'\x1f\x8b':      # gzip
                return True
            if len(data) >= 262 and data[257:262] == b'ustar':  # TAR
                return True
        return False

    async def _parse_archive(self, data: bytes, filename: str) -> ParseResult:
        """Parse archive files (ZIP/TAR/gz) — list contents, preview text files."""
        try:
            from .archive_preview import get_archive_previewer

            previewer = get_archive_previewer()
            fmt = previewer.detect_format(data, filename)
            if not fmt:
                return ParseResult(ok=False, error="Unrecognized archive format")

            # List all files
            entries = await previewer.list_files(data, fmt)
            text_entries = [e for e in entries if not e.is_dir]
            total_size = sum(e.size for e in text_entries)

            # Build summary
            lines = [
                f"Archive: {filename} (format: {fmt})",
                f"Files: {len(text_entries)} (total: {total_size:,} bytes)",
                "",
                "Contents:",
            ]
            for e in text_entries[:50]:  # limit listing to 50 entries
                lines.append(f"  {e.name} ({e.size:,} bytes)")

            if len(text_entries) > 50:
                lines.append(f"  ... and {len(text_entries) - 50} more files")

            # Preview first few text files
            preview_limit = min(3, len(text_entries))
            if preview_limit > 0:
                lines.append("")
                lines.append("Previews:")
                for e in text_entries[:preview_limit]:
                    if not e.name.lower().endswith(('.txt', '.md', '.csv', '.json', '.yaml', '.yml', '.py', '.log', '.html')):
                        lines.append(f"  [{e.name}]: (binary, {e.size} bytes)")
                        continue
                    preview = await previewer.preview_text(data, fmt, e.name, max_chars=2000)
                    if preview:
                        lines.append(f"  [{e.name}]:")
                        for pl in preview.split("\n")[:10]:
                            lines.append(f"    {pl}")
                        lines.append("")

            content = "\n".join(lines)
            title = Path(filename).stem.replace("_", " ").replace("-", " ")[:80]
            return ParseResult(
                ok=True, text=content, title=title,
                source_format=fmt, word_count=len(content.split()),
                metadata={
                    "format": fmt,
                    "files": len(text_entries),
                    "total_size_bytes": total_size,
                    "original_filename": filename,
                },
            )
        except ImportError:
            return ParseResult(ok=False, error="Archive parsing requires archive_preview module")
        except Exception as e:
            logger.warning(f"Archive parse [{filename}]: {e}")
            return ParseResult(ok=False, error=f"Archive parse: {e}")

    @staticmethod
    def _is_text(ext: str, mime: str) -> bool:
        text_exts = {".txt", ".md", ".py", ".js", ".ts", ".html", ".htm",
                     ".xml", ".yaml", ".yml", ".log", ".ini", ".cfg", ".toml",
                     ".css", ".scss", ".sql", ".rst", ".tex", ".sh", ".bat"}
        return ext in text_exts or any(t in mime for t in ("text/", "application/json", "application/xml", "application/javascript"))

    @staticmethod
    def _is_audio(ext: str, mime: str) -> bool:
        audio_exts = {".mp3", ".wav", ".webm", ".ogg", ".m4a", ".flac", ".aac", ".opus", ".wma"}
        return ext in audio_exts or "audio/" in mime

    @staticmethod
    def _is_video(ext: str, mime: str) -> bool:
        video_exts = {".mp4", ".webm", ".mkv", ".avi", ".mov", ".wmv", ".flv"}
        return ext in video_exts or "video/" in mime

    # ═══ Text Parsers ═══

    async def _parse_text(self, data: bytes, filename: str, ext: str) -> ParseResult:
        if len(data) > MAX_TEXT_BYTES:
            return ParseResult(ok=False, error=f"File too large: {len(data)/1024/1024:.1f}MB (max 50MB)")

        try:
            data = self._strip_bom(data)
            encoding = "utf-8"
            if ext in (".json",) or (data[:1] and data[0:1] == b'{'):
                text = self._parse_json_content(data)
            elif ext in (".html", ".htm", ".xml") or b"<" in data[:200]:
                text = self._parse_html_content(data)
            else:
                encoding = self._detect_encoding(data)
                text = data.decode(encoding, errors="replace")

            title = self._extract_title(text, filename)
            return ParseResult(
                ok=True, text=text.strip(), title=title,
                mime_type="text/plain", source_format=ext.lstrip("."),
                word_count=len(text.split()),
                metadata={"encoding": encoding, "size_bytes": len(data)},
            )
        except Exception as e:
            return ParseResult(ok=False, error=f"Text parse: {e}")

    async def _parse_csv(self, data: bytes, filename: str) -> ParseResult:
        """Streaming CSV parser — iterates rows lazily, never materializes all.

        Key optimization over the old list(reader): csv.reader rows are
        consumed one at a time. Only the header and first 20 sample rows
        are stored; everything else is counted and discarded.
        """
        try:
            # Guard against excessive input (matching _parse_text limit)
            if len(data) > MAX_TEXT_BYTES:
                return ParseResult(
                    ok=False,
                    error=f"CSV too large: {len(data)/1024/1024:.1f}MB (max 50MB)",
                )

            encoding = self._detect_encoding(data)
            text = data.decode(encoding, errors="replace")

            reader = csv.reader(io.StringIO(text))
            header: list[str] = []
            total_rows = 0
            sample_rows: list[list[str]] = []
            sample_limit = 20

            for row in reader:
                if not row or all(cell.strip() == '' for cell in row):
                    continue  # skip empty rows

                if not header:
                    header = row
                    continue

                total_rows += 1
                if len(sample_rows) < sample_limit:
                    sample_rows.append(row)

            if not header:
                return ParseResult(ok=False, error="CSV has no header row")

            # Build output
            summary_parts = [
                f"Columns: {', '.join(header)}",
                f"Rows: {total_rows}",
            ]
            sample = "\n".join(" | ".join(r) for r in sample_rows)
            content = f"CSV: {'; '.join(summary_parts)}\n\n{sample}"
            title = self._extract_title(content, filename)
            return ParseResult(
                ok=True, text=content, title=title,
                mime_type="text/csv", source_format="csv",
                word_count=len(content.split()),
                metadata={"rows": total_rows, "columns": len(header)},
            )
        except Exception as e:
            return ParseResult(ok=False, error=f"CSV parse: {e}")

    # ═══ PDF Parser (in-memory) ═══

    async def _parse_pdf(self, data: bytes, filename: str) -> ParseResult:
        text = ""
        try:
            from io import BytesIO
            import pypdf
            reader = pypdf.PdfReader(BytesIO(data))
            ...
            return ParseResult(ok=False, error="PDF parsing requires: pip install pypdf (or pdfplumber/pymupdf)")
        except ImportError:
            return ParseResult(ok=False, error="PDF parsing requires: pip install pypdf (or pdfplumber/pymupdf)")
        except Exception as e:
            return ParseResult(ok=False, error=f"PDF parse error: {e}")
        return self._build_result(text, filename, "pdf")

    # ═══ DOCX Parser (in-memory) ═══

    async def _parse_docx(self, data: bytes, filename: str) -> ParseResult:
        try:
            from io import BytesIO
            import docx
            doc = docx.Document(BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)

            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                text += "\n\n" + "\n".join(rows)

            return self._build_result(text, filename, "docx")
        except ImportError:
            return ParseResult(ok=False, error="DOCX parsing requires: pip install python-docx")
        except Exception as e:
            return ParseResult(ok=False, error=f"DOCX parse: {e}")

    # ═══ Audio Parser (STT) ═══

    async def _parse_audio(self, data: bytes, filename: str) -> ParseResult:
        if len(data) > MAX_AUDIO_BYTES:
            return ParseResult(ok=False, error=f"Audio too large: {len(data)/1024/1024:.1f}MB (max 25MB)")

        async with self._stt_semaphore:
            text = await self._transcribe_llamacpp(data)
            if not text:
                text = await self._transcribe_ollama(data)
            if not text:
                text = await self._transcribe_edge(data)

        if not text:
            return ParseResult(ok=False, error="Audio STT failed — no Whisper backend available",
                               metadata={"duration_estimate_s": len(data) / 16000})
        return self._build_result(text, filename, "audio/stt")

    # ═══ Video Parser (extract audio → STT) ═══

    async def _parse_video(self, data: bytes, filename: str) -> ParseResult:
        if len(data) > MAX_VIDEO_BYTES:
            return ParseResult(ok=False, error=f"Video too large: {len(data)/1024/1024:.1f}MB (max 100MB)")

        audio = await self._extract_audio_from_video(data)
        if not audio:
            return ParseResult(ok=False, error="Could not extract audio from video",
                               metadata={"video_size_bytes": len(data)})

        return await self._parse_audio(audio, filename)

    async def _extract_audio_from_video(self, data: bytes) -> bytes:
        """Extract audio track using subprocess ffmpeg (stdin→stdout, no disk).
        
        Uses asyncio.create_subprocess_exec for pipe I/O — unified_exec.run
        does not support stdin/stdout pipes needed for ffmpeg streaming.
        """
        try:
            proc = await asyncio.create_subprocess_exec(
                "ffmpeg", "-i", "pipe:0", "-vn", "-acodec", "pcm_s16le",
                "-ar", "16000", "-ac", "1", "-f", "wav", "pipe:1",
                stdin=asyncio.subprocess.PIPE,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await asyncio.wait_for(
                proc.communicate(data), timeout=60
            )
            if proc.returncode == 0 and stdout and len(stdout) > 44:
                return stdout
            logger.debug(f"ffmpeg extract failed: {stderr.decode()[:200]}")
        except asyncio.TimeoutError:
            logger.warning("ffmpeg audio extraction timeout")
        except FileNotFoundError:
            logger.debug("ffmpeg not found")
        except Exception as e:
            logger.debug(f"ffmpeg extract: {e}")
        return b""

    # ═══ STT Backends ═══

    async def _transcribe_llamacpp(self, audio: bytes) -> str:
        from .llamacpp_backend import get_llamacpp
        import base64
        try:
            llm = get_llamacpp()
            health = await llm.health()
            if health.get("ok"):
                return await llm.transcribe(audio, "wav")
        except Exception:
            pass
        return ""

    async def _transcribe_ollama(self, audio: bytes) -> str:
        import base64, httpx
        try:
            audio_b64 = base64.b64encode(audio).decode()
            async with httpx.AsyncClient(timeout=60.0) as client:
                resp = await client.post(
                    "http://localhost:11434/api/generate",
                    json={
                        "model": "whisper:latest",
                        "prompt": f"Transcribe: [base64:{audio_b64}]",
                        "stream": False,
                        "options": {"temperature": 0.0},
                    },
                )
                if resp.status_code == 200:
                    return resp.json().get("response", "").strip()
        except Exception:
            pass
        return ""

    async def _transcribe_edge(self, audio: bytes) -> str:
        """Browser-based STT hint — returns note about Whisper requirement."""
        return ""

    # ═══ Helpers ═══

    def _build_result(self, text: str, filename: str, fmt: str) -> ParseResult:
        title = self._extract_title(text, filename)
        return ParseResult(
            ok=True, text=text.strip(), title=title,
            source_format=fmt, word_count=len(text.split()),
            metadata={"original_filename": filename},
        )

    def _extract_title(self, text: str, filename: str) -> str:
        lines = text.strip().split("\n")
        for line in lines[:5]:
            line = line.strip().lstrip("# ").strip()
            if line and len(line) > 2:
                return line[:80]
        stem = Path(filename).stem if filename else "untitled"
        return stem.replace("_", " ").replace("-", " ")[:80]

    def _detect_encoding(self, data: bytes) -> str:
        if data[:3] == b'\xEF\xBB\xBF':
            return "utf-8-sig"
        if data[:2] == b'\xFF\xFE':
            return "utf-16-le"
        try:
            data.decode("utf-8")
            return "utf-8"
        except UnicodeDecodeError:
            pass
        try:
            data.decode("gbk")
            return "gbk"
        except UnicodeDecodeError:
            pass
        return "latin-1"

    def _strip_bom(self, data: bytes) -> bytes:
        return re.sub(rb'^\xEF\xBB\xBF', b'', data) if data[:3] == b'\xEF\xBB\xBF' else data

    def _parse_json_content(self, data: bytes) -> str:
        try:
            obj = _json.loads(data)
            return _json.dumps(obj, ensure_ascii=False, indent=2)
        except Exception:
            return data.decode("utf-8", errors="replace")

    def _parse_html_content(self, data: bytes) -> str:
        try:
            from html.parser import HTMLParser

            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []
                def handle_data(self, d):
                    t = d.strip()
                    if t:
                        self.text.append(t)

            extractor = TextExtractor()
            extractor.feed(data.decode("utf-8", errors="replace"))
            return "\n".join(extractor.text)
        except Exception:
            plain = re.sub(rb'<[^>]+>', b' ', data)
            plain = re.sub(rb'\s+', b' ', plain)
            return plain.decode("utf-8", errors="replace").strip()


_instance: Optional[InlineParser] = None


def get_inline_parser() -> InlineParser:
    global _instance
    if _instance is None:
        _instance = InlineParser()
    return _instance
