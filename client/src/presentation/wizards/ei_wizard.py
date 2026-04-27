"""
EIWizard - 环评报告生成向导

引导用户完成报告生成的完整流程
"""

from PySide6.QtWidgets import (
    QWizard, QWizardPage, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QLineEdit, QTextEdit, QFileDialog,
    QProgressBar, QMessageBox, QFormLayout, QComboBox
)
from PySide6.QtCore import Qt, QThread, Signal, Slot

# 导入 EIAgentAdapter
try:
    from client.src.business.ei_agent.ei_agent_adapter import (
        get_ei_agent_adapter,
        submit_ei_task,
        EIAgentAdapter
    )
    _HAS_ADAPTER = True
except ImportError:
    _HAS_ADAPTER = False
    print("[EWizard] 警告: EIAgentAdapter 未找到，将使用降级方案")

# 异步任务工作线程
class AsyncTaskWorker(QThread):
    """异步任务执行器（在独立线程中运行异步任务）"""
    finished = Signal(dict)  # 任务完成信号
    error = Signal(str)    # 错误信号
    
    def __init__(self, task_type: str, params: dict):
        super().__init__()
        self.task_type = task_type
        self.params = params
        self.result = None
    
    def run(self):
        """执行异步任务"""
        try:
            import asyncio
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            # 提交任务
            task_id = submit_ei_task(self.task_type, self.params)
            
            # 等待任务完成（简化实现：轮询状态）
            import time
            adapter = get_ei_agent_adapter()
            max_wait = 60  # 最多等待60秒
            for _ in range(max_wait):
                status = adapter.get_task_status(task_id)
                if status and status.value in ["completed", "failed", "cancelled"]:
                    result = adapter.get_task_result(task_id)
                    self.finished.emit(result or {"status": "failed", "error": "无结果"})
                    return
                time.sleep(1)
            
            self.error.emit("任务超时")
        except Exception as e:
            self.error.emit(str(e))


class EIWizard(QWizard):
    """环评报告生成向导
    
    引导用户完成报告生成的完整流程
    使用 EIAgentAdapter 接入系统 Agent 架构
    """
    
    def __init__(self, parent=None):
        """初始化向导
        
        Args:
            parent: 父组件
        """
        super().__init__(parent)
        
        # 1. 创建 EIAgentAdapter（如果可用）
        self.adapter = None
        self.use_adapter = False
        
        if _HAS_ADAPTER:
            try:
                self.adapter = get_ei_agent_adapter()
                self.use_adapter = True
                print("[EIWizard] 使用 EIAgentAdapter")
            except Exception as e:
                print(f"[EIWizard] EIAgentAdapter 初始化失败: {e}")
        
        # 降级方案：使用旧的 EIAgent
        if not self.use_adapter:
            try:
                from client.src.business.ei_agent import EIAgent
                self.agent = EIAgent()
                print("[EIWizard] 使用降级方案: EIAgent")
            except ImportError:
                self.agent = None
                print("[EIWizard] 警告: 无可用 Agent")
        
        # 2. 设置向导属性
        self.setWindowTitle("环评报告生成向导")
        self.setWizardStyle(QWizard.ModernStyle)
        
        # 3. 创建6个步骤
        self.addPage(Step1UploadDocument())
        self.addPage(Step2InputProjectInfo())
        self.addPage(Step2_5SpatialData())  # 新增：坐标和敏感点输入
        self.addPage(Step3AutoLearning())
        self.addPage(Step4GenerateReport())
        self.addPage(Step5ExportReport())
        
        # 4. 当前状态
        self.document_path = ""
        self.project_info = {}
        self.report_content = ""
    
    # ========== 与 Agent 集成 ==========
    
    def _call_agent(self, tool_name: str, **kwargs):
        """调用 Agent 的工具
        
        优先使用 EIAgentAdapter，降级使用 EIAgent
        
        Args:
            tool_name: 工具名称
            **kwargs: 工具参数
            
        Returns:
            dict: 工具返回结果
        """
        
        # 使用 EIAgentAdapter
        if self.use_adapter and self.adapter:
            try:
                # 映射工具名到任务类型
                tool_to_task = {
                    "get_coordinates": "environmental_impact_assessment",
                    "analyze_sensitive_zones": "risk_assessment",
                    "analyze_document": "report_generation",
                    "generate_report": "report_generation",
                    "query_regulation": "regulation_retrieval",
                    "query_pollution_coefficient": "pollution_coefficient",
                }
                
                task_type = tool_to_task.get(tool_name, "report_generation")
                
                # 提交任务
                task_id = submit_ei_task(task_type, kwargs)
                
                # 等待任务完成（简化实现：轮询）
                import time
                max_wait = 60
                for _ in range(max_wait):
                    from client.src.business.ei_agent.ei_agent_adapter import get_ei_agent_adapter
                    adapter = get_ei_agent_adapter()
                    status = adapter.get_task_status(task_id)
                    if status and status.value in ["completed", "failed", "cancelled"]:
                        result = adapter.get_task_result(task_id)
                        return result or {"error": "无结果"}
                    time.sleep(1)
                
                return {"error": "任务超时"}
                
            except Exception as e:
                print(f"[EIWizard] Adapter 调用失败: {e}")
                # 降级到旧方案
        
        # 降级方案：使用旧的 EIAgent
        if hasattr(self, 'agent') and self.agent:
            try:
                result = self.agent.execute_tool(tool_name, **kwargs)
                return result
            except Exception as e:
                QMessageBox.critical(self, "错误", f"调用 Agent 失败: {str(e)}")
                return {"error": str(e)}
        
        return {"error": "无可用 Agent"}


class Step1UploadDocument(QWizardPage):
    """步骤1: 上传文档"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("步骤1: 上传文档")
        self.setSubTitle("请上传环评报告和法律法规文件")
        
        # 创建UI
        layout = QVBoxLayout()
        
        # 说明标签
        label = QLabel("请上传以下文件：\n1. 环评报告（PDF/Word/TXT）\n2. 法律法规文件（PDF/Word/TXT）")
        layout.addWidget(label)
        
        # 上传按钮
        self.upload_btn = QPushButton("选择文件")
        self.upload_btn.clicked.connect(self._upload_file)
        layout.addWidget(self.upload_btn)
        
        # 已上传文件列表
        self.file_label = QLabel("已选择文件：无")
        layout.addWidget(self.file_label)
        
        self.setLayout(layout)
    
    def _upload_file(self):
        """上传文件"""
        
        # 打开文件对话框
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择文件",
            "",
            "所有支持的文件 (*.pdf *.docx *.doc *.txt);;PDF 文件 (*.pdf);;Word 文件 (*.docx *.doc);;文本文件 (*.txt)"
        )
        
        if file_path:
            # 保存文件路径
            wizard = self.wizard()
            wizard.document_path = file_path
            
            # 显示文件路径
            self.file_label.setText(f"已选择文件：{file_path}")
    
    def validatePage(self):
        """验证页面
        
        Returns:
            bool: 是否通过验证
        """
        
        wizard = self.wizard()
        
        if not wizard.document_path:
            QMessageBox.warning(self, "警告", "请先上传文件")
            return False
        
        return True


class Step2InputProjectInfo(QWizardPage):
    """步骤2: 输入项目信息"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("步骤2: 输入项目信息")
        self.setSubTitle("请输入项目的基本信息")
        
        # 创建UI
        layout = QFormLayout()
        
        # 项目名称
        self.project_name_edit = QLineEdit()
        layout.addRow("项目名称:", self.project_name_edit)
        
        # 建设地点
        self.location_edit = QLineEdit()
        layout.addRow("建设地点:", self.location_edit)
        
        # 行业类别
        self.industry_combo = QComboBox()
        self.industry_combo.addItems(["制造业", "化工业", "电力业", "建筑业", "其他"])
        layout.addRow("行业类别:", self.industry_combo)
        
        # 建设单位
        self.builder_edit = QLineEdit()
        layout.addRow("建设单位:", self.builder_edit)
        
        # 项目规模
        self.scale_combo = QComboBox()
        self.scale_combo.addItems(["大型", "中型", "小型"])
        layout.addRow("项目规模:", self.scale_combo)
        
        self.setLayout(layout)
    
    def validatePage(self):
        """验证页面
        
        Returns:
            bool: 是否通过验证
        """
        
        # 检查必填字段
        if not self.project_name_edit.text():
            QMessageBox.warning(self, "警告", "请输入项目名称")
            return False
        
        if not self.location_edit.text():
            QMessageBox.warning(self, "警告", "请输入建设地点")
            return False
        
        # 保存项目信息
        wizard = self.wizard()
        wizard.project_info = {
            "project_name": self.project_name_edit.text(),
            "location": self.location_edit.text(),
            "industry": self.industry_combo.currentText(),
            "builder": self.builder_edit.text(),
            "scale": self.scale_combo.currentText()
        }
        
        return True
    

class Step2_5SpatialData(QWizardPage):
    """步骤2.5: 输入坐标和敏感点"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("步骤2.5: 坐标和敏感点")
        self.setSubTitle("请输入项目坐标，系统将自动分析环境敏感点")
        
        # 创建UI
        layout = QFormLayout()
        
        # 坐标输入 - 纬度
        self.lat_edit = QLineEdit()
        self.lat_edit.setPlaceholderText("例如: 31.2304")
        layout.addRow("纬度:", self.lat_edit)
        
        # 坐标输入 - 经度
        self.lon_edit = QLineEdit()
        self.lon_edit.setPlaceholderText("例如: 121.4737")
        layout.addRow("经度:", self.lon_edit)
        
        # 获取坐标按钮
        self.get_coord_btn = QPushButton("自动获取坐标（根据建设地点）")
        self.get_coord_btn.clicked.connect(self._auto_get_coordinates)
        layout.addRow("", self.get_coord_btn)
        
        # 搜索半径
        self.radius_edit = QLineEdit("5000")
        self.radius_edit.setPlaceholderText("单位：米")
        layout.addRow("搜索半径（米）:", self.radius_edit)
        
        # 分析敏感点按钮
        self.analyze_btn = QPushButton("分析环境敏感点")
        self.analyze_btn.clicked.connect(self._analyze_sensitive_zones)
        layout.addRow("", self.analyze_btn)
        
        # 敏感点显示
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        self.result_text.setMaximumHeight(200)
        layout.addRow("敏感点:", self.result_text)
        
        self.setLayout(layout)
    
    def initializePage(self):
        """初始化页面"""
        # 如果已经有坐标，显示它们
        wizard = self.wizard()
        if hasattr(wizard, 'project_coords'):
            self.lat_edit.setText(str(wizard.project_coords.get('latitude', '')))
            self.lon_edit.setText(str(wizard.project_coords.get('longitude', '')))
        
        # 如果已经有敏感点分析结果，显示它们
        if hasattr(wizard, 'sensitive_zones'):
            self._display_sensitive_zones(wizard.sensitive_zones)
    
    def _auto_get_coordinates(self):
        """自动获取坐标（根据建设地点）"""
        wizard = self.wizard()
        location = wizard.project_info.get('location', '')
        
        if not location:
            QMessageBox.warning(self, "警告", "请先在步骤2中输入建设地点")
            return
        
        try:
            # 调用 EIAgent 的 get_coordinates 工具
            result = wizard._call_agent(
                "get_coordinates",
                address=location
            )
            
            if result.get("success"):
                lat = result.get("latitude", 0)
                lon = result.get("longitude", 0)
                
                self.lat_edit.setText(str(lat))
                self.lon_edit.setText(str(lon))
                
                # 保存到 wizard
                wizard.project_coords = {"latitude": lat, "longitude": lon}
                
                QMessageBox.information(self, "成功", f"已获取坐标：\n纬度: {lat}\n经度: {lon}")
            else:
                QMessageBox.warning(self, "警告", f"获取坐标失败: {result.get('error', '未知错误')}")
        
        except Exception as e:
            QMessageBox.critical(self, "错误", f"获取坐标失败: {str(e)}")
    
    def _analyze_sensitive_zones(self):
        """分析环境敏感点"""
        # 获取坐标
        try:
            lat = float(self.lat_edit.text())
            lon = float(self.lon_edit.text())
            radius = float(self.radius_edit.text())
        except ValueError:
            QMessageBox.warning(self, "警告", "请输入有效的坐标和半径")
            return
        
        try:
            # 调用 EIAgent 的 analyze_sensitive_zones 工具
            wizard = self.wizard()
            result = wizard._call_agent(
                "analyze_sensitive_zones",
                latitude=lat,
                longitude=lon,
                radius=radius
            )
            
            if result.get("success"):
                sensitive_zones = result.get("sensitive_zones", [])
                
                # 保存到 wizard
                wizard.sensitive_zones = sensitive_zones
                wizard.project_coords = {"latitude": lat, "longitude": lon}
                
                # 显示结果
                self._display_sensitive_zones(sensitive_zones)
                
                QMessageBox.information(self, "成功", f"已识别 {len(sensitive_zones)} 个环境敏感点")
            else:
                QMessageBox.warning(self, "警告", f"分析敏感点失败: {result.get('error', '未知错误')}")
        
        except Exception as e:
            QMessageBox.critical(self, "错误", f"分析敏感点失败: {str(e)}")
    
    def _display_sensitive_zones(self, sensitive_zones):
        """显示敏感点"""
        text = ""
        for i, zone in enumerate(sensitive_zones, 1):
            text += f"{i}. {zone.get('name', '未知')} ({zone.get('type', '未知类型')})\n"
            text += f"   距离: {zone.get('distance', '未知')} 米\n"
            if 'direction' in zone:
                text += f"   方向: {zone.get('direction')}\n"
            text += "\n"
        
        self.result_text.setText(text)
    
    def validatePage(self):
        """验证页面"""
        # 检查是否输入了坐标
        try:
            lat = float(self.lat_edit.text())
            lon = float(self.lon_edit.text())
        except ValueError:
            QMessageBox.warning(self, "警告", "请输入有效的坐标")
            return False
        
        # 检查是否分析了敏感点
        wizard = self.wizard()
        if not hasattr(wizard, 'sensitive_zones') or not wizard.sensitive_zones:
            reply = QMessageBox.question(
                self, "确认",
                "尚未分析环境敏感点，是否继续？",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.No:
                return False
        
        # 保存坐标到 wizard
        wizard.project_coords = {"latitude": lat, "longitude": lon}
        
        return True
    

class Step3AutoLearning(QWizardPage):
    """步骤3: 自动学习"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("步骤3: 自动学习")
        self.setSubTitle("系统正在自动学习，请稍候...")
        
        # 创建UI
        layout = QVBoxLayout()
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 0)  # 无限循环模式
        layout.addWidget(self.progress_bar)
        
        # 状态标签
        self.status_label = QLabel("正在分析文档...")
        layout.addWidget(self.status_label)
        
        # 结果显示
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        layout.addWidget(self.result_text)
        
        self.setLayout(layout)
        
        # 标记是否已完成
        self.is_completed = False
    
    def initializePage(self):
        """初始化页面"""
        
        # 重置状态
        self.is_completed = False
        self.status_label.setText("正在分析文档...")
        self.result_text.clear()
        
        # 调用 EIAgent 的工具
        wizard = self.wizard()
        
        try:
            # 分析文档
            result = wizard._call_agent(
                "analyze_document",
                file_path=wizard.document_path
            )
            
            # 显示结果
            self.result_text.setText(str(result))
            self.status_label.setText("文档分析完成")
            
            # 标记已完成
            self.is_completed = True
        
        except Exception as e:
            self.status_label.setText(f"文档分析失败: {str(e)}")
            QMessageBox.critical(self, "错误", f"文档分析失败: {str(e)}")
    
    def isComplete(self):
        """检查是否已完成
        
        Returns:
            bool: 是否已完成
        """
        return self.is_completed


class Step4GenerateReport(QWizardPage):
    """步骤4: 生成报告"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("步骤4: 生成报告")
        self.setSubTitle("系统正在生成报告，请稍候...")
        
        # 创建UI
        layout = QVBoxLayout()
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 0)  # 无限循环模式
        layout.addWidget(self.progress_bar)
        
        # 状态标签
        self.status_label = QLabel("正在生成报告...")
        layout.addWidget(self.status_label)
        
        # 结果显示
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        layout.addWidget(self.result_text)
        
        self.setLayout(layout)
        
        # 标记是否已完成
        self.is_completed = False
    
    def initializePage(self):
        """初始化页面"""
        
        # 重置状态
        self.is_completed = False
        self.status_label.setText("正在生成报告...")
        self.result_text.clear()
        
        # 调用 EIAgent 的工具
        wizard = self.wizard()
        
        try:
            # 生成报告
            result = wizard._call_agent(
                "generate_report",
                project_info=wizard.project_info
            )
            
            # 保存报告内容
            if "result" in result:
                wizard.report_content = result["result"]
            elif "report_content" in result:
                wizard.report_content = result["report_content"]
            else:
                wizard.report_content = str(result)
            
            # 显示结果
            self.result_text.setText(wizard.report_content)
            self.status_label.setText("报告生成完成")
            
            # 标记已完成
            self.is_completed = True
        
        except Exception as e:
            self.status_label.setText(f"报告生成失败: {str(e)}")
            QMessageBox.critical(self, "错误", f"报告生成失败: {str(e)}")
    
    def isComplete(self):
        """检查是否已完成
        
        Returns:
            bool: 是否已完成
        """
        return self.is_completed


class Step5ExportReport(QWizardPage):
    """步骤5: 导出报告"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setTitle("步骤5: 导出报告")
        self.setSubTitle("请选择导出格式")
        
        # 创建UI
        layout = QVBoxLayout()
        
        # 说明标签
        label = QLabel("报告已生成，请选择导出格式：")
        layout.addWidget(label)
        
        # 导出按钮
        self.export_word_btn = QPushButton("导出为 Word")
        self.export_word_btn.clicked.connect(self._export_word)
        layout.addWidget(self.export_word_btn)
        
        self.export_pdf_btn = QPushButton("导出为 PDF")
        self.export_pdf_btn.clicked.connect(self._export_pdf)
        layout.addWidget(self.export_pdf_btn)
        
        # 预览区域
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        layout.addWidget(self.preview_text)
        
        self.setLayout(layout)
    
    def initializePage(self):
        """初始化页面"""
        
        # 显示报告预览
        wizard = self.wizard()
        self.preview_text.setText(wizard.report_content)
    
    def _export_word(self):
        """导出为 Word"""
        
        # 打开文件对话框
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出为 Word",
            "",
            "Word 文件 (*.docx)"
        )
        
        if file_path:
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                
                # 1. 创建 Word 文档
                doc = Document()
                
                # 2. 添加标题
                wizard = self.wizard()
                title = wizard.project_info.get('project_name', '未知项目')
                doc.add_heading(f"{title} 环境影响报告书", 0)
                
                # 3. 添加报告内容（简化实现）
                # 实际应该解析 wizard.report_content，这里只添加基本结构
                doc.add_heading("1. 项目概述", 1)
                doc.add_paragraph(f"项目名称：{wizard.project_info.get('project_name', '未知项目')}")
                doc.add_paragraph(f"建设地点：{wizard.project_info.get('location', '未知地点')}")
                doc.add_paragraph(f"行业类别：{wizard.project_info.get('industry', '未知行业')}")
                
                doc.add_heading("2. 环境现状", 1)
                doc.add_paragraph("（此处省略环境现状分析...）")
                
                doc.add_heading("3. 环境影响预测", 1)
                doc.add_paragraph("（此处省略环境影响预测...）")
                
                doc.add_heading("4. 环境保护措施", 1)
                doc.add_paragraph("（此处省略环境保护措施...）")
                
                doc.add_heading("5. 结论", 1)
                doc.add_paragraph("（此处省略结论...）")
                
                # 4. 如果有敏感点，添加敏感点章节
                if hasattr(wizard, 'sensitive_zones') and wizard.sensitive_zones:
                    doc.add_heading("6. 环境敏感点分析", 1)
                    for i, zone in enumerate(wizard.sensitive_zones, 1):
                        doc.add_paragraph(f"{i}. {zone.get('name', '未知')} ({zone.get('type', '未知类型')})")
                        doc.add_paragraph(f"   距离: {zone.get('distance', '未知')} 米")
                
                # 5. 保存文档
                doc.save(file_path)
                
                QMessageBox.information(self, "成功", f"报告已导出到:\n{file_path}")
                
            except ImportError:
                QMessageBox.critical(self, "错误", "python-docx 库未安装，请运行:\npip install python-docx")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"导出 Word 失败:\n{str(e)}")
    
    def _export_pdf(self):
        """导出为 PDF"""
        
        # 打开文件对话框
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出为 PDF",
            "",
            "PDF 文件 (*.pdf)"
        )
        
        if file_path:
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.pdfgen import canvas
                from reportlab.lib.units import cm
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                
                # 1. 注册中文字体（使用系统字体）
                try:
                    # 尝试使用 Windows 系统中的中文字体
                    pdfmetrics.registerFont(TTFont('SimSun', 'C:\\Windows\\Fonts\\simsun.ttc'))
                    font_name = 'SimSun'
                except:
                    # 如果没有中文字体，使用默认字体（可能无法显示中文）
                    font_name = 'Helvetica'
                
                # 2. 创建 PDF 文档
                c = canvas.Canvas(file_path, pagesize=A4)
                width, height = A4
                
                # 3. 添加标题
                wizard = self.wizard()
                title = wizard.project_info.get('project_name', '未知项目')
                
                c.setFont(font_name, 20)
                c.drawString(2*cm, height-3*cm, f"{title} 环境影响报告书")
                
                # 4. 添加报告内容（简化实现）
                y = height - 5*cm
                c.setFont(font_name, 12)
                
                # 项目概述
                c.drawString(2*cm, y, "1. 项目概述")
                y -= 1*cm
                
                c.setFont(font_name, 10)
                c.drawString(3*cm, y, f"项目名称：{wizard.project_info.get('project_name', '未知项目')}")
                y -= 0.5*cm
                c.drawString(3*cm, y, f"建设地点：{wizard.project_info.get('location', '未知地点')}")
                y -= 0.5*cm
                c.drawString(3*cm, y, f"行业类别：{wizard.project_info.get('industry', '未知行业')}")
                y -= 1*cm
                
                # 环境现状（简化）
                c.setFont(font_name, 12)
                c.drawString(2*cm, y, "2. 环境现状")
                y -= 1*cm
                
                c.setFont(font_name, 10)
                c.drawString(3*cm, y, "（此处省略环境现状分析...）")
                y -= 1*cm
                
                # 如果有敏感点，添加敏感点信息
                if hasattr(wizard, 'sensitive_zones') and wizard.sensitive_zones:
                    c.setFont(font_name, 12)
                    c.drawString(2*cm, y, "3. 环境敏感点分析")
                    y -= 1*cm
                    
                    c.setFont(font_name, 10)
                    for i, zone in enumerate(wizard.sensitive_zones, 1):
                        text = f"{i}. {zone.get('name', '未知')} ({zone.get('type', '未知类型')}) - 距离: {zone.get('distance', '未知')} 米"
                        c.drawString(3*cm, y, text)
                        y -= 0.5*cm
                    y -= 0.5*cm
                
                # 5. 保存 PDF
                c.save()
                
                QMessageBox.information(self, "成功", f"报告已导出到:\n{file_path}")
                
            except ImportError:
                QMessageBox.critical(self, "错误", "reportlab 库未安装，请运行:\npip install reportlab")
            except Exception as e:
                QMessageBox.critical(self, "错误", f"导出 PDF 失败:\n{str(e)}")


if __name__ == "__main__":
    # 测试 EIWizard
    
    import sys
    from PySide6.QtWidgets import QApplication
    
    app = QApplication(sys.argv)
    
    wizard = EIWizard()
    wizard.show()
    
    sys.exit(app.exec())
