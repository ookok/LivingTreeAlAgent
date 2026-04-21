"""
多模态文档解析器
=================

支持：
- 文本格式: Word/PDF/TXT/Excel
- 图片格式: 工艺流程图、设备照片
- 混合文档: 包含文字、表格、图片的PDF报告
- 非结构化输入: 用户简写、语音转文字

Author: Hermes Desktop Team
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Any, Tuple
from enum import Enum
import re
import os

# ============================================================
# 第一部分：文档类型和配置
# ============================================================

class DocumentType(Enum):
    PDF = "pdf"
    WORD = "word"
    EXCEL = "excel"
    TEXT = "text"
    IMAGE = "image"
    HTML = "html"
    MARKDOWN = "markdown"
    UNKNOWN = "unknown"


@dataclass
class DocumentConfig:
    enable_ocr: bool = True
    enable_table_parsing: bool = True
    enable_layout_analysis: bool = True
    ocr_language: str = "chi_sim+eng"
    max_image_size: int = 10 * 1024 * 1024
    table_detection_method: str = "edge"


@dataclass
class ParsedDocument:
    doc_type: DocumentType
    title: str = ""
    text_content: str = ""
    structured_content: Dict[str, Any] = field(default_factory=dict)
    tables: List[Dict] = field(default_factory=list)
    images: List[Dict] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    raw_content: Any = None


# ============================================================
# 第二部分：基础解析器
# ============================================================

class BaseParser(ABC):
    def __init__(self, config: Optional[DocumentConfig] = None):
        self.config = config or DocumentConfig()

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        pass

    def clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()


# ============================================================
# 第三部分：文本解析器
# ============================================================

class TextParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return ParsedDocument(
                doc_type=DocumentType.TEXT,
                text_content=self.clean_text(content),
                structured_content={"full_text": content},
                metadata={"encoding": "utf-8", "size": len(content)}
            )
        except:
            for encoding in ['gbk', 'gb2312', 'gb18030']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return ParsedDocument(
                        doc_type=DocumentType.TEXT,
                        text_content=self.clean_text(content),
                        structured_content={"full_text": content}
                    )
                except:
                    continue
        return ParsedDocument(doc_type=DocumentType.TEXT)


class MarkdownParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else ""

        sections = {}
        current_section = "introduction"
        for line in content.split('\n'):
            header_match = re.match(r'^##?\s+(.+)$', line)
            if header_match:
                current_section = header_match.group(1)
                sections[current_section] = []
            else:
                if current_section in sections:
                    sections[current_section].append(line)

        for sec in sections:
            sections[sec] = '\n'.join(sections[sec])

        return ParsedDocument(
            doc_type=DocumentType.MARKDOWN,
            title=title,
            text_content=self.clean_text(content),
            structured_content=sections
        )


# ============================================================
# 第四部分：表格解析器
# ============================================================

class TableParser:
    def __init__(self, config: Optional[DocumentConfig] = None):
        self.config = config or DocumentConfig()

    def parse_csv(self, file_path: str) -> List[Dict]:
        tables = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            if lines:
                headers = [h.strip() for h in lines[0].split(',')]
                rows = []
                for line in lines[1:]:
                    values = [v.strip() for v in line.split(',')]
                    if len(values) == len(headers):
                        rows.append(values)
                tables.append({"headers": headers, "rows": rows})
        except Exception as e:
            print(f"CSV解析失败: {e}")
        return tables

    def parse_excel(self, file_path: str) -> List[Dict]:
        tables = []
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                headers = []
                rows = []
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if i == 0:
                        headers = [str(cell) if cell else "" for cell in row]
                    else:
                        rows.append([str(cell) if cell else "" for cell in row])
                if headers:
                    tables.append({"sheet": sheet_name, "headers": headers, "rows": rows})
            wb.close()
        except ImportError:
            print("openpyxl未安装")
        except Exception as e:
            print(f"Excel解析失败: {e}")
        return tables

    def extract_tables_from_text(self, text: str) -> List[Dict]:
        tables = []
        md_table_pattern = r'\|(.+)\|\n\|[-:\s|]+\|\n((?:\|.+\|\n?)+)'
        matches = re.finditer(md_table_pattern, text)
        for match in matches:
            header_line = match.group(1)
            body_lines = match.group(2).strip().split('\n')
            headers = [h.strip() for h in header_line.split('|') if h.strip()]
            rows = []
            for line in body_lines:
                values = [v.strip() for v in line.split('|') if v.strip()]
                if values:
                    rows.append(values)
            if headers and rows:
                tables.append({"headers": headers, "rows": rows})
        return tables


# ============================================================
# 第五部分：PDF解析器
# ============================================================

class PDFParser(BaseParser):
    def __init__(self, config: Optional[DocumentConfig] = None):
        super().__init__(config)
        self.table_parser = TableParser(config)

    def parse(self, file_path: str) -> ParsedDocument:
        text_content = ""
        tables = []
        images = []

        try:
            import fitz
            doc = fitz.open(file_path)
            metadata = {
                "page_count": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", "")
            }
            for page_num in range(len(doc)):
                page = doc[page_num]
                text_content += page.get_text()
                text_content += f"\n--- 第{page_num + 1}页 ---\n"
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n - pix.alpha < 4:
                        img_data = pix.tobytes("png")
                    else:
                        img_data = fitz.Pixmap(fitz.csRGB, pix).tobytes("png")
                    img_path = f"{file_path}_img_{page_num}_{img_index}.png"
                    with open(img_path, 'wb') as img_file:
                        img_file.write(img_data)
                    images.append({"path": img_path, "page": page_num + 1, "index": img_index})
            doc.close()
            tables = self.table_parser.extract_tables_from_text(text_content)
        except ImportError:
            print("PyMuPDF未安装")
        except Exception as e:
            print(f"PDF解析失败: {e}")

        return ParsedDocument(
            doc_type=DocumentType.PDF,
            text_content=self.clean_text(text_content),
            tables=tables,
            images=images,
            metadata=metadata
        )


# ============================================================
# 第六部分：Word解析器
# ============================================================

class WordParser(BaseParser):
    def __init__(self, config: Optional[DocumentConfig] = None):
        super().__init__(config)
        self.table_parser = TableParser(config)

    def parse(self, file_path: str) -> ParsedDocument:
        text_content = ""
        tables = []

        try:
            from docx import Document
            doc = Document(file_path)
            title = doc.paragraphs[0].text if doc.paragraphs else ""
            for para in doc.paragraphs:
                text_content += para.text + "\n"
            for table in doc.tables:
                headers = []
                rows = []
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    if i == 0:
                        headers = cells
                    else:
                        rows.append(cells)
                if headers:
                    tables.append({"headers": headers, "rows": rows})
            metadata = {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
        except ImportError:
            print("python-docx未安装")
            metadata = {}
        except Exception as e:
            print(f"Word解析失败: {e}")
            metadata = {}

        return ParsedDocument(
            doc_type=DocumentType.WORD,
            title=title,
            text_content=self.clean_text(text_content),
            tables=tables,
            metadata=metadata
        )


# ============================================================
# 第七部分：图像解析器
# ============================================================

class ImageParser(BaseParser):
    def __init__(self, config: Optional[DocumentConfig] = None):
        super().__init__(config)

    def parse(self, file_path: str) -> ParsedDocument:
        text_content = ""
        description = ""
        metadata = {}

        try:
            from PIL import Image
            img = Image.open(file_path)
            metadata = {"format": img.format, "size": img.size, "mode": img.mode}
            if self.config.enable_ocr:
                text_content, description = self._ocr_image(file_path)
        except ImportError:
            print("PIL未安装")
        except Exception as e:
            print(f"图像解析失败: {e}")

        return ParsedDocument(
            doc_type=DocumentType.IMAGE,
            text_content=self.clean_text(text_content),
            images=[{"path": file_path, "description": description}],
            metadata=metadata
        )

    def _ocr_image(self, file_path: str) -> Tuple[str, str]:
        text = ""
        description = ""
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang=self.config.ocr_language)
            description = text
        except ImportError:
            try:
                import easyocr
                reader = easyocr.Reader(['ch_sim', 'en'])
                results = reader.readtext(file_path)
                text = " ".join([result[1] for result in results])
                description = text
            except ImportError:
                print("OCR库均未安装")
        return text, description


# ============================================================
# 第八部分：文档解析工厂
# ============================================================

class DocumentParserFactory:
    _parsers = {
        DocumentType.TEXT: TextParser,
        DocumentType.MARKDOWN: MarkdownParser,
        DocumentType.PDF: PDFParser,
        DocumentType.WORD: WordParser,
        DocumentType.IMAGE: ImageParser,
    }

    @classmethod
    def get_parser(cls, file_path: str, config: Optional[DocumentConfig] = None) -> BaseParser:
        ext = file_path.lower().split('.')[-1]
        type_map = {
            'txt': DocumentType.TEXT, 'md': DocumentType.MARKDOWN,
            'pdf': DocumentType.PDF, 'docx': DocumentType.WORD, 'doc': DocumentType.WORD,
            'xlsx': DocumentType.EXCEL, 'xls': DocumentType.EXCEL, 'csv': DocumentType.EXCEL,
            'png': DocumentType.IMAGE, 'jpg': DocumentType.IMAGE,
            'jpeg': DocumentType.IMAGE, 'gif': DocumentType.IMAGE, 'bmp': DocumentType.IMAGE,
        }
        doc_type = type_map.get(ext, DocumentType.UNKNOWN)
        parser_class = cls._parsers.get(doc_type, TextParser)
        return parser_class(config)

    @classmethod
    def parse(cls, file_path: str, config: Optional[DocumentConfig] = None) -> ParsedDocument:
        parser = cls.get_parser(file_path, config)
        return parser.parse(file_path)


# ============================================================
# 第九部分：统一文档理解接口
# ============================================================

class UnifiedDocumentParser:
    def __init__(self, config: Optional[DocumentConfig] = None):
        self.config = config or DocumentConfig()

    def parse(self, input_source: Any) -> ParsedDocument:
        if isinstance(input_source, str):
            if input_source.startswith(('http://', 'https://')):
                return self._parse_from_url(input_source)
            elif os.path.exists(input_source):
                return DocumentParserFactory.parse(input_source, self.config)
            else:
                return self._parse_from_text(input_source)
        return ParsedDocument(doc_type=DocumentType.UNKNOWN)

    def _parse_from_url(self, url: str) -> ParsedDocument:
        return ParsedDocument(doc_type=DocumentType.UNKNOWN)

    def _parse_from_text(self, text: str) -> ParsedDocument:
        return ParsedDocument(
            doc_type=DocumentType.TEXT,
            text_content=text,
            structured_content={"full_text": text}
        )

    def parse_batch(self, file_paths: List[str]) -> List[ParsedDocument]:
        results = []
        for path in file_paths:
            try:
                results.append(self.parse(path))
            except Exception as e:
                print(f"解析失败 {path}: {e}")
                results.append(ParsedDocument(doc_type=DocumentType.UNKNOWN))
        return results


__all__ = [
    'DocumentType', 'DocumentConfig', 'ParsedDocument',
    'BaseParser', 'TextParser', 'MarkdownParser', 'TableParser',
    'PDFParser', 'WordParser', 'ImageParser',
    'DocumentParserFactory', 'UnifiedDocumentParser'
]
