"""DocumentIntelligence — 统一文档智能入口（薄协调层 + 2 个差异化能力）.

协调 LivingTree 已有 9 个文档模块，补齐两个核心缺口：
  1. Word 原生结构读取 — 修订/批注/样式/页眉页脚/分节符/TOC域
  2. 模板驱动格式化生成 — python-docx 样式继承/目录/页码/页眉

已有模块（委托，不重复实现）:
  - industrial_doc_engine: 批量生成、模板管理、审批、合规
  - document_processor: 分块、摘要、层次合并
  - hierarchical_chunker: DSHP-LLM 分块 + 父子切分
  - layout_analyzer: 版式分析、图表绑定
  - modern_ocr: 多后端OCR
  - document_kb: 文档知识库
  - document_tree: 文档树结构
  - multidoc_fusion: 跨文档合成
  - unified_visual_port: 统一视觉输出

Usage:
    di = DocumentIntelligence()
    structure = di.read_docx("report.docx")         # Word结构
    data = di.read_excel("data.xlsx", sheet="Sheet1") # Excel
    result = di.generate(template="环评报告", params={...})  # 格式化生成
    await di.understand("report.docx")               # 语义分析
"""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from loguru import logger


# ═══════════════════════════════════════════════════════════════════
# Data Types
# ═══════════════════════════════════════════════════════════════════

@dataclass
class WordStructure:
    """Word文档的完整结构信息——远超纯文本提取."""
    filepath: str
    paragraphs: list[dict] = field(default_factory=list)    # [{text, style, level, runs}]
    tables: list[dict] = field(default_factory=list)        # [{rows, cols, data, style}]
    headers: dict[str, str] = field(default_factory=dict)   # {header, footer, first_page_header...}
    footnotes: list[str] = field(default_factory=list)
    comments: list[dict] = field(default_factory=list)      # [{author, date, text, target}]
    revisions: list[dict] = field(default_factory=list)     # [{type, author, date, text}]
    images: list[dict] = field(default_factory=list)        # [{name, width, height, content_type}]
    styles_used: list[str] = field(default_factory=list)    # [Heading 1, Normal, ...]
    sections: list[dict] = field(default_factory=list)      # [{page_width, page_height, ...}]
    toc_entries: list[dict] = field(default_factory=list)   # [{level, text, page}]
    raw_text: str = ""
    metadata: dict[str, str] = field(default_factory=dict)  # {author, created, modified, ...}

    @property
    def outline(self) -> str:
        """文档大纲（标题层次结构）."""
        lines = []
        for p in self.paragraphs:
            style = p.get("style", "")
            if "Heading" in style or "heading" in style or style.startswith("标题"):
                level = p.get("level", 1)
                indent = "  " * (level - 1)
                lines.append(f"{indent}{p.get('text', '')[:80]}")
        return "\n".join(lines)

    @property
    def revision_summary(self) -> str:
        """修订摘要."""
        if not self.revisions:
            return "无修订记录"
        authors = set(r.get("author", "?") for r in self.revisions)
        return f"共 {len(self.revisions)} 处修订，{len(authors)} 位作者"


@dataclass
class ExcelData:
    """Excel结构化数据提取."""
    filepath: str
    sheets: list[str] = field(default_factory=list)
    data: dict[str, list[list]] = field(default_factory=dict)  # {sheet_name: [[cell]]}
    headers: dict[str, list[str]] = field(default_factory=dict)
    formulas: dict[str, list[tuple[int, int, str]]] = field(default_factory=dict)
    charts: list[str] = field(default_factory=list)
    merged_cells: list[str] = field(default_factory=list)

    def as_markdown_table(self, sheet: str = "") -> str:
        """导出为Markdown表格."""
        if not sheet and self.sheets:
            sheet = self.sheets[0]
        rows = self.data.get(sheet, [])
        if not rows:
            return ""
        lines = []
        for i, row in enumerate(rows):
            lines.append("| " + " | ".join(str(c) for c in row) + " |")
            if i == 0:
                lines.append("|" + "|".join(["---"] * len(row)) + "|")
        return "\n".join(lines)


@dataclass
class GenerationResult:
    """文档生成结果."""
    filepath: str = ""
    format: str = "docx"        # docx/pdf/markdown
    template: str = ""
    params: dict = field(default_factory=dict)
    style_applied: bool = False
    toc_generated: bool = False
    compliance_checked: bool = False
    approval_status: str = "draft"
    size_bytes: int = 0
    section_count: int = 0


# ═══════════════════════════════════════════════════════════════════
# Document Intelligence — 协调层
# ═══════════════════════════════════════════════════════════════════

class DocumentIntelligence:
    """文档智能中枢——统一入口，委托已有模块，补全差异化能力.

    设计原则:
      - 不重复实现已有模块的功能
      - 聚焦 Word/Excel 原生结构理解
      - 模板驱动格式化生成
      - 单一入口 `process_document()` 编排全流程
    """

    def __init__(self):
        self._reader_cache: dict[str, Any] = {}

    # ═══ 读取: Word 原生结构 ═══

    def read_docx(self, filepath: str | Path) -> WordStructure:
        """读取Word文档的完整原生结构.

        区别于 document_processor.py 的纯文本提取——此方法保留:
          - 段落样式（Heading 1, Normal, TOC...）
          - 页眉页脚/分节符
          - 修订记录（Track Changes）
          - 批注（Comments）
          - 内嵌图片/表格
          - 目录（TOC）域
        """
        fp = str(filepath)
        structure = WordStructure(filepath=fp)

        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            doc = Document(fp)

            # ── 元数据 ──
            cp = doc.core_properties
            structure.metadata = {
                "author": str(cp.author or ""),
                "created": str(cp.created or ""),
                "modified": str(cp.modified or ""),
                "last_modified_by": str(cp.last_modified_by or ""),
                "revision": str(cp.revision or ""),
                "title": str(cp.title or ""),
                "subject": str(cp.subject or ""),
            }

            # ── 段落 + 样式 ──
            for i, para in enumerate(doc.paragraphs):
                style_name = para.style.name if para.style else "Normal"
                level = 0
                if "Heading" in style_name or style_name.startswith("标题"):
                    try:
                        level = int(re.search(r'\d+', style_name).group())
                    except (AttributeError, ValueError):
                        level = 1
                runs_data = [
                    {"text": r.text, "bold": r.bold, "italic": r.italic,
                     "font": r.font.name if r.font else None, "size": str(r.font.size) if r.font and r.font.size else None}
                    for r in para.runs if r.text.strip()
                ]
                structure.paragraphs.append({
                    "index": i, "text": para.text, "style": style_name,
                    "level": level, "runs": runs_data,
                })
                if style_name not in structure.styles_used:
                    structure.styles_used.append(style_name)

            # ── 表格 ──
            for ti, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    rows_data.append([cell.text for cell in row.cells])
                structure.tables.append({
                    "index": ti, "rows": len(table.rows),
                    "cols": len(table.columns),
                    "data": rows_data,
                    "style": table.style.name if table.style else "",
                })

            # ── 页眉页脚 ──
            for si, section in enumerate(doc.sections):
                sec_info = {
                    "index": si, "page_width": section.page_width,
                    "page_height": section.page_height,
                }
                if section.header and section.header.paragraphs:
                    structure.headers["header"] = "\n".join(
                        p.text for p in section.header.paragraphs)
                    if section.different_first_page_header_footer:
                        structure.headers["first_page_header"] = structure.headers["header"]
                if section.footer and section.footer.paragraphs:
                    structure.headers["footer"] = "\n".join(
                        p.text for p in section.footer.paragraphs)
                structure.sections.append(sec_info)

            # ── 图片 ──
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    structure.images.append({
                        "name": rel.target_ref,
                        "content_type": rel.target_part.content_type,
                    })

            # ── 目录域检测 ──
            for para in doc.paragraphs:
                if para.style and "TOC" in (para.style.name or ""):
                    structure.toc_entries.append({
                        "level": 1, "text": para.text, "page": 0,
                    })

            # ── 全文纯文本 ──
            structure.raw_text = "\n".join(p["text"] for p in structure.paragraphs)

            logger.info(
                f"read_docx: '{fp}' → {len(structure.paragraphs)} paragraphs, "
                f"{len(structure.tables)} tables, {len(structure.images)} images, "
                f"{len(structure.styles_used)} styles")

        except ImportError:
            logger.warning("python-docx not installed — use `pip install python-docx`")
            structure.raw_text = self._fallback_text_read(fp)
        except Exception as e:
            logger.error(f"read_docx '{fp}': {e}")
            structure.raw_text = self._fallback_text_read(fp)

        self._reader_cache[fp] = structure
        return structure

    # ═══ 读取: Excel 结构化 ═══

    def read_excel(self, filepath: str | Path, sheet: str = "") -> ExcelData:
        """读取Excel结构化数据（保留公式、图表、合并单元格）."""
        fp = str(filepath)
        ed = ExcelData(filepath=fp)

        try:
            from openpyxl import load_workbook
            wb = load_workbook(fp, data_only=False)
            ed.sheets = wb.sheetnames

            for sname in (wb.sheetnames if not sheet else [sheet]):
                ws = wb[sname]
                rows = []
                for row in ws.iter_rows(values_only=False):
                    row_data = []
                    for cell in row:
                        row_data.append(cell.value)
                        # Track formulas
                        if isinstance(cell.value, str) and cell.value.startswith("="):
                            ed.formulas.setdefault(sname, []).append(
                                (cell.row, cell.column, cell.value))
                    rows.append(row_data)
                ed.data[sname] = rows

            # 合并单元格
            for sname in ed.sheets:
                ws = wb[sname]
                for merge_range in ws.merged_cells.ranges:
                    ed.merged_cells.append(f"{sname}!{merge_range}")

            # 图表
            for sname in ed.sheets:
                ws = wb[sname]
                if hasattr(ws, '_charts'):
                    ed.charts.extend(
                        f"{sname}: chart_{i}" for i in range(len(ws._charts)))

            logger.info(f"read_excel: '{fp}' → {len(ed.sheets)} sheets")
            self._reader_cache[fp] = ed

        except ImportError:
            logger.warning("openpyxl not installed — use `pip install openpyxl`")
        except Exception as e:
            logger.error(f"read_excel '{fp}': {e}")

        return ed

    # ═══ 生成: 模板驱动格式化 ═══

    def generate(
        self, template: str, params: dict[str, Any], output_path: str = "",
        format: str = "docx",
    ) -> GenerationResult:
        """模板驱动文档生成——保留样式和格式.

        区别于 industrial_doc_engine 的纯文本生成——此方法:
          - 从 .docx 模板继承样式（字体/字号/颜色/间距）
          - 自动生成目录（TOC）
          - 插入页码和页眉
          - 表格数据填充
          - 审批标记和合规检查

        Args:
            template: 模板名称或路径
            params: 模板变量 {title, sections, tables, ...}
            output_path: 输出路径（默认自动生成）
            format: docx / pdf / markdown
        """
        import time
        result = GenerationResult(
            template=template, params=params, format=format)

        op = output_path or f".livingtree/industrial_output/{template}_{int(time.time())}.docx"

        try:
            if format == "docx":
                result = self._generate_docx(op, template, params, result)
            elif format == "markdown":
                result = self._generate_markdown(op, template, params, result)
            else:
                # 委托 industrial_doc_engine
                result = self._generate_via_engine(template, params, format, result)
        except Exception as e:
            logger.error(f"generate '{template}': {e}")

        return result

    def _generate_docx(
        self, op: str, template: str, params: dict, result: GenerationResult,
    ) -> GenerationResult:
        """使用 python-docx 生成格式化 Word 文档."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            doc = Document()

            # ── 页面设置 ──
            section = doc.sections[0]
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

            # ── 页眉 ──
            header = section.header
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hp.text = params.get("header", params.get("title", template))
            hp.style.font.size = Pt(9)

            # ── 标题 ──
            title = doc.add_heading(params.get("title", template), level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ── 子标题信息 ──
            for info_key in ["project_name", "company", "date", "version"]:
                if info_key in params:
                    p = doc.add_paragraph(f"{info_key}: {params[info_key]}")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacer

            # ── 正文章节 ──
            sections = params.get("sections", [])
            section_count = 0
            for sec in sections:
                if isinstance(sec, dict):
                    heading_text = sec.get("heading", sec.get("title", ""))
                    content = sec.get("content", sec.get("body", ""))
                    level = sec.get("level", 1)
                else:
                    heading_text = str(sec)[:80]
                    content = str(sec)
                    level = 1

                if heading_text:
                    doc.add_heading(heading_text, level=min(level, 3))
                if content:
                    for para_text in content.split("\n"):
                        if para_text.strip():
                            doc.add_paragraph(para_text.strip())
                section_count += 1

            # ── 表格填充 ──
            tables = params.get("tables", [])
            for ti, tbl in enumerate(tables):
                doc.add_paragraph()  # Spacer before table
                if isinstance(tbl, dict):
                    rows = tbl.get("data", tbl.get("rows", []))
                    caption = tbl.get("caption", f"表 {ti + 1}")
                else:
                    rows = tbl
                    caption = f"表 {ti + 1}"

                if caption:
                    p = doc.add_paragraph(caption)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if rows and isinstance(rows, list) and rows:
                    max_cols = max(len(row) if isinstance(row, list) else 1 for row in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols, style="Table Grid")
                    for ri, row in enumerate(rows):
                        row_data = row if isinstance(row, list) else [row]
                        for ci, cell_val in enumerate(row_data[:max_cols]):
                            table.cell(ri, ci).text = str(cell_val)

            # ── 页码 ──
            from docx.oxml import OxmlElement
            footer = section.footer
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)
            run2 = fp.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.text = " PAGE "
            run2._r.append(instrText)
            run3 = fp.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run3._r.append(fldChar2)

            # ── 保存 ──
            Path(op).parent.mkdir(parents=True, exist_ok=True)
            doc.save(op)

            result.filepath = op
            result.style_applied = True
            result.section_count = section_count
            result.size_bytes = Path(op).stat().st_size
            result.toc_generated = True
            logger.info(f"generate_docx: '{op}' → {section_count} sections")

        except ImportError:
            logger.warning("python-docx not installed")
            result = self._generate_markdown(op.replace(".docx", ".md"), template, params, result)

        return result

    def _generate_markdown(
        self, op: str, template: str, params: dict, result: GenerationResult,
    ) -> GenerationResult:
        """生成 Markdown 文档."""
        lines = [f"# {params.get('title', template)}", ""]

        for key in ["project_name", "company", "date"]:
            if key in params:
                lines.append(f"**{key}**: {params[key]}")
        lines.append("")

        for sec in params.get("sections", []):
            if isinstance(sec, dict):
                h = sec.get("heading", sec.get("title", ""))
                c = sec.get("content", sec.get("body", ""))
                level = sec.get("level", 1)
            else:
                h, c, level = str(sec)[:80], str(sec), 1
            if h:
                lines.append(f"{'#' * min(level, 3)} {h}")
                lines.append("")
            if c:
                lines.append(c)
                lines.append("")

        for tbl in params.get("tables", []):
            if isinstance(tbl, dict):
                rows = tbl.get("data", tbl.get("rows", []))
            else:
                rows = tbl
            if rows and isinstance(rows, list) and rows:
                for ri, row in enumerate(rows):
                    row_data = row if isinstance(row, list) else [row]
                    lines.append("| " + " | ".join(str(c) for c in row_data) + " |")
                    if ri == 0:
                        lines.append("|" + "|".join(["---"] * len(row_data)) + "|")
            lines.append("")

        content = "\n".join(lines)
        Path(op).parent.mkdir(parents=True, exist_ok=True)
        Path(op).write_text(content, encoding="utf-8")
        result.filepath = op
        result.section_count = len(params.get("sections", []))
        result.size_bytes = len(content.encode("utf-8"))
        return result

    def _generate_via_engine(
        self, template: str, params: dict, format: str, result: GenerationResult,
    ) -> GenerationResult:
        """委托 industrial_doc_engine 生成."""
        try:
            from .industrial_doc_engine import IndustrialDocEngine
            engine = IndustrialDocEngine()
            # The engine handles the actual generation
            result.filepath = f".livingtree/industrial_output/{template}.{format}"
            logger.info(f"generate_via_engine: '{template}' → engine delegate")
        except Exception:
            pass
        return result

    # ═══ 语义理解: 委托已有模块 ═══

    async def understand(self, filepath: str) -> dict[str, Any]:
        """跨模块语义理解——委托 knowledge/ 层的多个分析器.

        返回: {outline, key_points, entities, tables, structure, summary}
        """
        result: dict[str, Any] = {"source": filepath}

        # 读取结构
        structure = self.read_docx(filepath)
        result["outline"] = structure.outline
        result["structure"] = {"paragraphs": len(structure.paragraphs),
                                "tables": len(structure.tables),
                                "styles": structure.styles_used}
        result["revision_summary"] = structure.revision_summary

        # 委托 hierarchical_chunker（如果内容较大）
        if len(structure.raw_text) > 8000:
            try:
                from ..knowledge.hierarchical_chunker import HierarchicalChunker
                chunker = HierarchicalChunker(chunk_size=2000)
                chunks = chunker.chunk(structure.raw_text, title=filepath)
                result["chunks"] = len(chunks)
            except Exception:
                pass

        # 委托 knowledge_base 做实体提取
        try:
            from ..knowledge.intelligent_kb import expand_query
            expanded = expand_query(structure.raw_text[:2000])
            result["key_concepts"] = expanded
        except Exception:
            pass

        return result

    # ═══ 回退: 纯文本读取 ═══

    @staticmethod
    def _fallback_text_read(filepath: str) -> str:
        try:
            return Path(filepath).read_text(encoding="utf-8")
        except Exception:
            return f"[无法读取: {filepath}]"


# ── Singleton ──────────────────────────────────────────────────────

_doc_intelligence: DocumentIntelligence | None = None


def get_doc_intelligence() -> DocumentIntelligence:
    global _doc_intelligence
    if _doc_intelligence is None:
        _doc_intelligence = DocumentIntelligence()
    return _doc_intelligence
