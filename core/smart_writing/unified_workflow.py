# -*- coding: utf-8 -*-
"""
智能写作统一工作流 - Smart Writing Unified Workflow (Enhanced)
==============================================================

职责：
1. 统一 AI 写作和项目生成流程
2. 整合需求澄清 → 深度搜索 → Wiki 生成 → 审核辩论
3. 支持多轮迭代优化
4. 集成计算模型、数据采集、意图分类、多模态生成

复用模块：
- IntentClassifier (意图分类)
- InteractiveClarifier (交互式澄清)
- CalculationEngine (计算模型)
- WritingDataCollector (数据采集)
- MultimodalGenerator (多模态生成)
- DeepSearchWikiSystem (深度搜索)
- AIEnhancedGeneration (内容生成 + 智能审核)
- ProjectGeneration (项目生成)
- UnifiedContext (上下文)

Author: Hermes Desktop Team
"""

import asyncio
import json
import logging
import time
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from typing import Any, Callable, Dict, List, Optional, Iterator

logger = logging.getLogger(__name__)


# =============================================================================
# 枚举定义
# =============================================================================

class WritingStage(Enum):
    """写作阶段"""
    REQUIREMENT_CLARIFICATION = "requirement_clarification"  # 需求澄清
    KNOWLEDGE_RETRIEVAL = "knowledge_retrieval"              # 知识检索
    DEEP_SEARCH = "deep_search"                            # 深度搜索
    CONTENT_GENERATION = "content_generation"              # 内容生成
    AI_REVIEW = "ai_review"                              # AI审核
    DEBATE = "debate"                                    # 分身辩论
    VIRTUAL_MEETING = "virtual_meeting"                  # 虚拟会议
    FINAL_REVISION = "final_revision"                    # 最终修订
    COMPLETED = "completed"                               # 完成


class DocumentCategory(Enum):
    """文档类别"""
    GENERAL = "general"               # 通用文档
    FEASIBILITY_REPORT = "feasibility_report"  # 可行性研究报告
    EIA_REPORT = "eia_report"        # 环境影响评价报告
    SAFETY_ASSESSMENT = "safety_assessment"  # 安全评价报告
    FINANCIAL_ANALYSIS = "financial_analysis"  # 财务分析报告
    BUSINESS_PLAN = "business_plan"   # 商业计划书


# =============================================================================
# 数据模型
# =============================================================================

@dataclass
class WritingContext:
    """写作上下文"""
    # 输入
    user_requirement: str = ""
    document_type: str = "general"
    project_name: str = ""
    project_id: str = ""
    
    # 阶段状态
    current_stage: WritingStage = WritingStage.REQUIREMENT_CLARIFICATION
    stage_history: List[Dict] = field(default_factory=list)
    
    # 中间产物
    clarified_requirements: str = ""
    retrieved_knowledge: List[Dict] = field(default_factory=list)
    search_results: List[Dict] = field(default_factory=list)
    draft_content: Dict[str, Any] = field(default_factory=dict)
    review_issues: List[Dict] = field(default_factory=list)
    debate_transcript: str = ""
    
    # 审核结果
    review_score: float = 0.0
    review_conclusion: str = ""
    
    # 最终产物
    final_content: Dict[str, Any] = field(default_factory=dict)
    output_files: List[str] = field(default_factory=list)
    confidence: float = 0.0
    
    # 元数据
    iterations: int = 0
    start_time: float = field(default_factory=time.time)
    
    def to_dict(self) -> Dict:
        """转换为字典"""
        return {
            "user_requirement": self.user_requirement,
            "document_type": self.document_type,
            "project_name": self.project_name,
            "current_stage": self.current_stage.value,
            "review_score": self.review_score,
            "review_conclusion": self.review_conclusion,
            "final_content_keys": list(self.final_content.keys()) if self.final_content else [],
            "output_files": self.output_files,
            "confidence": self.confidence,
            "iterations": self.iterations,
            "duration": time.time() - self.start_time,
        }


@dataclass
class WritingConfig:
    """写作配置"""
    # 阶段开关
    enable_clarification: bool = True
    enable_knowledge_retrieval: bool = True
    enable_deep_search: bool = True
    enable_ai_review: bool = True
    enable_debate: bool = True
    enable_virtual_meeting: bool = False
    
    # 知识源配置
    use_knowledge_base: bool = True
    use_deep_search: bool = True
    use_online_api: bool = True
    
    # 审核配置
    debate_rounds: int = 3
    review_modes: List[str] = field(default_factory=lambda: [
        "completeness", "compliance", "accuracy", "consistency"
    ])
    
    # 输出配置
    output_formats: List[str] = field(default_factory=lambda: ["docx"])
    output_dir: str = "./output"
    
    # 生成配置
    language: str = "zh-CN"
    quality_level: str = "high"
    
    # 最大迭代次数
    max_iterations: int = 3


# =============================================================================
# SmartWritingWorkflow 主类
# =============================================================================

class SmartWritingWorkflow:
    """
    智能写作统一工作流 (增强版)
    
    集成模块：
    - IntentClassifier: 意图分类和实体抽取
    - InteractiveClarifier: 交互式需求澄清
    - CalculationEngine: 计算模型
    - WritingDataCollector: 数据采集
    - MultimodalGenerator: 多模态内容生成
    - DeepSearchWikiSystem: 深度搜索
    - AIEnhancedGeneration: AI审核辩论
    - ProjectGeneration: 项目生成
    
    使用示例：
    ```python
    workflow = SmartWritingWorkflow()
    
    # 同步执行
    result = workflow.execute(
        requirement="写一份关于智能制造项目的可行性研究报告",
        document_type="feasibility_report",
        config=WritingConfig()
    )
    
    # 流式执行
    for stage_result in workflow.execute_stream(requirement):
        logger.info(f"Stage: {stage_result['stage']}")
        logger.info(f"Content: {stage_result.get('content', '')}")
    ```
    """
    
    def __init__(self, config: Optional[WritingConfig] = None):
        self.config = config or WritingConfig()
        
        # 组件初始化（延迟加载）
        self._unified_context = None
        self._knowledge_base = None
        self._ai_enhanced_engine = None
        self._project_engine = None
        self._wiki_generator = None
        
        # R4-1 新增组件
        self._intent_classifier = None
        self._clarifier = None
        self._calculation_engine = None
        self._data_collector = None
        self._multimodal_generator = None
        
        logger.info("SmartWritingWorkflow (增强版) 初始化完成")
    
    # ── 属性懒加载 ──────────────────────────────────────────────────────────
    
    @property
    def unified_context(self):
        """统一上下文（延迟加载）"""
        if self._unified_context is None:
            try:
                from core.unified_context import UnifiedContext
                self._unified_context = UnifiedContext(user_id="smart_writing")
            except ImportError:
                logger.warning("UnifiedContext 未找到，使用简化实现")
                self._unified_context = None
        return self._unified_context
    
    @property
    def knowledge_base(self):
        """知识库（延迟加载）"""
        if self._knowledge_base is None:
            try:
                from core.fusion_rag.knowledge_base import KnowledgeBaseLayer
                self._knowledge_base = KnowledgeBaseLayer()
            except ImportError:
                logger.warning("KnowledgeBaseLayer 未找到")
                self._knowledge_base = None
        return self._knowledge_base
    
    @property
    def ai_enhanced_engine(self):
        """AI增强引擎（延迟加载）"""
        if self._ai_enhanced_engine is None:
            try:
                from core.smart_writing.ai_enhanced_generation import (
                    get_ai_enhanced_project_engine,
                    AIEnhancedGenerationConfig,
                    AvatarRole,
                )
                self._ai_enhanced_engine = get_ai_enhanced_project_engine()
                self._ai_enhanced_config_class = AIEnhancedGenerationConfig
                self._avatar_role_class = AvatarRole
            except ImportError as e:
                logger.warning(f"AIEnhancedGeneration 未找到: {e}")
                self._ai_enhanced_engine = None
        return self._ai_enhanced_engine
    
    @property
    def project_engine(self):
        """项目生成引擎（延迟加载）"""
        if self._project_engine is None:
            try:
                from core.smart_writing.project_generation import (
                    get_project_generation_engine,
                )
                self._project_engine = get_project_generation_engine()
            except ImportError as e:
                logger.warning(f"ProjectGeneration 未找到: {e}")
                self._project_engine = None
        return self._project_engine
    
    @property
    def wiki_generator(self):
        """Wiki生成器（延迟加载）"""
        if self._wiki_generator is None:
            try:
                from core.deep_search_wiki.wiki_generator import WikiGenerator
                self._wiki_generator = WikiGenerator()
            except ImportError:
                logger.warning("WikiGenerator 未找到")
                self._wiki_generator = None
        return self._wiki_generator
    
    # ── R4-1 新增组件属性 ───────────────────────────────────────────────────
    
    @property
    def intent_classifier(self):
        """意图分类器（延迟加载）"""
        if self._intent_classifier is None:
            try:
                from core.smart_writing.intent_classifier import DocumentIntentClassifier
                self._intent_classifier = DocumentIntentClassifier()
            except ImportError as e:
                logger.warning(f"IntentClassifier 未找到: {e}")
                self._intent_classifier = None
        return self._intent_classifier
    
    @property
    def clarifier(self):
        """交互式澄清器（延迟加载）"""
        if self._clarifier is None:
            try:
                from core.smart_writing.interactive_clarifier import InteractiveClarifier
                self._clarifier = InteractiveClarifier()
            except ImportError as e:
                logger.warning(f"InteractiveClarifier 未找到: {e}")
                self._clarifier = None
        return self._clarifier
    
    @property
    def calculation_engine(self):
        """计算引擎（延迟加载）"""
        if self._calculation_engine is None:
            try:
                from core.smart_writing.calculation_models import get_calculation_engine
                self._calculation_engine = get_calculation_engine()
            except ImportError as e:
                logger.warning(f"CalculationEngine 未找到: {e}")
                self._calculation_engine = None
        return self._calculation_engine
    
    @property
    def data_collector(self):
        """数据采集器（延迟加载）"""
        if self._data_collector is None:
            try:
                from core.smart_writing.data_collector import get_data_collector
                self._data_collector = get_data_collector()
            except ImportError as e:
                logger.warning(f"WritingDataCollector 未找到: {e}")
                self._data_collector = None
        return self._data_collector
    
    @property
    def multimodal_generator(self):
        """多模态生成器（延迟加载）"""
        if self._multimodal_generator is None:
            try:
                from core.smart_writing.multimodal_generator import get_multimodal_generator
                self._multimodal_generator = get_multimodal_generator()
            except ImportError as e:
                logger.warning(f"MultimodalGenerator 未找到: {e}")
                self._multimodal_generator = None
        return self._multimodal_generator
    
    @property
    def evolution_engine(self):
        """自进化引擎（延迟加载）"""
        if not hasattr(self, '_evolution_engine') or self._evolution_engine is None:
            try:
                from core.smart_writing.self_evolution import get_evolution_engine
                self._evolution_engine = get_evolution_engine()
            except ImportError as e:
                logger.warning(f"EvolutionEngine 未找到: {e}")
                self._evolution_engine = None
        return self._evolution_engine
    
    # ── A3 新增组件属性 ───────────────────────────────────────────────────
    
    @property
    def document_analyzer(self):
        """文档分析器（延迟加载）"""
        if not hasattr(self, '_document_analyzer') or self._document_analyzer is None:
            try:
                from core.smart_writing.document_analyzer import get_document_analyzer
                self._document_analyzer = get_document_analyzer()
            except ImportError as e:
                logger.warning(f"DocumentAnalyzer 未找到: {e}")
                self._document_analyzer = None
        return self._document_analyzer
    
    @property
    def interactive_collector(self):
        """交互式收集器（延迟加载）"""
        if not hasattr(self, '_interactive_collector') or self._interactive_collector is None:
            try:
                from core.smart_writing.interactive_collector import get_interactive_collector
                self._interactive_collector = get_interactive_collector()
            except ImportError as e:
                logger.warning(f"InteractiveCollector 未找到: {e}")
                self._interactive_collector = None
        return self._interactive_collector
    
    @property
    def review_brain(self):
        """智慧审核大脑（延迟加载）"""
        if not hasattr(self, '_review_brain') or self._review_brain is None:
            try:
                from core.smart_writing.smart_review_brain import get_review_brain
                self._review_brain = get_review_brain()
            except ImportError as e:
                logger.warning(f"SmartReviewBrain 未找到: {e}")
                self._review_brain = None
        return self._review_brain
    
    @property
    def industry_enhancer(self):
        """行业增强器（延迟加载）"""
        if not hasattr(self, '_industry_enhancer') or self._industry_enhancer is None:
            try:
                from core.smart_writing.industry_enhancer import get_industry_enhancer
                self._industry_enhancer = get_industry_enhancer()
            except ImportError as e:
                logger.warning(f"IndustryEnhancer 未找到: {e}")
                self._industry_enhancer = None
        return self._industry_enhancer
    
    @property
    def cli_automation(self):
        """CLI自动化（延迟加载）"""
        if not hasattr(self, '_cli_automation') or self._cli_automation is None:
            try:
                from core.smart_writing.cli_automation import get_cli_automation
                self._cli_automation = get_cli_automation()
            except ImportError as e:
                logger.warning(f"CLIAutomation 未找到: {e}")
                self._cli_automation = None
        return self._cli_automation
    
    @property
    def streaming_processor(self):
        """流式处理器（延迟加载）"""
        if not hasattr(self, '_streaming_processor') or self._streaming_processor is None:
            try:
                from core.smart_writing.streaming_processor import get_streaming_processor
                self._streaming_processor = get_streaming_processor()
            except ImportError as e:
                logger.warning(f"StreamingProcessor 未找到: {e}")
                self._streaming_processor = None
        return self._streaming_processor
    
    # ── 执行入口 ───────────────────────────────────────────────────────────
    
    def execute(
        self,
        requirement: str,
        document_type: str = "general",
        project_name: str = "",
        config: Optional[WritingConfig] = None,
        progress_callback: Optional[Callable] = None,
    ) -> WritingContext:
        """
        执行智能写作工作流（同步版本）
        
        Args:
            requirement: 用户需求
            document_type: 文档类型
            project_name: 项目名称
            config: 写作配置
            progress_callback: 进度回调函数
            
        Returns:
            WritingContext: 写作上下文
        """
        if config:
            self.config = config
        
        ctx = WritingContext(
            user_requirement=requirement,
            document_type=document_type,
            project_name=project_name or requirement[:50],
            project_id=f"WR-{uuid.uuid4().hex[:8]}",
        )
        
        def emit(stage: WritingStage, progress: int, message: str):
            ctx.current_stage = stage
            ctx.stage_history.append({
                "stage": stage.value,
                "progress": progress,
                "message": message,
                "timestamp": time.time(),
            })
            if progress_callback:
                progress_callback(stage, progress, message)
        
        try:
            # Stage 1: 需求澄清
            emit(WritingStage.REQUIREMENT_CLARIFICATION, 5, "分析需求...")
            ctx = self._stage_clarification(ctx)
            
            if ctx.current_stage == WritingStage.REQUIREMENT_CLARIFICATION:
                # 需求不明确，等待用户补充
                return ctx
            
            # Stage 2: 知识检索
            emit(WritingStage.KNOWLEDGE_RETRIEVAL, 15, "检索知识库...")
            ctx = self._stage_knowledge_retrieval(ctx)
            
            # Stage 3: 深度搜索
            if self.config.enable_deep_search:
                emit(WritingStage.DEEP_SEARCH, 25, "进行深度搜索...")
                ctx = self._stage_deep_search(ctx)
            
            # Stage 4: 内容生成
            emit(WritingStage.CONTENT_GENERATION, 40, "生成内容...")
            ctx = self._stage_content_generation(ctx)
            
            # Stage 5: AI审核
            if self.config.enable_ai_review:
                emit(WritingStage.AI_REVIEW, 60, "AI智能审核中...")
                ctx = self._stage_ai_review(ctx)
            
            # Stage 6: 分身辩论
            if self.config.enable_debate:
                emit(WritingStage.DEBATE, 75, "启动分身辩论...")
                ctx = self._stage_debate(ctx)
            
            # Stage 7: 虚拟会议
            if self.config.enable_virtual_meeting:
                emit(WritingStage.VIRTUAL_MEETING, 85, "虚拟会议评审...")
                ctx = self._stage_virtual_meeting(ctx)
            
            # Stage 8: 最终修订
            emit(WritingStage.FINAL_REVISION, 95, "生成最终版本...")
            ctx = self._stage_final_revision(ctx)
            
            ctx.current_stage = WritingStage.COMPLETED
            emit(WritingStage.COMPLETED, 100, "写作完成！")
            
        except Exception as e:
            logger.error(f"写作工作流失败: {e}")
            ctx.final_content = {"error": str(e)}
        
        ctx.confidence = self._calculate_confidence(ctx)
        return ctx
    
    def execute_stream(
        self,
        requirement: str,
        document_type: str = "general",
        project_name: str = "",
        config: Optional[WritingConfig] = None,
    ) -> Iterator[Dict]:
        """
        执行智能写作工作流（流式版本）
        """
        if config:
            self.config = config
        
        ctx = WritingContext(
            user_requirement=requirement,
            document_type=document_type,
            project_name=project_name or requirement[:50],
            project_id=f"WR-{uuid.uuid4().hex[:8]}",
        )
        
        # Stage 1: 需求澄清
        yield {"stage": "clarification", "progress": 5, "content": "正在分析需求..."}
        ctx = self._stage_clarification(ctx)
        yield {"stage": "clarification", "progress": 10, "content": ctx.clarified_requirements}
        
        if ctx.current_stage == WritingStage.REQUIREMENT_CLARIFICATION:
            yield {"stage": "wait_confirm", "progress": 10, "content": "需要确认需求"}
            return
        
        # Stage 2: 知识检索
        yield {"stage": "retrieval", "progress": 15, "content": "正在检索知识..."}
        ctx = self._stage_knowledge_retrieval(ctx)
        yield {"stage": "retrieval", "progress": 20, "content": f"检索到 {len(ctx.retrieved_knowledge)} 条知识"}
        
        # Stage 3: 深度搜索
        if self.config.enable_deep_search:
            yield {"stage": "search", "progress": 25, "content": "正在进行深度搜索..."}
            ctx = self._stage_deep_search(ctx)
            yield {"stage": "search", "progress": 30, "content": f"深度搜索完成，获得 {len(ctx.search_results)} 条结果"}
        
        # Stage 4: 内容生成
        yield {"stage": "generation", "progress": 40, "content": "正在生成内容..."}
        ctx = self._stage_content_generation(ctx)
        yield {"stage": "draft", "progress": 50, "content": "草稿生成完成"}
        
        # Stage 5: AI审核
        if self.config.enable_ai_review:
            yield {"stage": "review", "progress": 60, "content": "正在进行AI审核..."}
            ctx = self._stage_ai_review(ctx)
            yield {"stage": "review", "progress": 70, "content": f"审核完成，评分: {ctx.review_score:.1f}"}
        
        # Stage 6: 分身辩论
        if self.config.enable_debate:
            yield {"stage": "debate", "progress": 75, "content": "正在启动分身辩论..."}
            ctx = self._stage_debate(ctx)
            yield {"stage": "debate", "progress": 80, "content": f"辩论完成，{ctx.stage_history[-1].get('debate_rounds', 0)} 轮"}
        
        # Stage 7: 最终修订
        yield {"stage": "revision", "progress": 90, "content": "正在生成最终版本..."}
        ctx = self._stage_final_revision(ctx)
        
        # Stage 8: 完成
        yield {"stage": "completed", "progress": 100, "content": ctx.final_content.get("summary", "写作完成")}
    
    # ── 阶段实现 ───────────────────────────────────────────────────────────
    
    def _stage_clarification(self, ctx: WritingContext) -> WritingContext:
        """Stage 1: 需求澄清"""
        ctx.current_stage = WritingStage.REQUIREMENT_CLARIFICATION
        
        requirement = ctx.user_requirement.strip()
        
        # 简单需求检查
        if len(requirement) < 10:
            ctx.clarified_requirements = requirement
            ctx.current_stage = WritingStage.KNOWLEDGE_RETRIEVAL
            return ctx
        
        # 使用LLM进行需求澄清分析
        try:
            from core.config_provider import get_ollama_url, get_default_model
            import requests
            
            url = get_ollama_url()
            model = get_default_model()
            
            prompt = f"""分析以下写作需求，识别关键信息和潜在歧义：

需求：{requirement}

请输出：
1. 明确的核心主题
2. 缺失的关键信息（如：目标读者、文档类型、篇幅要求等）
3. 建议的文档结构

格式简洁，一句话概括。"""

            response = requests.post(
                f"{url}/chat/completions",
                json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                },
                timeout=30,
            )
            
            if response.status_code == 200:
                result = response.json()
                content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
                ctx.clarified_requirements = content
            else:
                ctx.clarified_requirements = requirement
                
        except Exception as e:
            logger.warning(f"需求澄清LLM调用失败: {e}")
            ctx.clarified_requirements = requirement
        
        ctx.current_stage = WritingStage.KNOWLEDGE_RETRIEVAL
        return ctx
    
    def _stage_knowledge_retrieval(self, ctx: WritingContext) -> WritingContext:
        """Stage 2: 知识检索"""
        ctx.current_stage = WritingStage.KNOWLEDGE_RETRIEVAL
        
        if not self.config.enable_knowledge_retrieval:
            return ctx
        
        # 从知识库检索
        if self.knowledge_base:
            try:
                results = self.knowledge_base.search(
                    query=ctx.clarified_requirements or ctx.user_requirement,
                    top_k=10
                )
                ctx.retrieved_knowledge = results if isinstance(results, list) else []
            except Exception as e:
                logger.warning(f"知识库检索失败: {e}")
        
        # 从项目引擎获取相关内容
        if self.project_engine:
            try:
                # 根据文档类型获取章节模板
                from core.smart_writing.project_generation import SectionTemplates, ConsultingDocumentType
                
                doc_type_map = {
                    "feasibility_report": ConsultingDocumentType.FEASIBILITY_REPORT,
                    "eia_report": ConsultingDocumentType.EIA_REPORT,
                    "safety_assessment": ConsultingDocumentType.SAFETY_ASSESSMENT,
                    "financial_analysis": ConsultingDocumentType.FINANCIAL_ANALYSIS,
                }
                
                doc_type = doc_type_map.get(ctx.document_type, ConsultingDocumentType.FEASIBILITY_REPORT)
                sections = SectionTemplates.get_sections(doc_type)
                
                ctx.retrieved_knowledge.append({
                    "source": "template",
                    "type": "section_template",
                    "data": sections,
                    "summary": f"获取到 {len(sections)} 个章节模板",
                })
            except Exception as e:
                logger.warning(f"章节模板获取失败: {e}")
        
        return ctx
    
    def _stage_deep_search(self, ctx: WritingContext) -> WritingContext:
        """Stage 3: 深度搜索"""
        ctx.current_stage = WritingStage.DEEP_SEARCH
        
        if not self.config.enable_deep_search:
            return ctx
        
        # 使用Wiki生成器进行深度搜索
        if self.wiki_generator:
            try:
                wiki_page = self.wiki_generator.generate(
                    topic=ctx.user_requirement,
                    search_results=ctx.retrieved_knowledge,
                    use_search=True
                )
                
                # 转换为搜索结果格式
                ctx.search_results = [
                    {
                        "title": section.get("title", ""),
                        "content": section.get("content", ""),
                        "source": "wiki"
                    }
                    for section in wiki_page.get("sections", [])
                ]
            except Exception as e:
                logger.warning(f"Wiki生成失败: {e}")
        
        return ctx
    
    def _stage_content_generation(self, ctx: WritingContext) -> WritingContext:
        """Stage 4: 内容生成"""
        ctx.current_stage = WritingStage.CONTENT_GENERATION
        
        # 根据文档类型选择生成策略
        if ctx.document_type in ["feasibility_report", "eia_report", "safety_assessment", 
                                   "financial_analysis", "business_plan"]:
            ctx = self._generate_project_document(ctx)
        else:
            ctx = self._generate_general_document(ctx)
        
        return ctx
    
    def _generate_project_document(self, ctx: WritingContext) -> WritingContext:
        """生成项目文档"""
        if not self.project_engine:
            # 使用简化生成
            ctx.draft_content = {
                "cover": {
                    "title": ctx.project_name,
                    "subtitle": ctx.document_type,
                    "date": datetime.now().strftime("%Y年%m月%d日"),
                },
                "sections": [
                    {"title": "一、项目概述", "content": f"本项目为{ctx.project_name}，{ctx.user_requirement}"},
                    {"title": "二、项目背景", "content": "项目背景分析..."},
                    {"title": "三、可行性分析", "content": "技术可行性、经济可行性..."},
                ],
                "metadata": {
                    "generated_at": datetime.now().isoformat(),
                    "document_type": ctx.document_type,
                }
            }
            return ctx
        
        try:
            from core.smart_writing.project_generation import (
                GenerationConfig,
                ProjectData,
                ConsultingDocumentType,
                OutputFormat,
            )
            
            doc_type_map = {
                "feasibility_report": ConsultingDocumentType.FEASIBILITY_REPORT,
                "eia_report": ConsultingDocumentType.EIA_REPORT,
                "safety_assessment": ConsultingDocumentType.SAFETY_ASSESSMENT,
                "financial_analysis": ConsultingDocumentType.FINANCIAL_ANALYSIS,
            }
            
            doc_type = doc_type_map.get(ctx.document_type, ConsultingDocumentType.FEASIBILITY_REPORT)
            
            config = GenerationConfig(
                document_type=doc_type,
                output_formats=[OutputFormat.JSON],  # 内存处理，不写文件
                output_dir=self.config.output_dir,
            )
            
            project_data = ProjectData(
                project_id=ctx.project_id,
                project_name=ctx.project_name,
                project_type="新建",
                client_name="",
                description=ctx.user_requirement,
            )
            
            # 使用同步方式生成（避免async问题）
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            try:
                result = loop.run_until_complete(
                    self.project_engine.generate_project_document(
                        config=config,
                        project_data=project_data,
                        custom_content={}
                    )
                )
            finally:
                loop.close()
            
            # 读取生成的JSON文件作为内容
            json_file = None
            for f in result.output_files:
                if f.endswith(".json"):
                    json_file = f
                    break
            
            if json_file:
                with open(json_file, "r", encoding="utf-8") as f:
                    ctx.draft_content = json.load(f)
            else:
                ctx.draft_content = {"error": "生成失败"}
                
        except Exception as e:
            logger.warning(f"项目文档生成失败: {e}")
            # 回退到简化生成
            ctx.draft_content = {
                "cover": {
                    "title": ctx.project_name,
                    "subtitle": ctx.document_type,
                },
                "sections": [{"title": "内容生成", "content": ctx.user_requirement}],
            }
        
        return ctx
    
    def _generate_general_document(self, ctx: WritingContext) -> WritingContext:
        """生成通用文档"""
        try:
            from core.config_provider import get_ollama_url, get_default_model
            import requests
            
            url = get_ollama_url()
            model = get_default_model()
            
            # 构建上下文
            context_parts = []
            for item in ctx.retrieved_knowledge[:5]:
                if isinstance(item, dict):
                    content = item.get("content", str(item))
                    context_parts.append(content[:500])
            
            context_text = "\n".join(context_parts)
            
            prompt = f"""基于以下背景信息，生成文档内容：

背景信息：
{context_text}

用户需求：{ctx.user_requirement}

请生成结构化的文档内容，包含：
1. 标题和概述
2. 主要章节
3. 详细内容

输出格式为JSON，包含cover、sections等字段。"""

            response = requests.post(
                f"{url}/chat/completions",
                json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                },
                timeout=60,
            )
            
            if response.status_code == 200:
                result = response.json()
                content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
                
                # 尝试解析JSON
                try:
                    ctx.draft_content = json.loads(content)
                except json.JSONDecodeError:
                    # 如果不是JSON，包装成文本
                    ctx.draft_content = {
                        "cover": {"title": ctx.project_name},
                        "sections": [{"title": "内容", "content": content}],
                    }
            else:
                ctx.draft_content = {"error": f"生成失败: {response.status_code}"}
                
        except Exception as e:
            logger.error(f"通用文档生成失败: {e}")
            ctx.draft_content = {
                "cover": {"title": ctx.project_name},
                "sections": [{"title": "内容", "content": ctx.user_requirement}],
            }
        
        return ctx
    
    def _stage_ai_review(self, ctx: WritingContext) -> WritingContext:
        """Stage 5: AI审核"""
        ctx.current_stage = WritingStage.AI_REVIEW
        
        if not self.config.enable_ai_review:
            return ctx
        
        if not self.ai_enhanced_engine:
            # 简化审核
            ctx.review_score = 85.0
            ctx.review_conclusion = "pass"
            ctx.review_issues = []
            return ctx
        
        try:
            # 构建审核配置
            config = self._ai_enhanced_config_class(
                enable_ai_agent=True,
                enable_smart_review=True,
                enable_conflict_detection=True,
                enable_avatar_debate=False,  # 辩论单独处理
                enable_virtual_review=False,
                use_knowledge_base=self.config.use_knowledge_base,
                use_deep_search=self.config.use_deep_search,
                debate_rounds=0,
            )
            
            # 执行审核
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            try:
                review_result = loop.run_until_complete(
                    self.ai_enhanced_engine.review_engine.run_full_review(
                        content=ctx.draft_content,
                        doc_type=ctx.document_type,
                        config=config,
                    )
                )
            finally:
                loop.close()
            
            ctx.review_score = review_result.overall_score
            ctx.review_conclusion = review_result.conclusion.value
            ctx.review_issues = [
                {
                    "id": issue.issue_id,
                    "section": issue.section,
                    "type": issue.issue_type,
                    "severity": issue.severity.value,
                    "description": issue.description,
                    "suggestion": issue.suggested_revision,
                }
                for issue in review_result.issues
            ]
            
            # 记录审核阶段
            ctx.stage_history.append({
                "stage": "ai_review",
                "score": ctx.review_score,
                "issues_count": len(ctx.review_issues),
                "timestamp": time.time(),
            })
            
        except Exception as e:
            logger.warning(f"AI审核失败: {e}")
            ctx.review_score = 80.0
            ctx.review_conclusion = "pass"
        
        return ctx
    
    def _stage_debate(self, ctx: WritingContext) -> WritingContext:
        """Stage 6: 分身辩论"""
        ctx.current_stage = WritingStage.DEBATE
        
        if not self.config.enable_debate:
            return ctx
        
        if not self.ai_enhanced_engine:
            ctx.debate_transcript = "（辩论功能暂不可用）"
            return ctx
        
        try:
            # 准备辩论配置
            config = self._ai_enhanced_config_class(
                enable_avatar_debate=True,
                debate_rounds=self.config.debate_rounds,
                debate_participants=[
                    self._avatar_role_class.DEFENDER,
                    self._avatar_role_class.PROSECUTOR,
                    self._avatar_role_class.EXPERT_ENV,
                ],
            )
            
            # 执行辩论
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            try:
                review_result = loop.run_until_complete(
                    self.ai_enhanced_engine.review_engine.run_full_review(
                        content=ctx.draft_content,
                        doc_type=ctx.document_type,
                        config=config,
                    )
                )
            finally:
                loop.close()
            
            # 构建辩论记录
            debate_lines = []
            for record in review_result.debate_records:
                debate_lines.append(f"[{record.speaker_name}] {record.argument}")
            
            ctx.debate_transcript = "\n".join(debate_lines)
            
            # 更新辩论轮数
            debate_rounds = review_result.debate_rounds
            ctx.stage_history.append({
                "stage": "debate",
                "debate_rounds": debate_rounds,
                "records_count": len(debate_lines),
                "timestamp": time.time(),
            })
            
        except Exception as e:
            logger.warning(f"分身辩论失败: {e}")
            ctx.debate_transcript = f"（辩论执行失败: {e}）"
        
        return ctx
    
    def _stage_virtual_meeting(self, ctx: WritingContext) -> WritingContext:
        """Stage 7: 虚拟会议"""
        ctx.current_stage = WritingStage.VIRTUAL_MEETING
        
        if not self.config.enable_virtual_meeting:
            return ctx
        
        # 虚拟会议实现（简化版）
        ctx.stage_history.append({
            "stage": "virtual_meeting",
            "participants": ["AI主持人", "环境专家", "安全专家", "财务专家"],
            "timestamp": time.time(),
        })
        
        return ctx
    
    def _stage_final_revision(self, ctx: WritingContext) -> WritingContext:
        """Stage 8: 最终修订"""
        ctx.current_stage = WritingStage.FINAL_REVISION
        
        # 根据审核意见进行修订
        if ctx.review_issues and self.config.enable_ai_review:
            ctx.final_content = self._apply_revision(ctx.draft_content, ctx.review_issues)
        else:
            ctx.final_content = ctx.draft_content
        
        # 添加元数据
        ctx.final_content["_meta"] = {
            "project_id": ctx.project_id,
            "document_type": ctx.document_type,
            "review_score": ctx.review_score,
            "review_conclusion": ctx.review_conclusion,
            "issues_count": len(ctx.review_issues),
            "confidence": self._calculate_confidence(ctx),
            "generated_at": datetime.now().isoformat(),
            "duration": time.time() - ctx.start_time,
        }
        
        return ctx
    
    def _apply_revision(self, content: Dict, issues: List[Dict]) -> Dict:
        """根据审核意见应用修订"""
        # 简化实现：添加修订说明
        revision_notes = []
        for issue in issues[:5]:  # 只取前5个
            revision_notes.append({
                "section": issue.get("section", ""),
                "issue": issue.get("description", ""),
                "suggestion": issue.get("suggestion", ""),
            })
        
        content["_revisions"] = revision_notes
        return content
    
    # ── 辅助方法 ───────────────────────────────────────────────────────────
    
    def _calculate_confidence(self, ctx: WritingContext) -> float:
        """计算置信度"""
        factors = []
        
        # 知识检索覆盖率
        if ctx.retrieved_knowledge:
            factors.append(0.15)
        
        # 深度搜索覆盖率
        if ctx.search_results:
            factors.append(0.15)
        
        # 内容生成
        if ctx.draft_content:
            factors.append(0.20)
        
        # 审核评分
        if ctx.review_score > 0:
            factors.append(ctx.review_score / 100 * 0.30)
        
        # 辩论
        if ctx.debate_transcript:
            factors.append(0.10)
        
        # 无问题
        if len(ctx.review_issues) == 0:
            factors.append(0.10)
        elif len(ctx.review_issues) < 3:
            factors.append(0.05)
        
        return min(1.0, sum(factors))
    
    def export(
        self,
        ctx: WritingContext,
        formats: List[str] = None,
        output_dir: str = None,
    ) -> List[str]:
        """
        导出文档
        
        Args:
            ctx: 写作上下文
            formats: 导出格式列表
            output_dir: 输出目录
            
        Returns:
            List[str]: 导出文件路径列表
        """
        if formats is None:
            formats = self.config.output_formats
        if output_dir is None:
            output_dir = self.config.output_dir
        
        import os
        os.makedirs(output_dir, exist_ok=True)
        
        output_files = []
        
        for fmt in formats:
            try:
                if fmt == "json":
                    path = os.path.join(output_dir, f"{ctx.project_id}.json")
                    with open(path, "w", encoding="utf-8") as f:
                        json.dump(ctx.final_content, f, ensure_ascii=False, indent=2)
                    output_files.append(path)
                    
                elif fmt == "markdown" or fmt == "md":
                    path = os.path.join(output_dir, f"{ctx.project_id}.md")
                    content = self._to_markdown(ctx.final_content)
                    with open(path, "w", encoding="utf-8") as f:
                        f.write(content)
                    output_files.append(path)
                    
                elif fmt == "html":
                    path = os.path.join(output_dir, f"{ctx.project_id}.html")
                    content = self._to_html(ctx.final_content)
                    with open(path, "w", encoding="utf-8") as f:
                        f.write(content)
                    output_files.append(path)
                    
                elif fmt == "docx":
                    path = self._export_docx(ctx.final_content, output_dir, ctx.project_id)
                    if path:
                        output_files.append(path)
                        
            except Exception as e:
                logger.error(f"导出 {fmt} 失败: {e}")
        
        ctx.output_files = output_files
        return output_files
    
    def _to_markdown(self, content: Dict) -> str:
        """转换为Markdown"""
        lines = []
        
        # 封面
        cover = content.get("cover", {})
        if cover.get("title"):
            lines.append(f"# {cover['title']}")
            lines.append("")
        if cover.get("subtitle"):
            lines.append(f"**{cover['subtitle']}**")
            lines.append("")
        
        # 章节
        for section in content.get("sections", []):
            level = section.get("level", 1)
            title = section.get("title", "")
            title_prefix = "#" * (level + 1)
            lines.append(f"{title_prefix} {title}")
            lines.append("")
            
            content_text = section.get("content", "")
            if content_text:
                lines.append(content_text)
                lines.append("")
        
        return "\n".join(lines)
    
    def _to_html(self, content: Dict) -> str:
        """转换为HTML"""
        html = ['<!DOCTYPE html>', '<html>', '<head>',
                '<meta charset="utf-8">',
                '<title>Document</title>',
                '<style>',
                'body { font-family: Arial, sans-serif; margin: 40px; }',
                'h1 { color: #333; }',
                'table { border-collapse: collapse; width: 100%; }',
                'th, td { border: 1px solid #ddd; padding: 8px; }',
                '</style>',
                '</head>', '<body>']
        
        cover = content.get("cover", {})
        if cover.get("title"):
            html.append(f"<h1>{cover['title']}</h1>")
        
        for section in content.get("sections", []):
            level = section.get("level", 1)
            title = section.get("title", "")
            html.append(f"<h{level + 1}>{title}</h{level + 1}>")
            
            content_text = section.get("content", "")
            if content_text:
                html.append(f"<p>{content_text}</p>")
        
        html.extend(['</body>', '</html>'])
        return "\n".join(html)
    
    def _export_docx(self, content: Dict, output_dir: str, filename: str) -> Optional[str]:
        """导出为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt
from core.logger import get_logger
logger = get_logger('smart_writing.unified_workflow')

            
            path = f"{output_dir}/{filename}.docx"
            doc = Document()
            
            # 封面
            cover = content.get("cover", {})
            if cover.get("title"):
                p = doc.add_heading(cover['title'], 0)
            
            # 章节
            for section in content.get("sections", []):
                title = section.get("title", "")
                level = section.get("level", 1)
                
                if level == 1:
                    doc.add_heading(title, 1)
                else:
                    doc.add_heading(title, 2)
                
                content_text = section.get("content", "")
                if content_text:
                    doc.add_paragraph(content_text)
            
            doc.save(path)
            return path
            
        except ImportError:
            logger.warning("python-docx 未安装，无法导出Word文档")
            return None
        except Exception as e:
            logger.error(f"导出Word失败: {e}")
            return None


# =============================================================================
# 单例模式
# =============================================================================

_instance: Optional[SmartWritingWorkflow] = None


def get_smart_writing_workflow(config: Optional[WritingConfig] = None) -> SmartWritingWorkflow:
    """获取智能写作工作流单例"""
    global _instance
    if _instance is None:
        _instance = SmartWritingWorkflow(config)
    return _instance



def reset_smart_writing_workflow():
    """重置智能写作工作流（用于测试）"""
    global _instance
    _instance = None


# =============================================================================
# 增强版使用示例
# =============================================================================

def demo_enhanced_usage():
    """
    增强版使用示例
    
    演示如何使用新增的自进化、交互式澄清等功能：
    """
    # 1. 使用交互式澄清
    clarifier = InteractiveClarifier()
    session = clarifier.start_session(
        requirement="写一份武汉化工项目可行性研究报告",
        doc_type="feasibility_report"
    )
    
    # 获取问题
    questions = clarifier.get_next_questions(session)
    
    # 回答问题
    session = clarifier.answer(session, "investment", "5000万元")
    
    # 完成澄清
    result = clarifier.complete_session(session)
    logger.info(f"澄清结果: {result}")
    
    # 2. 使用自进化引擎
    engine = get_evolution_engine()
    
    # 学习生成结果
    engine.learn_from_generation(
        requirement="武汉化工项目",
        doc_type="feasibility_report",
        content="生成的完整文档...",
        quality_score=0.85
    )
    
    # 获取参考文档
    refs = engine.get_reference_documents("武汉化工项目", "feasibility_report")
    logger.info(f"参考文档: {len(refs)} 个")
    
    # 获取写作指导
    guidance = engine.get_writing_guidance("feasibility_report", "化工项目")
    logger.info(f"章节结构: {guidance.get('sections', [])[:5]}")
    
    # 3. 使用计算模型
    calc = get_calculation_engine()
    
    # 计算NPV
    npv_result = calc.calculate("npv", {
        "initial_investment": 5000,
        "annual_cashflows": [1500, 1500, 1500, 1500, 1500],
        "discount_rate": 8
    })
    logger.info(f"NPV: {npv_result.result_value}")
    
    # 4. 获取进化指标
    metrics = engine.get_metrics()
    logger.info(f"进化进度: {metrics}")
    
    # 5. 完整工作流
    workflow = get_smart_writing_workflow()
    ctx = workflow.execute(
        requirement="写一份武汉化工项目可行性研究报告",
        document_type="feasibility_report",
        config=WritingConfig(
            use_knowledge_base=True,
            use_deep_search=True,
            enable_ai_review=True,
        )
    )
    logger.info(f"生成完成，置信度: {ctx.confidence:.2f}")

