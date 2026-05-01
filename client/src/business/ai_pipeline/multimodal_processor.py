"""
多模态输入处理引擎 - MultimodalProcessor

核心功能：
1. 文本输入处理（自然语言、Markdown、表格）
2. 文件上传处理（PDF、Word、图片、原型文件）
3. 可视化输入处理（流程图、ER图、线框图）
4. 外部集成（JIRA、Confluence、API文档）
"""

from typing import Dict, Any, Optional, List
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
import asyncio
from loguru import logger


class InputType(Enum):
    """输入类型"""
    TEXT = "text"
    FILE = "file"
    IMAGE = "image"
    PDF = "pdf"
    DOCUMENT = "document"
    FIGMA = "figma"
    URL = "url"
    STRUCTURED = "structured"


class FileType(Enum):
    """文件类型"""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"
    MD = "md"
    JPG = "jpg"
    PNG = "png"
    SVG = "svg"
    FIGMA = "figma"


@dataclass
class InputData:
    """输入数据"""
    id: str
    type: InputType
    content: str = ""
    file_path: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    parsed_content: Optional[Dict[str, Any]] = None


@dataclass
class ParsedResult:
    """解析结果"""
    success: bool
    content: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None
    confidence: float = 0.0


class MultimodalProcessor:
    """
    多模态输入处理引擎
    
    核心特性：
    1. 文本输入处理 - 自然语言描述、结构化数据、Markdown格式
    2. 文件上传处理 - PDF/Word/Excel解析、图片分析、原型文件
    3. 可视化输入处理 - 流程图、ER图、界面线框图
    4. 外部集成 - JIRA/Confluence数据、API文档导入、竞品分析链接
    """

    def __init__(self):
        self._logger = logger.bind(component="MultimodalProcessor")
        self._input_handlers = {
            InputType.TEXT: self._handle_text,
            InputType.FILE: self._handle_file,
            InputType.IMAGE: self._handle_image,
            InputType.PDF: self._handle_pdf,
            InputType.URL: self._handle_url,
            InputType.STRUCTURED: self._handle_structured
        }

    async def process(self, input_data: InputData) -> ParsedResult:
        """处理输入数据"""
        handler = self._input_handlers.get(input_data.type)
        if not handler:
            return ParsedResult(
                success=False,
                error=f"不支持的输入类型: {input_data.type}"
            )
        
        try:
            result = await handler(input_data)
            self._logger.info(f"输入处理成功: {input_data.type}")
            return result
        except Exception as e:
            self._logger.error(f"输入处理失败: {e}")
            return ParsedResult(success=False, error=str(e))

    async def _handle_text(self, input_data: InputData) -> ParsedResult:
        """处理文本输入"""
        content = input_data.content
        
        # 提取关键词
        keywords = self._extract_keywords(content)
        
        # 检测结构化数据
        structured_data = self._detect_structured(content)
        
        # 分类需求类型
        category = self._classify_requirement(content)
        
        return ParsedResult(
            success=True,
            content={
                "raw": content,
                "keywords": keywords,
                "structured": structured_data,
                "category": category,
                "length": len(content)
            },
            confidence=0.95
        )

    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        # 简单的关键词提取（可以扩展为使用NLP模型）
        common_keywords = ["用户", "功能", "需求", "系统", "模块", "页面", "API", "接口"]
        found = []
        
        for keyword in common_keywords:
            if keyword in text:
                found.append(keyword)
        
        return found

    def _detect_structured(self, text: str) -> Optional[Dict[str, Any]]:
        """检测结构化数据"""
        # 检查是否包含表格
        if "|" in text and "\n" in text:
            lines = text.strip().split("\n")
            if any("|" in line for line in lines[:5]):
                return {"type": "table", "hint": "可能包含表格数据"}
        
        # 检查是否为JSON
        if text.strip().startswith("{") or text.strip().startswith("["):
            try:
                import json
                parsed = json.loads(text)
                return {"type": "json", "data": parsed}
            except:
                pass
        
        # 检查是否为Markdown列表
        if text.strip().startswith("- ") or text.strip().startswith("* "):
            return {"type": "list", "hint": "可能包含列表"}
        
        return None

    def _classify_requirement(self, text: str) -> str:
        """分类需求类型"""
        if any(word in text for word in ["性能", "响应时间", "并发", "速度"]):
            return "非功能需求-性能"
        if any(word in text for word in ["安全", "权限", "加密"]):
            return "非功能需求-安全"
        if any(word in text for word in ["页面", "界面", "UI"]):
            return "功能需求-前端"
        if any(word in text for word in ["API", "接口", "后端"]):
            return "功能需求-后端"
        return "功能需求-通用"

    async def _handle_file(self, input_data: InputData) -> ParsedResult:
        """处理文件上传"""
        if not input_data.file_path:
            return ParsedResult(success=False, error="文件路径为空")
        
        file_path = Path(input_data.file_path)
        if not file_path.exists():
            return ParsedResult(success=False, error="文件不存在")
        
        # 获取文件类型
        file_ext = file_path.suffix.lower()
        file_type = self._get_file_type(file_ext)
        
        if file_type == FileType.PDF:
            return await self._parse_pdf(input_data.file_path)
        elif file_type in [FileType.DOCX]:
            return await self._parse_document(input_data.file_path)
        elif file_type in [FileType.XLSX, FileType.CSV]:
            return await self._parse_spreadsheet(input_data.file_path)
        elif file_type in [FileType.JPG, FileType.PNG]:
            return await self._parse_image(input_data.file_path)
        elif file_type == FileType.JSON:
            return await self._parse_json(input_data.file_path)
        elif file_type == FileType.MD:
            return await self._parse_markdown(input_data.file_path)
        
        return ParsedResult(
            success=True,
            content={
                "file_type": file_type.value,
                "size": file_path.stat().st_size,
                "name": file_path.name,
                "hint": "文件已上传，等待进一步处理"
            },
            confidence=0.8
        )

    def _get_file_type(self, ext: str) -> FileType:
        """获取文件类型"""
        mapping = {
            ".pdf": FileType.PDF,
            ".docx": FileType.DOCX,
            ".xlsx": FileType.XLSX,
            ".csv": FileType.CSV,
            ".json": FileType.JSON,
            ".md": FileType.MD,
            ".jpg": FileType.JPG,
            ".jpeg": FileType.JPG,
            ".png": FileType.PNG,
            ".svg": FileType.SVG
        }
        return mapping.get(ext, FileType.DOCUMENT)

    async def _handle_image(self, input_data: InputData) -> ParsedResult:
        """处理图片输入"""
        if input_data.file_path:
            return await self._parse_image(input_data.file_path)
        
        return ParsedResult(
            success=True,
            content={
                "type": "image",
                "hint": "图片已接收，可进行OCR识别"
            },
            confidence=0.7
        )

    async def _handle_pdf(self, input_data: InputData) -> ParsedResult:
        """处理PDF输入"""
        if input_data.file_path:
            return await self._parse_pdf(input_data.file_path)
        
        return ParsedResult(success=False, error="PDF文件路径为空")

    async def _handle_url(self, input_data: InputData) -> ParsedResult:
        """处理URL输入"""
        url = input_data.content
        if not url:
            return ParsedResult(success=False, error="URL为空")
        
        # 识别URL类型
        url_type = self._classify_url(url)
        
        return ParsedResult(
            success=True,
            content={
                "url": url,
                "type": url_type,
                "hint": "URL已记录，可进行内容抓取"
            },
            confidence=0.85
        )

    def _classify_url(self, url: str) -> str:
        """分类URL类型"""
        if "jira" in url.lower():
            return "jira"
        if "confluence" in url.lower():
            return "confluence"
        if "github" in url.lower():
            return "github"
        if "figma" in url.lower():
            return "figma"
        if "api" in url.lower() or "swagger" in url.lower():
            return "api_documentation"
        return "general"

    async def _handle_structured(self, input_data: InputData) -> ParsedResult:
        """处理结构化输入"""
        try:
            import json
            data = json.loads(input_data.content)
            return ParsedResult(
                success=True,
                content={
                    "type": "structured",
                    "data": data,
                    "schema": self._infer_schema(data)
                },
                confidence=0.95
            )
        except json.JSONDecodeError as e:
            return ParsedResult(success=False, error=f"JSON解析失败: {e}")

    def _infer_schema(self, data: Dict[str, Any]) -> Dict[str, str]:
        """推断数据结构"""
        schema = {}
        for key, value in data.items():
            schema[key] = type(value).__name__
        return schema

    async def _parse_pdf(self, file_path: str) -> ParsedResult:
        """解析PDF文件"""
        try:
            # 尝试使用PyPDF2或其他库
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                
                return ParsedResult(
                    success=True,
                    content={
                        "type": "pdf",
                        "pages": len(reader.pages),
                        "text": text[:5000],
                        "word_count": len(text.split())
                    },
                    confidence=0.9
                )
            except ImportError:
                return ParsedResult(
                    success=True,
                    content={
                        "type": "pdf",
                        "hint": "PDF解析库未安装，使用其他方式处理",
                        "file_path": file_path
                    },
                    confidence=0.6
                )
        except Exception as e:
            return ParsedResult(success=False, error=f"PDF解析失败: {e}")

    async def _parse_document(self, file_path: str) -> ParsedResult:
        """解析文档文件"""
        try:
            try:
                from docx import Document
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                
                return ParsedResult(
                    success=True,
                    content={
                        "type": "docx",
                        "text": text[:5000],
                        "paragraphs": len(doc.paragraphs)
                    },
                    confidence=0.9
                )
            except ImportError:
                return ParsedResult(
                    success=True,
                    content={
                        "type": "docx",
                        "hint": "python-docx未安装",
                        "file_path": file_path
                    },
                    confidence=0.6
                )
        except Exception as e:
            return ParsedResult(success=False, error=f"文档解析失败: {e}")

    async def _parse_spreadsheet(self, file_path: str) -> ParsedResult:
        """解析电子表格"""
        try:
            try:
                import pandas as pd
                df = pd.read_excel(file_path) if file_path.endswith('.xlsx') else pd.read_csv(file_path)
                
                return ParsedResult(
                    success=True,
                    content={
                        "type": "spreadsheet",
                        "rows": len(df),
                        "columns": len(df.columns),
                        "columns_list": df.columns.tolist(),
                        "sample": df.head(5).to_dict('records')
                    },
                    confidence=0.95
                )
            except ImportError:
                return ParsedResult(
                    success=True,
                    content={
                        "type": "spreadsheet",
                        "hint": "pandas未安装",
                        "file_path": file_path
                    },
                    confidence=0.6
                )
        except Exception as e:
            return ParsedResult(success=False, error=f"表格解析失败: {e}")

    async def _parse_image(self, file_path: str) -> ParsedResult:
        """解析图片"""
        try:
            from PIL import Image
            img = Image.open(file_path)
            
            return ParsedResult(
                success=True,
                content={
                    "type": "image",
                    "format": img.format,
                    "size": img.size,
                    "mode": img.mode,
                    "hint": "可进行OCR识别或图像分析"
                },
                confidence=0.85
            )
        except ImportError:
            return ParsedResult(
                success=True,
                content={
                    "type": "image",
                    "hint": "PIL未安装，无法提取元数据",
                    "file_path": file_path
                },
                confidence=0.6
            )
        except Exception as e:
            return ParsedResult(success=False, error=f"图片解析失败: {e}")

    async def _parse_json(self, file_path: str) -> ParsedResult:
        """解析JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            return ParsedResult(
                success=True,
                content={
                    "type": "json",
                    "data": data,
                    "schema": self._infer_schema(data)
                },
                confidence=0.95
            )
        except Exception as e:
            return ParsedResult(success=False, error=f"JSON解析失败: {e}")

    async def _parse_markdown(self, file_path: str) -> ParsedResult:
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return ParsedResult(
                success=True,
                content={
                    "type": "markdown",
                    "content": content,
                    "headings": self._extract_headings(content)
                },
                confidence=0.95
            )
        except Exception as e:
            return ParsedResult(success=False, error=f"Markdown解析失败: {e}")

    def _extract_headings(self, markdown: str) -> List[str]:
        """提取Markdown标题"""
        headings = []
        for line in markdown.split("\n"):
            if line.startswith("#"):
                level = len(line.split()[0])
                text = line.lstrip("#").strip()
                headings.append({"level": level, "text": text})
        return headings


def get_multimodal_processor() -> MultimodalProcessor:
    """获取多模态处理器单例"""
    global _processor_instance
    if _processor_instance is None:
        _processor_instance = MultimodalProcessor()
    return _processor_instance


_processor_instance = None