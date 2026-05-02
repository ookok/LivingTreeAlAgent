"""
Word渲染引擎 - 使用reference.docx模板系统 + 样式学习

核心功能：
1. 加载reference.docx模板
2. 从用户上传的文档中自动学习样式
3. 解析Markdown DSL并映射到学习到的样式
4. 生成格式正确的Word文档
5. 支持样式槽位填充
6. 支持实时样式调整
"""
import os
from pathlib import Path
from typing import List, Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .markdown_dsl_parser import DSLNode, NodeType, StyleType
from .style_learner import StyleLearner, StyleDefinition


class WordRenderer:
    """
    Word渲染引擎
    
    使用reference.docx模板或从用户文档学习到的样式生成工业级格式的Word文档。
    
    核心特性：
    - 支持从用户上传的文档自动学习样式
    - 支持样式融合（学习多个文档后融合成统一样式）
    - 支持实时样式调整
    - 支持增量学习
    """
    
    # 样式名称映射
    STYLE_MAP = {
        StyleType.HEADING_1: 'Heading 1',
        StyleType.HEADING_2: 'Heading 2',
        StyleType.HEADING_3: 'Heading 3',
        StyleType.NORMAL_TEXT: 'Normal',
        StyleType.ENV_TABLE: 'Environment Table',
        StyleType.CODE_BLOCK: 'Code',
    }
    
    def __init__(self, template_path: Optional[str] = None):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx库未安装，请运行: pip install python-docx")
        
        self.template_path = template_path or self._get_default_template()
        self.document = None
        self.styles_cache = {}
        
        # 样式学习器
        self.style_learner = StyleLearner()
        
        # 学习到的样式映射
        self.learned_style_mapping = {}
        
        # 当前样式覆盖（用于实时调整）
        self.style_overrides = {}
    
    def _get_default_template(self) -> str:
        """获取默认模板路径"""
        default_path = Path(__file__).parent / 'templates' / 'reference.docx'
        return str(default_path)
    
    def load_template(self, template_path: Optional[str] = None):
        """加载模板文档"""
        path = template_path or self.template_path
        
        if os.path.exists(path):
            self.document = Document(path)
            self._cache_styles()
            print(f"✅ 已加载模板: {path}")
        else:
            # 创建新文档作为基础模板
            self.document = Document()
            self._create_default_styles()
            print(f"⚠️ 模板不存在，创建新文档: {path}")
    
    def _cache_styles(self):
        """缓存样式"""
        if self.document:
            self.styles_cache = {style.name: style for style in self.document.styles}
    
    def _create_default_styles(self):
        """创建默认工业级样式"""
        if not self.document:
            return
        
        # 获取或创建样式
        styles = self.document.styles
        
        # Heading 1 - 黑体/居中/16pt
        heading1 = styles['Heading 1']
        heading1.font.name = '黑体'
        heading1.font.size = Pt(16)
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading1.paragraph_format.space_before = Pt(12)
        heading1.paragraph_format.space_after = Pt(12)
        
        # Heading 2 - 黑体/14pt
        heading2 = styles['Heading 2']
        heading2.font.name = '黑体'
        heading2.font.size = Pt(14)
        heading2.paragraph_format.space_before = Pt(10)
        heading2.paragraph_format.space_after = Pt(8)
        
        # Heading 3 - 黑体/12pt
        heading3 = styles['Heading 3']
        heading3.font.name = '黑体'
        heading3.font.size = Pt(12)
        heading3.paragraph_format.space_before = Pt(8)
        heading3.paragraph_format.space_after = Pt(6)
        
        # Normal - 仿宋/小四(12pt)
        normal = styles['Normal']
        normal.font.name = '仿宋'
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        
        # Code - 等线/10pt
        if 'Code' not in styles:
            from docx.enum.style import WD_STYLE_TYPE
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = '等线'
            code_style.font.size = Pt(10)
            code_style.paragraph_format.font.name = '等线'
        
        # Environment Table style
        if 'Environment Table' not in styles:
            from docx.enum.style import WD_STYLE_TYPE
            env_table_style = styles.add_style('Environment Table', WD_STYLE_TYPE.TABLE)
    
    def render(self, nodes: List[DSLNode], output_path: str) -> bool:
        """
        渲染文档
        
        Args:
            nodes: DSL节点列表
            output_path: 输出文件路径
        
        Returns:
            是否成功
        """
        if not self.document:
            self.load_template()
        
        try:
            for node in nodes:
                self._render_node(node)
            
            self.document.save(output_path)
            print(f"✅ 文档已保存: {output_path}")
            return True
        
        except Exception as e:
            print(f"❌ 渲染失败: {e}")
            return False
    
    def _render_node(self, node: DSLNode):
        """渲染单个节点"""
        if node.node_type == NodeType.HEADING:
            self._render_heading(node)
        
        elif node.node_type == NodeType.PARAGRAPH:
            self._render_paragraph(node)
        
        elif node.node_type == NodeType.LIST:
            self._render_list(node)
        
        elif node.node_type == NodeType.CODE:
            self._render_code(node)
        
        elif node.node_type == NodeType.TABLE:
            self._render_table(node)
        
        elif node.node_type == NodeType.STYLE:
            # 样式指令，更新当前样式状态
            pass
    
    def _render_heading(self, node: DSLNode):
        """渲染标题"""
        style_name = self.STYLE_MAP.get(node.style, 'Heading 1')
        
        # 检查样式是否存在
        if style_name not in self.styles_cache:
            style_name = 'Heading 1'
        
        self.document.add_heading(node.content, level=self._get_heading_level(node))
    
    def _get_heading_level(self, node: DSLNode) -> int:
        """获取标题级别"""
        mapping = {
            StyleType.HEADING_1: 1,
            StyleType.HEADING_2: 2,
            StyleType.HEADING_3: 3,
        }
        return mapping.get(node.style, 1)
    
    def _render_paragraph(self, node: DSLNode):
        """渲染段落"""
        style_name = self.STYLE_MAP.get(node.style, 'Normal')
        
        if style_name not in self.styles_cache:
            style_name = 'Normal'
        
        paragraph = self.document.add_paragraph(node.content)
        paragraph.style = self.document.styles[style_name]
    
    def _render_list(self, node: DSLNode):
        """渲染列表"""
        items = node.metadata.get('items', [])
        
        for item in items:
            paragraph = self.document.add_paragraph(item, style='List Bullet')
    
    def _render_code(self, node: DSLNode):
        """渲染代码块"""
        style_name = self.STYLE_MAP.get(node.style, 'Code')
        
        if style_name not in self.styles_cache:
            style_name = 'Normal'
        
        paragraph = self.document.add_paragraph(node.content)
        paragraph.style = self.document.styles[style_name]
    
    def _render_table(self, node: DSLNode):
        """渲染表格"""
        headers = node.metadata.get('headers', [])
        rows = node.metadata.get('rows', [])
        
        if not headers:
            return
        
        table = self.document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Grid Table 4 Accent 1'  # 使用内置样式
        
        # 设置表格对齐
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 填充表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # 填充数据行
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell in enumerate(row):
                if j < len(row_cells):
                    row_cells[j].text = str(cell)
                    row_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def generate_template(self, output_path: str):
        """
        生成reference.docx模板文件
        
        Args:
            output_path: 输出路径
        """
        self.document = Document()
        self._create_default_styles()
        
        # 添加说明内容
        self.document.add_heading('工业文档模板', level=1)
        self.document.add_paragraph('这是一个工业级文档模板，包含以下样式：')
        
        self.document.add_heading('样式说明', level=2)
        styles = [
            ('Heading 1', '黑体/居中/16pt - 用于章节标题'),
            ('Heading 2', '黑体/14pt - 用于小节标题'),
            ('Heading 3', '黑体/12pt - 用于子节标题'),
            ('Normal', '仿宋/小四/1.5倍行距 - 用于正文'),
            ('Code', '等线/10pt - 用于代码块'),
            ('Environment Table', '三线表样式 - 用于环境数据表格'),
        ]
        
        for style_name, description in styles:
            self.document.add_paragraph(f'• {style_name}: {description}')
        
        self.document.save(output_path)
        print(f"✅ 模板已生成: {output_path}")
    
    # ==================== 样式学习与调整方法 ====================
    
    def learn_styles_from_document(self, docx_path: str, document_type: str = "unknown") -> Dict[str, Any]:
        """
        从用户上传的文档学习样式
        
        Args:
            docx_path: Word文档路径
            document_type: 文档类型
        
        Returns:
            学习结果
        """
        result = self.style_learner.learn_from_document(docx_path, document_type)
        
        if result['success']:
            # 更新样式映射
            self.learned_style_mapping = self.style_learner.get_style_mapping()
            print(f"✅ 已从文档学习到 {len(self.learned_style_mapping)} 个样式")
        
        return result
    
    def learn_styles_from_multiple(self, doc_paths: List[str], document_type: str = "unknown") -> Dict[str, Any]:
        """
        从多个文档学习样式并融合
        
        Args:
            doc_paths: Word文档路径列表
            document_type: 文档类型
        
        Returns:
            学习结果
        """
        result = self.style_learner.learn_from_multiple_documents(doc_paths, document_type)
        
        if result['success']:
            # 融合样式
            self.style_learner.fuse_styles()
            self.learned_style_mapping = self.style_learner.get_style_mapping()
            print(f"✅ 已从 {len(doc_paths)} 个文档学习并融合样式")
        
        return result
    
    def apply_learned_styles(self):
        """应用学习到的样式到当前文档"""
        if not self.document:
            self.load_template()
        
        # 根据学习到的样式创建文档样式
        for style_def in self.style_learner.paragraph_styles.values():
            self._apply_style_definition(style_def)
    
    def _apply_style_definition(self, style_def: StyleDefinition):
        """应用样式定义"""
        if not self.document:
            return
        
        styles = self.document.styles
        
        # 检查样式是否已存在
        if style_def.name not in styles:
            # 创建新样式
            style = styles.add_style(style_def.name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_def.name]
        
        # 设置字体属性
        style.font.name = style_def.font_name
        style.font.size = Pt(style_def.font_size)
        style.font.bold = style_def.font_bold
        style.font.italic = style_def.font_italic
        
        # 设置对齐方式
        align_map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
        }
        style.paragraph_format.alignment = align_map.get(style_def.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # 设置段落间距
        style.paragraph_format.line_spacing = style_def.line_spacing
        style.paragraph_format.space_before = Pt(style_def.space_before)
        style.paragraph_format.space_after = Pt(style_def.space_after)
        
        # 更新缓存
        self.styles_cache[style_def.name] = style
    
    def set_style_override(self, style_type: StyleType, **kwargs):
        """
        设置样式覆盖（用于实时调整）
        
        Args:
            style_type: 样式类型
            kwargs: 样式属性（font_name, font_size, font_bold等）
        """
        self.style_overrides[style_type] = kwargs
        print(f"✅ 已设置样式覆盖: {style_type.value} -> {kwargs}")
    
    def clear_style_overrides(self):
        """清除所有样式覆盖"""
        self.style_overrides.clear()
        print("✅ 已清除所有样式覆盖")
    
    def _get_effective_style(self, style_type: StyleType) -> Dict[str, Any]:
        """获取有效样式（考虑覆盖）"""
        # 基础样式
        base_style = {
            'font_name': '仿宋',
            'font_size': 12,
            'font_bold': False,
            'alignment': 'left',
            'line_spacing': 1.5,
        }
        
        # 根据样式类型调整
        if style_type == StyleType.HEADING_1:
            base_style.update({
                'font_name': '黑体',
                'font_size': 16,
                'font_bold': True,
                'alignment': 'center',
            })
        elif style_type == StyleType.HEADING_2:
            base_style.update({
                'font_name': '黑体',
                'font_size': 14,
                'font_bold': True,
            })
        elif style_type == StyleType.HEADING_3:
            base_style.update({
                'font_name': '黑体',
                'font_size': 12,
                'font_bold': True,
            })
        
        # 应用学习到的样式
        if self.learned_style_mapping:
            mapped_name = self.learned_style_mapping.get(self.STYLE_MAP.get(style_type))
            if mapped_name and mapped_name in self.style_learner.paragraph_styles:
                learned = self.style_learner.paragraph_styles[mapped_name]
                base_style.update({
                    'font_name': learned.font_name or base_style['font_name'],
                    'font_size': learned.font_size or base_style['font_size'],
                    'font_bold': learned.font_bold,
                    'alignment': learned.alignment,
                    'line_spacing': learned.line_spacing,
                })
        
        # 应用覆盖
        if style_type in self.style_overrides:
            base_style.update(self.style_overrides[style_type])
        
        return base_style
    
    def export_learned_styles(self, output_path: str):
        """导出学习到的样式"""
        self.style_learner.export_styles(output_path)
    
    def import_learned_styles(self, input_path: str):
        """导入学习到的样式"""
        self.style_learner.import_styles(input_path)
        self.learned_style_mapping = self.style_learner.get_style_mapping()
    
    def get_learned_styles_summary(self) -> Dict[str, Any]:
        """获取学习到的样式摘要"""
        return {
            'paragraph_styles': list(self.style_learner.paragraph_styles.keys()),
            'table_styles': list(self.style_learner.table_styles.keys()),
            'character_styles': list(self.style_learner.character_styles.keys()),
            'style_mapping': self.learned_style_mapping,
            'learning_history_count': len(self.style_learner.learning_history)
        }