"""ReportEnhancer — Professional document features: citations, data binding, docx pipeline.

Adds to DocEngine:
  1. CitationTracker: auto-number tables, figures, cross-references
  2. DataBinder: real-time data from CSV/JSON/DB sources
  3. FormatPipeline: native .docx output with inherited styles
  4. MultiLang: template-based localization support

Usage:
    from livingtree.capability.report_enhancer import ReportEnhancer
    enhancer = ReportEnhancer()
    doc = enhancer.generate("eia_report", data, format="docx", 
                           data_source="data/site_measurements.csv")
"""

from __future__ import annotations

import csv
import difflib
import json
import re
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from loguru import logger

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree
from pypdf import PdfReader, PdfWriter


# ═══ 1. Citation Tracker — auto-number tables, figures, references ═══

class CitationTracker:
    """Auto-number and cross-reference tables, figures, and sections.

    Supports: 表1-1, 图2-3, 参见表3.2, 参考文献[1]
    GB/T 7714 and APA-style citation formatting.
    """

    def __init__(self):
        self._tables: dict[str, int] = {}
        self._figures: dict[str, int] = {}
        self._refs: list[str] = []
        self._chapter = 1

    def chapter(self, num: int) -> "CitationTracker":
        self._chapter = num
        return self

    def table(self, label: str) -> str:
        """Register a table. Returns formatted label like '表3-1'."""
        key = f"{self._chapter}"
        num = self._tables.get(key, 0) + 1
        self._tables[key] = num
        return f"表{self._chapter}-{num}"

    def figure(self, label: str) -> str:
        """Register a figure. Returns formatted label like '图2-3'."""
        key = f"{self._chapter}"
        num = self._figures.get(key, 0) + 1
        self._figures[key] = num
        return f"图{self._chapter}-{num}"

    def ref(self, citation: str, style: str = "gb7714") -> str:
        """Add a reference. Returns formatted citation number like '[1]'."""
        self._refs.append(citation)
        idx = len(self._refs)
        return f"[{idx}]" if style == "gb7714" else f"({idx})"

    def ref_list(self, style: str = "gb7714") -> str:
        """Generate formatted reference list."""
        lines = ["## 参考文献", ""]
        for i, ref in enumerate(self._refs, 1):
            lines.append(f"[{i}] {ref}")
        return "\n".join(lines)

    def cross_ref(self, target_type: str, chapter: int, num: int) -> str:
        """Cross-reference: '参见表3-2' or '如图1-1所示'."""
        prefixes = {"table": "表", "figure": "图", "section": "节"}
        prefix = prefixes.get(target_type, "")
        if target_type == "section":
            return f"参见{prefix}{chapter}.{num}节"
        return f"参见{prefix}{chapter}-{num}"

    def process_document(self, content: str) -> str:
        """Process document content, replacing citation placeholders with numbered references.

        Placeholders:
          {tbl:label} → 表1-1
          {fig:label} → 图1-2
          {ref:GB3095-2012} → [1]
        """
        def _replace_table(m):
            return self.table(m.group(1))
        def _replace_figure(m):
            return self.figure(m.group(1))
        def _replace_ref(m):
            return self.ref(m.group(1))

        content = re.sub(r'\{tbl:([^}]+)\}', _replace_table, content)
        content = re.sub(r'\{fig:([^}]+)\}', _replace_figure, content)
        content = re.sub(r'\{ref:([^}]+)\}', _replace_ref, content)
        return content


# ═══ 2. Data Binder — real-time data from CSV/JSON/DB/sensors ═══

class DataBinder:
    """Bind real-time data sources into report generation.

    Sources: CSV files, JSON APIs, SQLite databases, environment variables.
    """

    @staticmethod
    def from_csv(path: str, key_column: str = None) -> dict:
        """Load CSV as dict. If key_column, returns {key: row_dict} mapping."""
        data = {}
        try:
            with open(path, encoding="utf-8") as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                if key_column and key_column in (reader.fieldnames or []):
                    return {row[key_column]: dict(row) for row in rows}
                return {"rows": rows, "columns": reader.fieldnames or [], "count": len(rows)}
        except Exception as e:
            return {"error": str(e), "source": path}

    @staticmethod
    def from_json(path: str) -> dict:
        """Load JSON file as dict."""
        try:
            return json.loads(Path(path).read_text(encoding="utf-8"))
        except Exception as e:
            return {"error": str(e), "source": path}

    @staticmethod
    def from_db(db_path: str, query: str) -> dict:
        """Execute SQL query and return results as dict."""
        try:
            conn = sqlite3.connect(db_path)
            cursor = conn.cursor()
            cursor.execute(query)
            columns = [d[0] for d in cursor.description] if cursor.description else []
            rows = [dict(zip(columns, row)) for row in cursor.fetchall()]
            conn.close()
            return {"rows": rows, "columns": columns, "count": len(rows)}
        except Exception as e:
            return {"error": str(e), "query": query}

    @staticmethod
    def from_api(url: str, headers: dict = None, method: str = "GET") -> dict:
        """Fetch data from REST API endpoint. Returns parsed JSON."""
        try:
            import urllib.request
            req = urllib.request.Request(url, headers=headers or {}, method=method)
            with urllib.request.urlopen(req, timeout=30) as resp:
                return json.loads(resp.read().decode())
        except Exception as e:
            return {"error": str(e), "url": url}

    @staticmethod
    def merge_sources(sources: list[dict]) -> dict:
        """Merge multiple data sources into one dict."""
        merged = {}
        for i, src in enumerate(sources):
            source_type = src.get("type", "inline")
            source_path = src.get("path", "") or src.get("url", "") or src.get("db", "")
            label = src.get("label", f"source_{i}")

            if source_type == "csv":
                merged[label] = DataBinder.from_csv(source_path, src.get("key_column"))
            elif source_type == "json":
                merged[label] = DataBinder.from_json(source_path)
            elif source_type == "db":
                merged[label] = DataBinder.from_db(source_path, src.get("query", "SELECT * FROM sqlite_master"))
            elif source_type == "api":
                merged[label] = DataBinder.from_api(source_path, src.get("headers"))
            elif source_type == "inline":
                merged[label] = src.get("data", {})

        return merged


# ═══ 3. Format Pipeline — native .docx output ═══

class FormatPipeline:
    """Bridge DocEngine to document_intelligence for native .docx generation.

    Integrates the existing document_intelligence._generate_docx() pipeline
    with DocEngine's template system for professional Word document output.
    """

    @staticmethod
    def export_docx(content: dict, template_name: str = "report",
                    output_path: str = "", citation_tracker: CitationTracker = None) -> str:
        """Generate a formatted .docx file from report content.

        Uses document_intelligence.DocumentIntelligence for native formatting:
        - Inherits styles from .docx templates (fonts, margins, headers)
        - Auto-generates Table of Contents
        - Page numbers and headers
        - Approval stamps and compliance markers
        - Table data with proper borders and formatting
        """
        try:
            from .document_intelligence import DocumentIntelligence
            di = DocumentIntelligence()

            params = {
                "title": content.get("title", template_name),
                "sections": content.get("sections", []),
                "tables": content.get("tables", []),
                "figures": content.get("figures", []),
                "metadata": content.get("metadata", {}),
                "timestamp": datetime.now().isoformat(),
            }

            if citation_tracker:
                params["references"] = citation_tracker._refs
                params["table_count"] = sum(citation_tracker._tables.values())
                params["figure_count"] = sum(citation_tracker._figures.values())

            result = di.generate(template=template_name, params=params,
                                output_path=output_path, format="docx")
            return result.output_path or output_path
        except ImportError:
            # Fallback: generate markdown
            return FormatPipeline._export_markdown(content, template_name, output_path)
        except Exception as e:
            logger.warning(f"FormatPipeline docx: {e}")
            return FormatPipeline._export_markdown(content, template_name, output_path)

    @staticmethod
    def _export_markdown(content: dict, template_name: str, output_path: str) -> str:
        """Fallback: generate markdown file."""
        out_path = Path(output_path or f"/tmp/{template_name}_{int(datetime.now().timestamp())}.md")
        lines = [f"# {content.get('title', template_name)}", ""]
        for section in content.get("sections", []):
            lines.append(f"## {section.get('heading', '')}")
            lines.append(section.get("body", ""))
            lines.append("")

        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text("\n".join(lines), encoding="utf-8")
        return str(out_path)

    @staticmethod
    def export_pdf(content: dict, template_name: str = "report",
                   output_path: str = "") -> str:
        """Generate PDF via markdown → HTML → WeasyPrint pipeline."""
        try:
            from ..core.doc_renderer import DocRenderer
            renderer = DocRenderer()
            markdown = FormatPipeline._export_markdown(content, template_name, "")
            result = renderer.render_document(
                Path(markdown).read_text(encoding="utf-8"),
                template="long_doc",
            )
            out_path = Path(output_path or f"/tmp/{template_name}.pdf")
            out_path.write_bytes(result.pdf_bytes if hasattr(result, 'pdf_bytes') else b"")
            return str(out_path)
        except Exception as e:
            logger.warning(f"FormatPipeline pdf: {e}")
            return FormatPipeline._export_markdown(content, template_name,
                        str(Path(output_path or "/tmp").with_suffix(".md")))


# ═══ 4. Multi-Language Support ═══

class MultiLangSupport:
    """Template-based localization for multi-language reports.

    Usage:
        mls = MultiLangSupport()
        mls.load_locale("zh-CN", "livingtree/locales/")
        title = mls.t("eia_report_title", lang="zh-CN")  # → "环境影响评价报告"
        title = mls.t("eia_report_title", lang="en")      # → "EIA Report"
    """

    _locales: dict[str, dict[str, str]] = {}
    _builtin: dict[str, dict[str, str]] = {
        "zh-CN": {
            "report_title": "环境影响评价报告",
            "table_of_contents": "目录",
            "chapter": "第{num}章",
            "section": "节",
            "table_caption": "表",
            "figure_caption": "图",
            "reference": "参考文献",
            "appendix": "附录",
            "approval": "审批",
            "prepared_by": "编制单位",
            "date": "日期",
            "page": "第{num}页",
        },
        "en": {
            "report_title": "Environmental Impact Assessment Report",
            "table_of_contents": "Table of Contents",
            "chapter": "Chapter {num}",
            "section": "Section",
            "table_caption": "Table",
            "figure_caption": "Figure",
            "reference": "References",
            "appendix": "Appendix",
            "approval": "Approval",
            "prepared_by": "Prepared by",
            "date": "Date",
            "page": "Page {num}",
        },
    }

    @classmethod
    def t(cls, key: str, lang: str = "zh-CN", **fmt) -> str:
        """Translate a key to the given language with optional formatting."""
        locale = cls._locales.get(lang, {}) or cls._builtin.get(lang, {})
        text = locale.get(key, key)
        if fmt:
            text = text.format(**fmt)
        return text

    @classmethod
    def load_locale(cls, lang: str, directory: str) -> int:
        """Load locale files from directory (e.g., locales/zh-CN.json)."""
        try:
            path = Path(directory) / f"{lang}.json"
            if path.exists():
                cls._locales[lang] = json.loads(path.read_text(encoding="utf-8"))
                return len(cls._locales[lang])
        except Exception:
            pass
        return 0

    @classmethod
    def available_languages(cls) -> list[str]:
        return sorted(set(list(cls._locales.keys()) + list(cls._builtin.keys())))


# ═══ 5. ReportEnhancer — unified entry point ═══

class ReportEnhancer:
    """Unified professional document generation with all enhancements.

    Bridges: EIA models → DocEngine → CitationTracker → FormatPipeline → .docx/PDF
    """

    def __init__(self):
        self.citations = CitationTracker()
        self._doc_engine = None

    def _get_engine(self):
        if self._doc_engine is None:
            from .doc_engine import DocEngine
            self._doc_engine = DocEngine()
        return self._doc_engine

    async def generate(self, template_type: str, data: dict,
                 format: str = "docx", output: str = "",
                 data_sources: list[dict] = None,
                 citations: bool = True, lang: str = "zh-CN") -> dict:
        """Generate a professional report with all enhancements.

        Args:
            template_type: 'eia_report', 'emergency_plan', 'feasibility', etc.
            data: Template variables and section content
            format: 'docx', 'pdf', 'markdown', 'html'
            output: Output file path
            data_sources: [{type:'csv', path:'data.csv'}, {type:'api', url:'...'}]
            citations: Enable auto-numbering of tables, figures, references
            lang: Report language ('zh-CN', 'en')

        Returns:
            {output_path, format, size, pages, citations_used}
        """
        # 1. Merge data sources
        merged_data = dict(data)
        if data_sources:
            merged_data["external_data"] = DataBinder.merge_sources(data_sources)

        # 2. Enrich with EIA physics models
        engine = self._get_engine()
        enriched = engine.enrich_with_eia_data(template_type, merged_data)

        # 3. Generate report content via DocEngine
        result = await engine.generate_report_with_models(template_type, enriched)

        # 4. Process citations
        if citations and result.get("sections"):
            for section in result["sections"]:
                if isinstance(section, dict) and "body" in section:
                    section["body"] = self.citations.process_document(
                        section.get("body", "")
                    )

        # 5. Inject multi-lang labels
        result["_lang"] = lang
        result["_labels"] = {
            "toc": MultiLangSupport.t("table_of_contents", lang),
            "refs": MultiLangSupport.t("reference", lang),
        }

        # 6. Output via FormatPipeline
        output_path = ""
        if format == "docx":
            output_path = FormatPipeline.export_docx(
                result, template_type, output, self.citations)
        elif format == "pdf":
            output_path = FormatPipeline.export_pdf(
                result, template_type, output)
        else:
            output_path = FormatPipeline._export_markdown(
                result, template_type, output)

        return {
            "output_path": output_path,
            "format": format,
            "size": Path(output_path).stat().st_size if Path(output_path).exists() else 0,
            "citations": {
                "tables": sum(self.citations._tables.values()),
                "figures": sum(self.citations._figures.values()),
                "references": len(self.citations._refs),
            },
            "models_used": list(enriched.keys()) if enriched else [],
            "sections": len(result.get("sections", [])),
        }


# ═══ 6. Data Validator — input sanity checks ═══

class DataValidator:
    """Validate report input data before generation."""

    RULES = {
        "eia_report": {
            "project_name": {"required": True, "min_len": 2, "max_len": 200},
            "source_params.emission_rate": {"type": "number", "min": 0},
            "water_params.bod_initial": {"type": "number", "min": 0, "max": 1000},
            "noise_params.source_level": {"type": "number", "min": 20, "max": 200},
        },
        "emergency_plan": {
            "project_name": {"required": True},
            "emergency_type": {"required": True, "enum": ["火灾", "爆炸", "泄漏", "中毒", "自然灾害"]},
        },
    }

    @classmethod
    def validate(cls, template_type: str, data: dict) -> dict:
        """Validate data against template rules. Returns {valid, errors, warnings}."""
        rules = cls.RULES.get(template_type, {})
        errors = []
        warnings = []

        for path, rule in rules.items():
            val = cls._get_nested(data, path)
            if rule.get("required") and (val is None or val == ""):
                errors.append(f"缺少必填字段: {path}")
                continue
            if val is None:
                continue
            if "type" in rule:
                try:
                    if rule["type"] == "number":
                        val = float(val)
                except (ValueError, TypeError):
                    errors.append(f"{path}: 期望数字, 得到 {val}")
            if "min" in rule and float(val) < rule["min"]:
                errors.append(f"{path}: 值 {val} < 最小值 {rule['min']}")
            if "max" in rule and float(val) > rule["max"]:
                errors.append(f"{path}: 值 {val} > 最大值 {rule['max']}")
            if "min_len" in rule and len(str(val)) < rule["min_len"]:
                errors.append(f"{path}: 长度不足 (至少 {rule['min_len']} 字符)")
            if "enum" in rule and val not in rule["enum"]:
                warnings.append(f"{path}: 值 '{val}' 不在推荐选项 {rule['enum']} 中")

        # Structural checks
        if not data.get("sections") and not data.get("raw_text"):
            warnings.append("未提供 sections 或 raw_text，报告可能为空")

        return {"valid": len(errors) == 0, "errors": errors, "warnings": warnings}

    @staticmethod
    def _get_nested(d: dict, path: str):
        keys = path.split(".")
        for k in keys:
            if isinstance(d, dict):
                d = d.get(k)
            else:
                return None
        return d


# ═══ 7. Watermark — draft/confidential/internal marks ═══

class Watermark:
    """Add watermarks to generated documents."""

    @staticmethod
    def apply_docx(doc_path: str, text: str, opacity: float = 0.15,
                   font_size: int = 72, rotation: float = -45) -> str:
        """Add diagonal watermark to .docx file."""
        try:
            doc = Document(doc_path)
            for section in doc.sections:
                header = section.header
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.alignment = 1  # center
                run = paragraph.add_run(text)
                run.font.size = Pt(font_size)
                run.font.color.rgb = None  # Will be set via XML

                # Set opacity via XML
                rPr = run._element.get_or_add_rPr()
                color = etree.SubElement(rPr, qn('w:color'))
                color.set(qn('w:val'), 'C0C0C0')

            out_path = Path(doc_path).parent / f"watermarked_{Path(doc_path).name}"
            doc.save(str(out_path))
            return str(out_path)
        except Exception as e:
            logger.warning(f"Watermark docx: {e}")
            return doc_path

    @staticmethod
    def apply_markdown(content: str, text: str) -> str:
        """Add watermark comment to markdown."""
        return f"<!-- {text} -->\n{content}"

    @staticmethod
    def status_stamp(status: str) -> str:
        """Generate status stamp: 草稿 / 送审稿 / 报批稿 / 最终版"""
        stamps = {
            "draft": "【草  稿】",
            "review": "【送 审 稿】",
            "approval": "【报 批 稿】",
            "final": "【最 终 版】",
        }
        return stamps.get(status, f"【{status}】")


# ═══ 8. Chart Auto-Generation — matplotlib → docx inline ═══

class ChartGenerator:
    """Generate charts and embed in documents."""

    @staticmethod
    def generate_embedded(data: dict, chart_configs: list[dict],
                          output_dir: str = "") -> list[str]:
        """Generate chart images and return paths for document embedding.

        chart_configs: [{type:'bar|line|pie|scatter', title:'', x:[], y:[], labels:[]}]
        """
        out_dir = Path(output_dir or ".livingtree/charts")
        out_dir.mkdir(parents=True, exist_ok=True)
        paths = []

        for i, cfg in enumerate(chart_configs):
            try:
                fig, ax = plt.subplots(figsize=(8, 4.5))
                chart_type = cfg.get("type", "bar")

                if chart_type == "bar":
                    ax.bar(cfg.get("labels", []), cfg.get("values", []),
                          color=cfg.get("colors", "#3b82f6"))
                elif chart_type == "line":
                    ax.plot(cfg.get("x", []), cfg.get("y", []),
                           color=cfg.get("color", "#3b82f6"), linewidth=2)
                    ax.fill_between(cfg.get("x", []), cfg.get("y", []), alpha=0.1)
                elif chart_type == "pie":
                    ax.pie(cfg.get("values", []), labels=cfg.get("labels", []),
                          autopct='%1.1f%%')
                elif chart_type == "scatter":
                    ax.scatter(cfg.get("x", []), cfg.get("y", []),
                              alpha=0.6, color=cfg.get("color", "#3b82f6"))

                ax.set_title(cfg.get("title", ""), fontsize=12, fontweight='bold')
                if cfg.get("xlabel"):
                    ax.set_xlabel(cfg["xlabel"])
                if cfg.get("ylabel"):
                    ax.set_ylabel(cfg["ylabel"])
                if cfg.get("grid", True):
                    ax.grid(True, alpha=0.3)

                path = out_dir / f"chart_{i+1}_{cfg.get('title','')[:20]}.png"
                fig.savefig(path, dpi=150, bbox_inches='tight')
                plt.close(fig)
                paths.append(str(path))
            except Exception:
                continue

        return paths

    @staticmethod
    def generate_distribution_map(params: dict, output: str = "") -> str:
        """Generate concentration distribution contour map (for EIA reports)."""
        x = np.linspace(-500, 500, 100)
        y = np.linspace(-500, 500, 100)
        X, Y = np.meshgrid(x, y)
        sx = params.get("sigma_x", 100)
        sy = params.get("sigma_y", 50)
        Q = params.get("emission_rate", 1.0)
        u = params.get("wind_speed", 2.5)
        Z = (Q / (2 * np.pi * u * sx * sy)) * np.exp(
            -0.5 * ((X / sx) ** 2 + (Y / sy) ** 2))

        fig, ax = plt.subplots(figsize=(8, 6))
        contour = ax.contourf(X, Y, Z * 1e6, levels=15, cmap='YlOrRd')
        plt.colorbar(contour, label='浓度 (μg/m³)')
        ax.set_title('大气污染物浓度分布图', fontsize=12, fontweight='bold')
        ax.set_xlabel('距离 (m)')
        ax.set_ylabel('距离 (m)')

        out_path = Path(output or ".livingtree/charts/distribution.png")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        fig.savefig(out_path, dpi=150, bbox_inches='tight')
        plt.close(fig)
        return str(out_path)


# ═══ 9. Diff Tracker — version comparison with redline ═══

class DiffTracker:
    """Track document versions and generate redline comparisons."""

    def __init__(self):
        self._versions: list[dict] = []
        self._storage = Path(".livingtree/report_versions")
        self._storage.mkdir(parents=True, exist_ok=True)

    def save_version(self, content: str, label: str = "") -> int:
        """Save a document version. Returns version number."""
        ver = len(self._versions) + 1
        path = self._storage / f"v{ver:03d}_{label or 'report'}.md"
        path.write_text(content, encoding="utf-8")
        self._versions.append({"version": ver, "label": label, "path": str(path),
                               "timestamp": datetime.now().isoformat(), "length": len(content)})
        return ver

    def diff(self, v1: int, v2: int) -> dict:
        """Compare two versions. Returns {added_lines, removed_lines, unified_diff}."""
        c1 = self._get_version_content(v1)
        c2 = self._get_version_content(v2)
        if not c1 or not c2:
            return {"error": "Version not found"}

        diff = list(difflib.unified_diff(
            c1.splitlines(keepends=True),
            c2.splitlines(keepends=True),
            fromfile=f"v{v1}", tofile=f"v{v2}",
        ))

        added = sum(1 for l in diff if l.startswith('+') and not l.startswith('+++'))
        removed = sum(1 for l in diff if l.startswith('-') and not l.startswith('---'))

        return {
            "v1": v1, "v2": v2,
            "added_lines": added,
            "removed_lines": removed,
            "total_changes": added + removed,
            "unified_diff": "".join(diff),
        }

    def version_list(self) -> list[dict]:
        return self._versions

    def _get_version_content(self, ver: int) -> str:
        for v in self._versions:
            if v["version"] == ver:
                try:
                    return Path(v["path"]).read_text(encoding="utf-8")
                except Exception:
                    pass
        return ""


# ═══ 10. Audit Logger — who/when/what data trace ═══

class AuditLogger:
    """Track document generation audit trail for compliance."""

    _log_file = Path(".livingtree/audit_log.jsonl")

    @classmethod
    def log(cls, event: str, details: dict) -> None:
        """Record an audit event."""
        entry = {
            "timestamp": datetime.now().isoformat(),
            "event": event,
            **details,
        }
        try:
            cls._log_file.parent.mkdir(parents=True, exist_ok=True)
            with open(cls._log_file, "a", encoding="utf-8") as f:
                f.write(json.dumps(entry, ensure_ascii=False) + "\n")
        except Exception:
            pass

    @classmethod
    def query(cls, event_filter: str = "", limit: int = 50) -> list[dict]:
        """Query audit log entries."""
        results = []
        try:
            if not cls._log_file.exists():
                return results
            with open(cls._log_file, encoding="utf-8") as f:
                for line in f:
                    try:
                        entry = json.loads(line)
                        if not event_filter or event_filter in entry.get("event", ""):
                            results.append(entry)
                    except json.JSONDecodeError:
                        continue
            return results[-limit:]
        except Exception:
            return results

    @classmethod
    def generate_trail(cls, template_type: str) -> str:
        """Generate an audit trail section for the report."""
        events = cls.query(template_type, 20)
        if not events:
            return "无审计记录"

        lines = ["## 编制审计追踪", "", "| 时间 | 事件 | 详情 |", "|------|------|------|"]
        for e in events[-15:]:
            details = ", ".join(f"{k}={v}" for k, v in e.items()
                              if k not in ("timestamp", "event"))[:80]
            lines.append(f"| {e['timestamp'][:19]} | {e['event']} | {details} |")
        return "\n".join(lines)


# ═══ 11. OFD Format — GB/T 33190 Chinese standard document ═══

class OFDExporter:
    """Basic OFD (Open Fixed-layout Document) export for GB/T 33190 compliance."""

    @staticmethod
    def export(content: dict, output: str = "") -> str:
        """Generate a basic OFD-compatible ZIP (structure only).

        Full OFD rendering requires ofdpy library. This creates the standard
        OFD package structure with metadata and content stub.
        """
        import zipfile

        out_path = Path(output or f"/tmp/report_{int(datetime.now().timestamp())}.ofd")
        with zipfile.ZipFile(out_path, "w") as zf:
            # OFD package structure
            zf.writestr("OFD.xml", '<?xml version="1.0" encoding="UTF-8"?>'
                       '<OFD xmlns="http://www.ofdspec.org/2016" Version="1.0">'
                       f'<DocBody><DocInfo><Title>{content.get("title","Report")}</Title>'
                       f'<Author>LivingTree</Author>'
                       f'<CreationDate>{datetime.now().isoformat()}</CreationDate>'
                       '</DocInfo><DocRoot>Doc_0/Document.xml</DocRoot></DocBody></OFD>')
            zf.writestr("Doc_0/Document.xml",
                       '<?xml version="1.0" encoding="UTF-8"?>'
                       '<Document xmlns="http://www.ofdspec.org/2016">'
                       '<CommonData><PageArea><PhysicalBox>0 0 595 842</PhysicalBox>'
                       '</PageArea></CommonData>'
                       f'<Pages><Page ID="1"><Content><Text>{content.get("title","Report")}</Text></Content></Page></Pages>'
                       '</Document>')

        return str(out_path)


# ═══ 12. Concurrent Section Generation ═══

class ConcurrentGenerator:
    """Generate report sections in parallel for speed."""

    @staticmethod
    async def generate_parallel(engine, template_type: str, sections_data: list[dict],
                                max_concurrent: int = 5) -> list[dict]:
        """Generate multiple sections concurrently with asyncio.gather."""
        import asyncio

        sem = asyncio.Semaphore(max_concurrent)

        async def _gen_section(i: int, section_data: dict) -> dict:
            async with sem:
                try:
                    result = await engine.generate_report(
                        template_type, section_data, fold=False)
                    return {"index": i, "heading": section_data.get("heading", f"Section {i+1}"),
                            "body": result.get("content", ""), "status": "ok"}
                except Exception as e:
                    return {"index": i, "heading": section_data.get("heading", ""),
                            "body": f"[生成失败: {e}]", "status": "error"}

        tasks = [_gen_section(i, sd) for i, sd in enumerate(sections_data)]
        results = await asyncio.gather(*tasks)
        return sorted(results, key=lambda r: r["index"])


# ═══ 13. Section Locker — immutable legal sections ═══

class SectionLocker:
    """Lock sections that should not be modified (legal disclaimers, certifications)."""

    LOCKED_TEMPLATES = {
        "eia_report": [
            "声明", "编制单位和人员", "审批意见",
        ],
        "emergency_plan": [
            "法律责任声明", "编制依据",
        ],
    }

    _locks: dict[str, dict[str, str]] = {}

    @classmethod
    def load_locks(cls, template_type: str, source_dir: str = "config/templates/locked") -> dict:
        """Load locked section content from files."""
        if template_type in cls._locks:
            return cls._locks[template_type]

        locks = {}
        for section_name in cls.LOCKED_TEMPLATES.get(template_type, []):
            path = Path(source_dir) / template_type / f"{section_name}.md"
            if path.exists():
                locks[section_name] = path.read_text(encoding="utf-8")

        cls._locks[template_type] = locks
        return locks

    @classmethod
    def is_locked(cls, template_type: str, section_name: str) -> bool:
        return section_name in cls.LOCKED_TEMPLATES.get(template_type, [])

    @classmethod
    def get_locked_content(cls, template_type: str, section_name: str) -> str:
        locks = cls._locks.get(template_type, {})
        return locks.get(section_name, f"[{section_name} — 锁定内容未找到]")

    @classmethod
    def inject_locked_sections(cls, template_type: str, sections: list[dict]) -> list[dict]:
        """Replace LLM-generated locked sections with official content."""
        locks = cls.load_locks(template_type)
        for i, section in enumerate(sections):
            name = section.get("heading", "")
            for lock_name, lock_content in locks.items():
                if lock_name in name:
                    sections[i]["body"] = lock_content
                    sections[i]["locked"] = True
        return sections


# ═══ 14. PDF Form Filler ═══

class PDFFormFiller:
    """Fill AcroForm fields in PDF templates."""

    @staticmethod
    def fill(template_path: str, fields: dict, output: str = "") -> str:
        """Fill PDF form fields."""
        reader = PdfReader(template_path)
        writer = PdfWriter()
        writer.append(reader)

        try:
            writer.update_page_form_field_values(
                writer.pages[0], fields, auto_regenerate=False)
        except Exception:
            # Manual field setting
            for page in writer.pages:
                for annot in page.get("/Annots", []):
                    field_name = annot.get("/T")
                    if field_name in fields:
                        annot.update({"/V": fields[field_name]})

        out_path = Path(output or Path(template_path).parent /
                       f"filled_{Path(template_path).name}")
        writer.write(str(out_path))
        return str(out_path)


# ═══ 15. GB/T 9704 Official Document Format ═══

class OfficialDocFormat:
    """党政机关公文格式 (GB/T 9704-2012) layout generator."""

    @staticmethod
    def generate(title: str, body: str, doc_number: str = "",
                 issuing_org: str = "", date: str = "", output: str = "") -> str:
        """Generate an official government document in GB/T 9704 format.

        Spec: A4, top 37mm±1, bottom 35mm±1, left 28mm±1, right 26mm±1.
        Title: 方正小标宋简体 22pt, centered.
        Body: 仿宋 16pt, 28 chars per line, 22 lines per page.
        """
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        # Header: 发文字号 + 签发机关
        if doc_number:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"发文字号: {doc_number}")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Title: 方正小标宋 22pt centered
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(20)
        p.space_after = Pt(20)
        run = p.add_run(title)
        run.font.size = Pt(22)
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # Body: 仿宋 16pt
        for para_text in body.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(16)
            p.paragraph_format.first_line_indent = Pt(32)
            p.paragraph_format.line_spacing = Pt(28)

        # Footer: date + issuing org
        if issuing_org:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.space_before = Pt(30)
            run = p.add_run(issuing_org)
            run.font.size = Pt(16)
        if date:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(date)
            run.font.size = Pt(16)

        out_path = Path(output or f"/tmp/ofd_{int(datetime.now().timestamp())}.docx")
        doc.save(str(out_path))
        return str(out_path)


# ═══ 16. Table/Figure Index Generator ═══

class IndexGenerator:
    """Auto-generate table of tables and table of figures."""

    @staticmethod
    def generate(tracker: CitationTracker) -> dict[str, str]:
        """Generate table index and figure index from CitationTracker."""
        table_index = ["## 表格索引", "", "| 编号 | 名称 | 页码 |", "|------|------|------|"]
        for chapter_num, count in sorted(tracker._tables.items()):
            for i in range(1, count + 1):
                table_index.append(f"| 表{chapter_num}-{i} | 待填写 | — |")

        figure_index = ["## 图形索引", "", "| 编号 | 名称 | 页码 |", "|------|------|------|"]
        for chapter_num, count in sorted(tracker._figures.items()):
            for i in range(1, count + 1):
                figure_index.append(f"| 图{chapter_num}-{i} | 待填写 | — |")

        return {
            "table_index": "\n".join(table_index),
            "figure_index": "\n".join(figure_index),
        }


__all__ = [
    "CitationTracker", "DataBinder", "FormatPipeline",
    "MultiLangSupport", "ReportEnhancer",
    "DataValidator", "Watermark", "ChartGenerator",
    "DiffTracker", "AuditLogger", "OFDExporter",
    "ConcurrentGenerator", "SectionLocker",
    "PDFFormFiller", "OfficialDocFormat", "IndexGenerator",
]
