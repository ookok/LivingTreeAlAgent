"""Tools — PDF parsing, translation, flowcharts, enhanced maps.

Unified utility tab combining document processing, language tools,
diagram generation, and map visualization.
"""

from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Optional

from textual import on, work
from textual.app import ComposeResult
from textual.containers import Horizontal, Vertical, VerticalScroll
from textual.screen import Screen
from textual.widgets import (
    Button, Input, Label, RichLog, Select, Static, TabbedContent, TabPane, TextArea,
)

from ..widgets.model_selector import ModelSelector
from ..widgets.pkg_manager_widget import PackageManagerWidget


class ToolsScreen(Screen):
    BINDINGS = [("escape", "app.pop_screen", "返回")]

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self._hub = None
        self._current_file: Optional[str] = None

    def set_hub(self, hub) -> None:
        self._hub = hub

    def compose(self) -> ComposeResult:
        yield Static("[dim]← 返回首页 (Esc)[/dim]", id="back-link")
        with TabbedContent(id="tools-tabs"):
            with TabPane("📄 PDF/Doc", id="pdf-tab"):
                yield Vertical(
                    Horizontal(
                        Input(placeholder="PDF or document path...", id="pdf-path"),
                        Button("📖 Extract", variant="primary", id="pdf-extract-btn"),
                        Button("📝 Summarize", variant="default", id="pdf-summarize-btn"),
                        id="pdf-toolbar",
                    ),
                    TextArea("", id="pdf-content", read_only=True),
                    RichLog(id="pdf-output", highlight=True, markup=True, wrap=True),
                )
            with TabPane("🌐 翻译", id="translate-tab"):
                yield Vertical(
                    Horizontal(
                        Select([("中→英","zh→en"),("英→中","en→zh"),("中→日","zh→ja"),
                                ("日→中","ja→zh"),("自动检测","auto")], value="zh→en", id="translate-lang"),
                        Button("🌐 Translate", variant="primary", id="translate-btn"),
                        id="translate-toolbar",
                    ),
                    TextArea("", id="translate-input"),
                    Label("[dim]Translation output:[/dim]", id="translate-label"),
                    TextArea("", id="translate-output", read_only=True),
                )
            with TabPane("📊 流程图", id="flowchart-tab"):
                yield Vertical(
                    Input(placeholder="Describe the diagram... e.g. 用户登录流程", id="flowchart-desc"),
                    Button("📊 Generate", variant="primary", id="flowchart-gen-btn"),
                    TextArea("", id="flowchart-output", read_only=True),
                )
            with TabPane("🗺️ 地图", id="map-tab"):
                yield Vertical(
                    Horizontal(
                        Input(placeholder="🔍 Location (city, address, coordinates)...", id="map-query"),
                        Button("🗺️ Search", variant="primary", id="map-search-btn"),
                        Select([("矢量","vec"),("卫星","img"),("地形","ter"),("标注","cva")],
                               value="vec", id="map-layer"),
                        Button("Beijing", variant="default", id="map-bj"),
                        Button("Shanghai", variant="default", id="map-sh"),
                        id="map-toolbar",
                    ),
                    RichLog(id="map-display", highlight=True, markup=True, wrap=True),
                    id="map-container",
                )

            with TabPane("🔬 模型", id="models-tab"):
                yield ModelSelector()

            with TabPane("📦 包管理", id="packages-tab"):
                yield PackageManagerWidget()

    def on_mount(self) -> None:
        hub = getattr(self.app, '_hub', None)
        if hub and hasattr(self, 'set_hub'):
            self.set_hub(hub)
        output = self.query_one("#pdf-output", RichLog)
        output.write("[bold]📄 Document Processor[/bold]")
        output.write("[dim]Extract text from PDF/DOCX, then summarize with AI.[/dim]")

    # ── PDF/Doc tab ──

    @work(exclusive=False)
    async def on_button_pressed(self, event: Button.Pressed) -> None:
        bid = event.button.id
        if bid == "pdf-extract-btn":
            await self._extract_pdf()
        elif bid == "pdf-summarize-btn":
            await self._summarize_doc()
        elif bid == "translate-btn":
            await self._translate()
        elif bid == "flowchart-gen-btn":
            await self._gen_flowchart()
        elif bid == "map-search-btn":
            await self._map_search()

    async def _extract_pdf(self) -> None:
        path = self.query_one("#pdf-path", Input).value.strip()
        output = self.query_one("#pdf-output", RichLog)
        editor = self.query_one("#pdf-content", TextArea)
        if not path:
            output.write("[yellow]Enter file path[/yellow]")
            return

        p = Path(path)
        if not p.exists():
            output.write(f"[red]File not found: {path}[/red]")
            return

        suffix = p.suffix.lower()
        text = ""
        output.clear()

        try:
            if suffix == ".pdf":
                try:
                    import PyPDF2
                    reader = PyPDF2.PdfReader(str(p))
                    pages = []
                    for i, page in enumerate(reader.pages[:20]):
                        t = page.extract_text() or ""
                        pages.append(f"--- Page {i+1} ---\n{t}")
                    text = "\n\n".join(pages)
                    output.write(f"[green]Extracted {len(reader.pages)} pages[/green]")
                except ImportError:
                    output.write("[yellow]PyPDF2 not installed (pip install PyPDF2)[/yellow]")
                    return
            elif suffix == ".docx":
                try:
                    from docx import Document
                    doc = Document(str(p))
                    text = "\n".join(p.text for p in doc.paragraphs)
                    output.write(f"[green]Extracted {len(doc.paragraphs)} paragraphs[/green]")
                except ImportError:
                    output.write("[yellow]python-docx not installed[/yellow]")
                    return
            elif suffix in (".md", ".txt", ".py", ".json", ".yaml", ".html", ".csv"):
                text = p.read_text(encoding="utf-8", errors="replace")
                output.write(f"[green]Loaded {len(text)} chars[/green]")
            else:
                try:
                    text = p.read_text(encoding="utf-8", errors="replace")
                    output.write(f"[green]Loaded as text: {len(text)} chars[/green]")
                except Exception:
                    output.write(f"[red]Cannot read {suffix} files[/red]")
                    return

            editor.text = text[:50000]
            self._current_file = path
        except Exception as e:
            output.write(f"[red]Error: {e}[/red]")

    @work(exclusive=False)
    async def _summarize_doc(self) -> None:
        editor = self.query_one("#pdf-content", TextArea)
        output = self.query_one("#pdf-output", RichLog)
        text = editor.text.strip()
        if not text:
            output.write("[yellow]No content to summarize[/yellow]")
            return
        if not self._hub:
            return

        api_key = self._hub.config.model.deepseek_api_key
        if not api_key:
            output.write("[yellow]API key not configured[/yellow]")
            return

        output.write("[bold]📝 Summarizing...[/bold]")
        try:
            import aiohttp
            headers = {"Content-Type":"application/json","Authorization":f"Bearer {api_key}"}
            payload = {
                "model": self._hub.config.model.flash_model,
                "messages": [{"role":"system","content":"Summarize in Chinese. Key points + main content + conclusion."},
                             {"role":"user","content": f"Summarize:\n\n{text[:8000]}"}],
                "temperature":0.2,"max_tokens":1024,
            }
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self._hub.config.model.deepseek_base_url}/v1/chat/completions",
                    headers=headers,json=payload,timeout=aiohttp.ClientTimeout(total=60),
                ) as resp:
                    data = await resp.json()
                    summary = data["choices"][0]["message"]["content"]
            output.write(f"\n[bold green]Summary:[/bold green]\n{summary}")
        except Exception as e:
            output.write(f"[red]{e}[/red]")

    # ── Translate tab ──

    @work(exclusive=False)
    async def _translate(self) -> None:
        inp = self.query_one("#translate-input", TextArea)
        out = self.query_one("#translate-output", TextArea)
        lang = str(self.query_one("#translate-lang", Select).value or "zh→en")
        text = inp.text.strip()
        if not text:
            return
        if not self._hub:
            return

        api_key = self._hub.config.model.deepseek_api_key
        if not api_key:
            out.text = "API key not configured"
            return

        lang_map = {"zh→en":"Chinese to English","en→zh":"English to Chinese",
                    "zh→ja":"Chinese to Japanese","ja→zh":"Japanese to Chinese",
                    "auto":"the target language"}
        direction = lang_map.get(lang, "Chinese to English")

        try:
            import aiohttp
            headers = {"Content-Type":"application/json","Authorization":f"Bearer {api_key}"}
            payload = {
                "model": self._hub.config.model.flash_model,
                "messages": [{"role":"system","content":f"Translate to {direction}. Output ONLY the translation, no explanations."},
                             {"role":"user","content": text}],
                "temperature":0.1,"max_tokens":4096,
            }
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self._hub.config.model.deepseek_base_url}/v1/chat/completions",
                    headers=headers,json=payload,timeout=aiohttp.ClientTimeout(total=60),
                ) as resp:
                    data = await resp.json()
                    result = data["choices"][0]["message"]["content"]
            out.text = result
        except Exception as e:
            out.text = f"Error: {e}"

    # ── Flowchart tab ──

    @work(exclusive=False)
    async def _gen_flowchart(self) -> None:
        desc = self.query_one("#flowchart-desc", Input).value.strip()
        out = self.query_one("#flowchart-output", TextArea)
        if not desc:
            return
        if not self._hub:
            return

        api_key = self._hub.config.model.deepseek_api_key
        if not api_key:
            out.text = "API key not configured"
            return

        try:
            import aiohttp
            headers = {"Content-Type":"application/json","Authorization":f"Bearer {api_key}"}
            payload = {
                "model": self._hub.config.model.flash_model,
                "messages": [{"role":"system","content":(
                    "Generate a text flowchart using ASCII art (box-drawing characters).\n"
                    "Use ┌─┐│└─┘├┤┬┴┼ for boxes and → for arrows.\n"
                    "Output ONLY the diagram, no explanations.\n"
                    "Make it visually clear and well-structured."
                )},{"role":"user","content": f"Generate an ASCII flowchart for: {desc}"}],
                "temperature":0.3,"max_tokens":2048,
            }
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self._hub.config.model.deepseek_base_url}/v1/chat/completions",
                    headers=headers,json=payload,timeout=aiohttp.ClientTimeout(total=60),
                ) as resp:
                    data = await resp.json()
                    diagram = data["choices"][0]["message"]["content"]
            out.text = diagram
        except Exception as e:
            out.text = f"Error: {e}"

    # ── Map tab ──

    @work(exclusive=False)
    async def _map_search(self) -> None:
        query = self.query_one("#map-query", Input).value.strip()
        layer = str(self.query_one("#map-layer", Select).value or "vec")
        display = self.query_one("#map-display", RichLog)
        display.clear()

        if not query:
            # Default: Beijing
            query = "北京"

        display.write(f"[bold]🗺️ 天地图 · {query}[/bold]")

        try:
            from ..capability.tianditu import geocode, static_map, map_url, reverse_geocode

            # Try geocoding
            result = geocode(query)
            if result.get("found"):
                lat, lon = float(result["lat"]), float(result["lon"])
                display.write(f"  📍 {result.get('name',query)} | {result.get('admin','')}")
                display.write(f"  🔗 {map_url(lon, lat)}")

                # Fetch tile info
                from ..capability.tianditu import fetch_tile, lonlat_to_tile
                cx, cy, px, py = lonlat_to_tile(lon, lat, 12)
                display.write(f"  🧩 Tile: ({cx},{cy}) zoom=12 pixel=({px},{py})")

                # Try fetch a tile to verify connectivity
                tile = fetch_tile(layer, cx, cy, 12)
                if tile:
                    display.write(f"  ✅ Tile fetched: {len(tile)} bytes ({layer})")
                    # Show ASCII block art
                    map_art = static_map(lon, lat, zoom=12, layer=layer)
                    display.write(f"\n{map_art}")
                else:
                    display.write(f"  ⚠️ Tile fetch failed — check network/key")

                # Try reverse geocode for detail
                rev = reverse_geocode(lat, lon)
                if rev.get("found"):
                    display.write(f"  🏠 {rev.get('address','')}")
            else:
                # Direct coordinate parse?
                try:
                    parts = query.split(",")
                    if len(parts) == 2:
                        lat, lon = float(parts[0]), float(parts[1])
                        display.write(f"  📍 Coordinates: ({lat:.4f}, {lon:.4f})")
                        display.write(f"  🔗 {map_url(lon, lat)}")
                    else:
                        display.write(f"  ❌ Location not found: {query}")
                        display.write(f"  [dim]Try: 北京, 上海, 39.9,116.4[/dim]")
                except ValueError:
                    display.write(f"  ❌ Location not found: {query}")

        except ImportError:
            display.write("[yellow]Tianditu module not available[/yellow]")
        except Exception as e:
            display.write(f"[red]Error: {e}[/red]")

    async def refresh(self, **kwargs) -> None:
        pass
