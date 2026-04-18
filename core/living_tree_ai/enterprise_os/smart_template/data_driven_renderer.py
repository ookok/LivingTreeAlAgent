"""
数据驱动渲染器 (Data-Driven Renderer)

将新项目数据注入模板配置，动态组装生成新的 Word/PDF 文档。

核心流程：
1. 加载 template_config.json
2. 注入 project_data
3. 渲染生成完整文档
4. 支持多种输出格式（DOCX/HTML/Markdown）

解决的问题：
- 格式100%一致：骨架来自定稿文档
- 嵌套表格自动处理：AI 理解表格语义
- 变量自动替换：data_key → 实际值
"""

import json
import logging
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Any, Callable
from enum import Enum
from pathlib import Path

logger = logging.getLogger(__name__)


# ==================== 数据模型 ====================

class OutputFormat(Enum):
    """输出格式"""
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "markdown"
    JSON = "json"  # 仅结构化输出


@dataclass
class RenderContext:
    """渲染上下文"""
    project_id: str = ""
    project_name: str = ""
    generated_at: str = ""
    generated_by: str = "Hermes AI"
    version: str = "1.0.0"
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class RenderResult:
    """渲染结果"""
    success: bool
    output_path: Optional[str] = None
    content: Optional[str] = None  # 非文件输出的内容
    format: str = ""
    message: str = ""
    stats: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "output_path": self.output_path,
            "content": self.content,
            "format": self.format,
            "message": self.message,
            "stats": self.stats
        }


# ==================== 数据驱动渲染器 ====================

class DataDrivenRenderer:
    """
    数据驱动渲染器

    将结构化数据注入模板，生成目标文档。
    支持多种输出格式。
    """

    def __init__(self):
        self._custom_formatters: Dict[str, Callable] = {}

    def register_formatter(self, data_key: str, formatter: Callable):
        """
        注册自定义格式化函数

        Args:
            data_key: 数据键名
            formatter: 格式化函数，签名为 (value: Any, context: RenderContext) -> str
        """
        self._custom_formatters[data_key] = formatter

    async def render(
        self,
        template_config: Dict[str, Any],
        project_data: Dict[str, Any],
        output_path: Optional[str] = None,
        output_format: OutputFormat = OutputFormat.DOCX,
        context: Optional[RenderContext] = None
    ) -> RenderResult:
        """
        渲染文档

        Args:
            template_config: 模板配置（dict 或 JSON 路径）
            project_data: 项目数据
            output_path: 输出文件路径
            output_format: 输出格式
            context: 渲染上下文

        Returns:
            RenderResult: 渲染结果
        """
        # 加载模板配置
        if isinstance(template_config, str):
            if Path(template_config).exists():
                with open(template_config, 'r', encoding='utf-8') as f:
                    template_config = json.load(f)
            else:
                template_config = json.loads(template_config)

        if context is None:
            context = RenderContext()

        try:
            # 根据输出格式选择渲染方法
            if output_format == OutputFormat.DOCX:
                return await self._render_docx(template_config, project_data, output_path, context)
            elif output_format == OutputFormat.HTML:
                return await self._render_html(template_config, project_data, output_path, context)
            elif output_format == OutputFormat.MARKDOWN:
                return await self._render_markdown(template_config, project_data, output_path, context)
            elif output_format == OutputFormat.JSON:
                return await self._render_json(template_config, project_data, context)
            else:
                return RenderResult(
                    success=False,
                    message=f"不支持的输出格式: {output_format}"
                )

        except Exception as e:
            logger.error(f"渲染失败: {e}")
            return RenderResult(
                success=False,
                message=f"渲染失败: {str(e)}"
            )

    async def _render_docx(
        self,
        template_config: Dict[str, Any],
        project_data: Dict[str, Any],
        output_path: Optional[str],
        context: RenderContext
    ) -> RenderResult:
        """渲染为 DOCX 格式"""
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            return RenderResult(
                success=False,
                message="需要安装 python-docx 库: pip install python-docx"
            )

        doc = Document()

        # 设置默认字体（支持中文）
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        blocks = template_config.get('blocks', [])
        tables = template_config.get('tables', [])

        # 构建 data_key → data 的映射
        data_map = self._build_data_map(tables, project_data)

        stats = {"blocks_rendered": 0, "tables_rendered": 0, "variables_replaced": 0}

        # 渲染块
        for block in blocks:
            block_type = block.get('block_type', 'static')
            content = block.get('content', '')

            if block_type == 'static':
                # 固定文本
                self._add_paragraph(doc, content, block)
                stats["blocks_rendered"] += 1

            elif block_type == 'dynamic':
                # 可变数据
                data_key = block.get('data_key', '')
                value = self._resolve_value(data_key, project_data, data_map, context)

                if value is not None:
                    rendered_content = self._apply_custom_formatter(data_key, value, context)
                    self._add_paragraph(doc, rendered_content, block)
                    stats["variables_replaced"] += 1
                else:
                    # 使用原始内容作为占位符提示
                    self._add_paragraph(doc, f"[{data_key}: 待填充]", block)

        # 渲染表格
        for table in tables:
            data_key = table.get('data_key', '')
            table_data = self._resolve_table_data(data_key, project_data, data_map)

            if table_data:
                self._render_table(doc, table, table_data)
                stats["tables_rendered"] += 1

        # 保存文件
        if output_path:
            doc.save(output_path)
            return RenderResult(
                success=True,
                output_path=output_path,
                format="docx",
                message="DOCX 文档生成成功",
                stats=stats
            )
        else:
            return RenderResult(
                success=True,
                content="DOCX 需要文件输出，请提供 output_path",
                format="docx",
                message="请提供 output_path",
                stats=stats
            )

    async def _render_html(
        self,
        template_config: Dict[str, Any],
        project_data: Dict[str, Any],
        output_path: Optional[str],
        context: RenderContext
    ) -> RenderResult:
        """渲染为 HTML 格式"""
        blocks = template_config.get('blocks', [])
        tables = template_config.get('tables', [])

        data_map = self._build_data_map(tables, project_data)
        stats = {"blocks_rendered": 0, "tables_rendered": 0, "variables_replaced": 0}

        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            "<meta charset='utf-8'>",
            "<title>报告文档</title>",
            "<style>",
            "body { font-family: 'Microsoft YaHei', sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }",
            "h1, h2, h3 { color: #333; }",
            "table { border-collapse: collapse; width: 100%; margin: 20px 0; }",
            "th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }",
            "th { background-color: #f5f5f5; }",
            ".placeholder { color: #999; font-style: italic; }",
            "</style>",
            "</head>",
            "<body>"
        ]

        for block in blocks:
            block_type = block.get('block_type', 'static')
            content = block.get('content', '')
            level = block.get('level', 0)

            if block_type == 'static':
                if level == 1:
                    html_parts.append(f"<h1>{content}</h1>")
                elif level == 2:
                    html_parts.append(f"<h2>{content}</h2>")
                elif level == 3:
                    html_parts.append(f"<h3>{content}</h3>")
                else:
                    html_parts.append(f"<p>{content}</p>")
                stats["blocks_rendered"] += 1

            elif block_type == 'dynamic':
                data_key = block.get('data_key', '')
                value = self._resolve_value(data_key, project_data, data_map, context)

                if value is not None:
                    rendered = self._apply_custom_formatter(data_key, value, context)
                    html_parts.append(f"<p><strong>{rendered}</strong></p>")
                    stats["variables_replaced"] += 1
                else:
                    html_parts.append(f"<p class='placeholder'>[{data_key}: 待填充]</p>")

        # 渲染表格
        for table in tables:
            data_key = table.get('data_key', '')
            table_data = self._resolve_table_data(data_key, project_data, data_map)

            if table_data:
                html_parts.append(self._render_table_html(table, table_data))
                stats["tables_rendered"] += 1

        html_parts.extend([
            "</body>",
            "</html>"
        ])

        html_content = "\n".join(html_parts)

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return RenderResult(
                success=True,
                output_path=output_path,
                content=html_content,
                format="html",
                message="HTML 文档生成成功",
                stats=stats
            )
        else:
            return RenderResult(
                success=True,
                content=html_content,
                format="html",
                message="HTML 渲染成功",
                stats=stats
            )

    async def _render_markdown(
        self,
        template_config: Dict[str, Any],
        project_data: Dict[str, Any],
        output_path: Optional[str],
        context: RenderContext
    ) -> RenderResult:
        """渲染为 Markdown 格式"""
        blocks = template_config.get('blocks', [])
        tables = template_config.get('tables', [])

        data_map = self._build_data_map(tables, project_data)
        stats = {"blocks_rendered": 0, "tables_rendered": 0, "variables_replaced": 0}

        md_parts = []
        current_heading_level = 0

        for block in blocks:
            block_type = block.get('block_type', 'static')
            content = block.get('content', '')
            level = block.get('level', 0)

            if block_type == 'static':
                if level > 0:
                    md_parts.append(f"{'#' * level} {content}")
                    current_heading_level = level
                else:
                    md_parts.append(content)
                stats["blocks_rendered"] += 1

            elif block_type == 'dynamic':
                data_key = block.get('data_key', '')
                value = self._resolve_value(data_key, project_data, data_map, context)

                if value is not None:
                    rendered = self._apply_custom_formatter(data_key, value, context)
                    md_parts.append(rendered)
                    stats["variables_replaced"] += 1
                else:
                    md_parts.append(f"**[{data_key}: 待填充]**")

        # 渲染表格
        for table in tables:
            data_key = table.get('data_key', '')
            table_data = self._resolve_table_data(data_key, project_data, data_map)

            if table_data:
                md_parts.append(self._render_table_markdown(table, table_data))
                stats["tables_rendered"] += 1

        md_content = "\n\n".join(md_parts)

        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            return RenderResult(
                success=True,
                output_path=output_path,
                content=md_content,
                format="markdown",
                message="Markdown 文档生成成功",
                stats=stats
            )
        else:
            return RenderResult(
                success=True,
                content=md_content,
                format="markdown",
                message="Markdown 渲染成功",
                stats=stats
            )

    async def _render_json(
        self,
        template_config: Dict[str, Any],
        project_data: Dict[str, Any],
        context: RenderContext
    ) -> RenderResult:
        """渲染为结构化 JSON"""
        blocks = template_config.get('blocks', [])
        tables = template_config.get('tables', [])

        data_map = self._build_data_map(tables, project_data)

        result_data = {
            "metadata": {
                "template_id": template_config.get('template_id'),
                "template_name": template_config.get('template_name'),
                "project_id": context.project_id,
                "generated_at": context.generated_at,
                "generated_by": context.generated_by
            },
            "content": [],
            "tables": []
        }

        # 渲染块
        for block in blocks:
            block_type = block.get('block_type', 'static')
            content = block.get('content', '')

            if block_type == 'static':
                result_data["content"].append({
                    "type": "static",
                    "text": content
                })
            elif block_type == 'dynamic':
                data_key = block.get('data_key', '')
                value = self._resolve_value(data_key, project_data, data_map, context)
                result_data["content"].append({
                    "type": "dynamic",
                    "data_key": data_key,
                    "value": value
                })

        # 渲染表格
        for table in tables:
            data_key = table.get('data_key', '')
            table_data = self._resolve_table_data(data_key, project_data, data_map)
            result_data["tables"].append({
                "data_key": data_key,
                "data": table_data
            })

        json_content = json.dumps(result_data, ensure_ascii=False, indent=2)

        return RenderResult(
            success=True,
            content=json_content,
            format="json",
            message="JSON 结构化输出成功",
            stats={"blocks": len(result_data["content"]), "tables": len(result_data["tables"])}
        )

    def _add_paragraph(self, doc, content: str, block: Dict[str, Any]):
        """添加段落到 DOCX"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        style = block.get('style', {})
        level = block.get('level', 0)

        if level > 0:
            # 标题
            para = doc.add_heading(content, level=min(level, 9))
        else:
            para = doc.add_paragraph(content)

        # 设置对齐
        alignment = style.get('alignment', 'left')
        if alignment == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 设置字体样式
        for run in para.runs:
            if style.get('bold'):
                run.bold = True
            if style.get('italic'):
                run.font.italic = True

        return para

    def _build_data_map(
        self,
        tables: List[Dict[str, Any]],
        project_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """构建数据映射"""
        data_map = {}

        # 表格数据
        for table in tables:
            data_key = table.get('data_key', '')
            if data_key and data_key in project_data:
                data_map[data_key] = project_data[data_key]

        # 扁平化处理
        self._flatten_data(project_data, "", data_map)

        return data_map

    def _flatten_data(
        self,
        data: Any,
        prefix: str,
        result: Dict[str, Any]
    ):
        """扁平化嵌套数据"""
        if isinstance(data, dict):
            for key, value in data.items():
                new_key = f"{prefix}.{key}" if prefix else key
                if isinstance(value, (dict, list)):
                    self._flatten_data(value, new_key, result)
                else:
                    result[new_key] = value
        elif isinstance(data, list):
            for idx, item in enumerate(data):
                new_key = f"{prefix}[{idx}]"
                if isinstance(item, (dict, list)):
                    self._flatten_data(item, new_key, result)
                else:
                    result[new_key] = item

    def _resolve_value(
        self,
        data_key: str,
        project_data: Dict[str, Any],
        data_map: Dict[str, Any],
        context: RenderContext
    ) -> Optional[Any]:
        """解析变量值"""
        # 直接匹配
        if data_key in project_data:
            return project_data[data_key]

        # 嵌套匹配
        if data_key in data_map:
            return data_map[data_key]

        # 尝试点号路径
        keys = data_key.split('.')
        value = project_data
        for k in keys:
            if isinstance(value, dict) and k in value:
                value = value[k]
            else:
                return None
        return value

    def _resolve_table_data(
        self,
        data_key: str,
        project_data: Dict[str, Any],
        data_map: Dict[str, Any]
    ) -> Optional[List[Dict[str, Any]]]:
        """解析表格数据"""
        if data_key in project_data:
            data = project_data[data_key]
            if isinstance(data, list):
                return data
        return None

    def _apply_custom_formatter(
        self,
        data_key: str,
        value: Any,
        context: RenderContext
    ) -> str:
        """应用自定义格式化"""
        if data_key in self._custom_formatters:
            formatter = self._custom_formatters[data_key]
            return formatter(value, context)
        return str(value)

    def _render_table(
        self,
        doc,
        table_config: Dict[str, Any],
        table_data: List[Dict[str, Any]]
    ):
        """渲染表格到 DOCX"""
        from docx.shared import Pt

        headers = table_config.get('headers', [])
        if not headers:
            return

        # 创建表格
        num_rows = len(table_data) + 1
        num_cols = len(headers)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        # 填充表头
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            if idx < num_cols:
                header_cells[idx].text = header.get('name', header.get('key', ''))

        # 填充数据行
        for row_idx, row_data in enumerate(table_data):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, header in enumerate(headers):
                if col_idx < num_cols:
                    key = header.get('key', '')
                    value = row_data.get(key, '')
                    row_cells[col_idx].text = str(value) if value is not None else ''

    def _render_table_html(
        self,
        table_config: Dict[str, Any],
        table_data: List[Dict[str, Any]]
    ) -> str:
        """渲染表格为 HTML"""
        headers = table_config.get('headers', [])
        if not headers:
            return ""

        html_parts = ["<table>"]

        # 表头
        html_parts.append("<thead><tr>")
        for header in headers:
            html_parts.append(f"<th>{header.get('name', header.get('key', ''))}</th>")
        html_parts.append("</tr></thead>")

        # 数据行
        html_parts.append("<tbody>")
        for row_data in table_data:
            html_parts.append("<tr>")
            for header in headers:
                key = header.get('key', '')
                value = row_data.get(key, '')
                html_parts.append(f"<td>{value if value is not None else ''}</td>")
            html_parts.append("</tr>")
        html_parts.append("</tbody>")

        html_parts.append("</table>")
        return "\n".join(html_parts)

    def _render_table_markdown(
        self,
        table_config: Dict[str, Any],
        table_data: List[Dict[str, Any]]
    ) -> str:
        """渲染表格为 Markdown"""
        headers = table_config.get('headers', [])
        if not headers:
            return ""

        md_parts = []

        # 表头
        header_row = "| " + " | ".join(h.get('name', h.get('key', '')) for h in headers) + " |"
        separator = "| " + " | ".join("---" for _ in headers) + " |"

        md_parts.append(header_row)
        md_parts.append(separator)

        # 数据行
        for row_data in table_data:
            data_row = "| " + " | ".join(str(row_data.get(h.get('key', ''), '')) for h in headers) + " |"
            md_parts.append(data_row)

        return "\n".join(md_parts)


# ==================== 便捷函数 ====================

_renderer: Optional[DataDrivenRenderer] = None


def get_data_driven_renderer() -> DataDrivenRenderer:
    """获取渲染器单例"""
    global _renderer
    if _renderer is None:
        _renderer = DataDrivenRenderer()
    return _renderer


async def render_document(
    template_config: Dict[str, Any],
    project_data: Dict[str, Any],
    output_path: Optional[str] = None,
    output_format: OutputFormat = OutputFormat.DOCX,
    context: Optional[RenderContext] = None
) -> RenderResult:
    """渲染文档"""
    renderer = get_data_driven_renderer()
    return await renderer.render(template_config, project_data, output_path, output_format, context)